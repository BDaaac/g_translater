
import sys
import subprocess
import importlib.util

def ensure_package(package_name, import_name=None, extras=None):
    """–ü—Ä–æ–≤–µ—Ä—è–µ—Ç –Ω–∞–ª–∏—á–∏–µ –ø–∞–∫–µ—Ç–∞ –∏ —É—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ—Ç –µ–≥–æ –ø—Ä–∏ –Ω–µ–æ–±—Ö–æ–¥–∏–º–æ—Å—Ç–∏."""
    import_name = import_name or package_name
    if importlib.util.find_spec(import_name) is None:
        print(f"–ü–∞–∫–µ—Ç '{package_name}' –Ω–µ –Ω–∞–π–¥–µ–Ω. –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞—é...")
        try:
            install_target = package_name + extras if extras else package_name
            subprocess.check_call([sys.executable, "-m", "pip", "install", install_target])
        except Exception as e:
            print(f"–ù–µ —É–¥–∞–ª–æ—Å—å —É—Å—Ç–∞–Ω–æ–≤–∏—Ç—å –ø–∞–∫–µ—Ç '{package_name}': {e}")
            return False
    return True

ensure_package("bs4", "bs4")
ensure_package("PySocks", "socks") # For SOCKS proxy support
ensure_package("PyQt6", "PyQt6")
ensure_package("google-generativeai", "google")
ensure_package("python-docx", "docx")
ensure_package("lxml", "lxml")
ensure_package("ebooklib", "ebooklib")
ensure_package("Pillow", "PIL")

DOCX_AVAILABLE = importlib.util.find_spec("docx") is not None
LXML_AVAILABLE = importlib.util.find_spec("lxml") is not None
EBOOKLIB_AVAILABLE = importlib.util.find_spec("ebooklib") is not None
PILLOW_AVAILABLE = importlib.util.find_spec("PIL") is not None
BS4_AVAILABLE = importlib.util.find_spec("bs4") is not None

import os
import sys
import glob
import argparse
import traceback
import xml.etree.ElementTree as ET
import time
import math
from pathlib import Path
import re
import uuid
import tempfile
from io import BytesIO
import zipfile
import configparser
import base64
import imghdr
import html
import json
import threading  # <<< –ù–û–í–ò–ù–ö–ê: –¥–ª—è ApiKeyManager
import shutil  # <<< –ù–û–í–ò–ù–ö–ê: –¥–ª—è TranslatedChaptersManagerDialog
from urllib.parse import urlparse, urljoin, unquote
import warnings

from bs4 import BeautifulSoup, Tag, NavigableString, XMLParsedAsHTMLWarning
warnings.filterwarnings("ignore", category=XMLParsedAsHTMLWarning)

from PyQt6 import QtWidgets, QtCore, QtGui
from PyQt6.QtWidgets import (
    QApplication, QDialog, QVBoxLayout, QListWidget, QPushButton,
    QDialogButtonBox, QLabel, QWidget, QLineEdit, QComboBox, QSpinBox,
    QCheckBox, QPlainTextEdit, QDoubleSpinBox, QProgressBar, QTextEdit,
    QGridLayout, QGroupBox, QHBoxLayout, QMessageBox, QFileDialog, QScrollArea,
    QListWidgetItem, QTableWidget, QTableWidgetItem, QFormLayout  # <<< –ù–û–í–ò–ù–ö–ê: –¥–ª—è –Ω–æ–≤–æ–≥–æ –¥–∏–∞–ª–æ–≥–∞
)
from PyQt6.QtCore import QStandardPaths, Qt

# --- –ù–ê–ß–ê–õ–û –ë–õ–û–ö–ê –î–õ–Ø –ó–ê–ú–ï–ù–´ (—ç—Ç–æ—Ç –±–ª–æ–∫ –ø—Ä–∞–≤–∏–ª—å–Ω—ã–π, –º–µ–Ω—è—Ç—å –Ω–µ –Ω—É–∂–Ω–æ) ---
from google.api_core import exceptions as google_exceptions
from google import generativeai as genai
# –ü—Ä–∞–≤–∏–ª—å–Ω—ã–π –∏–º–ø–æ—Ä—Ç –º–æ–¥—É–ª—è 'types', –∫–æ—Ç–æ—Ä—ã–π —Å–æ–¥–µ—Ä–∂–∏—Ç –∏ —Ç–∏–ø—ã, –∏ –∏—Å–∫–ª—é—á–µ–Ω–∏—è –∫–æ–Ω—Ç–µ–Ω—Ç–∞
import google.generativeai.types as genai_types
# --- –ö–û–ù–ï–¶ –ë–õ–û–ö–ê –î–õ–Ø –ó–ê–ú–ï–ù–´ ---

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn

from lxml import etree
from ebooklib import epub
from PIL import Image

from functools import partial
from concurrent.futures import ThreadPoolExecutor, as_completed, Future, wait, CancelledError

MODELS = {




    "Gemini 2.5 Pro": { # From user list / original code
        "id": "models/gemini-2.5-pro",
        "rpm": 5, # Moderate RPM
        "needs_chunking": True, # Assume requires chunking
        "post_request_delay": 2 # –£–º–µ–Ω—å—à–µ–Ω–Ω–∞—è –∑–∞–¥–µ—Ä–∂–∫–∞ –¥–ª—è –±—ã—Å—Ç—Ä–æ–≥–æ –ø–µ—Ä–µ–≤–æ–¥–∞
    },

    "Gemini 2.5 Flash": { # From user list / original code
        "id": "models/gemini-2.5-flash",
        "rpm": 10, # Moderate RPM
        "needs_chunking": True, # Assume requires chunking
        "post_request_delay": 2 # –£–º–µ–Ω—å—à–µ–Ω–Ω–∞—è –∑–∞–¥–µ—Ä–∂–∫–∞ –¥–ª—è –±—ã—Å—Ç—Ä–æ–≥–æ –ø–µ—Ä–µ–≤–æ–¥–∞
    },

    "Gemini 2.5 Flash-Lite Preview": { # From user list / original code
        "id": "models/gemini-2.5-flash-lite-preview-06-17",
        "rpm": 15, # Moderate RPM
        "needs_chunking": True, # Assume requires chunking
        "post_request_delay": 2 # –£–º–µ–Ω—å—à–µ–Ω–Ω–∞—è –∑–∞–¥–µ—Ä–∂–∫–∞ –¥–ª—è –±—ã—Å—Ç—Ä–æ–≥–æ –ø–µ—Ä–µ–≤–æ–¥–∞
    },

    "Gemini 2.5 Pro Experimental 03-25": { # From user list / original code
        "id": "models/gemini-2.5-pro-preview-03-25",
        "rpm": 10, # Moderate RPM
        "needs_chunking": True, # Assume requires chunking
        "post_request_delay": 2 # –£–º–µ–Ω—å—à–µ–Ω–Ω–∞—è –∑–∞–¥–µ—Ä–∂–∫–∞ –¥–ª—è –±—ã—Å—Ç—Ä–æ–≥–æ –ø–µ—Ä–µ–≤–æ–¥–∞
    },


    "Gemini 2.0 Flash": { # From user list / original code
        "id": "models/gemini-2.0-flash",
        "rpm": 15, # Higher RPM for Flash
        "needs_chunking": True, # Requires chunking for large inputs
        "post_request_delay": 2 # –£–º–µ–Ω—å—à–µ–Ω–Ω–∞—è –∑–∞–¥–µ—Ä–∂–∫–∞ –¥–ª—è –±—ã—Å—Ç—Ä–æ–≥–æ –ø–µ—Ä–µ–≤–æ–¥–∞
    },
    "Gemini 2.0 Flash Experimental": { # From user list / original code
        "id": "models/gemini-2.0-flash-exp",
        "rpm": 10, # Higher RPM for Flash
        "needs_chunking": True, # Requires chunking for large inputs
        "post_request_delay": 2 # –£–º–µ–Ω—å—à–µ–Ω–Ω–∞—è –∑–∞–¥–µ—Ä–∂–∫–∞ –¥–ª—è –±—ã—Å—Ç—Ä–æ–≥–æ –ø–µ—Ä–µ–≤–æ–¥–∞
    },
    "Gemini 2.0 Flash-Lite": { # From user list
        "id": "models/gemini-2.0-flash-lite",
        "rpm": 30, # Guess: Higher than standard Flash
        "needs_chunking": True, # Assume needs chunking like other Flash
        "post_request_delay": 2 # –£–º–µ–Ω—å—à–µ–Ω–Ω–∞—è –∑–∞–¥–µ—Ä–∂–∫–∞ –¥–ª—è –±—ã—Å—Ç—Ä–æ–≥–æ –ø–µ—Ä–µ–≤–æ–¥–∞
    },

    "Gemini 1.5 Flash": { # From user list (using recommended 'latest' tag)
        "id": "models/gemini-1.5-flash-latest",
        "rpm": 20, # Guess: Higher RPM for Flash models
        "needs_chunking": True, # Assume needs chunking
        "post_request_delay": 2 # –£–º–µ–Ω—å—à–µ–Ω–Ω–∞—è –∑–∞–¥–µ—Ä–∂–∫–∞ –¥–ª—è –±—ã—Å—Ç—Ä–æ–≥–æ –ø–µ—Ä–µ–≤–æ–¥–∞
    },
    "gemma-3-27b-it": { # From user list / original code
        "id": "models/gemma-3-27b-it",
        "rpm": 30, # Moderate RPM
        "needs_chunking": True, # Assume requires chunking
        "post_request_delay": 2 # –£–º–µ–Ω—å—à–µ–Ω–Ω–∞—è –∑–∞–¥–µ—Ä–∂–∫–∞ –¥–ª—è –±—ã—Å—Ç—Ä–æ–≥–æ –ø–µ—Ä–µ–≤–æ–¥–∞
    },


}

DEFAULT_MODEL_NAME = "Gemini 2.5 Flash Preview" if "Gemini 2.0 Flash" in MODELS else list(MODELS.keys())[0]

MAX_RETRIES = 3
RETRY_DELAY_SECONDS = 25
API_TIMEOUT_SECONDS = 600 # 10 –º–∏–Ω—É—Ç

DEFAULT_CHARACTER_LIMIT_FOR_CHUNK = 900_000 # Default limit (can be adjusted in GUI)
DEFAULT_CHUNK_SEARCH_WINDOW = 500 # Default window (can be adjusted in GUI)
MIN_CHUNK_SIZE = 500 # Minimum size to avoid tiny chunks
CHUNK_HTML_SOURCE = True # Keep False: HTML chunking with embedded images is complex and disabled by default

SETTINGS_FILE = 'translator_settings.ini'

OUTPUT_FORMATS = {
    "–¢–µ–∫—Å—Ç–æ–≤—ã–π —Ñ–∞–π–ª (.txt)": "txt",
    "–î–æ–∫—É–º–µ–Ω—Ç Word (.docx)": "docx",
    "Markdown (.md)": "md",
    "EPUB (.epub)": "epub", # Triggers EPUB rebuild logic if input is also EPUB
    "FictionBook2 (.fb2)": "fb2",
    "HTML (.html)": "html",
}
DEFAULT_OUTPUT_FORMAT_DISPLAY = "–¢–µ–∫—Å—Ç–æ–≤—ã–π —Ñ–∞–π–ª (.txt)" # Default display name for format dropdown

IMAGE_PLACEHOLDER_PREFIX = "img_placeholder_"
def create_image_placeholder(img_uuid):
    return f"<||{IMAGE_PLACEHOLDER_PREFIX}{img_uuid}||>"

def find_image_placeholders(text):
    pattern = re.compile(r"<\|\|(" + IMAGE_PLACEHOLDER_PREFIX + r"([a-f0-9]{32}))\|\|>")
    return [(match.group(0), match.group(2)) for match in pattern.finditer(text)]

TRANSLATED_SUFFIX = "_translated"

def add_translated_suffix(filename):
    """Adds _translated before the file extension, handling multiple suffixes."""
    if not filename: return filename
    path = Path(filename)

    suffixes = "".join(path.suffixes)
    if not suffixes: # Handle case with no extension (e.g., just "myfile")
        stem = path.name

        return str(path.parent / f"{stem}{TRANSLATED_SUFFIX}")
    else:

        stem = path.name.replace(suffixes, "")

        return str(path.parent / f"{stem}{TRANSLATED_SUFFIX}{suffixes}")

def format_size(size_bytes):
   """Converts bytes to a human-readable format (KB, MB, GB)."""
   if size_bytes == 0: return "0 B"
   size_name = ("B", "KB", "MB", "GB", "TB")
   i = int(math.floor(math.log(size_bytes, 1024))) if size_bytes > 0 else 0
   i = min(i, len(size_name) - 1)
   p = math.pow(1024, i)
   s = round(size_bytes / p, 2)
   return f"{s} {size_name[i]}"




def split_text_into_chunks(text, limit_chars, search_window, min_chunk_size):
    """Splits text into chunks, respecting paragraphs and sentences where possible."""
    chunks = []
    start_index = 0
    text_len = len(text)
    target_size = max(min_chunk_size, limit_chars - search_window // 2)

    while start_index < text_len:
        if text_len - start_index <= limit_chars:
            chunks.append(text[start_index:])
            break

        ideal_end_index = min(start_index + target_size, text_len)
        search_start = max(start_index + min_chunk_size, ideal_end_index - search_window)
        search_end = min(ideal_end_index + search_window, text_len)
        split_index = -1

        potential_splits = []

        search_slice = text[search_start:search_end]
        if search_slice:
            for match in re.finditer(r'\n\n', search_slice):
                 potential_splits.append((abs((search_start + match.end()) - ideal_end_index), search_start + match.end(), 1))
            for match in re.finditer(r"[.!?]\s+", search_slice):
                 potential_splits.append((abs((search_start + match.end()) - ideal_end_index), search_start + match.end(), 2))
            for match in re.finditer(r'\n', search_slice):
                  potential_splits.append((abs((search_start + match.end()) - ideal_end_index), search_start + match.end(), 3))

            for match in re.finditer(r' ', search_slice):
                current_split_pos = search_start + match.end()
                preceding_text = text[max(0, current_split_pos - 50):current_split_pos]
                following_text = text[current_split_pos:min(text_len, current_split_pos + 5)]
                if f"<||{IMAGE_PLACEHOLDER_PREFIX}" in preceding_text and "||>" not in following_text:
                     continue # Likely inside a placeholder, don't split here
                potential_splits.append((abs(current_split_pos - ideal_end_index), current_split_pos, 4))


        potential_splits.sort()

        if potential_splits:
             split_index = potential_splits[0][1]

             if split_index <= start_index + min_chunk_size:
                 split_index = -1 # Ignore this split point

        if split_index == -1:
             if ideal_end_index > start_index + min_chunk_size:
                 split_index = ideal_end_index
             else: # Force split at limit or end of text
                 split_index = min(start_index + limit_chars, text_len)

        split_index = min(split_index, text_len)
        if split_index <= start_index:

             split_index = min(start_index + limit_chars, text_len)
             if split_index <= start_index: # Final fallback if limit is tiny or zero
                 split_index = text_len

        chunks.append(text[start_index:split_index])
        start_index = split_index

    return [chunk for chunk in chunks if chunk.strip()]

def get_image_extension_from_data(image_data, fallback_ext="jpeg"):
    """Determines image extension from binary data."""
    if not image_data: return fallback_ext
    ext = imghdr.what(None, image_data)
    if ext == 'jpeg': return 'jpg'
    if ext is None and PILLOW_AVAILABLE:
        try:
            with Image.open(BytesIO(image_data)) as img:
                img_format = img.format
                if img_format:
                    fmt_lower = img_format.lower()
                    if fmt_lower == 'jpeg': return 'jpg'
                    if fmt_lower in ['png', 'gif', 'bmp', 'tiff', 'webp']: return fmt_lower
        except Exception: pass # Ignore Pillow errors if imghdr failed
    return ext if ext else fallback_ext


def convert_emf_to_png(emf_data):
    """Converts EMF image data to PNG using Pillow."""
    if not PILLOW_AVAILABLE:
        print("[WARN] Pillow library not found, cannot convert EMF image. Skipping.")
        return None
    try:

        with Image.open(BytesIO(emf_data)) as img:

            if img.mode == 'CMYK': img = img.convert('RGB')
            elif img.mode == 'P': img = img.convert('RGBA') # Convert palette to RGBA for transparency
            elif img.mode == '1': img = img.convert('L') # Convert bilevel to grayscale

            png_bytes_io = BytesIO()
            img.save(png_bytes_io, format='PNG')
            return png_bytes_io.getvalue()
    except ImportError: # Might happen if EMF plugin for Pillow is missing
         print("[ERROR] Failed to convert EMF: Pillow EMF support might be missing or incomplete on this system.")
         return None
    except Exception as e:
        print(f"[ERROR] Failed to convert EMF to PNG: {e}")
        return None

def extract_number_from_path(path):
    """–ò–∑–≤–ª–µ–∫–∞–µ—Ç –Ω–æ–º–µ—Ä –∏–∑ –∏–º–µ–Ω–∏ —Ñ–∞–π–ª–∞ –¥–ª—è —Å–æ—Ä—Ç–∏—Ä–æ–≤–∫–∏."""
    filename = os.path.basename(path)
    match = re.search(
        r"(?:chapter|part|section|page|item|file|ch|pt|pg|_)(\d+)",
        filename,
        re.IGNORECASE,
    )
    if not match:
        match = re.search(r"(\d+)", filename)
    if match:
        try:
            return int(match.group(1))
        except (ValueError, IndexError):
            return float("inf")
    return float("inf")

class ApiKeyManager:
    """–£–ø—Ä–∞–≤–ª—è–µ—Ç –ø—É–ª–æ–º API –∫–ª—é—á–µ–π —Å –æ—Ç—Å–ª–µ–∂–∏–≤–∞–Ω–∏–µ–º –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏—è –∏ —Ä–æ—Ç–∞—Ü–∏–µ–π."""
    def __init__(self, api_keys):
        if not api_keys:
            raise ValueError("–°–ø–∏—Å–æ–∫ API –∫–ª—é—á–µ–π –Ω–µ –º–æ–∂–µ—Ç –±—ã—Ç—å –ø—É—Å—Ç—ã–º.")
        self.api_keys = list(set(api_keys))  # –£–±–∏—Ä–∞–µ–º –¥—É–±–ª–∏–∫–∞—Ç—ã
        self.current_index = 0
        self.usage_counts = {key: 0 for key in self.api_keys}
        self.usage_limits = {key: 1000 for key in self.api_keys}
        self.exhausted_keys = set()  # –ö–ª—é—á–∏ —Å –∏—Å—á–µ—Ä–ø–∞–Ω–Ω–æ–π –∫–≤–æ—Ç–æ–π
        self.lock = threading.Lock()

    def get_next_available_key(self):
        """–í–æ–∑–≤—Ä–∞—â–∞–µ—Ç —Å–ª–µ–¥—É—é—â–∏–π –¥–æ—Å—Ç—É–ø–Ω—ã–π –∫–ª—é—á —Å —É—á–µ—Ç–æ–º –ª–∏–º–∏—Ç–æ–≤."""
        with self.lock:
            # –ü—Ä–æ–±—É–µ–º –Ω–∞–π—Ç–∏ —Ä–∞–±–æ—á–∏–π –∫–ª—é—á
            attempts = 0
            while attempts < len(self.api_keys):
                key = self.api_keys[self.current_index]
                self.current_index = (self.current_index + 1) % len(self.api_keys)
                attempts += 1
                
                # –ü—Ä–æ–ø—É—Å–∫–∞–µ–º –∏—Å—á–µ—Ä–ø–∞–Ω–Ω—ã–µ –∫–ª—é—á–∏
                if key in self.exhausted_keys:
                    continue
                    
                if self.usage_counts[key] < self.usage_limits[key]:
                    self.usage_counts[key] += 1
                    return key
                    
            # –í—Å–µ –∫–ª—é—á–∏ –∏—Å—á–µ—Ä–ø–∞–Ω—ã –∏–ª–∏ –¥–æ—Å—Ç–∏–≥–ª–∏ –ª–∏–º–∏—Ç–∞
            return None
            
    def mark_key_exhausted(self, key):
        """–ü–æ–º–µ—á–∞–µ—Ç –∫–ª—é—á –∫–∞–∫ –∏—Å—á–µ—Ä–ø–∞–Ω–Ω—ã–π"""
        with self.lock:
            if key in self.api_keys:
                self.exhausted_keys.add(key)
                print(f"[API KEY] –ö–ª—é—á ...{key[-4:]} –ø–æ–º–µ—á–µ–Ω –∫–∞–∫ –∏—Å—á–µ—Ä–ø–∞–Ω–Ω—ã–π")
                
    def has_available_keys(self):
        """–ü—Ä–æ–≤–µ—Ä—è–µ—Ç, –µ—Å—Ç—å –ª–∏ –¥–æ—Å—Ç—É–ø–Ω—ã–µ –∫–ª—é—á–∏"""
        with self.lock:
            return len(self.exhausted_keys) < len(self.api_keys)

    def reset_usage(self):
        """–°–±—Ä–∞—Å—ã–≤–∞–µ—Ç —Å—á–µ—Ç—á–∏–∫–∏ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏—è."""
        with self.lock:
            for key in self.api_keys:
                self.usage_counts[key] = 0
            # –ù–ï —Å–±—Ä–∞—Å—ã–≤–∞–µ–º exhausted_keys - –æ–Ω–∏ –æ—Å—Ç–∞—é—Ç—Å—è –∏—Å—á–µ—Ä–ø–∞–Ω–Ω—ã–º–∏

    def get_usage_report(self):
        """–í–æ–∑–≤—Ä–∞—â–∞–µ—Ç –æ—Ç—á–µ—Ç –æ–± –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏–∏ –∫–ª—é—á–µ–π."""
        with self.lock:
            reports = []
            for key in self.api_keys:
                key_short = f"...{key[-4:]}"
                status = "–ò–°–ß–ï–†–ü–ê–ù" if key in self.exhausted_keys else f"{self.usage_counts[key]}/{self.usage_limits[key]}"
                reports.append(f"{key_short}: {status}")
            return ", ".join(reports)

class RateLimitTracker:
    """–û—Ç—Å–ª–µ–∂–∏–≤–∞–µ—Ç –ª–∏–º–∏—Ç—ã API –Ω–∞ –æ—Å–Ω–æ–≤–µ –∑–∞–≥–æ–ª–æ–≤–∫–æ–≤ –æ—Ç–≤–µ—Ç–∞"""
    def __init__(self):
        self.limits = {}  # {api_key: {'limit': X, 'remaining': Y, 'reset': Z}}
        self.lock = threading.Lock()
        
    def update_from_headers(self, api_key, headers):
        """–û–±–Ω–æ–≤–ª—è–µ—Ç –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –æ –ª–∏–º–∏—Ç–∞—Ö –∏–∑ –∑–∞–≥–æ–ª–æ–≤–∫–æ–≤ –æ—Ç–≤–µ—Ç–∞"""
        with self.lock:
            if api_key not in self.limits:
                self.limits[api_key] = {}
                
            # –ü—Ä–æ–±—É–µ–º —Ä–∞–∑–Ω—ã–µ –≤–∞—Ä–∏–∞–Ω—Ç—ã –Ω–∞–∑–≤–∞–Ω–∏–π –∑–∞–≥–æ–ª–æ–≤–∫–æ–≤
            rate_limit_headers = {
                'limit': ['x-ratelimit-limit', 'ratelimit-limit', 'x-rate-limit-limit'],
                'remaining': ['x-ratelimit-remaining', 'ratelimit-remaining', 'x-rate-limit-remaining'],
                'reset': ['x-ratelimit-reset', 'ratelimit-reset', 'x-rate-limit-reset']
            }
            
            for key, possible_names in rate_limit_headers.items():
                for header_name in possible_names:
                    if header_name in headers:
                        try:
                            if key == 'reset':
                                self.limits[api_key][key] = int(headers[header_name])
                            else:
                                self.limits[api_key][key] = int(headers[header_name])
                        except (ValueError, TypeError):
                            pass
                        break
                        
    def get_remaining_requests(self, api_key):
        """–í–æ–∑–≤—Ä–∞—â–∞–µ—Ç –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ –æ—Å—Ç–∞–≤—à–∏—Ö—Å—è –∑–∞–ø—Ä–æ—Å–æ–≤"""
        with self.lock:
            return self.limits.get(api_key, {}).get('remaining', None)
            
    def should_wait(self, api_key, threshold=2):
        """–û–ø—Ä–µ–¥–µ–ª—è–µ—Ç, –Ω—É–∂–Ω–æ –ª–∏ –∂–¥–∞—Ç—å –ø–µ—Ä–µ–¥ —Å–ª–µ–¥—É—é—â–∏–º –∑–∞–ø—Ä–æ—Å–æ–º"""
        remaining = self.get_remaining_requests(api_key)
        if remaining is not None and remaining <= threshold:
            reset_time = self.limits.get(api_key, {}).get('reset', 0)
            if reset_time:
                wait_time = max(0, reset_time - time.time())
                return True, wait_time
        return False, 0
        
    def get_status(self, api_key):
        """–í–æ–∑–≤—Ä–∞—â–∞–µ—Ç —Å—Ç—Ä–æ–∫—É —Å–æ —Å—Ç–∞—Ç—É—Å–æ–º –ª–∏–º–∏—Ç–æ–≤"""
        with self.lock:
            if api_key not in self.limits:
                return "–ù–µ—Ç –¥–∞–Ω–Ω—ã—Ö –æ –ª–∏–º–∏—Ç–∞—Ö"
            
            info = self.limits[api_key]
            if 'remaining' in info and 'limit' in info:
                return f"{info['remaining']}/{info['limit']} –∑–∞–ø—Ä–æ—Å–æ–≤ –æ—Å—Ç–∞–ª–æ—Å—å"
            return "–ß–∞—Å—Ç–∏—á–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ –æ –ª–∏–º–∏—Ç–∞—Ö"

class InitialSetupDialog(QDialog):
    """–ù–∞—á–∞–ª—å–Ω—ã–π –¥–∏–∞–ª–æ–≥ –¥–ª—è –≤–≤–æ–¥–∞ –≤—Å–µ—Ö –Ω–∞—Å—Ç—Ä–æ–µ–∫ –ø–µ—Ä–µ–¥ –∑–∞–ø—É—Å–∫–æ–º –ø–µ—Ä–µ–≤–æ–¥—á–∏–∫–∞ —Å –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–æ–π —Ä–æ—Ç–∞—Ü–∏–µ–π"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("–ù–∞—Å—Ç—Ä–æ–π–∫–∞ –ø–µ—Ä–µ–≤–æ–¥—á–∏–∫–∞ —Å –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–æ–π —Ä–æ—Ç–∞—Ü–∏–µ–π")
        self.setMinimumSize(700, 750)
        self.selected_file = None
        self.output_folder = None
        self.api_keys = []
        self.glossary_dict = {}
        self.selected_model = DEFAULT_MODEL_NAME
        # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –∑–Ω–∞—á–µ–Ω–∏–µ –ø–æ —É–º–æ–ª—á–∞–Ω–∏—é –Ω–∞ –æ—Å–Ω–æ–≤–µ –º–æ–¥–µ–ª–∏
        self.concurrent_requests = MODELS.get(DEFAULT_MODEL_NAME, {}).get("rpm", 10)
        self.init_ui()
        
    def init_ui(self):
        layout = QVBoxLayout(self)
    
        # –ò–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è –æ —Ä–µ–∂–∏–º–µ
        info_label = QLabel(
            "üîÑ –†–µ–∂–∏–º –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–æ–π —Ä–æ—Ç–∞—Ü–∏–∏ –∫–ª—é—á–µ–π\n"
            "–ü—Ä–∏ –¥–æ—Å—Ç–∏–∂–µ–Ω–∏–∏ –ª–∏–º–∏—Ç–∞ –ø—Ä–æ–≥—Ä–∞–º–º–∞ –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏ –ø–µ—Ä–µ–∫–ª—é—á–∏—Ç—Å—è –Ω–∞ —Å–ª–µ–¥—É—é—â–∏–π –∫–ª—é—á"
        )
        info_label.setStyleSheet("background-color: #e8f4f8; padding: 10px; border-radius: 5px;")
        layout.addWidget(info_label)
    
        # 1. –í—ã–±–æ—Ä —Ñ–∞–π–ª–∞ –¥–ª—è –ø–µ—Ä–µ–≤–æ–¥–∞
        file_group = QGroupBox("1. –§–∞–π–ª –¥–ª—è –ø–µ—Ä–µ–≤–æ–¥–∞")
        file_layout = QVBoxLayout(file_group)
        file_btn_layout = QHBoxLayout()
        self.file_btn = QPushButton("–í—ã–±—Ä–∞—Ç—å —Ñ–∞–π–ª...")
        self.file_btn.clicked.connect(self.select_file)
        self.file_label = QLabel("–§–∞–π–ª –Ω–µ –≤—ã–±—Ä–∞–Ω")
        file_btn_layout.addWidget(self.file_btn)
        file_btn_layout.addWidget(self.file_label, 1)
        file_layout.addLayout(file_btn_layout)
        layout.addWidget(file_group)
    
        # 2. –ü—É—Ç—å —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è
        output_group = QGroupBox("2. –ü–∞–ø–∫–∞ –¥–ª—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è –ø–µ—Ä–µ–≤–æ–¥–∞")
        output_layout = QVBoxLayout(output_group)
        output_btn_layout = QHBoxLayout()
        self.output_btn = QPushButton("–í—ã–±—Ä–∞—Ç—å –ø–∞–ø–∫—É...")
        self.output_btn.clicked.connect(self.select_output_folder)
        self.output_label = QLabel("–ü–∞–ø–∫–∞ –Ω–µ –≤—ã–±—Ä–∞–Ω–∞")
        output_btn_layout.addWidget(self.output_btn)
        output_btn_layout.addWidget(self.output_label, 1)
        output_layout.addLayout(output_btn_layout)
        layout.addWidget(output_group)
    
        # 3. API –∫–ª—é—á–∏
        keys_group = QGroupBox("3. API –∫–ª—é—á–∏ Gemini")
        keys_layout = QVBoxLayout(keys_group)
        keys_layout.addWidget(QLabel("–í–≤–µ–¥–∏—Ç–µ –∫–ª—é—á–∏ (–ø–æ –æ–¥–Ω–æ–º—É –Ω–∞ —Å—Ç—Ä–æ–∫—É):"))
        self.keys_edit = QTextEdit()
        self.keys_edit.setMaximumHeight(100)
        self.keys_edit.setPlaceholderText("–ö–ª—é—á1\n–ö–ª—é—á2\n–ö–ª—é—á3...")
        self.keys_edit.textChanged.connect(self.update_keys_count)
        keys_layout.addWidget(self.keys_edit)
    
        # –°—á–µ—Ç—á–∏–∫ –∫–ª—é—á–µ–π
        self.keys_count_label = QLabel("–ö–ª—é—á–µ–π: 0")
        self.keys_count_label.setStyleSheet("color: blue; font-size: 10px;")
        keys_layout.addWidget(self.keys_count_label)
    
        keys_btn_layout = QHBoxLayout()
        load_keys_btn = QPushButton("üìÅ –ó–∞–≥—Ä—É–∑–∏—Ç—å –∏–∑ —Ñ–∞–π–ª–∞")
        load_keys_btn.clicked.connect(self.load_keys_from_file)
        keys_btn_layout.addWidget(load_keys_btn)
        keys_btn_layout.addStretch()
        keys_layout.addLayout(keys_btn_layout)
        layout.addWidget(keys_group)
    
        # 4. –ì–ª–æ—Å—Å–∞—Ä–∏–π (—Å –ø–æ–¥–¥–µ—Ä–∂–∫–æ–π JSON)
        glossary_group = QGroupBox("4. –ì–ª–æ—Å—Å–∞—Ä–∏–π (–æ–ø—Ü–∏–æ–Ω–∞–ª—å–Ω–æ)")
        glossary_layout = QVBoxLayout(glossary_group)
    
        # –ò–Ω—Å—Ç—Ä—É–∫—Ü–∏—è
        glossary_info = QLabel(
            "–í–≤–µ–¥–∏—Ç–µ —Ç–µ—Ä–º–∏–Ω—ã –≤ —Ñ–æ—Ä–º–∞—Ç–µ:\n"
            "‚Ä¢ –û–±—ã—á–Ω—ã–π: –û—Ä–∏–≥–∏–Ω–∞–ª = –ü–µ—Ä–µ–≤–æ–¥\n"
            "‚Ä¢ JSON: {\"term\": \"–ø–µ—Ä–µ–≤–æ–¥\", ...}"
        )
        glossary_info.setStyleSheet("color: #666; font-size: 10px;")
        glossary_layout.addWidget(glossary_info)
    
        # –¢–µ–∫—Å—Ç–æ–≤–æ–µ –ø–æ–ª–µ –¥–ª—è –≥–ª–æ—Å—Å–∞—Ä–∏—è
        self.glossary_text_edit = QPlainTextEdit()
        self.glossary_text_edit.setMaximumHeight(120)
        self.glossary_text_edit.setPlaceholderText(
            "Son Goku = –°–æ–Ω –ì–æ–∫—É\n"
            "Kamehameha = –ö–∞–º–µ—Ö–∞–º–µ—Ö–∞\n"
            "–ò–õ–ò JSON:\n"
            '{\"Lin An\": \"–õ–∏–Ω—å –ê–Ω—å\", \"Makima\": \"–ú–∞–∫–∏–º–∞\"}'
        )
        from PyQt6.QtGui import QFont
        self.glossary_text_edit.setFont(QFont("Consolas", 9))
        self.glossary_text_edit.textChanged.connect(self.update_glossary_count)
        glossary_layout.addWidget(self.glossary_text_edit)
    
        # –°—á–µ—Ç—á–∏–∫ —Ç–µ—Ä–º–∏–Ω–æ–≤
        self.glossary_count_label = QLabel("–¢–µ—Ä–º–∏–Ω–æ–≤: 0")
        self.glossary_count_label.setStyleSheet("color: blue; font-size: 10px;")
        glossary_layout.addWidget(self.glossary_count_label)
    
        # –ö–Ω–æ–ø–∫–∏ –¥–ª—è —Ä–∞–±–æ—Ç—ã —Å –≥–ª–æ—Å—Å–∞—Ä–∏–µ–º
        glossary_buttons_layout = QHBoxLayout()
    
        load_glossary_btn = QPushButton("üìÅ –ó–∞–≥—Ä—É–∑–∏—Ç—å")
        load_glossary_btn.clicked.connect(self.load_glossary_from_file)
        load_glossary_btn.setToolTip("–ó–∞–≥—Ä—É–∑–∏—Ç—å –≥–ª–æ—Å—Å–∞—Ä–∏–π –∏–∑ —Ñ–∞–π–ª–∞ (JSON –∏–ª–∏ TXT)")
        glossary_buttons_layout.addWidget(load_glossary_btn)
    
        save_glossary_btn = QPushButton("üíæ –°–æ—Ö—Ä–∞–Ω–∏—Ç—å")
        save_glossary_btn.clicked.connect(self.save_glossary_to_file)
        save_glossary_btn.setToolTip("–°–æ—Ö—Ä–∞–Ω–∏—Ç—å —Ç–µ–∫—É—â–∏–π –≥–ª–æ—Å—Å–∞—Ä–∏–π –≤ —Ñ–∞–π–ª")
        glossary_buttons_layout.addWidget(save_glossary_btn)
    
        clear_glossary_btn = QPushButton("üóëÔ∏è –û—á–∏—Å—Ç–∏—Ç—å")
        clear_glossary_btn.clicked.connect(lambda: self.glossary_text_edit.clear())
        clear_glossary_btn.setToolTip("–û—á–∏—Å—Ç–∏—Ç—å –≥–ª–æ—Å—Å–∞—Ä–∏–π")
        glossary_buttons_layout.addWidget(clear_glossary_btn)
    
        # –ü—Ä–µ–¥—É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω–Ω—ã–µ –≥–ª–æ—Å—Å–∞—Ä–∏–∏
        preset_combo = QComboBox()
        preset_combo.addItems([
            "–í—ã–±—Ä–∞—Ç—å –ø—Ä–µ—Å–µ—Ç...",
            "–ê–Ω–∏–º–µ/–ú–∞–Ω–≥–∞",
            "–ö–∏—Ç–∞–π—Å–∫–∏–µ –Ω–æ–≤–µ–ª–ª—ã", 
            "–ö–æ—Ä–µ–π—Å–∫–∏–µ –Ω–æ–≤–µ–ª–ª—ã",
            "–§—ç–Ω—Ç–µ–∑–∏",
            "–ù–∞—É—á–Ω–∞—è —Ñ–∞–Ω—Ç–∞—Å—Ç–∏–∫–∞"
        ])
        preset_combo.currentTextChanged.connect(self.load_preset_glossary)
        glossary_buttons_layout.addWidget(preset_combo)
    
        glossary_buttons_layout.addStretch()
        glossary_layout.addLayout(glossary_buttons_layout)
    
        layout.addWidget(glossary_group)
    
        # –û–ø—Ü–∏–∏ –≥–ª–æ—Å—Å–∞—Ä–∏—è
        glossary_options_layout = QHBoxLayout()

        self.dynamic_glossary_checkbox = QCheckBox("–î–∏–Ω–∞–º–∏—á–µ—Å–∫–∏–π –≥–ª–æ—Å—Å–∞—Ä–∏–π")
        self.dynamic_glossary_checkbox.setToolTip(
            "–ê–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏ —Ñ–∏–ª—å—Ç—Ä—É–µ—Ç –≥–ª–æ—Å—Å–∞—Ä–∏–π –¥–ª—è –∫–∞–∂–¥–æ–π –≥–ª–∞–≤—ã,\n"
            "–æ—Å—Ç–∞–≤–ª—è—è —Ç–æ–ª—å–∫–æ —Ä–µ–ª–µ–≤–∞–Ω—Ç–Ω—ã–µ —Ç–µ—Ä–º–∏–Ω—ã"
        )
        self.dynamic_glossary_checkbox.setChecked(True)
        glossary_options_layout.addWidget(self.dynamic_glossary_checkbox)

        glossary_options_layout.addStretch()
        glossary_layout.addLayout(glossary_options_layout)
    
        # 5. –ú–æ–¥–µ–ª—å –∏ –ø–æ—Ç–æ–∫–∏
        model_group = QGroupBox("5. –ù–∞—Å—Ç—Ä–æ–π–∫–∏ –º–æ–¥–µ–ª–∏")
        model_layout = QFormLayout(model_group)
        self.model_combo = QComboBox()
        self.model_combo.addItems(MODELS.keys())
        self.model_combo.setCurrentText(self.selected_model)
        self.concurrency_spin = QSpinBox()
        self.concurrency_spin.setMinimum(1)
        self.concurrency_spin.setMaximum(100)
        self.concurrency_spin.setValue(self.concurrent_requests)
        # –û–±–Ω–æ–≤–ª—è–µ–º –∑–Ω–∞—á–µ–Ω–∏–µ –ø—Ä–∏ —Å–º–µ–Ω–µ –º–æ–¥–µ–ª–∏
        self.model_combo.currentTextChanged.connect(self.update_concurrency_for_model)
        model_layout.addRow("–ú–æ–¥–µ–ª—å:", self.model_combo)
        model_layout.addRow("–ü–∞—Ä–∞–ª–ª–µ–ª—å–Ω—ã–µ –∑–∞–ø—Ä–æ—Å—ã:", self.concurrency_spin)
        layout.addWidget(model_group)
    
        # 6. –ö–∞—Å—Ç–æ–º–Ω—ã–π –ø—Ä–æ–º–ø—Ç (–ù–û–í–û–ï!)
        prompt_group = QGroupBox("6. –ü—Ä–æ–º–ø—Ç (–æ–ø—Ü–∏–æ–Ω–∞–ª—å–Ω–æ)")
        prompt_layout = QVBoxLayout(prompt_group)
        prompt_layout.addWidget(QLabel("–û—Å—Ç–∞–≤—å—Ç–µ –ø—É—Å—Ç—ã–º –¥–ª—è —Å—Ç–∞–Ω–¥–∞—Ä—Ç–Ω–æ–≥–æ –ø—Ä–æ–º–ø—Ç–∞:"))
        self.custom_prompt_edit = QPlainTextEdit()
        self.custom_prompt_edit.setMaximumHeight(150)
        self.custom_prompt_edit.setPlaceholderText(
            "–í–≤–µ–¥–∏—Ç–µ —Å–≤–æ–π –ø—Ä–æ–º–ø—Ç —Å –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä–æ–º {text}\n"
            "–ü—Ä–∏–º–µ—Ä: –ü–µ—Ä–µ–≤–µ–¥–∏ –Ω–∞ —Ä—É—Å—Å–∫–∏–π —è–∑—ã–∫: {text}"
        )
        prompt_layout.addWidget(self.custom_prompt_edit)
    
        # –ö–Ω–æ–ø–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ —Å—Ç–∞–Ω–¥–∞—Ä—Ç–Ω–æ–≥–æ –ø—Ä–æ–º–ø—Ç–∞
        load_default_btn = QPushButton("üìã –ó–∞–≥—Ä—É–∑–∏—Ç—å —Å—Ç–∞–Ω–¥–∞—Ä—Ç–Ω—ã–π –ø—Ä–æ–º–ø—Ç")
        load_default_btn.clicked.connect(self.load_default_prompt)
        prompt_layout.addWidget(load_default_btn)
    
        layout.addWidget(prompt_group)
    
        # –ö–Ω–æ–ø–∫–∏
        button_box = QDialogButtonBox()
        self.start_btn = QPushButton("üöÄ –°—Ç–∞—Ä—Ç")
        self.start_btn.clicked.connect(self.validate_and_start)
        button_box.addButton(self.start_btn, QDialogButtonBox.ButtonRole.AcceptRole)
        button_box.addButton(QDialogButtonBox.StandardButton.Cancel)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)

    def update_concurrency_for_model(self, model_name):
        """–û–±–Ω–æ–≤–ª—è–µ—Ç –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ –ø–∞—Ä–∞–ª–ª–µ–ª—å–Ω—ã—Ö –∑–∞–ø—Ä–æ—Å–æ–≤ –ø—Ä–∏ —Å–º–µ–Ω–µ –º–æ–¥–µ–ª–∏"""
        if model_name in MODELS:
            rpm = MODELS[model_name].get("rpm", 10)
            self.concurrency_spin.setValue(rpm)
            self.concurrent_requests = rpm

    def load_default_prompt(self):
        """–ó–∞–≥—Ä—É–∂–∞–µ—Ç —É–ª—É—á—à–µ–Ω–Ω—ã–π —Å—Ç–∞–Ω–¥–∞—Ä—Ç–Ω—ã–π –ø—Ä–æ–º–ø—Ç —Å –ø—Ä–∞–≤–∏–ª–∞–º–∏ –¥–∏–∞–ª–æ–≥–æ–≤"""
        default_prompt = """--- PROMPT START ---

**I. –ö–û–ù–¢–ï–ö–°–¢ –ò –ó–ê–î–ê–ß–ê**

**–í–∞—à–∞ –†–æ–ª—å:** –í—ã ‚Äî —ç–ª–∏—Ç–Ω—ã–π –ø–µ—Ä–µ–≤–æ–¥—á–∏–∫ –∏ —Ä–µ–¥–∞–∫—Ç–æ—Ä, **–º–∞—Å—Ç–µ—Ä —Ö—É–¥–æ–∂–µ—Å—Ç–≤–µ–Ω–Ω–æ–π –∞–¥–∞–ø—Ç–∞—Ü–∏–∏**, —Å–ø–µ—Ü–∏–∞–ª–∏–∑–∏—Ä—É—é—â–∏–π—Å—è –Ω–∞ **–ª–∏—Ç–µ—Ä–∞—Ç—É—Ä–Ω–æ–º –ø–µ—Ä–µ–≤–æ–¥–µ —Å–æ–¥–µ—Ä–∂–∏–º–æ–≥–æ EPUB-–∫–Ω–∏–≥** (–≤–µ–±-–Ω–æ–≤–µ–ª–ª, —Ä–∞–Ω–æ–±—ç, —Ä–æ–º–∞–Ω–æ–≤ –∏ —Ç.–¥.) —Å —è–∑—ã–∫–∞ –æ—Ä–∏–≥–∏–Ω–∞–ª–∞ –Ω–∞ —Ä—É—Å—Å–∫–∏–π —è–∑—ã–∫. –í—ã –æ–±–ª–∞–¥–∞–µ—Ç–µ –≥–ª—É–±–æ–∫–∏–º –ø–æ–Ω–∏–º–∞–Ω–∏–µ–º —è–∑—ã–∫–∞ –æ—Ä–∏–≥–∏–Ω–∞–ª–∞, **–µ–≥–æ –∫—É–ª—å—Ç—É—Ä–Ω—ã—Ö –∫–æ–¥–æ–≤**, —Ä–µ—á–µ–≤—ã—Ö –æ–±–æ—Ä–æ—Ç–æ–≤, **–ª–∏—Ç–µ—Ä–∞—Ç—É—Ä–Ω—ã—Ö –ø—Ä–∏–µ–º–æ–≤, –∞ —Ç–∞–∫–∂–µ** —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏—Ö –∞—Å–ø–µ–∫—Ç–æ–≤ —Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏—è XHTML. **–í–∞—à–∞ —Ü–µ–ª—å ‚Äì —Å–æ–∑–¥–∞—Ç—å —Ç–µ–∫—Å—Ç, –∫–æ—Ç–æ—Ä—ã–π —á–∏—Ç–∞–µ—Ç—Å—è —Ç–∞–∫, –±—É–¥—Ç–æ –∏–∑–Ω–∞—á–∞–ª—å–Ω–æ –±—ã–ª –Ω–∞–ø–∏—Å–∞–Ω –Ω–∞ —Ä—É—Å—Å–∫–æ–º —è–∑—ã–∫–µ –¥–ª—è —Ä—É—Å—Å–∫–æ—è–∑—ã—á–Ω–æ–≥–æ —á–∏—Ç–∞—Ç–µ–ª—è, –ø–æ–ª–Ω–æ—Å—Ç—å—é –∑–∞–º–µ–Ω—è—è –æ—Ä–∏–≥–∏–Ω–∞–ª –∏ —Å–æ—Ö—Ä–∞–Ω—è—è –ø—Ä–∏ —ç—Ç–æ–º –≤—Å—é –µ–≥–æ —Å—É—Ç—å, –¥—É—Ö –∏ —É–Ω–∏–∫–∞–ª—å–Ω–æ—Å—Ç—å.**

**–í–∞—à–∞ –ó–∞–¥–∞—á–∞:** –ü–µ—Ä–µ–¥ –≤–∞–º–∏ —Ñ—Ä–∞–≥–º–µ–Ω—Ç –æ—Ä–∏–≥–∏–Ω–∞–ª—å–Ω–æ–≥–æ —Ç–µ–∫—Å—Ç–∞ –∏–∑ —Ñ–∞–π–ª–∞ EPUB (–ø—Ä–µ–¥–æ—Å—Ç–∞–≤–ª–µ–Ω–Ω—ã–π –∫–∞–∫ `{text}` –≤ —Ñ–æ—Ä–º–∞—Ç–µ XHTML/HTML). –í–∞—à–∞ —Ü–µ–ª—å ‚Äî –≤—ã–ø–æ–ª–Ω–∏—Ç—å **–≤—ã—Å–æ–∫–æ–∫–ª–∞—Å—Å–Ω—É—é, –≥–ª—É–±–æ–∫—É—é –ª–∏—Ç–µ—Ä–∞—Ç—É—Ä–Ω—É—é –∞–¥–∞–ø—Ç–∞—Ü–∏—é** –Ω–∞ —Ä—É—Å—Å–∫–∏–π —è–∑—ã–∫, **–≤–∏—Ä—Ç—É–æ–∑–Ω–æ** —Å–æ—Ö—Ä–∞–Ω—è—è —Å–º—ã—Å–ª, —Å—Ç–∏–ª—å, **—ç–º–æ—Ü–∏–æ–Ω–∞–ª—å–Ω—ã–π –Ω–∞–∫–∞–ª, –¥–∏–Ω–∞–º–∏–∫—É –ø–æ–≤–µ—Å—Ç–≤–æ–≤–∞–Ω–∏—è** –∏ –∏—Å—Ö–æ–¥–Ω–æ–µ XHTML-—Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ. **–ö—Ä–∏—Ç–∏—á–µ—Å–∫–∏ –≤–∞–∂–Ω–æ, —á—Ç–æ–±—ã –≤ –∏—Ç–æ–≥–æ–≤–æ–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç–µ –ù–ï –û–°–¢–ê–õ–û–°–¨ –ù–ò –û–î–ù–û–ì–û –°–õ–û–í–ê –∏–ª–∏ –§–†–ê–ì–ú–ï–ù–¢–ê —Ç–µ–∫—Å—Ç–∞ –Ω–∞ —è–∑—ã–∫–µ –æ—Ä–∏–≥–∏–Ω–∞–ª–∞ (–∑–∞ –∏—Å–∫–ª—é—á–µ–Ω–∏–µ–º –Ω–µ–∏–∑–º–µ–Ω—è–µ–º—ã—Ö —á–∞—Å—Ç–µ–π XHTML, —É–∫–∞–∑–∞–Ω–Ω—ã—Ö –Ω–∏–∂–µ).**

**II. –û–ë–©–ò–ï –ü–†–ò–ù–¶–ò–ü–´ –ê–î–ê–ü–¢–ê–¶–ò–ò**

1.  **–ï—Å—Ç–µ—Å—Ç–≤–µ–Ω–Ω–æ—Å—Ç—å –∏ –•—É–¥–æ–∂–µ—Å—Ç–≤–µ–Ω–Ω–æ—Å—Ç—å –†—É—Å—Å–∫–æ–≥–æ –Ø–∑—ã–∫–∞:** –ü–µ—Ä–µ–≤–æ–¥ –¥–æ–ª–∂–µ–Ω –∑–≤—É—á–∞—Ç—å –∞–±—Å–æ–ª—é—Ç–Ω–æ –æ—Ä–≥–∞–Ω–∏—á–Ω–æ –∏ **–ª–∏—Ç–µ—Ä–∞—Ç—É—Ä–Ω–æ** –ø–æ-—Ä—É—Å—Å–∫–∏. –ò–∑–±–µ–≥–∞–π—Ç–µ –±—É–∫–≤–∞–ª—å–Ω–æ–≥–æ —Å–ª–µ–¥–æ–≤–∞–Ω–∏—è –≥—Ä–∞–º–º–∞—Ç–∏–∫–µ –∏–ª–∏ –∏–¥–∏–æ–º–∞–º –æ—Ä–∏–≥–∏–Ω–∞–ª–∞, –µ—Å–ª–∏ –æ–Ω–∏ —Å–æ–∑–¥–∞—é—Ç –Ω–µ–µ—Å—Ç–µ—Å—Ç–≤–µ–Ω–Ω—ã–µ –∏–ª–∏ –∫–æ—Å–Ω–æ—è–∑—ã—á–Ω—ã–µ –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏. –ù–∞—Ö–æ–¥–∏—Ç–µ —ç–∫–≤–∏–≤–∞–ª–µ–Ω—Ç–Ω—ã–µ —Ä—É—Å—Å–∫–∏–µ –≤—ã—Ä–∞–∂–µ–Ω–∏—è, **–∏–¥–∏–æ–º—ã –∏ —Ä–µ—á–µ–≤—ã–µ –æ–±–æ—Ä–æ—Ç—ã, –∫–æ—Ç–æ—Ä—ã–µ —Ç–æ—á–Ω–æ –ø–µ—Ä–µ–¥–∞—é—Ç –∑–∞–º—ã—Å–µ–ª –∞–≤—Ç–æ—Ä–∞.** **–°—Ç—Ä–µ–º–∏—Ç–µ—Å—å –∫ –±–æ–≥–∞—Ç—Å—Ç–≤—É, –æ–±—Ä–∞–∑–Ω–æ—Å—Ç–∏ –∏ –≤—ã—Ä–∞–∑–∏—Ç–µ–ª—å–Ω–æ—Å—Ç–∏ —Ä—É—Å—Å–∫–æ–≥–æ —è–∑—ã–∫–∞.**
2.  **–°–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ –°–º—ã—Å–ª–∞, –¢–æ–Ω–∞ –∏ –ê—Ç–º–æ—Å—Ñ–µ—Ä—ã:** –¢–æ—á–Ω–æ –ø–µ—Ä–µ–¥–∞–≤–∞–π—Ç–µ –æ—Å–Ω–æ–≤–Ω–æ–π —Å–º—ã—Å–ª, –∞—Ç–º–æ—Å—Ñ–µ—Ä—É (—é–º–æ—Ä, —Å–∞—Å–ø–µ–Ω—Å, –¥—Ä–∞–º—É, —Ä–æ–º–∞–Ω—Ç–∏–∫—É –∏ —Ç.–¥.) –∏ –∞–≤—Ç–æ—Ä—Å–∫–∏–π —Å—Ç–∏–ª—å –æ—Ä–∏–≥–∏–Ω–∞–ª–∞. **–£–¥–µ–ª—è–π—Ç–µ –æ—Å–æ–±–æ–µ –≤–Ω–∏–º–∞–Ω–∏–µ –ø–µ—Ä–µ–¥–∞—á–µ —ç–º–æ—Ü–∏–π –ø–µ—Ä—Å–æ–Ω–∞–∂–µ–π, –∏—Ö –≤–Ω—É—Ç—Ä–µ–Ω–Ω–∏—Ö –ø–µ—Ä–µ–∂–∏–≤–∞–Ω–∏–π, –º–æ—Ç–∏–≤–∞—Ü–∏–π –∏ —Ö–∞—Ä–∞–∫—Ç–µ—Ä–∞ —á–µ—Ä–µ–∑ –∏—Ö —Ä–µ—á—å –∏ –º—ã—Å–ª–∏.**
3.  **–ö—É–ª—å—Ç—É—Ä–Ω–∞—è –∏ –°—Ç–∏–ª–∏—Å—Ç–∏—á–µ—Å–∫–∞—è –ê–¥–∞–ø—Ç–∞—Ü–∏—è:**
    *   **–•–æ–Ω–æ—Ä–∏—Ñ–∏–∫–∏ (-—Å–∞–Ω, -–∫—É–Ω, -–Ω–∏–º, –≥—ç–≥—ç, —à–∏—Å—é–Ω, —Å—ç–º–ø–∞–π –∏ —Ç.–¥.):** –ö–∞–∫ –ø—Ä–∞–≤–∏–ª–æ, **–æ–ø—É—Å–∫–∞–π—Ç–µ** –∏–ª–∏ –∑–∞–º–µ–Ω—è–π—Ç–µ –µ—Å—Ç–µ—Å—Ç–≤–µ–Ω–Ω—ã–º–∏ —Ä—É—Å—Å–∫–∏–º–∏ —Ñ–æ—Ä–º–∞–º–∏ –æ–±—Ä–∞—â–µ–Ω–∏—è (–ø–æ –∏–º–µ–Ω–∏, "–≥–æ—Å–ø–æ–¥–∏–Ω/–≥–æ—Å–ø–æ–∂–∞", "–±—Ä–∞—Ç–µ—Ü/—Å–µ—Å—Ç—Ä–∏—Ü–∞", "—É—á–∏—Ç–µ–ª—å", "—Å—Ç–∞—Ä—à–∏–π" ‚Äì –≤ –∑–∞–≤–∏—Å–∏–º–æ—Å—Ç–∏ –æ—Ç –∫–æ–Ω—Ç–µ–∫—Å—Ç–∞ –∏ –æ—Ç–Ω–æ—à–µ–Ω–∏–π –º–µ–∂–¥—É –ø–µ—Ä—Å–æ–Ω–∞–∂–∞–º–∏).
    *   **–†–µ–∞–ª–∏–∏:** –ê–¥–∞–ø—Ç–∏—Ä—É–π—Ç–µ –Ω–µ–ø–æ–Ω—è—Ç–Ω—ã–µ –¥–ª—è —Ä—É—Å—Å–∫–æ–≥–æ —á–∏—Ç–∞—Ç–µ–ª—è –∫—É–ª—å—Ç—É—Ä–Ω—ã–µ –∏–ª–∏ –±—ã—Ç–æ–≤—ã–µ —Ä–µ–∞–ª–∏–∏: –Ω–∞–π–¥–∏—Ç–µ —Ä—É—Å—Å–∫–∏–π —ç–∫–≤–∏–≤–∞–ª–µ–Ω—Ç, –¥–∞–π—Ç–µ –∫—Ä–∞—Ç–∫–æ–µ, **–æ—Ä–≥–∞–Ω–∏—á–Ω–æ –≤–ø–ª–µ—Ç–µ–Ω–Ω–æ–µ –≤ –ø–æ–≤–µ—Å—Ç–≤–æ–≤–∞–Ω–∏–µ –ø–æ—è—Å–Ω–µ–Ω–∏–µ** (–Ω–∞–ø—Ä–∏–º–µ—Ä, "–æ–Ω –¥–æ—Å—Ç–∞–ª —Ü–∑—è–Ω—å ‚Äì –ø—Ä—è–º–æ–π –∫–∏—Ç–∞–π—Å–∫–∏–π –º–µ—á"), –∏–ª–∏ –∑–∞–º–µ–Ω–∏—Ç–µ –Ω–∞ –±–ª–∏–∑–∫—É—é –ø–æ —Å–º—ã—Å–ª—É –ø–æ–Ω—è—Ç–Ω—É—é –¥–µ—Ç–∞–ª—å. *–ò–∑–±–µ–≥–∞–π—Ç–µ —Å–Ω–æ—Å–æ–∫ –∏ –∫–æ–º–º–µ–Ω—Ç–∞—Ä–∏–µ–≤ –ø–µ—Ä–µ–≤–æ–¥—á–∏–∫–∞ –≤ —Ç–µ–∫—Å—Ç–µ.*
    *   **–û–Ω–æ–º–∞—Ç–æ–ø–µ—è (–∑–≤—É–∫–æ–ø–æ–¥—Ä–∞–∂–∞–Ω–∏—è):** –ó–∞–º–µ–Ω—è–π—Ç–µ —Ä—É—Å—Å–∫–∏–º–∏ –∑–≤—É–∫–æ–ø–æ–¥—Ä–∞–∂–∞–Ω–∏—è–º–∏ –∏–ª–∏ **—è—Ä–∫–∏–º–∏, –æ–±—Ä–∞–∑–Ω—ã–º–∏ –æ–ø–∏—Å–∞–Ω–∏—è–º–∏ –∑–≤—É–∫–∞/–¥–µ–π—Å—Ç–≤–∏—è** (–Ω–∞–ø—Ä–∏–º–µ—Ä, –≤–º–µ—Å—Ç–æ "–±–∞—Ö" –º–æ–∂–Ω–æ –Ω–∞–ø–∏—Å–∞—Ç—å "—Ä–∞–∑–¥–∞–ª—Å—è –≥–ª—É—Ö–æ–π —É–¥–∞—Ä").
    *   **–ò–º–µ–Ω–∞ —Å–æ–±—Å—Ç–≤–µ–Ω–Ω—ã–µ –∏ –ù–∞–∑–≤–∞–Ω–∏—è:** –ü—Ä–∏ –æ—Ç—Å—É—Ç—Å—Ç–≤–∏–∏ –≥–ª–æ—Å—Å–∞—Ä–∏—è, —Å—Ç—Ä–µ–º–∏—Ç–µ—Å—å –∫ –±–ª–∞–≥–æ–∑–≤—É—á–Ω–æ–π –∏ –ø–æ–Ω—è—Ç–Ω–æ–π –∞–¥–∞–ø—Ç–∞—Ü–∏–∏. –ï—Å–ª–∏ –≤–æ–∑–º–æ–∂–µ–Ω –æ—Å–º—ã—Å–ª–µ–Ω–Ω—ã–π –ø–µ—Ä–µ–≤–æ–¥ –Ω–∞–∑–≤–∞–Ω–∏—è (–Ω–∞–ø—Ä–∏–º–µ—Ä, —Ç–µ—Ö–Ω–∏–∫–∏ –∏–ª–∏ –∞—Ä—Ç–µ—Ñ–∞–∫—Ç–∞), –æ—Ç–¥–∞–≤–∞–π—Ç–µ –µ–º—É –ø—Ä–µ–¥–ø–æ—á—Ç–µ–Ω–∏–µ –ø–µ—Ä–µ–¥ —Ç—Ä–∞–Ω—Å–ª–∏—Ç–µ—Ä–∞—Ü–∏–µ–π. **–ò–∑–±–µ–≥–∞–π—Ç–µ –Ω–∞–≥—Ä–æ–º–æ–∂–¥–µ–Ω–∏—è —Ç—Ä—É–¥–Ω–æ–ø—Ä–æ–∏–∑–Ω–æ—Å–∏–º—ã—Ö —Ç—Ä–∞–Ω—Å–ª–∏—Ç–µ—Ä–∞—Ü–∏–π.**
    *   **–°—Ç–∏–ª–∏—Å—Ç–∏–∫–∞ —Ä–µ—á–∏ –ø–µ—Ä—Å–æ–Ω–∞–∂–µ–π:** –ï—Å–ª–∏ –æ—Ä–∏–≥–∏–Ω–∞–ª –ø—Ä–µ–¥–ø–æ–ª–∞–≥–∞–µ—Ç —Ä–∞–∑–ª–∏—á–∏—è –≤ –º–∞–Ω–µ—Ä–µ —Ä–µ—á–∏ —Ä–∞–∑–Ω—ã—Ö –ø–µ—Ä—Å–æ–Ω–∞–∂–µ–π (–ø—Ä–æ—Å—Ç–æ—Ä–µ—á–∏—è, –≤—ã—Å–æ–∫–∏–π —Å—Ç–∏–ª—å, –∞—Ä—Ö–∞–∏–∑–º—ã, –∂–∞—Ä–≥–æ–Ω, –¥–µ—Ç—Å–∫–∞—è —Ä–µ—á—å), **—Å—Ç–∞—Ä–∞–π—Ç–µ—Å—å –ø–µ—Ä–µ–¥–∞—Ç—å —ç—Ç–∏ —Ä–∞–∑–ª–∏—á–∏—è —Å—Ä–µ–¥—Å—Ç–≤–∞–º–∏ —Ä—É—Å—Å–∫–æ–≥–æ —è–∑—ã–∫–∞.**

**III. –¢–ï–•–ù–ò–ß–ï–°–ö–ò–ï –ò –°–¢–ò–õ–ò–°–¢–ò–ß–ï–°–ö–ò–ï –¢–†–ï–ë–û–í–ê–ù–ò–Ø**

**1. –†–∞–±–æ—Ç–∞ —Å XHTML/HTML-—Å—Ç—Ä—É–∫—Ç—É—Ä–æ–π EPUB:**
*   **–í–ê–® –ì–õ–ê–í–ù–´–ô –ü–†–ò–û–†–ò–¢–ï–¢ ‚Äî –ü–û–õ–ù–û–ï –°–û–•–†–ê–ù–ï–ù–ò–ï –ò–°–•–û–î–ù–û–ô XHTML/HTML-–°–¢–†–£–ö–¢–£–†–´.** –ü–æ–º–Ω–∏—Ç–µ, —á—Ç–æ EPUB-–∫–Ω–∏–≥–∞ —Å–æ—Å—Ç–æ–∏—Ç –∏–∑ XHTML-—Ñ–∞–π–ª–æ–≤. –í–∞—à–∞ –∑–∞–¥–∞—á–∞ ‚Äî —Ä–∞–±–æ—Ç–∞—Ç—å —Å –∫–æ–¥–æ–º —ç—Ç–∏—Ö —Ñ–∞–π–ª–æ–≤, –ø–µ—Ä–µ–≤–æ–¥—è —Ç–æ–ª—å–∫–æ —Ç–µ–∫—Å—Ç–æ–≤–æ–µ –Ω–∞–ø–æ–ª–Ω–µ–Ω–∏–µ.
*   **–°–û–•–†–ê–ù–Ø–ô–¢–ï –í–°–ï HTML-–¢–ï–ì–ò!** –ü–µ—Ä–µ–≤–æ–¥–∏—Ç–µ **–¢–û–õ–¨–ö–û –≤–∏–¥–∏–º—ã–π –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—é —Ç–µ–∫—Å—Ç** –≤–Ω—É—Ç—Ä–∏ —Ç–µ–≥–æ–≤ (–Ω–∞–ø—Ä–∏–º–µ—Ä, —Ç–µ–∫—Å—Ç –≤–Ω—É—Ç—Ä–∏ `<p>`, `<h1>`, `<li>`, `<td>`, `<span>`, `<a>` –∏ —Ç.–¥., –∞ —Ç–∞–∫–∂–µ –∑–Ω–∞—á–µ–Ω–∏—è –∞—Ç—Ä–∏–±—É—Ç–æ–≤ `title` –∏ `alt`, –µ—Å–ª–∏ –æ–Ω–∏ —Å–æ–¥–µ—Ä–∂–∞—Ç –æ—Å–º—ã—Å–ª–µ–Ω–Ω—ã–π —Ç–µ–∫—Å—Ç).
*   **–ù–ï –ú–ï–ù–Ø–ô–¢–ï, –ù–ï –£–î–ê–õ–Ø–ô–¢–ï –∏ –ù–ï –î–û–ë–ê–í–õ–Ø–ô–¢–ï** –Ω–∏–∫–∞–∫–∏–µ HTML-—Ç–µ–≥–∏ (`<p>`, `<div>`, `<img>`, `<a>` –∏ —Ç.–¥.), –∞—Ç—Ä–∏–±—É—Ç—ã (`class`, `id`, `href`, `src` –∏ —Ç.–¥.) –∏–ª–∏ —Å—Ç—Ä—É–∫—Ç—É—Ä—É –¥–æ–∫—É–º–µ–Ω—Ç–∞.
*   **–ö–æ–º–º–µ–Ω—Ç–∞—Ä–∏–∏ HTML (`<!-- ... -->`), —Å–∫—Ä–∏–ø—Ç—ã (`<script>...</script>`) –∏ —Å—Ç–∏–ª–∏ (`<style>...</style>`) –¥–æ–ª–∂–Ω—ã –æ—Å—Ç–∞–≤–∞—Ç—å—Å—è –ë–ï–ó –ò–ó–ú–ï–ù–ï–ù–ò–ô.** –°–æ–¥–µ—Ä–∂–∏–º–æ–µ —ç—Ç–∏—Ö —Ç–µ–≥–æ–≤ **–ù–ï –ü–ï–†–ï–í–û–î–ò–¢–°–Ø**.
*   **–¶–µ–ª—å:** –í—ã—Ö–æ–¥–Ω–æ–π –∫–æ–¥ –¥–æ–ª–∂–µ–Ω –±—ã—Ç—å –≤–∞–ª–∏–¥–Ω—ã–º XHTML/HTML —Å —Ç–æ–π –∂–µ —Å—Ç—Ä—É–∫—Ç—É—Ä–æ–π –∏ —Ç–µ–≥–∞–º–∏, —á—Ç–æ –∏ –≤—Ö–æ–¥–Ω–æ–π, –Ω–æ —Å **–ø–æ–ª–Ω–æ—Å—Ç—å—é –ø–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω—ã–º —Ç–µ–∫—Å—Ç–æ–≤—ã–º —Å–æ–¥–µ—Ä–∂–∏–º—ã–º** (–∫—Ä–æ–º–µ —É–∫–∞–∑–∞–Ω–Ω—ã—Ö –∏—Å–∫–ª—é—á–µ–Ω–∏–π).

**2. –°—Ç–∏–ª–∏—Å—Ç–∏—á–µ—Å–∫–∏–µ –¢—Ä–µ–±–æ–≤–∞–Ω–∏—è –∫ –¢–µ–∫—Å—Ç—É (–ü—Ä–∞–≤–∏–ª–∞ –ê–¥–∞–ø—Ç–∞—Ü–∏–∏):**
*   **2.1. (–û—Ñ–æ—Ä–º–ª–µ–Ω–∏–µ –ø—Ä—è–º–æ–π —Ä–µ—á–∏ –∏ —Ü–∏—Ç–∞—Ç):**
    *   –ö–≤–∞–¥—Ä–∞—Ç–Ω—ã–µ —Å–∫–æ–±–∫–∏ `[]`, –æ–±–æ–∑–Ω–∞—á–∞—é—â–∏–µ **–ø—Ä—è–º—É—é —Ä–µ—á—å –ø–µ—Ä—Å–æ–Ω–∞–∂–µ–π**, –∑–∞–º–µ–Ω—è–π—Ç–µ –Ω–∞ —Å—Ç–∞–Ω–¥–∞—Ä—Ç–Ω–æ–µ –æ—Ñ–æ—Ä–º–ª–µ–Ω–∏–µ –ø—Ä—è–º–æ–π —Ä–µ—á–∏ –≤ —Ä—É—Å—Å–∫–æ–º —è–∑—ã–∫–µ —Å –ø–æ–º–æ—â—å—é —Ç–∏—Ä–µ: `‚Äî –†–µ–ø–ª–∏–∫–∞.`
    *   –ö–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏–∏ –≤–∏–¥–∞ `„Äé–¶–∏—Ç–∞—Ç–∞/–†–µ–ø–ª–∏–∫–∞„Äè` –∏–ª–∏ `„Äå–¶–∏—Ç–∞—Ç–∞/–†–µ–ø–ª–∏–∫–∞„Äç` –∑–∞–º–µ–Ω—è–π—Ç–µ –Ω–∞ —Ä—É—Å—Å–∫–∏–µ –∫–∞–≤—ã—á–∫–∏-¬´—ë–ª–æ—á–∫–∏¬ª (`¬´–¶–∏—Ç–∞—Ç–∞/–†–µ–ø–ª–∏–∫–∞¬ª`), –µ—Å–ª–∏ —ç—Ç–æ –≤—ã–¥–µ–ª–µ–Ω–Ω–∞—è –º—ã—Å–ª—å, –Ω–∞–∑–≤–∞–Ω–∏–µ, —Ü–∏—Ç–∞—Ç–∞. –ï—Å–ª–∏ —ç—Ç–æ –ø–æ–ª–Ω–æ—Ü–µ–Ω–Ω–∞—è –ø—Ä—è–º–∞—è —Ä–µ—á—å, –æ—Ñ–æ—Ä–º–ª—è–π—Ç–µ –µ—ë —Å —Ç–∏—Ä–µ: `‚Äî –†–µ–ø–ª–∏–∫–∞.`
*   **2.2. (–û—Ñ–æ—Ä–º–ª–µ–Ω–∏–µ –º—ã—Å–ª–µ–π):** –í—Å–µ **–º—ã—Å–ª–∏ –≥–µ—Ä–æ–µ–≤** –æ—Ñ–æ—Ä–º–ª—è–π—Ç–µ —Ä—É—Å—Å–∫–∏–º–∏ –∫–∞–≤—ã—á–∫–∞–º–∏-¬´—ë–ª–æ—á–∫–∞–º–∏¬ª: `¬´–ú—ã—Å–ª—å –ø–µ—Ä—Å–æ–Ω–∞–∂–∞.¬ª`
*   **2.3. (–ü–ª–∞–≤–Ω–æ—Å—Ç—å –∏ —á–∏—Ç–∞–µ–º–æ—Å—Ç—å):** –£–¥–µ–ª–∏—Ç–µ –æ—Å–æ–±–æ–µ –≤–Ω–∏–º–∞–Ω–∏–µ **–ø–ª–∞–≤–Ω–æ—Å—Ç–∏ –∏ —Ä–∏—Ç–º—É —Ç–µ–∫—Å—Ç–∞**. –û–Ω –¥–æ–ª–∂–µ–Ω —á–∏—Ç–∞—Ç—å—Å—è –µ—Å—Ç–µ—Å—Ç–≤–µ–Ω–Ω–æ –∏ —É–≤–ª–µ–∫–∞—Ç–µ–ª—å–Ω–æ. –ü—Ä–∏ –Ω–µ–æ–±—Ö–æ–¥–∏–º–æ—Å—Ç–∏, –¥–µ–ª–∏—Ç–µ —Å–ª–∏—à–∫–æ–º –¥–ª–∏–Ω–Ω—ã–µ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è –Ω–∞ –±–æ–ª–µ–µ –∫–æ—Ä–æ—Ç–∫–∏–µ –¥–ª—è –ª—É—á—à–µ–π —á–∏—Ç–∞–µ–º–æ—Å—Ç–∏, –Ω–µ —Ç–µ—Ä—è—è —Å–≤—è–∑–∏ –∏ —Å–º—ã—Å–ª–∞.
*   **2.4. (–ü–µ—Ä–µ–¥–∞—á–∞ –ø—Ä–æ—Ç—è–∂–Ω—ã—Ö –∑–≤—É–∫–æ–≤/–∑–∞–∏–∫–∞–Ω–∏—è):** –î–ª—è –ø–µ—Ä–µ–¥–∞—á–∏ –ø—Ä–æ—Ç—è–∂–Ω—ã—Ö –∑–≤—É–∫–æ–≤ –∏–ª–∏ –∑–∞–∏–∫–∞–Ω–∏—è –æ–≥—Ä–∞–Ω–∏—á–∏–≤–∞–π—Ç–µ—Å—å **—Ç—Ä–µ–º—è-—á–µ—Ç—ã—Ä—å–º—è –ø–æ–≤—Ç–æ—Ä–µ–Ω–∏—è–º–∏ –±—É–∫–≤—ã**, —Ä–∞–∑–¥–µ–ª–µ–Ω–Ω—ã–º–∏ –¥–µ—Ñ–∏—Å–æ–º: `–ê-–∞-–∞—Ö...`, `–ù-–Ω-–Ω–µ—Ç...`.
*   **2.5. (–ó–Ω–∞–∫–∏ –ø—Ä–µ–ø–∏–Ω–∞–Ω–∏—è –≤ –∫–æ–Ω—Ü–µ —Ñ—Ä–∞–∑—ã):** –ï—Å–ª–∏ —Ñ—Ä–∞–∑–∞ –æ–∫–∞–Ω—á–∏–≤–∞–µ—Ç—Å—è –Ω–∞ `...!` –∏–ª–∏ `...?`, **—Å–æ—Ö—Ä–∞–Ω—è–π—Ç–µ —ç—Ç–æ—Ç –ø–æ—Ä—è–¥–æ–∫**. –î–ª—è —Å–æ—á–µ—Ç–∞–Ω–∏—è –≤–æ–ø—Ä–æ—Å–∏—Ç–µ–ª—å–Ω–æ–≥–æ –∏ –≤–æ—Å–∫–ª–∏—Ü–∞—Ç–µ–ª—å–Ω–æ–≥–æ –∑–Ω–∞–∫–æ–≤ –∏—Å–ø–æ–ª—å–∑—É–π—Ç–µ `?!` –∏–ª–∏ `!?`.
*   **2.6. (–û—Ñ–æ—Ä–º–ª–µ–Ω–∏–µ –º—ã—Å–ª–µ–π –±–µ–∑ —Ç–∏—Ä–µ):** –ú—ã—Å–ª–∏ –≤ –∫–∞–≤—ã—á–∫–∞—Ö –¥–æ–ª–∂–Ω—ã –±—ã—Ç—å —Å–∞–º–æ—Å—Ç–æ—è—Ç–µ–ª—å–Ω—ã–º–∏ –∫–æ–Ω—Å—Ç—Ä—É–∫—Ü–∏—è–º–∏. –ù–µ —Å—Ç–∞–≤—å—Ç–µ –ø–µ—Ä–µ–¥ –Ω–∏–º–∏ —Ç–∏—Ä–µ, –∫–∞–∫ –ø–µ—Ä–µ–¥ –ø—Ä—è–º–æ–π —Ä–µ—á—å—é.
    *   –ö–æ—Ä—Ä–µ–∫—Ç–Ω–æ: `–û–Ω –ø–æ–¥—É–º–∞–ª: ¬´–≠—Ç–æ –∏–Ω—Ç–µ—Ä–µ—Å–Ω–æ¬ª.` –∏–ª–∏ `¬´–≠—Ç–æ –∏–Ω—Ç–µ—Ä–µ—Å–Ω–æ¬ª, ‚Äî –º–µ–ª—å–∫–Ω—É–ª–æ —É –Ω–µ–≥–æ –≤ –≥–æ–ª–æ–≤–µ.`
    *   –ù–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω–æ: `‚Äî ¬´–ú—ã—Å–ª—å...¬ª`
*   **2.7. (–ö–æ–ª–∏—á–µ—Å—Ç–≤–æ –∑–Ω–∞–∫–æ–≤ –ø—Ä–µ–ø–∏–Ω–∞–Ω–∏—è):** –ß—Ä–µ–∑–º–µ—Ä–Ω–æ–µ –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ –æ–¥–∏–Ω–∞–∫–æ–≤—ã—Ö –∑–Ω–∞–∫–æ–≤ (`!!!!`, `????`) –∑–∞–º–µ–Ω—è–π—Ç–µ **–æ–¥–Ω–∏–º, –¥–≤—É–º—è (`!!`, `??`) –∏–ª–∏ —Å–æ—á–µ—Ç–∞–Ω–∏–µ–º `?!` / `!?`**.
*   **2.8. (–ü–µ—Ä–µ–¥–∞—á–∞ –∑–∞–∏–∫–∞–Ω–∏—è/—Ä–∞–∑–¥–µ–ª—å–Ω–æ–≥–æ –ø—Ä–æ–∏–∑–Ω–æ—à–µ–Ω–∏—è):** –°–æ—Ö—Ä–∞–Ω—è–π—Ç–µ —Ä–∞–∑–¥–µ–ª–µ–Ω–∏–µ –±—É–∫–≤ –¥–µ—Ñ–∏—Å–æ–º –¥–ª—è –ø–µ—Ä–µ–¥–∞—á–∏ –∑–∞–∏–∫–∞–Ω–∏—è –∏–ª–∏ –ø—Ä–æ—Ç—è–∂–Ω–æ–≥–æ –ø—Ä–æ–∏–∑–Ω–µ—Å–µ–Ω–∏—è: `–ü-–ø-–ø—Ä–∏–≤–µ—Ç...`, `–ß—Ç-—Ç-—Ç–æ-–æ?!`

    1. **–ö–ê–ñ–î–ê–Ø –†–ï–ü–õ–ò–ö–ê –î–ò–ê–õ–û–ì–ê –ù–ê–ß–ò–ù–ê–ï–¢–°–Ø –° –ù–û–í–û–ô –°–¢–†–û–ö–ò (–ù–û–í–û–ì–û –ê–ë–ó–ê–¶–ê)**
    2. **–ü—Ä–∞–≤–∏–ª—å–Ω–æ–µ –æ—Ñ–æ—Ä–º–ª–µ–Ω–∏–µ –¥–∏–∞–ª–æ–≥–æ–≤:**
       - –ü—Ä–æ—Å—Ç–∞—è —Ä–µ–ø–ª–∏–∫–∞: `‚Äî –¢–µ–∫—Å—Ç —Ä–µ–ø–ª–∏–∫–∏.`
       - –†–µ–ø–ª–∏–∫–∞ —Å –∞–≤—Ç–æ—Ä—Å–∫–æ–π —Ä–µ–º–∞—Ä–∫–æ–π –ü–û–°–õ–ï: `‚Äî –¢–µ–∫—Å—Ç —Ä–µ–ø–ª–∏–∫–∏, ‚Äî —Å–∫–∞–∑–∞–ª –ø–µ—Ä—Å–æ–Ω–∞–∂.`
       - –†–µ–ø–ª–∏–∫–∞ —Å –∞–≤—Ç–æ—Ä—Å–∫–æ–π —Ä–µ–º–∞—Ä–∫–æ–π –î–û: `–ü–µ—Ä—Å–æ–Ω–∞–∂ —Å–∫–∞–∑–∞–ª:` (–Ω–æ–≤–∞—è —Å—Ç—Ä–æ–∫–∞) `‚Äî –¢–µ–∫—Å—Ç —Ä–µ–ø–ª–∏–∫–∏.`
       - –ù–ï —Ä–∞–∑—Ä—ã–≤–∞–π—Ç–µ —Ä–µ–ø–ª–∏–∫—É –∏ –µ—ë –∞–≤—Ç–æ—Ä—Å–∫—É—é —Ä–µ–º–∞—Ä–∫—É –Ω–∞ —Ä–∞–∑–Ω—ã–µ –∞–±–∑–∞—Ü—ã!

    3. **–ó–ê–ü–†–ï–©–ï–ù–û:**
       ‚ùå –ù–µ–ø—Ä–∞–≤–∏–ª—å–Ω–æ:
       ```
       ‚Äî –†–µ–ø–ª–∏–∫–∞.
   
       ‚Äî –°–ª–µ–¥—É—é—â–∞—è —Ä–µ–ø–ª–∏–∫–∞.
   
       —Å–∫–∞–∑–∞–ª –æ–Ω.
       ```
   
       ‚úÖ –ü—Ä–∞–≤–∏–ª—å–Ω–æ:
       ```
       ‚Äî –†–µ–ø–ª–∏–∫–∞.
   
       ‚Äî –°–ª–µ–¥—É—é—â–∞—è —Ä–µ–ø–ª–∏–∫–∞, ‚Äî —Å–∫–∞–∑–∞–ª –æ–Ω.
       ```

    4. **–ú—ã—Å–ª–∏ –ø–µ—Ä—Å–æ–Ω–∞–∂–µ–π:** –æ—Ñ–æ—Ä–º–ª—è–π—Ç–µ –≤ –∫–∞–≤—ã—á–∫–∞—Ö-¬´—ë–ª–æ—á–∫–∞—Ö¬ª: ¬´–ú—ã—Å–ª—å –ø–µ—Ä—Å–æ–Ω–∞–∂–∞¬ª

**3. –û–ë–Ø–ó–ê–¢–ï–õ–¨–ù–û –û–§–û–†–ú–õ–Ø–¢–¨ –ù–ê–ó–í–ê–ù–ò–Ø –ì–õ–ê–í –í –í–ò–î–ï: –ì–ª–∞–≤–∞ X. –ù–∞–∑–≤–∞–Ω–∏–µ –≥–ª–∞–≤—ã**
–ï—Å–ª–∏ –ï–°–¢–¨ –≥–ª–∞–≤–∞, –Ω–æ –Ω–µ—Ç –Ω–∞–∑–≤–∞–Ω–∏—è, —Ç–æ –ø—Ä–æ—Å—Ç–æ: –ì–ª–∞–≤–∞ X
–ê –µ—Å–ª–∏ –Ω–µ—Ç –≥–ª–∞–≤—ã, –Ω–æ –µ—Å—Ç—å –Ω–∞–∑–≤–∞–Ω–∏–µ, —Ç–æ –ø—Ä–æ—Å—Ç–æ –Ω–∞–∑–≤–∞–Ω–∏–µ –Ω–∞–¥–æ —Å –ø–µ—Ä–µ–≤–æ–¥–æ–º

**IV. –ì–õ–û–°–°–ê–†–ò–ô (–ï—Å–ª–∏ –ø—Ä–∏–º–µ–Ω–∏–º–æ)**

*   –ï—Å–ª–∏ –ø—Ä–µ–¥–æ—Å—Ç–∞–≤–ª–µ–Ω –≥–ª–æ—Å—Å–∞—Ä–∏–π –∏–º–µ–Ω, —Ç–µ—Ä–º–∏–Ω–æ–≤, –Ω–∞–∑–≤–∞–Ω–∏–π —Ç–µ—Ö–Ω–∏–∫ –∏ —Ç.–¥. ‚Äî **—Å—Ç—Ä–æ–≥–æ –ø—Ä–∏–¥–µ—Ä–∂–∏–≤–∞–π—Ç–µ—Å—å –µ–≥–æ**. –ü–æ—Å–ª–µ–¥–æ–≤–∞—Ç–µ–ª—å–Ω–æ—Å—Ç—å –∏ –µ–¥–∏–Ω–æ–æ–±—Ä–∞–∑–∏–µ –∫—Ä–∏—Ç–∏—á–Ω—ã.

**V. –ò–¢–û–ì–û–í–´–ô –†–ï–ó–£–õ–¨–¢–ê–¢**

*   –ü—Ä–µ–¥–æ—Å—Ç–∞–≤—å—Ç–µ **–ò–°–ö–õ–Æ–ß–ò–¢–ï–õ–¨–ù–û –ø–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω—ã–π –∏ –∞–¥–∞–ø—Ç–∏—Ä–æ–≤–∞–Ω–Ω—ã–π XHTML/HTML-–∫–æ–¥.**
*   **–ö–ê–¢–ï–ì–û–†–ò–ß–ï–°–ö–ò –ó–ê–ü–†–ï–©–ï–ù–û –≤–∫–ª—é—á–∞—Ç—å –≤ –≤—ã–≤–æ–¥ –æ—Ä–∏–≥–∏–Ω–∞–ª—å–Ω—ã–π —Ç–µ–∫—Å—Ç –∏–ª–∏ –ª—é–±—ã–µ –µ–≥–æ —Ñ—Ä–∞–≥–º–µ–Ω—Ç—ã.**
*   **–ù–ï –¥–æ–±–∞–≤–ª—è–π—Ç–µ –Ω–∏–∫–∞–∫–∏—Ö –≤–≤–æ–¥–Ω—ã—Ö —Ñ—Ä–∞–∑** —Ç–∏–ø–∞ "–í–æ—Ç –ø–µ—Ä–µ–≤–æ–¥:", "–ê–¥–∞–ø—Ç–∞—Ü–∏—è:", **–∞ —Ç–∞–∫–∂–µ –Ω–∏–∫–∞–∫–∏—Ö –∑–∞–∫–ª—é—á–∏—Ç–µ–ª—å–Ω—ã—Ö —Ñ—Ä–∞–∑ –∏–ª–∏ –∫–æ–º–º–µ–Ω—Ç–∞—Ä–∏–µ–≤** (–∫—Ä–æ–º–µ –Ω–µ–∏–∑–º–µ–Ω–µ–Ω–Ω—ã—Ö –∫–æ–º–º–µ–Ω—Ç–∞—Ä–∏–µ–≤ HTML).

**VI. –§–ò–ù–ê–õ–¨–ù–ê–Ø –ü–†–û–í–ï–†–ö–ê (–ú—ã—Å–ª–µ–Ω–Ω–æ –ø–µ—Ä–µ–¥ –≤—ã–≤–æ–¥–æ–º):**
*   –¢–µ–∫—Å—Ç –≤–Ω—É—Ç—Ä–∏ HTML-–∫–æ–¥–∞ –∑–≤—É—á–∏—Ç **–µ—Å—Ç–µ—Å—Ç–≤–µ–Ω–Ω–æ, —Ö—É–¥–æ–∂–µ—Å—Ç–≤–µ–Ω–Ω–æ –∏ —É–≤–ª–µ–∫–∞—Ç–µ–ª—å–Ω–æ** –ø–æ-—Ä—É—Å—Å–∫–∏?
*   –°–º—ã—Å–ª, —Ç–æ–Ω, **—ç–º–æ—Ü–∏–∏ –∏ –∞—Ç–º–æ—Å—Ñ–µ—Ä–∞** –æ—Ä–∏–≥–∏–Ω–∞–ª–∞ –ø–µ—Ä–µ–¥–∞–Ω—ã —Ç–æ—á–Ω–æ?
*   **XHTML-—Ç–µ–≥–∏ –∏ —Å—Ç—Ä—É–∫—Ç—É—Ä–∞ –¥–æ–∫—É–º–µ–Ω—Ç–∞** —Å–æ—Ö—Ä–∞–Ω–µ–Ω—ã –≤ —Ç–æ—á–Ω–æ—Å—Ç–∏?
*   –¢–æ–ª—å–∫–æ –≤–∏–¥–∏–º—ã–π –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—é —Ç–µ–∫—Å—Ç –ø–µ—Ä–µ–≤–µ–¥–µ–Ω, –∞ —Ç–µ–≥–∏, –∞—Ç—Ä–∏–±—É—Ç—ã, —Å–∫—Ä–∏–ø—Ç—ã, —Å—Ç–∏–ª–∏ –∏ –∫–æ–º–º–µ–Ω—Ç–∞—Ä–∏–∏ –Ω–µ —Ç—Ä–æ–Ω—É—Ç—ã?
*   **–í—Å–µ —Å—Ç–∏–ª–∏—Å—Ç–∏—á–µ—Å–∫–∏–µ –∏ –∫—É–ª—å—Ç—É—Ä–Ω—ã–µ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è (—Ä–∞–∑–¥–µ–ª—ã II –∏ III.2) —É—á—Ç–µ–Ω—ã?**
*   –í –∏—Ç–æ–≥–æ–≤–æ–º –∫–æ–¥–µ **–ü–û–õ–ù–û–°–¢–¨–Æ –û–¢–°–£–¢–°–¢–í–£–ï–¢ —Ç–µ–∫—Å—Ç –Ω–∞ —è–∑—ã–∫–µ –æ—Ä–∏–≥–∏–Ω–∞–ª–∞** (–∑–∞ –∏—Å–∫–ª—é—á–µ–Ω–∏–µ–º —Ä–∞–∑—Ä–µ—à–µ–Ω–Ω—ã—Ö –Ω–µ–∏–∑–º–µ–Ω—è–µ–º—ã—Ö —ç–ª–µ–º–µ–Ω—Ç–æ–≤)? **–ü–†–û–í–ï–†–ï–ù–û?**

--- PROMPT END ---"""
    
        self.custom_prompt_edit.setPlainText(default_prompt)
        QMessageBox.information(
            self,
            "–ü—Ä–æ–º–ø—Ç –∑–∞–≥—Ä—É–∂–µ–Ω",
            "–ó–∞–≥—Ä—É–∂–µ–Ω —É–ª—É—á—à–µ–Ω–Ω—ã–π —Å—Ç–∞–Ω–¥–∞—Ä—Ç–Ω—ã–π –ø—Ä–æ–º–ø—Ç —Å –ø—Ä–∞–≤–∏–ª–∞–º–∏ —Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏—è –¥–∏–∞–ª–æ–≥–æ–≤"
        )
        
    def update_keys_count(self):
        """–û–±–Ω–æ–≤–ª—è–µ—Ç —Å—á–µ—Ç—á–∏–∫ API –∫–ª—é—á–µ–π"""
        keys_text = self.keys_edit.toPlainText()
        keys = [k.strip() for k in keys_text.splitlines() if k.strip()]
        unique_keys = list(set(keys))
        
        if len(keys) != len(unique_keys):
            self.keys_count_label.setText(f"–ö–ª—é—á–µ–π: {len(unique_keys)} (—É–Ω–∏–∫–∞–ª—å–Ω—ã—Ö –∏–∑ {len(keys)})")
            self.keys_count_label.setStyleSheet("color: orange; font-size: 10px;")
        else:
            self.keys_count_label.setText(f"–ö–ª—é—á–µ–π: {len(keys)}")
            self.keys_count_label.setStyleSheet("color: blue; font-size: 10px;")
            
    def update_glossary_count(self):
        """–û–±–Ω–æ–≤–ª—è–µ—Ç —Å—á–µ—Ç—á–∏–∫ —Ç–µ—Ä–º–∏–Ω–æ–≤ –≤ –≥–ª–æ—Å—Å–∞—Ä–∏–∏"""
        glossary_dict = self.parse_glossary_text()
        self.glossary_count_label.setText(f"–¢–µ—Ä–º–∏–Ω–æ–≤: {len(glossary_dict)}")
        
    def parse_glossary_text(self):
        """–ü–∞—Ä—Å–∏—Ç —Ç–µ–∫—Å—Ç –≥–ª–æ—Å—Å–∞—Ä–∏—è –≤ —Å–ª–æ–≤–∞—Ä—å (–ø–æ–¥–¥–µ—Ä–∂–∏–≤–∞–µ—Ç JSON –∏ –æ–±—ã—á–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç)"""
        glossary_dict = {}
        text = self.glossary_text_edit.toPlainText().strip()
    
        if not text:
            return glossary_dict
    
        # –ü—Ä–æ–≤–µ—Ä—è–µ–º, —è–≤–ª—è–µ—Ç—Å—è –ª–∏ —ç—Ç–æ JSON
        if text.startswith('{') and text.endswith('}'):
            try:
                glossary_dict = json.loads(text)
                return glossary_dict
            except json.JSONDecodeError:
                # –ù–µ JSON, –æ–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º –∫–∞–∫ –æ–±—ã—á–Ω—ã–π —Ç–µ–∫—Å—Ç
                pass
    
        # –û–±—ã—á–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç: Term = Translation
        for line in text.splitlines():
            line = line.strip()
            if not line or '=' not in line:
                continue
            
            parts = line.split('=', 1)
            if len(parts) == 2:
                original = parts[0].strip()
                translation = parts[1].strip()
            
                if original and translation:
                    glossary_dict[original] = translation
                
        return glossary_dict
        
    def load_preset_glossary(self, preset_name):
        """–ó–∞–≥—Ä—É–∂–∞–µ—Ç –ø—Ä–µ–¥—É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω–Ω—ã–π –≥–ª–æ—Å—Å–∞—Ä–∏–π"""
        presets = {
            "–ê–Ω–∏–º–µ/–ú–∞–Ω–≥–∞": {
                "San": "—Å–∞–Ω",
                "Chan": "—á–∞–Ω", 
                "Kun": "–∫—É–Ω",
                "Sama": "—Å–∞–º–∞",
                "Senpai": "—Å—ç–º–ø–∞–π",
                "Kouhai": "–∫–æ—Ö–∞–π",
                "Sensei": "—Å—ç–Ω—Å—ç–π",
                "Onii-chan": "–±—Ä–∞—Ç–∏–∫",
                "Onee-chan": "—Å–µ—Å—Ç—Ä–∏—á–∫–∞",
                "Baka": "–¥—É—Ä–∞–∫",
                "Kawaii": "–º–∏–ª—ã–π",
                "Sugoi": "–ø–æ—Ç—Ä—è—Å–∞—é—â–µ"
            },
            "–ö–∏—Ç–∞–π—Å–∫–∏–µ –Ω–æ–≤–µ–ª–ª—ã": {
                "Cultivator": "–∫—É–ª—å—Ç–∏–≤–∞—Ç–æ—Ä",
                "Dao": "–î–∞–æ",
                "Qi": "—Ü–∏",
                "Immortal": "–±–µ—Å—Å–º–µ—Ä—Ç–Ω—ã–π",
                "Sect": "—Å–µ–∫—Ç–∞",
                "Elder": "—Å—Ç–∞—Ä–µ–π—à–∏–Ω–∞",
                "Junior": "–º–ª–∞–¥—à–∏–π",
                "Senior": "—Å—Ç–∞—Ä—à–∏–π",
                "Young Master": "–º–æ–ª–æ–¥–æ–π –≥–æ—Å–ø–æ–¥–∏–Ω",
                "Jade Beauty": "–Ω–µ—Ñ—Ä–∏—Ç–æ–≤–∞—è –∫—Ä–∞—Å–∞–≤–∏—Ü–∞"
            },
            "–ö–æ—Ä–µ–π—Å–∫–∏–µ –Ω–æ–≤–µ–ª–ª—ã": {
                "Oppa": "–æ–ø–ø–∞",
                "Hyung": "—Ö—ë–Ω",
                "Noona": "–Ω—É–Ω–∞",
                "Unnie": "–æ–Ω–Ω–∏",
                "Sunbae": "—Å–æ–Ω–±—ç",
                "Hoobae": "—Ö—É–±—ç",
                "Ahjussi": "–∞–¥–∂–æ—Å—Å–∏",
                "Ahjumma": "–∞–¥–∂—É–º–º–∞"
            },
            "–§—ç–Ω—Ç–µ–∑–∏": {
                "Mage": "–º–∞–≥",
                "Knight": "—Ä—ã—Ü–∞—Ä—å",
                "Dragon": "–¥—Ä–∞–∫–æ–Ω",
                "Elf": "—ç–ª—å—Ñ",
                "Dwarf": "–≥–Ω–æ–º",
                "Orc": "–æ—Ä–∫",
                "Spell": "–∑–∞–∫–ª–∏–Ω–∞–Ω–∏–µ",
                "Sword": "–º–µ—á",
                "Shield": "—â–∏—Ç",
                "Armor": "–¥–æ—Å–ø–µ—Ö–∏"
            },
            "–ù–∞—É—á–Ω–∞—è —Ñ–∞–Ω—Ç–∞—Å—Ç–∏–∫–∞": {
                "AI": "–ò–ò",
                "Cyborg": "–∫–∏–±–æ—Ä–≥",
                "Android": "–∞–Ω–¥—Ä–æ–∏–¥",
                "Spaceship": "–∫–æ—Å–º–∏—á–µ—Å–∫–∏–π –∫–æ—Ä–∞–±–ª—å",
                "Laser": "–ª–∞–∑–µ—Ä",
                "Quantum": "–∫–≤–∞–Ω—Ç–æ–≤—ã–π",
                "Warp": "–≤–∞—Ä–ø",
                "Hyperspace": "–≥–∏–ø–µ—Ä–ø—Ä–æ—Å—Ç—Ä–∞–Ω—Å—Ç–≤–æ"
            }
        }
        
        if preset_name in presets:
            # –î–æ–±–∞–≤–ª—è–µ–º –∫ —Å—É—â–µ—Å—Ç–≤—É—é—â–µ–º—É –≥–ª–æ—Å—Å–∞—Ä–∏—é
            current_text = self.glossary_text_edit.toPlainText()
            if current_text and not current_text.endswith('\n'):
                current_text += '\n'
                
            new_lines = []
            for original, translation in presets[preset_name].items():
                new_lines.append(f"{original} = {translation}")
                
            self.glossary_text_edit.setPlainText(current_text + '\n'.join(new_lines))
            
    def load_glossary_from_file(self):
        """–ó–∞–≥—Ä—É–∂–∞–µ—Ç –≥–ª–æ—Å—Å–∞—Ä–∏–π –∏–∑ —Ñ–∞–π–ª–∞"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "–í—ã–±–µ—Ä–∏—Ç–µ —Ñ–∞–π–ª –≥–ª–æ—Å—Å–∞—Ä–∏—è",
            "",
            "All supported (*.json *.txt);;JSON files (*.json);;Text files (*.txt)"
        )
        
        if not file_path:
            return
            
        try:
            if file_path.lower().endswith('.json'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    glossary_data = json.load(f)
                    
                lines = []
                for key, value in glossary_data.items():
                    lines.append(f"{key} = {value}")
                    
                self.glossary_text_edit.setPlainText('\n'.join(lines))
                QMessageBox.information(
                    self,
                    "–£—Å–ø–µ—Ö",
                    f"–ó–∞–≥—Ä—É–∂–µ–Ω–æ {len(glossary_data)} —Ç–µ—Ä–º–∏–Ω–æ–≤ –∏–∑ JSON"
                )
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                self.glossary_text_edit.setPlainText(content)
                
                lines = [line.strip() for line in content.splitlines() if '=' in line.strip()]
                QMessageBox.information(
                    self,
                    "–£—Å–ø–µ—Ö", 
                    f"–ó–∞–≥—Ä—É–∂–µ–Ω–æ {len(lines)} —Ç–µ—Ä–º–∏–Ω–æ–≤ –∏–∑ —Ç–µ–∫—Å—Ç–æ–≤–æ–≥–æ —Ñ–∞–π–ª–∞"
                )
                
        except Exception as e:
            QMessageBox.critical(self, "–û—à–∏–±–∫–∞", f"–ù–µ —É–¥–∞–ª–æ—Å—å –∑–∞–≥—Ä—É–∑–∏—Ç—å —Ñ–∞–π–ª: {e}")
            
    def save_glossary_to_file(self):
        """–°–æ—Ö—Ä–∞–Ω—è–µ—Ç –≥–ª–æ—Å—Å–∞—Ä–∏–π –≤ —Ñ–∞–π–ª"""
        glossary_dict = self.parse_glossary_text()
        
        if not glossary_dict:
            QMessageBox.warning(
                self,
                "–ü—Ä–µ–¥—É–ø—Ä–µ–∂–¥–µ–Ω–∏–µ",
                "–ì–ª–æ—Å—Å–∞—Ä–∏–π –ø—É—Å—Ç –∏–ª–∏ –Ω–µ–ø—Ä–∞–≤–∏–ª—å–Ω–æ –æ—Ç—Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω"
            )
            return
            
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "–°–æ—Ö—Ä–∞–Ω–∏—Ç—å –≥–ª–æ—Å—Å–∞—Ä–∏–π",
            "glossary.json",
            "JSON files (*.json);;Text files (*.txt)"
        )
        
        if not file_path:
            return
            
        try:
            if file_path.lower().endswith('.json'):
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(glossary_dict, f, ensure_ascii=False, indent=2)
            else:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(self.glossary_text_edit.toPlainText())
                    
            QMessageBox.information(
                self,
                "–£—Å–ø–µ—Ö",
                f"–ì–ª–æ—Å—Å–∞—Ä–∏–π —Å–æ—Ö—Ä–∞–Ω–µ–Ω ({len(glossary_dict)} —Ç–µ—Ä–º–∏–Ω–æ–≤)"
            )
            
        except Exception as e:
            QMessageBox.critical(self, "–û—à–∏–±–∫–∞", f"–ù–µ —É–¥–∞–ª–æ—Å—å —Å–æ—Ö—Ä–∞–Ω–∏—Ç—å —Ñ–∞–π–ª: {e}")
            
    def select_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "–í—ã–±–µ—Ä–∏—Ç–µ —Ñ–∞–π–ª –¥–ª—è –ø–µ—Ä–µ–≤–æ–¥–∞",
            "",
            "EPUB —Ñ–∞–π–ª—ã (*.epub);;All files (*)"
        )
        if file_path:
            self.selected_file = file_path
            self.file_label.setText(os.path.basename(file_path))
            
    def select_output_folder(self):
        folder = QFileDialog.getExistingDirectory(
            self,
            "–í—ã–±–µ—Ä–∏—Ç–µ –ø–∞–ø–∫—É –¥–ª—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è –ø–µ—Ä–µ–≤–æ–¥–∞"
        )
        if folder:
            self.output_folder = folder
            self.output_label.setText(folder)
            
    def load_keys_from_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "–í—ã–±–µ—Ä–∏—Ç–µ —Ñ–∞–π–ª —Å –∫–ª—é—á–∞–º–∏",
            "",
            "Text files (*.txt)"
        )
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    keys = [line.strip() for line in f if line.strip()]
                self.keys_edit.setPlainText('\n'.join(keys))
            except Exception as e:
                QMessageBox.warning(self, "–û—à–∏–±–∫–∞", f"–ù–µ —É–¥–∞–ª–æ—Å—å –∑–∞–≥—Ä—É–∑–∏—Ç—å —Ñ–∞–π–ª: {e}")
                
    def validate_and_start(self):
        # –ü—Ä–æ–≤–µ—Ä–∫–∞ –æ–±—è–∑–∞—Ç–µ–ª—å–Ω—ã—Ö –ø–æ–ª–µ–π
        if not self.selected_file:
            QMessageBox.warning(self, "–û—à–∏–±–∫–∞", "–í—ã–±–µ—Ä–∏—Ç–µ —Ñ–∞–π–ª –¥–ª—è –ø–µ—Ä–µ–≤–æ–¥–∞")
            return
            
        if not self.output_folder:
            QMessageBox.warning(self, "–û—à–∏–±–∫–∞", "–í—ã–±–µ—Ä–∏—Ç–µ –ø–∞–ø–∫—É –¥–ª—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è")
            return
            
        keys_text = self.keys_edit.toPlainText()
        self.api_keys = [k.strip() for k in keys_text.splitlines() if k.strip()]
        
        if not self.api_keys:
            QMessageBox.warning(self, "–û—à–∏–±–∫–∞", "–í–≤–µ–¥–∏—Ç–µ —Ö–æ—Ç—è –±—ã –æ–¥–∏–Ω API –∫–ª—é—á")
            return
            
        # –ü–∞—Ä—Å–∏–º –≥–ª–æ—Å—Å–∞—Ä–∏–π
        self.glossary_dict = self.parse_glossary_text()
        
        self.selected_model = self.model_combo.currentText()
        self.concurrent_requests = self.concurrency_spin.value()
        
        self.accept()
        
    def get_settings(self):
        """–í–æ–∑–≤—Ä–∞—â–∞–µ—Ç –Ω–∞—Å—Ç—Ä–æ–π–∫–∏ –≤ –≤–∏–¥–µ —Å–ª–æ–≤–∞—Ä—è"""
        return {
            'file_path': self.selected_file,
            'output_folder': self.output_folder,
            'api_keys': self.api_keys,
            'glossary_dict': self.glossary_dict,
            'model': self.selected_model,
            'dynamic_glossary': self.dynamic_glossary_checkbox.isChecked(),
            'custom_prompt': self.custom_prompt_edit.toPlainText().strip() if hasattr(self, 'custom_prompt_edit') else None,
            'concurrent_requests': self.concurrent_requests
        }

class TranslationSessionManager:
    """–ú–µ–Ω–µ–¥–∂–µ—Ä —Å–µ—Å—Å–∏–∏ –ø–µ—Ä–µ–≤–æ–¥–∞ –¥–ª—è –æ—Ç—Å–ª–µ–∂–∏–≤–∞–Ω–∏—è –ø—Ä–æ–≥—Ä–µ—Å—Å–∞ –º–µ–∂–¥—É –ø–µ—Ä–µ–∑–∞–ø—É—Å–∫–∞–º–∏."""
    def __init__(self, session_file_path):
        self.session_file_path = session_file_path
        self.session_data = {
            'original_file': None,
            'output_folder': None,
            'total_files': 0,
            'completed_files': [],
            'failed_files': [],
            'content_filtered_files': [],  # –û—Ç–¥–µ–ª—å–Ω—ã–π —Å–ø–∏—Å–æ–∫ –¥–ª—è –∑–∞–±–ª–æ–∫–∏—Ä–æ–≤–∞–Ω–Ω—ã—Ö —Ñ–∏–ª—å—Ç—Ä–∞–º–∏
            'current_key_index': 0,
            'api_keys': [],
            'model': DEFAULT_MODEL_NAME,
            'concurrent_requests': 10,
            'glossary_dict': {},
            'glossary_path': None,  # –î–ª—è –æ–±—Ä–∞—Ç–Ω–æ–π —Å–æ–≤–º–µ—Å—Ç–∏–º–æ—Å—Ç–∏
            'file_type': None,
            'epub_html_files': []
        }
        
    def init_new_session(self, settings):
        """–ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä—É–µ—Ç –Ω–æ–≤—É—é —Å–µ—Å—Å–∏—é —Å –Ω–∞—á–∞–ª—å–Ω—ã–º–∏ –Ω–∞—Å—Ç—Ä–æ–π–∫–∞–º–∏"""
        self.session_data['original_file'] = settings['file_path']
        self.session_data['output_folder'] = settings['output_folder']
        self.session_data['api_keys'] = settings['api_keys']
        self.session_data['current_key_index'] = 0
        self.session_data['model'] = settings['model']
        self.session_data['concurrent_requests'] = settings['concurrent_requests']
    
        # –°–æ—Ö—Ä–∞–Ω—è–µ–º –í–°–ï –Ω–∞—Å—Ç—Ä–æ–π–∫–∏
        if 'glossary_dict' in settings:
            self.session_data['glossary_dict'] = settings['glossary_dict']
        elif 'glossary_path' in settings:
            self.session_data['glossary_path'] = settings['glossary_path']
    
        # –°–æ—Ö—Ä–∞–Ω—è–µ–º –¥–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω—ã–µ –Ω–∞—Å—Ç—Ä–æ–π–∫–∏
        if 'dynamic_glossary' in settings:
            self.session_data['dynamic_glossary'] = settings['dynamic_glossary']
        if 'custom_prompt' in settings:
            self.session_data['custom_prompt'] = settings['custom_prompt']
        
        self.session_data['completed_files'] = []
        self.session_data['failed_files'] = []
        self.session_data['content_filtered_files'] = []
    
        # –û–ø—Ä–µ–¥–µ–ª—è–µ–º —Ç–∏–ø —Ñ–∞–π–ª–∞ –∏ —Å–ø–∏—Å–æ–∫ –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏
        file_ext = os.path.splitext(settings['file_path'])[1].lower()
        self.session_data['file_type'] = file_ext[1:]  # –±–µ–∑ —Ç–æ—á–∫–∏
    
        if file_ext == '.epub':
            # –î–ª—è EPUB –Ω—É–∂–Ω–æ –ø–æ–ª—É—á–∏—Ç—å —Å–ø–∏—Å–æ–∫ HTML —Ñ–∞–π–ª–æ–≤
            html_files = self._get_epub_html_files(settings['file_path'])
            self.session_data['epub_html_files'] = html_files
            self.session_data['total_files'] = len(html_files)
        else:
            self.session_data['total_files'] = 1
        
        self.save_session()
        
    def _get_epub_html_files(self, epub_path):
        """–ò–∑–≤–ª–µ–∫–∞–µ—Ç —Å–ø–∏—Å–æ–∫ HTML —Ñ–∞–π–ª–æ–≤ –∏–∑ EPUB."""
        html_files = []
        try:
            with zipfile.ZipFile(epub_path, 'r') as epub_zip:
                for name in epub_zip.namelist():
                    if name.endswith(('.html', '.xhtml', '.htm')) and not name.startswith('META-INF/'):
                        html_files.append(name)
            return sorted(html_files, key=extract_number_from_path)
        except Exception as e:
            print(f"–û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è EPUB: {e}")
            return []
            
    def save_session(self):
        """–°–æ—Ö—Ä–∞–Ω—è–µ—Ç —Ç–µ–∫—É—â–µ–µ —Å–æ—Å—Ç–æ—è–Ω–∏–µ —Å–µ—Å—Å–∏–∏."""
        try:
            with open(self.session_file_path, 'w', encoding='utf-8') as f:
                json.dump(self.session_data, f, ensure_ascii=False, indent=2)
            return True
        except Exception as e:
            print(f"–û—à–∏–±–∫–∞ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è —Å–µ—Å—Å–∏–∏: {e}")
            return False
            
    def load_session(self):
        """–ó–∞–≥—Ä—É–∂–∞–µ—Ç —Å–æ—Ö—Ä–∞–Ω–µ–Ω–Ω—É—é —Å–µ—Å—Å–∏—é."""
        try:
            if os.path.exists(self.session_file_path):
                with open(self.session_file_path, 'r', encoding='utf-8') as f:
                    loaded_data = json.load(f)
                    self.session_data.update(loaded_data)
                return True
            return False
        except Exception as e:
            print(f"–û—à–∏–±–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ —Å–µ—Å—Å–∏–∏: {e}")
            return False
        
    def mark_file_completed(self, file_tuple):
        """–û—Ç–º–µ—á–∞–µ—Ç —Ñ–∞–π–ª –∫–∞–∫ —É—Å–ø–µ—à–Ω–æ –æ–±—Ä–∞–±–æ—Ç–∞–Ω–Ω—ã–π."""
        if file_tuple not in self.session_data['completed_files']:
            self.session_data['completed_files'].append(file_tuple)
            self.save_session()
            
    def mark_file_failed(self, file_tuple, error_msg):
        """–û—Ç–º–µ—á–∞–µ—Ç —Ñ–∞–π–ª –∫–∞–∫ –Ω–µ—É–¥–∞—á–Ω–æ –æ–±—Ä–∞–±–æ—Ç–∞–Ω–Ω—ã–π."""
        self.session_data['failed_files'].append({
            'file': file_tuple,
            'error': error_msg,
            'timestamp': time.strftime('%Y-%m-%d %H:%M:%S')
        })
        self.save_session()
        
    def mark_file_content_filtered(self, file_tuple, error_msg):
        """–û—Ç–º–µ—á–∞–µ—Ç —Ñ–∞–π–ª –∫–∞–∫ –∑–∞–±–ª–æ–∫–∏—Ä–æ–≤–∞–Ω–Ω—ã–π —Ñ–∏–ª—å—Ç—Ä–∞–º–∏ –∫–æ–Ω—Ç–µ–Ω—Ç–∞"""
        # –ü—Ä–æ–≤–µ—Ä—è–µ–º, –Ω–µ –¥–æ–±–∞–≤–ª–µ–Ω –ª–∏ —É–∂–µ
        if not any(f['file'] == file_tuple for f in self.session_data['content_filtered_files']):
            self.session_data['content_filtered_files'].append({
                'file': file_tuple,
                'error': error_msg,
                'timestamp': time.strftime('%Y-%m-%d %H:%M:%S')
            })
            self.save_session()
            
    def is_content_filtered(self, file_tuple):
        """–ü—Ä–æ–≤–µ—Ä—è–µ—Ç, –±—ã–ª–∞ –ª–∏ –≥–ª–∞–≤–∞ –∑–∞–±–ª–æ–∫–∏—Ä–æ–≤–∞–Ω–∞ —Ñ–∏–ª—å—Ç—Ä–∞–º–∏"""
        return any(f['file'] == file_tuple for f in self.session_data['content_filtered_files'])
        
    def get_pending_files(self):
        """–í–æ–∑–≤—Ä–∞—â–∞–µ—Ç —Å–ø–∏—Å–æ–∫ —Ñ–∞–π–ª–æ–≤ –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏ (–µ—â–µ –Ω–µ –∑–∞–≤–µ—Ä—à–µ–Ω–Ω—ã—Ö –∏ –Ω–µ –∑–∞–±–ª–æ–∫–∏—Ä–æ–≤–∞–Ω–Ω—ã—Ö)"""
        if self.session_data['file_type'] == 'epub':
            pending = []
            for html_file in self.session_data['epub_html_files']:
                file_tuple = ('epub', self.session_data['original_file'], html_file)
            
                # –ü—Ä–æ–ø—É—Å–∫–∞–µ–º —É–∂–µ —É—Å–ø–µ—à–Ω–æ –æ–±—Ä–∞–±–æ—Ç–∞–Ω–Ω—ã–µ
                if file_tuple in self.session_data['completed_files']:
                    continue
                
                # –ü—Ä–æ–ø—É—Å–∫–∞–µ–º –∑–∞–±–ª–æ–∫–∏—Ä–æ–≤–∞–Ω–Ω—ã–µ —Ñ–∏–ª—å—Ç—Ä–∞–º–∏
                if self.is_content_filtered(file_tuple):
                    continue
                
                # === –ù–û–í–û–ï: –ü—Ä–æ–ø—É—Å–∫–∞–µ–º –≥–ª–∞–≤—ã —Å —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏–º–∏ –æ—à–∏–±–∫–∞–º–∏ (500/503) ===
                # –û–Ω–∏ –ù–ï –¥–æ–ª–∂–Ω—ã –ø–æ–≤—Ç–æ—Ä—è—Ç—å—Å—è –ø—Ä–∏ —Å–º–µ–Ω–µ –∫–ª—é—á–∞
                is_technical_error = False
                for failed_entry in self.session_data['failed_files']:
                    if failed_entry.get('file') == file_tuple:
                        error = failed_entry.get('error', '').lower()
                        if '500' in error or '503' in error or 'internal server error' in error:
                            is_technical_error = True
                            break
                        
                if is_technical_error:
                    continue
                
                # –î–æ–±–∞–≤–ª—è–µ–º –≤ —Å–ø–∏—Å–æ–∫ –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏
                pending.append(file_tuple)
            
            return pending
        else:
            file_type = self.session_data['file_type']
            file_tuple = (file_type, self.session_data['original_file'], None)
        
            if file_tuple in self.session_data['completed_files']:
                return []
            if self.is_content_filtered(file_tuple):
                return []
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º –Ω–∞ —Ç–µ—Ö–Ω–∏—á–µ—Å–∫–∏–µ –æ—à–∏–±–∫–∏
            for failed_entry in self.session_data['failed_files']:
                if failed_entry.get('file') == file_tuple:
                    error = failed_entry.get('error', '').lower()
                    if '500' in error or '503' in error or 'internal server error' in error:
                        return []
                    
            return [file_tuple]
            
    def is_rate_limited(self, error_msg):
        """–ü—Ä–æ–≤–µ—Ä—è–µ—Ç, —è–≤–ª—è–µ—Ç—Å—è –ª–∏ –æ—à–∏–±–∫–∞ –ø—Ä–µ–≤—ã—à–µ–Ω–∏–µ–º –ª–∏–º–∏—Ç–∞ –∑–∞–ø—Ä–æ—Å–æ–≤."""
        rate_limit_indicators = [
            '429',
            'rate limit',
            'ResourceExhausted',
            'too many requests'
        ]
        # –ù–ï –≤–∫–ª—é—á–∞–µ–º 'quota exceeded' –∏ 'exceeded your current quota' —Å—é–¥–∞
        # —Ç–∞–∫ –∫–∞–∫ —ç—Ç–æ –æ—Ç–¥–µ–ª—å–Ω—ã–π —Å–ª—É—á–∞–π - –ø–æ–ª–Ω–æ–µ –∏—Å—á–µ—Ä–ø–∞–Ω–∏–µ –∫–≤–æ—Ç—ã
        error_lower = str(error_msg).lower()
        return any(indicator in error_lower for indicator in rate_limit_indicators) and \
               'exceeded your current quota' not in error_lower and \
               'quota exceeded' not in error_lower

    def is_quota_exceeded(self, error_msg):
        """–ü—Ä–æ–≤–µ—Ä—è–µ—Ç, –ø—Ä–µ–≤—ã—à–µ–Ω–∞ –ª–∏ –∫–≤–æ—Ç–∞ API –∫–ª—é—á–∞ –ø–æ–ª–Ω–æ—Å—Ç—å—é"""
        quota_indicators = [
            'exceeded your current quota',
            'quota exceeded',
            'out of quota'
        ]
        error_lower = str(error_msg).lower()
        return any(indicator in error_lower for indicator in quota_indicators)
        
    def is_content_filter_error(self, error_msg):
        """–ü—Ä–æ–≤–µ—Ä—è–µ—Ç, —è–≤–ª—è–µ—Ç—Å—è –ª–∏ –æ—à–∏–±–∫–∞ –±–ª–æ–∫–∏—Ä–æ–≤–∫–æ–π –∫–æ–Ω—Ç–µ–Ω—Ç–∞"""
        content_filter_indicators = [
            'PROHIBITED_CONTENT',
            'block_reason',
            'content filter',
            'blocked prompt',
            'safety',
            'harmful',
            'inappropriate',
            'BlockedPromptException',
            'StopCandidateException'
        ]
        error_lower = str(error_msg).lower()
        return any(indicator.lower() in error_lower for indicator in content_filter_indicators)
        
    def get_progress(self):
        """–í–æ–∑–≤—Ä–∞—â–∞–µ—Ç —Ç–µ–∫—É—â–∏–π –ø—Ä–æ–≥—Ä–µ—Å—Å."""
        completed = len(self.session_data['completed_files'])
        filtered = len(self.session_data['content_filtered_files'])
        total = self.session_data['total_files']
        return completed, filtered, total

class EpubCreator:
    """–°–æ–∑–¥–∞–µ—Ç EPUB —Ñ–∞–π–ª –≤–µ—Ä—Å–∏–∏ 2 –∏–∑ HTML –≥–ª–∞–≤."""
    def __init__(self, title, author="Unknown", language="ru"):
        self.title = title
        self.author = author
        self.language = language
        self.chapters = []
        self.uuid = str(uuid.uuid4())

    def add_chapter(self, filename, content, title):
        """–î–æ–±–∞–≤–ª—è–µ—Ç –≥–ª–∞–≤—É –≤ –∫–Ω–∏–≥—É."""
        self.chapters.append({
            'filename': filename,
            'content': content,
            'title': title,
            'id': f'chapter{len(self.chapters) + 1}'
        })

    def create_epub(self, output_path):
        """–°–æ–∑–¥–∞–µ—Ç EPUB —Ñ–∞–π–ª."""
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as epub:
            # –û–±—è–∑–∞—Ç–µ–ª—å–Ω—ã–π —Ñ–∞–π–ª mimetype
            epub.writestr('mimetype', 'application/epub+zip', compress_type=zipfile.ZIP_STORED)
            
            # META-INF/container.xml
            epub.writestr('META-INF/container.xml', self._create_container())
            
            # OEBPS/content.opf
            epub.writestr('OEBPS/content.opf', self._create_opf())
            
            # OEBPS/toc.ncx
            epub.writestr('OEBPS/toc.ncx', self._create_ncx())
            
            # OEBPS/styles.css
            epub.writestr('OEBPS/styles.css', self._create_styles())
            
            # –ì–ª–∞–≤—ã
            for chapter in self.chapters:
                chapter_content = f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>{chapter['title']}</title>
    <link rel="stylesheet" type="text/css" href="styles.css"/>
</head>
<body>
{chapter['content']}
</body>
</html>'''
                epub.writestr(f"OEBPS/{chapter['filename']}", chapter_content)

    def _create_container(self):
        return '''<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
    <rootfiles>
        <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
    </rootfiles>
</container>'''

    def _create_opf(self):
        opf = f'''<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="BookID" version="2.0">
    <metadata xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:opf="http://www.idpf.org/2007/opf">
        <dc:title>{self.title}</dc:title>
        <dc:creator>{self.author}</dc:creator>
        <dc:language>{self.language}</dc:language>
        <dc:identifier id="BookID">urn:uuid:{self.uuid}</dc:identifier>
    </metadata>
    <manifest>
        <item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>
        <item id="styles" href="styles.css" media-type="text/css"/>'''
        
        for chapter in self.chapters:
            opf += f'\n        <item id="{chapter["id"]}" href="{chapter["filename"]}" media-type="application/xhtml+xml"/>'
            
        opf += '\n    </manifest>\n    <spine toc="ncx">'
        
        for chapter in self.chapters:
            opf += f'\n        <itemref idref="{chapter["id"]}"/>'
            
        opf += '\n    </spine>\n</package>'
        return opf

    def _create_ncx(self):
        ncx = f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE ncx PUBLIC "-//NISO//DTD ncx 2005-1//EN" "http://www.daisy.org/z3986/2005/ncx-2005-1.dtd">
<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
    <head>
        <meta name="dtb:uid" content="urn:uuid:{self.uuid}"/>
        <meta name="dtb:depth" content="1"/>
        <meta name="dtb:totalPageCount" content="0"/>
        <meta name="dtb:maxPageNumber" content="0"/>
    </head>
    <docTitle><text>{self.title}</text></docTitle>
    <navMap>'''
    
        for i, chapter in enumerate(self.chapters):
            ncx += f'''
        <navPoint id="navPoint-{i+1}" playOrder="{i+1}">
            <navLabel><text>{chapter['title']}</text></navLabel>
            <content src="{chapter['filename']}"/>
        </navPoint>'''
        
        ncx += '\n    </navMap>\n</ncx>'
        return ncx

    def _create_styles(self):
        return '''body { font-family: Georgia, serif; } p { text-indent: 1.5em; margin: 0; }'''

class TranslatedChaptersManagerDialog(QDialog):
    """–î–∏–∞–ª–æ–≥ –¥–ª—è —É–ø—Ä–∞–≤–ª–µ–Ω–∏—è –ø–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω—ã–º–∏ –≥–ª–∞–≤–∞–º–∏ –∏ —Å–æ–∑–¥–∞–Ω–∏—è EPUB."""
    def __init__(self, translated_folder, parent=None):
        super().__init__(parent)
        self.translated_folder = translated_folder
        self.chapters_data = []
        self.setWindowTitle("–ú–µ–Ω–µ–¥–∂–µ—Ä EPUB")
        self.setMinimumSize(800, 600)
        self.init_ui()
        self.load_chapters()

    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.addWidget(QLabel("–£–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –ø–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω—ã–º–∏ –≥–ª–∞–≤–∞–º–∏ –¥–ª—è —Å–±–æ—Ä–∫–∏ EPUB:"))
        
        self.table = QTableWidget()
        self.table.setColumnCount(3)
        self.table.setHorizontalHeaderLabels(["–ü–æ—Ä—è–¥–æ–∫", "–ò–º—è —Ñ–∞–π–ª–∞", "–î–µ–π—Å—Ç–≤–∏—è"])
        
        # –†–∞—Å—Ç—è–≥–∏–≤–∞–µ–º –∫–æ–ª–æ–Ω–∫—É —Å –∏–º–µ–Ω–µ–º —Ñ–∞–π–ª–∞
        header = self.table.horizontalHeader()
        header.setSectionResizeMode(1, header.ResizeMode.Stretch)
        
        self.table.setSelectionBehavior(self.table.SelectionBehavior.SelectRows)
        self.table.setEditTriggers(self.table.EditTrigger.NoEditTriggers)
        layout.addWidget(self.table)
        
        # –ö–Ω–æ–ø–∫–∏ —É–ø—Ä–∞–≤–ª–µ–Ω–∏—è
        buttons_layout = QHBoxLayout()
        
        add_btn = QPushButton("‚ûï –î–æ–±–∞–≤–∏—Ç—å —Ñ–∞–π–ª")
        add_btn.clicked.connect(self.add_external_file)
        buttons_layout.addWidget(add_btn)
        
        replace_btn = QPushButton("üîÑ –ó–∞–º–µ–Ω–∏—Ç—å")
        replace_btn.clicked.connect(self.replace_selected_file)
        buttons_layout.addWidget(replace_btn)
        
        delete_btn = QPushButton("üóëÔ∏è –£–¥–∞–ª–∏—Ç—å")
        delete_btn.clicked.connect(self.delete_selected_file)
        buttons_layout.addWidget(delete_btn)
        
        view_btn = QPushButton("üëÅÔ∏è –ü—Ä–æ—Å–º–æ—Ç—Ä")
        view_btn.clicked.connect(lambda: self.view_chapter(self.table.currentRow()))
        buttons_layout.addWidget(view_btn)
        
        buttons_layout.addStretch()
        
        create_epub_btn = QPushButton("üìö –°–æ–∑–¥–∞—Ç—å EPUB")
        create_epub_btn.clicked.connect(self.create_epub)
        create_epub_btn.setStyleSheet("font-weight: bold; background-color: #4CAF50; color: white;")
        buttons_layout.addWidget(create_epub_btn)
        
        layout.addLayout(buttons_layout)
        
        # –ö–Ω–æ–ø–∫–∏ –¥–∏–∞–ª–æ–≥–∞
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)

    def load_chapters(self):
        """–ó–∞–≥—Ä—É–∂–∞–µ—Ç —Å–ø–∏—Å–æ–∫ –ø–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω—ã—Ö —Ñ–∞–π–ª–æ–≤"""
        if not os.path.exists(self.translated_folder):
            return
            
        # –ò—â–µ–º HTML —Ñ–∞–π–ª—ã
        html_files = []
        for ext in ['*.html', '*.xhtml', '*.htm']:
            html_files.extend(glob.glob(os.path.join(self.translated_folder, ext)))
        
        # –°–æ—Ä—Ç–∏—Ä—É–µ–º –ø–æ –Ω–æ–º–µ—Ä—É –≥–ª–∞–≤—ã
        html_files.sort(key=extract_number_from_path)
        
        self.chapters_data = []
        for i, file_path in enumerate(html_files):
            filename = os.path.basename(file_path)
            title = f"–ì–ª–∞–≤–∞ {i+1}"
            
            # –ü—ã—Ç–∞–µ–º—Å—è –∏–∑–≤–ª–µ—á—å –∑–∞–≥–æ–ª–æ–≤–æ–∫ –∏–∑ —Ñ–∞–π–ª–∞
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    # –ò—â–µ–º –∑–∞–≥–æ–ª–æ–≤–æ–∫ –≤ —Ç–µ–≥–∞—Ö h1, h2 –∏–ª–∏ title
                    import re
                    title_match = re.search(r'<(?:h[1-6]|title)[^>]*>([^<]+)</(?:h[1-6]|title)>', content, re.IGNORECASE)
                    if title_match:
                        title = title_match.group(1).strip()
            except:
                pass
                
            self.chapters_data.append({
                'file_path': file_path,
                'filename': filename,
                'title': title,
                'order': i
            })
            
        self.update_table()
        
    def update_table(self):
        """–û–±–Ω–æ–≤–ª—è–µ—Ç —Ç–∞–±–ª–∏—Ü—É"""
        self.table.setRowCount(len(self.chapters_data))
        
        for row, chapter in enumerate(self.chapters_data):
            # –ü–æ—Ä—è–¥–∫–æ–≤—ã–π –Ω–æ–º–µ—Ä
            order_item = QTableWidgetItem(str(chapter['order'] + 1))
            self.table.setItem(row, 0, order_item)
            
            # –ò–º—è —Ñ–∞–π–ª–∞
            filename_item = QTableWidgetItem(f"{chapter['title']} ({chapter['filename']})")
            self.table.setItem(row, 1, filename_item)
            
            # –ö–Ω–æ–ø–∫–∏ –¥–µ–π—Å—Ç–≤–∏–π
            actions_widget = QWidget()
            actions_layout = QHBoxLayout(actions_widget)
            actions_layout.setContentsMargins(5, 0, 5, 0)
            
            up_btn = QPushButton("‚Üë")
            up_btn.setMaximumWidth(30)
            up_btn.clicked.connect(lambda checked, r=row: self.move_chapter(r, -1))
            actions_layout.addWidget(up_btn)
            
            down_btn = QPushButton("‚Üì") 
            down_btn.setMaximumWidth(30)
            down_btn.clicked.connect(lambda checked, r=row: self.move_chapter(r, 1))
            actions_layout.addWidget(down_btn)
            
            actions_layout.addStretch()
            self.table.setCellWidget(row, 2, actions_widget)

    def move_chapter(self, row, direction):
        """–ü–µ—Ä–µ–º–µ—â–∞–µ—Ç –≥–ª–∞–≤—É –≤–≤–µ—Ä—Ö –∏–ª–∏ –≤–Ω–∏–∑"""
        if direction == -1 and row > 0:
            # –ü–µ—Ä–µ–º–µ—â–∞–µ–º –≤–≤–µ—Ä—Ö
            self.chapters_data[row], self.chapters_data[row-1] = self.chapters_data[row-1], self.chapters_data[row]
        elif direction == 1 and row < len(self.chapters_data) - 1:
            # –ü–µ—Ä–µ–º–µ—â–∞–µ–º –≤–Ω–∏–∑
            self.chapters_data[row], self.chapters_data[row+1] = self.chapters_data[row+1], self.chapters_data[row]
        
        # –û–±–Ω–æ–≤–ª—è–µ–º –ø–æ—Ä—è–¥–∫–æ–≤—ã–µ –Ω–æ–º–µ—Ä–∞
        for i, chapter in enumerate(self.chapters_data):
            chapter['order'] = i
            
        self.update_table()

    def add_external_file(self):
        """–î–æ–±–∞–≤–ª—è–µ—Ç –≤–Ω–µ—à–Ω–∏–π —Ñ–∞–π–ª"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "–í—ã–±–µ—Ä–∏—Ç–µ HTML —Ñ–∞–π–ª",
            "",
            "HTML files (*.html *.xhtml *.htm)"
        )
        
        if file_path:
            filename = os.path.basename(file_path)
            title = f"–ì–ª–∞–≤–∞ {len(self.chapters_data) + 1}"
            
            self.chapters_data.append({
                'file_path': file_path,
                'filename': filename,
                'title': title,
                'order': len(self.chapters_data)
            })
            
            self.update_table()

    def replace_selected_file(self):
        """–ó–∞–º–µ–Ω—è–µ—Ç –≤—ã–±—Ä–∞–Ω–Ω—ã–π —Ñ–∞–π–ª"""
        current_row = self.table.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "–ü—Ä–µ–¥—É–ø—Ä–µ–∂–¥–µ–Ω–∏–µ", "–í—ã–±–µ—Ä–∏—Ç–µ –≥–ª–∞–≤—É –¥–ª—è –∑–∞–º–µ–Ω—ã")
            return
            
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "–í—ã–±–µ—Ä–∏—Ç–µ –Ω–æ–≤—ã–π HTML —Ñ–∞–π–ª",
            "",
            "HTML files (*.html *.xhtml *.htm)"
        )
        
        if file_path:
            self.chapters_data[current_row]['file_path'] = file_path
            self.chapters_data[current_row]['filename'] = os.path.basename(file_path)
            self.update_table()

    def delete_selected_file(self):
        """–£–¥–∞–ª—è–µ—Ç –≤—ã–±—Ä–∞–Ω–Ω—ã–π —Ñ–∞–π–ª"""
        current_row = self.table.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "–ü—Ä–µ–¥—É–ø—Ä–µ–∂–¥–µ–Ω–∏–µ", "–í—ã–±–µ—Ä–∏—Ç–µ –≥–ª–∞–≤—É –¥–ª—è —É–¥–∞–ª–µ–Ω–∏—è")
            return
            
        reply = QMessageBox.question(
            self,
            "–ü–æ–¥—Ç–≤–µ—Ä–∂–¥–µ–Ω–∏–µ",
            "–£–¥–∞–ª–∏—Ç—å –≤—ã–±—Ä–∞–Ω–Ω—É—é –≥–ª–∞–≤—É –∏–∑ —Å–ø–∏—Å–∫–∞?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            del self.chapters_data[current_row]
            # –û–±–Ω–æ–≤–ª—è–µ–º –ø–æ—Ä—è–¥–∫–æ–≤—ã–µ –Ω–æ–º–µ—Ä–∞
            for i, chapter in enumerate(self.chapters_data):
                chapter['order'] = i
            self.update_table()

    def view_chapter(self, index):
        """–ü—Ä–æ—Å–º–∞—Ç—Ä–∏–≤–∞–µ—Ç —Å–æ–¥–µ—Ä–∂–∏–º–æ–µ –≥–ª–∞–≤—ã"""
        if index < 0 or index >= len(self.chapters_data):
            return
            
        chapter = self.chapters_data[index]
        try:
            with open(chapter['file_path'], 'r', encoding='utf-8') as f:
                content = f.read()
                
            # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º –¥–∏–∞–ª–æ–≥ —Å —Å–æ–¥–µ—Ä–∂–∏–º—ã–º
            dialog = QDialog(self)
            dialog.setWindowTitle(f"–ü—Ä–æ—Å–º–æ—Ç—Ä: {chapter['title']}")
            dialog.setMinimumSize(600, 400)
            
            layout = QVBoxLayout(dialog)
            text_edit = QPlainTextEdit()
            text_edit.setPlainText(content)
            text_edit.setReadOnly(True)
            layout.addWidget(text_edit)
            
            button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
            button_box.rejected.connect(dialog.reject)
            layout.addWidget(button_box)
            
            dialog.exec()
            
        except Exception as e:
            QMessageBox.critical(self, "–û—à–∏–±–∫–∞", f"–ù–µ —É–¥–∞–ª–æ—Å—å –ø—Ä–æ—á–∏—Ç–∞—Ç—å —Ñ–∞–π–ª: {e}")

    def create_epub(self):
        """–°–æ–∑–¥–∞–µ—Ç EPUB —Ñ–∞–π–ª"""
        if not self.chapters_data:
            QMessageBox.warning(self, "–ü—Ä–µ–¥—É–ø—Ä–µ–∂–¥–µ–Ω–∏–µ", "–ù–µ—Ç –≥–ª–∞–≤ –¥–ª—è —Å–æ–∑–¥–∞–Ω–∏—è EPUB")
            return
            
        # –ó–∞–ø—Ä–∞—à–∏–≤–∞–µ–º –∏–º—è —Ñ–∞–π–ª–∞
        output_path, _ = QFileDialog.getSaveFileName(
            self,
            "–°–æ—Ö—Ä–∞–Ω–∏—Ç—å EPUB –∫–∞–∫",
            "translated_book.epub",
            "EPUB files (*.epub)"
        )
        
        if not output_path:
            return
            
        try:
            # –ó–∞–ø—Ä–∞—à–∏–≤–∞–µ–º –º–µ—Ç–∞–¥–∞–Ω–Ω—ã–µ
            title, ok = QtWidgets.QInputDialog.getText(self, "–ù–∞–∑–≤–∞–Ω–∏–µ –∫–Ω–∏–≥–∏", "–í–≤–µ–¥–∏—Ç–µ –Ω–∞–∑–≤–∞–Ω–∏–µ –∫–Ω–∏–≥–∏:")
            if not ok or not title:
                title = "–ü–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω–∞—è –∫–Ω–∏–≥–∞"
                
            author, ok = QtWidgets.QInputDialog.getText(self, "–ê–≤—Ç–æ—Ä", "–í–≤–µ–¥–∏—Ç–µ –∞–≤—Ç–æ—Ä–∞:")
            if not ok or not author:
                author = "Unknown"
            
            # –°–æ–∑–¥–∞–µ–º EPUB
            epub_creator = EpubCreator(title, author)
            
            for chapter in self.chapters_data:
                with open(chapter['file_path'], 'r', encoding='utf-8') as f:
                    content = f.read()
                    
                # –ò–∑–≤–ª–µ–∫–∞–µ–º —Å–æ–¥–µ—Ä–∂–∏–º–æ–µ body –∏–ª–∏ –∏—Å–ø–æ–ª—å–∑—É–µ–º –≤–µ—Å—å –∫–æ–Ω—Ç–µ–Ω—Ç
                body_match = re.search(r'<body[^>]*>(.*?)</body>', content, re.DOTALL | re.IGNORECASE)
                if body_match:
                    chapter_content = body_match.group(1)
                else:
                    chapter_content = content
                    
                filename = f"chapter_{chapter['order']:03d}.xhtml"
                epub_creator.add_chapter(filename, chapter_content, chapter['title'])
            
            epub_creator.create_epub(output_path)
            
            QMessageBox.information(
                self,
                "–£—Å–ø–µ—Ö",
                f"EPUB —Ñ–∞–π–ª —É—Å–ø–µ—à–Ω–æ —Å–æ–∑–¥–∞–Ω:\n{output_path}"
            )
            
        except Exception as e:
            QMessageBox.critical(self, "–û—à–∏–±–∫–∞", f"–ù–µ —É–¥–∞–ª–æ—Å—å —Å–æ–∑–¥–∞—Ç—å EPUB: {e}")

class ContextManager:
    """–£–ø—Ä–∞–≤–ª—è–µ—Ç –¥–æ–ª–≥–æ—Å—Ä–æ—á–Ω—ã–º –∫–æ–Ω—Ç–µ–∫—Å—Ç–æ–º: –≥–ª–æ—Å—Å–∞—Ä–∏–µ–º, —Ä–µ–∑—é–º–µ –∏ —Ç.–¥."""
    def __init__(self, output_folder):
        self.output_folder = output_folder
        self.glossary_file = os.path.join(output_folder, 'context_glossary.json')
        self.glossary = {}
        self.load_glossary()

    def load_glossary(self):
        """–ó–∞–≥—Ä—É–∂–∞–µ—Ç –≥–ª–æ—Å—Å–∞—Ä–∏–π –∏–∑ —Ñ–∞–π–ª–∞"""
        try:
            if os.path.exists(self.glossary_file):
                with open(self.glossary_file, 'r', encoding='utf-8') as f:
                    self.glossary = json.load(f)
        except Exception as e:
            print(f"–û—à–∏–±–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ –≥–ª–æ—Å—Å–∞—Ä–∏—è: {e}")
            self.glossary = {}

    def save_glossary(self):
        """–°–æ—Ö—Ä–∞–Ω—è–µ—Ç –≥–ª–æ—Å—Å–∞—Ä–∏–π –≤ —Ñ–∞–π–ª"""
        try:
            os.makedirs(self.output_folder, exist_ok=True)
            with open(self.glossary_file, 'w', encoding='utf-8') as f:
                json.dump(self.glossary, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"–û—à–∏–±–∫–∞ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è –≥–ª–æ—Å—Å–∞—Ä–∏—è: {e}")

    def get_glossary_as_json_str(self):
        """–í–æ–∑–≤—Ä–∞—â–∞–µ—Ç –≥–ª–æ—Å—Å–∞—Ä–∏–π –∫–∞–∫ —Å—Ç—Ä–æ–∫—É JSON"""
        return json.dumps(self.glossary, ensure_ascii=False, indent=2)
        
    def set_glossary_from_json_str(self, json_str):
        """–£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ—Ç –≥–ª–æ—Å—Å–∞—Ä–∏–π –∏–∑ —Å—Ç—Ä–æ–∫–∏ JSON"""
        try:
            self.glossary = json.loads(json_str)
            self.save_glossary()
        except json.JSONDecodeError as e:
            raise ValueError(f"–ù–µ–≤–µ—Ä–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç JSON: {e}")

    def format_glossary_for_prompt(self, text_content=None):
        """–§–æ—Ä–º–∞—Ç–∏—Ä—É–µ—Ç –≥–ª–æ—Å—Å–∞—Ä–∏–π –¥–ª—è –ø—Ä–æ–º–ø—Ç–∞ —Å –≤–æ–∑–º–æ–∂–Ω–æ–π —Ñ–∏–ª—å—Ç—Ä–∞—Ü–∏–µ–π"""
        if not self.glossary:
            return ""
            
        # –ï—Å–ª–∏ —É–∫–∞–∑–∞–Ω –∫–æ–Ω—Ç–µ–Ω—Ç, –ø—Ä–∏–º–µ–Ω—è–µ–º –¥–∏–Ω–∞–º–∏—á–µ—Å–∫—É—é —Ñ–∏–ª—å—Ç—Ä–∞—Ü–∏—é
        glossary_to_use = self.glossary
        if text_content:
            glossary_to_use = DynamicGlossaryFilter.filter_glossary(
                text_content, self.glossary
            )
            
        if not glossary_to_use:
            return ""
            
        glossary_lines = []
        for original, translation in glossary_to_use.items():
            glossary_lines.append(f"  {original} = {translation}")
            
        return f"\n\n**–ì–õ–û–°–°–ê–†–ò–ô:**\n" + "\n".join(glossary_lines)

def run_translation_with_auto_restart(initial_settings=None):
    """–ì–ª–∞–≤–Ω–∞—è —Ñ—É–Ω–∫—Ü–∏—è –¥–ª—è –∑–∞–ø—É—Å–∫–∞ –ø–µ—Ä–µ–≤–æ–¥–∞ —Å –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏–º –ø–µ—Ä–µ–∑–∞–ø—É—Å–∫–æ–º –ø—Ä–∏ rate limit."""
    
    # –û–ø—Ä–µ–¥–µ–ª—è–µ–º –ø—É—Ç—å –∫ —Ñ–∞–π–ª—É —Å–µ—Å—Å–∏–∏
    session_file = os.path.join(
        initial_settings['output_folder'] if initial_settings else os.getcwd(),
        'translation_session.json'
    )
    
    # –°–æ–∑–¥–∞–µ–º –º–µ–Ω–µ–¥–∂–µ—Ä —Å–µ—Å—Å–∏–∏
    session_manager = TranslationSessionManager(session_file)
    
    # –ï—Å–ª–∏ –µ—Å—Ç—å –Ω–∞—á–∞–ª—å–Ω—ã–µ –Ω–∞—Å—Ç—Ä–æ–π–∫–∏, –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä—É–µ–º –Ω–æ–≤—É—é —Å–µ—Å—Å–∏—é
    if initial_settings:
        session_manager.init_new_session(initial_settings)
        print(f"–ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–∞ –Ω–æ–≤–∞—è —Å–µ—Å—Å–∏—è –ø–µ—Ä–µ–≤–æ–¥–∞")
    else:
        # –ü—ã—Ç–∞–µ–º—Å—è –∑–∞–≥—Ä—É–∑–∏—Ç—å —Å—É—â–µ—Å—Ç–≤—É—é—â—É—é —Å–µ—Å—Å–∏—é
        if not session_manager.load_session():
            print("–ù–µ —É–¥–∞–ª–æ—Å—å –∑–∞–≥—Ä—É–∑–∏—Ç—å —Å–µ—Å—Å–∏—é")
            return False
    
    # –°–æ–∑–¥–∞–µ–º –µ–¥–∏–Ω—ã–π ApiKeyManager —Å–æ –≤—Å–µ–º–∏ –∫–ª—é—á–∞–º–∏
    all_api_keys = session_manager.session_data['api_keys']
    shared_api_key_manager = ApiKeyManager(all_api_keys)
    print(f"–ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω –º–µ–Ω–µ–¥–∂–µ—Ä —Å {len(all_api_keys)} –∫–ª—é—á–∞–º–∏")
    
    # === –î–û–ë–ê–í–¨–¢–ï –≠–¢–£ –°–¢–†–û–ö–£ ===
    translator_window = None  # –û–±—ä—è–≤–ª—è–µ–º –ø–µ—Ä–µ–º–µ–Ω–Ω—É—é –∑–∞—Ä–∞–Ω–µ–µ
    
    # –û—Å–Ω–æ–≤–Ω–æ–π —Ü–∏–∫–ª —Å –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏–º –ø–µ—Ä–µ–∑–∞–ø—É—Å–∫–æ–º
    continue_translation = True
    while continue_translation:
        # –ü–æ–ª—É—á–∞–µ–º —Å–ø–∏—Å–æ–∫ —Ñ–∞–π–ª–æ–≤ –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏ (–∏—Å–∫–ª—é—á–∞—è –∑–∞–±–ª–æ–∫–∏—Ä–æ–≤–∞–Ω–Ω—ã–µ —Ñ–∏–ª—å—Ç—Ä–∞–º–∏)
        pending_files = session_manager.get_pending_files()
        
        if not pending_files:
            print("–í—Å–µ —Ñ–∞–π–ª—ã –æ–±—Ä–∞–±–æ—Ç–∞–Ω—ã!")
            
            # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º —Ñ–∏–Ω–∞–ª—å–Ω—É—é —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫—É
            completed = len(session_manager.session_data['completed_files'])
            filtered = len(session_manager.session_data['content_filtered_files'])
            failed = len(session_manager.session_data['failed_files'])
            
            message = f"–ü–µ—Ä–µ–≤–æ–¥ –∑–∞–≤–µ—Ä—à–µ–Ω!\n\n"
            message += f"‚úÖ –£—Å–ø–µ—à–Ω–æ –ø–µ—Ä–µ–≤–µ–¥–µ–Ω–æ: {completed}\n"
            
            if filtered > 0:
                message += f"üö´ –ó–∞–±–ª–æ–∫–∏—Ä–æ–≤–∞–Ω–æ —Ñ–∏–ª—å—Ç—Ä–∞–º–∏: {filtered}\n"
                message += "(–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ –æ—Å–Ω–æ–≤–Ω–æ–µ –æ–∫–Ω–æ –ø—Ä–æ–≥—Ä–∞–º–º—ã –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏ —á–µ—Ä–µ–∑ OpenRouter)\n"
                
            if failed > 0:
                message += f"‚ùå –û—à–∏–±–∫–∏: {failed}\n"
                
            QMessageBox.information(
                None,
                "–ü–µ—Ä–µ–≤–æ–¥ –∑–∞–≤–µ—Ä—à–µ–Ω",
                message
            )
            
            # –ù–ï —É–¥–∞–ª—è–µ–º —Ñ–∞–π–ª —Å–µ—Å—Å–∏–∏ –µ—Å–ª–∏ –µ—Å—Ç—å –∑–∞–±–ª–æ–∫–∏—Ä–æ–≤–∞–Ω–Ω—ã–µ –≥–ª–∞–≤—ã
            if filtered == 0:
                try:
                    os.remove(session_file)
                    print("–§–∞–π–ª —Å–µ—Å—Å–∏–∏ —É–¥–∞–ª–µ–Ω")
                except:
                    pass
            else:
                print(f"–°–µ—Å—Å–∏—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∞ –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏ {filtered} –∑–∞–±–ª–æ–∫–∏—Ä–æ–≤–∞–Ω–Ω—ã—Ö –≥–ª–∞–≤")
                
            return True
        
        # –ü—Ä–æ–≤–µ—Ä—è–µ–º, –µ—Å—Ç—å –ª–∏ –¥–æ—Å—Ç—É–ø–Ω—ã–µ –∫–ª—é—á–∏
        if not shared_api_key_manager.has_available_keys():
            print("–í—Å–µ API –∫–ª—é—á–∏ –∏—Å—á–µ—Ä–ø–∞–Ω—ã!")
            completed, filtered, total = session_manager.get_progress()
            QMessageBox.critical(
                None,
                "–ö–ª—é—á–∏ –∏—Å—á–µ—Ä–ø–∞–Ω—ã",
                f"–í—Å–µ –¥–æ—Å—Ç—É–ø–Ω—ã–µ API –∫–ª—é—á–∏ –±—ã–ª–∏ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω—ã.\n\n"
                f"–û–±—Ä–∞–±–æ—Ç–∞–Ω–æ: {completed}/{total} —Ñ–∞–π–ª–æ–≤\n"
                f"–ó–∞–±–ª–æ–∫–∏—Ä–æ–≤–∞–Ω–æ —Ñ–∏–ª—å—Ç—Ä–∞–º–∏: {filtered}\n\n"
                f"–î–æ–±–∞–≤—å—Ç–µ –Ω–æ–≤—ã–µ –∫–ª—é—á–∏ –∏ –ø–µ—Ä–µ–∑–∞–ø—É—Å—Ç–∏—Ç–µ –ø—Ä–æ–≥—Ä–∞–º–º—É."
            )
            return False
        
        print(f"–û—Å—Ç–∞–ª–æ—Å—å –æ–±—Ä–∞–±–æ—Ç–∞—Ç—å —Ñ–∞–π–ª–æ–≤: {len(pending_files)}")
        print(f"–°—Ç–∞—Ç—É—Å –∫–ª—é—á–µ–π: {shared_api_key_manager.get_usage_report()}")
        
        # –°–æ–∑–¥–∞–µ–º –ø—Ä–∏–ª–æ–∂–µ–Ω–∏–µ Qt
        app = QApplication.instance()
        if not app:
            app = QApplication(sys.argv)
            
        try:
            # === –ù–û–í–û–ï: –§–ª–∞–≥ –¥–ª—è –æ—Ç—Å–ª–µ–∂–∏–≤–∞–Ω–∏—è –Ω–µ–æ–±—Ö–æ–¥–∏–º–æ—Å—Ç–∏ –ø—Ä–æ–¥–æ–ª–∂–µ–Ω–∏—è ===
            need_restart = False
            
            # –°–æ–∑–¥–∞–µ–º –∫–∞—Å—Ç–æ–º–Ω—ã–π –∫–ª–∞—Å—Å —Å –ø–µ—Ä–µ–æ–ø—Ä–µ–¥–µ–ª–µ–Ω–Ω—ã–º –º–µ—Ç–æ–¥–æ–º
            class AutoRestartTranslatorApp(TranslatorApp):
                def __init__(self, api_key_manager, session_mgr):
                    self.api_key_manager = api_key_manager
                    self.session_manager = session_mgr
                    super().__init__()
                    
                def show_rate_limit_restart_dialog(self):
                    """–ü–æ–∫–∞–∑—ã–≤–∞–µ—Ç –¥–∏–∞–ª–æ–≥ –ø–µ—Ä–µ–∑–∞–ø—É—Å–∫–∞ –ø—Ä–∏ rate limit"""
                    nonlocal need_restart
                    need_restart = True
                    self.close()
                    
            # –ó–∞–ø—É—Å–∫–∞–µ–º –ø–µ—Ä–µ–≤–æ–¥—á–∏–∫ —Å –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏–º –ø–µ—Ä–µ–∑–∞–ø—É—Å–∫–æ–º
            translator_window = AutoRestartTranslatorApp(shared_api_key_manager, session_manager)
            
            # –ù–∞—Å—Ç—Ä–∞–∏–≤–∞–µ–º –ø–∞—Ä–∞–º–µ—Ç—Ä—ã –∏–∑ —Å–µ—Å—Å–∏–∏
            if hasattr(translator_window, 'api_key_edit'):
                # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –ø–µ—Ä–≤—ã–π –¥–æ—Å—Ç—É–ø–Ω—ã–π –∫–ª—é—á
                available_key = shared_api_key_manager.get_next_available_key()
                if available_key:
                    translator_window.api_key_edit.setText(available_key)
            
            # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –º–æ–¥–µ–ª—å
            if hasattr(translator_window, 'model_combo'):
                model_name = session_manager.session_data.get('model', DEFAULT_MODEL_NAME)
                index = translator_window.model_combo.findText(model_name)
                if index >= 0:
                    translator_window.model_combo.setCurrentIndex(index)
            
            # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ –ø–æ—Ç–æ–∫–æ–≤
            if hasattr(translator_window, 'concurrent_requests_spin'):
                concurrent = session_manager.session_data.get('concurrent_requests', 10)
                translator_window.concurrent_requests_spin.setValue(concurrent)
            
            translator_window.show()
            app.exec()
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º, –Ω—É–∂–µ–Ω –ª–∏ –ø–µ—Ä–µ–∑–∞–ø—É—Å–∫
            if need_restart:
                print("–ü–µ—Ä–µ–∑–∞–ø—É—Å–∫ –∏–∑-–∑–∞ –¥–æ—Å—Ç–∏–∂–µ–Ω–∏—è –ª–∏–º–∏—Ç–æ–≤...")
                continue_translation = True
            else:
                continue_translation = False
                
        except Exception as e:
            print(f"–û—à–∏–±–∫–∞ –≤ —Ü–∏–∫–ª–µ –ø–µ—Ä–µ–≤–æ–¥–∞: {e}")
            continue_translation = False
            
    return True

class DynamicGlossaryFilter:
    """–ö–ª–∞—Å—Å –¥–ª—è –¥–∏–Ω–∞–º–∏—á–µ—Å–∫–æ–π —Ñ–∏–ª—å—Ç—Ä–∞—Ü–∏–∏ –≥–ª–æ—Å—Å–∞—Ä–∏—è –ø–æ —Å–æ–¥–µ—Ä–∂–∏–º–æ–º—É —Ç–µ–∫—Å—Ç–∞"""
    
    @staticmethod
    def filter_glossary(text_content, full_glossary, min_relevance_score=0.1):
        """
        –§–∏–ª—å—Ç—Ä—É–µ—Ç –≥–ª–æ—Å—Å–∞—Ä–∏–π, –æ—Å—Ç–∞–≤–ª—è—è —Ç–æ–ª—å–∫–æ —Ä–µ–ª–µ–≤–∞–Ω—Ç–Ω—ã–µ –¥–ª—è –¥–∞–Ω–Ω–æ–≥–æ —Ç–µ–∫—Å—Ç–∞ —Ç–µ—Ä–º–∏–Ω—ã
        """
        if not full_glossary or not text_content:
            return full_glossary
            
        # –ü—Ä–∏–≤–æ–¥–∏–º —Ç–µ–∫—Å—Ç –∫ –Ω–∏–∂–Ω–µ–º—É —Ä–µ–≥–∏—Å—Ç—Ä—É –¥–ª—è –ø–æ–∏—Å–∫–∞
        text_lower = text_content.lower()
        
        filtered_glossary = {}
        
        for original, translation in full_glossary.items():
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º –Ω–∞–ª–∏—á–∏–µ —Ç–µ—Ä–º–∏–Ω–∞ –≤ —Ç–µ–∫—Å—Ç–µ (—Ä–∞–∑–ª–∏—á–Ω—ã–µ –≤–∞—Ä–∏–∞–Ω—Ç—ã)
            original_lower = original.lower()
            
            # –ü—Ä—è–º–æ–µ –≤—Ö–æ–∂–¥–µ–Ω–∏–µ
            if original_lower in text_lower:
                filtered_glossary[original] = translation
                continue
                
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º —á–∞—Å—Ç–∏ —Å–æ—Å—Ç–∞–≤–Ω—ã—Ö —Ç–µ—Ä–º–∏–Ω–æ–≤
            if ' ' in original_lower:
                words = original_lower.split()
                if any(word in text_lower for word in words if len(word) > 2):
                    filtered_glossary[original] = translation
                    continue
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º —Å—Ö–æ–∂–∏–µ —Ñ–æ—Ä–º—ã (–¥–ª—è –∏–º–µ–Ω —Å–æ–±—Å—Ç–≤–µ–Ω–Ω—ã—Ö)
            if len(original) > 3:
                # –ü—Ä–æ–≤–µ—Ä—è–µ–º –≤—Ö–æ–∂–¥–µ–Ω–∏—è –ø–æ–¥—Å—Ç—Ä–æ–∫ –¥–ª–∏–Ω–æ–π –Ω–µ –º–µ–Ω–µ–µ 4 —Å–∏–º–≤–æ–ª–æ–≤
                for i in range(len(original_lower) - 3):
                    substring = original_lower[i:i+4]
                    if substring in text_lower:
                        filtered_glossary[original] = translation
                        break
        
        return filtered_glossary

def read_docx_with_images(filepath, temp_dir, image_map):
    """Reads DOCX, extracts text, replaces images with placeholders, saves images."""
    if not DOCX_AVAILABLE: raise ImportError("python-docx library is required.")
    if not os.path.exists(filepath): raise FileNotFoundError(f"DOCX file not found: {filepath}")

    doc = docx.Document(filepath)
    output_lines = []

    is_bold_chapter = re.compile(r'^\s*(–ì–ª–∞–≤–∞|Chapter|Part)\s+([0-9IVXLCDM]+|[a-zA-Z–∞-—è–ê-–Ø]+)\b.*', re.IGNORECASE)
    doc_rels = doc.part.rels
    processed_image_rids = set()
    processed_rid_to_uuid = {}

    for element in doc.element.body:

        if element.tag.endswith('p'):
            para = docx.text.paragraph.Paragraph(element, doc)
            para_text_parts = []
            contains_image = False
            for run in para.runs:

                drawing_elems = run.element.xpath('.//w:drawing')
                if drawing_elems:
                    for drawing in drawing_elems:

                        inline_elems = drawing.xpath('.//wp:inline | .//wp:anchor')
                        if inline_elems:
                            for inline in inline_elems:

                                blip_fill = inline.xpath('.//a:blip')
                                if blip_fill:
                                    rId = blip_fill[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')

                                    if rId and rId in doc_rels and "image" in doc_rels[rId].target_ref:
                                        if rId not in processed_image_rids:
                                            try:
                                                img_part = doc_rels[rId].target_part
                                                img_data = img_part.blob
                                                original_filename = os.path.basename(img_part.partname)

                                                img_ext_original = os.path.splitext(original_filename)[-1].lower().strip('.')
                                                img_ext_detected = get_image_extension_from_data(img_data, fallback_ext=img_ext_original or "png")

                                                if img_ext_original == 'emf' or img_ext_detected == 'emf':
                                                    png_data = convert_emf_to_png(img_data)
                                                    if png_data:
                                                        img_data = png_data; img_ext_final = 'png'; content_type = 'image/png'
                                                        print(f"[INFO] DOCX: Converted EMF image '{original_filename}' to PNG.")
                                                    else:
                                                        print(f"[WARN] DOCX: Failed to convert EMF '{original_filename}', skipping."); continue # Skip if conversion failed
                                                else:
                                                    img_ext_final = img_ext_detected; content_type = f"image/{img_ext_final}"

                                                width, height = None, None
                                                try:
                                                    extent = inline.xpath('.//wp:extent');
                                                    if extent:
                                                        emu_per_px = 9525 # Approx conversion factor
                                                        width = int(extent[0].get('cx')) // emu_per_px
                                                        height = int(extent[0].get('cy')) // emu_per_px
                                                except Exception: pass # Ignore errors getting dimensions

                                                img_uuid = uuid.uuid4().hex
                                                saved_filename = f"{img_uuid}.{img_ext_final}"; saved_path = os.path.join(temp_dir, saved_filename)
                                                with open(saved_path, 'wb') as img_file: img_file.write(img_data)
                                                image_map[img_uuid] = {'saved_path': saved_path, 'original_filename': original_filename, 'content_type': content_type, 'width': width, 'height': height}
                                                processed_image_rids.add(rId); processed_rid_to_uuid[rId] = img_uuid

                                                placeholder = create_image_placeholder(img_uuid)
                                                para_text_parts.append(placeholder); contains_image = True
                                            except Exception as e:
                                                print(f"[WARN] DOCX: Error processing image rId {rId}: {e}"); para_text_parts.append(run.text) # Append run text on error
                                        else: # Image already processed (e.g., copy-pasted image)
                                            if rId in processed_rid_to_uuid:
                                                para_text_parts.append(create_image_placeholder(processed_rid_to_uuid[rId])); contains_image = True
                                            else: # Should not happen if processed correctly
                                                print(f"[WARN] DOCX: rId {rId} processed but not in UUID map."); para_text_parts.append(run.text)
                                    else: # Not an image relationship or rId invalid
                                        para_text_parts.append(run.text)
                                else: # No blip fill found
                                     para_text_parts.append(run.text)
                        else: # No inline/anchor element
                             para_text_parts.append(run.text)
                else: # No drawing element in run
                     para_text_parts.append(run.text)

            full_para_text = "".join(para_text_parts).strip()
            style_name = para.style.name.lower() if para.style and para.style.name else ''
            is_heading_style = False

            is_run_bold = all(r.bold for r in para.runs if r.text.strip()) # Check if all text runs are bold
            if style_name.startswith('heading 1') or (style_name == 'normal' and is_bold_chapter.match(full_para_text) and is_run_bold):
                output_lines.append(f"# {full_para_text}"); is_heading_style = True
            elif style_name.startswith('heading 2'):
                output_lines.append(f"## {full_para_text}"); is_heading_style = True
            elif style_name.startswith('heading 3'):
                output_lines.append(f"### {full_para_text}"); is_heading_style = True

            elif not full_para_text.strip() and not contains_image:

                 if output_lines and output_lines[-1] != "": output_lines.append("")
                 continue

            elif not is_heading_style and (style_name.startswith('list paragraph') or (para.paragraph_format and para.paragraph_format.left_indent and full_para_text)):
                 list_marker = "*"; # Default marker

                 num_match = re.match(r'^\s*(\d+\.|\([a-z]\)|\([A-Z]\)|[a-z]\.|[A-Z]\.)\s+', para.text) # Numbered or lettered lists
                 bullet_match = re.match(r'^\s*([\*\-\‚Ä¢\‚ÅÉ])\s+', para.text) # Common bullet chars
                 if num_match: list_marker = num_match.group(1)
                 elif bullet_match: list_marker = bullet_match.group(1)

                 clean_list_text = re.sub(r'^\s*(\d+\.|\([a-z]\)|\([A-Z]\)|[a-z]\.|[A-Z]\.|[\*\-\‚Ä¢\‚ÅÉ])\s*', '', full_para_text)
                 output_lines.append(f"{list_marker} {clean_list_text}")

            elif not is_heading_style and (full_para_text or contains_image):
                 output_lines.append(full_para_text)

        elif element.tag.endswith('tbl'):

            if output_lines and output_lines[-1]: output_lines.append("")
            output_lines.append("[--- –¢–ê–ë–õ–ò–¶–ê (–Ω–µ –æ–±—Ä–∞–±–æ—Ç–∞–Ω–æ) ---]")
            output_lines.append("")

    final_text = "";
    for i, line in enumerate(output_lines):
        final_text += line

        if i < len(output_lines) - 1:
             final_text += "\n"

             is_current_placeholder_line = IMAGE_PLACEHOLDER_PREFIX in line
             is_next_placeholder_line = IMAGE_PLACEHOLDER_PREFIX in output_lines[i+1]
             is_current_heading = line.startswith('#')
             is_current_list = re.match(r'^([\*\-\‚Ä¢\‚ÅÉ]|\d+\.|\([a-z]\)|\([A-Z]\)|[a-z]\.|[A-Z]\.)\s', line)
             is_current_table = "[--- –¢–ê–ë–õ–ò–¶–ê" in line
             is_next_table = "[--- –¢–ê–ë–õ–ò–¶–ê" in output_lines[i+1]

             if (output_lines[i+1] != "" and line != "" and # Both lines have content
                 not is_current_heading and not is_current_list and # Not headings or lists
                 not is_current_table and not is_next_table and # Not tables
                 not (is_current_placeholder_line and is_next_placeholder_line)): # Not two image lines together
                     final_text += "\n"

    print(f"[INFO] DOCX Read: Extracted {len(image_map)} images.")
    return final_text.strip()


def process_html_images(html_content, source_context, temp_dir, image_map):
    """
    Parses HTML, extracts images, replaces with placeholders, converts Hx/title to Markdown-like,
    and then extracts text content for translation.
    `source_context` can be a tuple (zipfile.ZipFile, html_path_in_zip) or a base directory path.
    """
    if not BS4_AVAILABLE: raise ImportError("BeautifulSoup4 is required for HTML processing.")

    if "<svg" in html_content.lower() or "xmlns:" in html_content.lower() or \
       html_content.strip().startswith("<?xml"):
        parser_type = 'lxml-xml'
    else:
        parser_type = 'lxml'

    try:
        soup = BeautifulSoup(html_content, parser_type)
    except Exception as e_parse:
        print(f"DEBUG process_html_images: Parse failed with '{parser_type}': {e_parse}. Trying 'html.parser'.")
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
        except Exception as e_parse_fallback:
            print(f"[ERROR] BeautifulSoup failed to parse HTML content with primary parser '{parser_type}' and fallback 'html.parser'. Error: {e_parse_fallback}")
            raise ValueError(f"Failed to parse HTML content after trying multiple parsers: {e_parse_fallback}")

    zip_file_obj = None
    source_html_path = None
    base_path = ""
    if isinstance(source_context, tuple) and len(source_context) == 2 and isinstance(source_context[0], zipfile.ZipFile):
         zip_file_obj = source_context[0]
         source_html_path = source_context[1]
         if source_html_path:
             base_path = os.path.dirname(source_html_path).replace('\\', '/')
             if base_path == '.': base_path = ""
    elif isinstance(source_context, str) and os.path.isdir(source_context):
        base_path = source_context
        source_html_path = "unknown.html"
        zip_file_obj = None
    else:
        zip_file_obj = None
        base_path = ""
        source_html_path = "unknown.html"

    image_processing_context = zip_file_obj if zip_file_obj else base_path
    
    # --- Image Processing ---
    for tag in soup.find_all(['img', 'svg']):
        if not tag.parent: continue
        img_uuid = None
        tag_name = tag.name.lower()
        
        try:
            if tag_name == 'img':
                img_uuid = _process_single_image(tag, image_processing_context, base_path, source_html_path, temp_dir, image_map, is_svg_image=False)
                tag.replace_with(NavigableString(create_image_placeholder(img_uuid)) if img_uuid else "")
            elif tag_name == 'svg':
                svg_image_tag = tag.find(lambda t: t.name.lower() == 'image', recursive=False)
                if svg_image_tag:
                    img_uuid = _process_single_image(svg_image_tag, image_processing_context, base_path, source_html_path, temp_dir, image_map, is_svg_image=True)
                tag.replace_with(NavigableString(create_image_placeholder(img_uuid)) if img_uuid else "")
        except Exception as e:
            print(f"[ERROR] process_html_images: Error replacing tag <{tag_name}>: {e}")
            try: tag.replace_with("")
            except Exception as remove_err: print(f"[ERROR] Failed to remove tag after error: {remove_err}")
            
    # --- –ù–ê–ß–ê–õ–û –ë–õ–û–ö–ê –ù–û–†–ú–ê–õ–ò–ó–ê–¶–ò–ò HTML (–ù–û–í–´–ô –ë–õ–û–ö) ---
    # –ü—Ä–µ–æ–±—Ä–∞–∑—É–µ–º —Å—Ç–∏–ª–∏–∑–æ–≤–∞–Ω–Ω—ã–µ span –≤ —Å–µ–º–∞–Ω—Ç–∏—á–µ—Å–∫–∏–µ —Ç–µ–≥–∏ em/strong
    for span in soup.find_all('span', style=True):
        style = span['style'].lower()
        if 'font-style' in style and 'italic' in style:
            span.name = 'em'
            del span['style'] # –£–¥–∞–ª—è–µ–º –∞—Ç—Ä–∏–±—É—Ç style –ø–æ—Å–ª–µ –ø—Ä–µ–æ–±—Ä–∞–∑–æ–≤–∞–Ω–∏—è
        elif 'font-weight' in style and ('bold' in style or any(w.strip() in ['700', '800', '900'] for w in style.split('font-weight:')[1].split(';')[0].split())):
            span.name = 'strong'
            del span['style']

    # –û–±—ä–µ–¥–∏–Ω—è–µ–º –ø–æ—Å–ª–µ–¥–æ–≤–∞—Ç–µ–ª—å–Ω—ã–µ —Ç–µ–≥–∏ em –∏ strong
    for tag_name in ['em', 'strong']:
        for tag in soup.find_all(tag_name):
            next_sibling = tag.next_sibling
            while next_sibling and isinstance(next_sibling, NavigableString) and not next_sibling.strip():
                next_sibling = next_sibling.next_sibling # –ü—Ä–æ–ø—É—Å–∫–∞–µ–º –ø—É—Å—Ç—ã–µ —Å—Ç—Ä–æ–∫–∏
            if next_sibling and next_sibling.name == tag.name:
                # –ï—Å–ª–∏ —É —Å–ª–µ–¥—É—é—â–µ–≥–æ —Ç–µ–≥–∞ –µ—Å—Ç—å –∞—Ç—Ä–∏–±—É—Ç—ã, –æ—Ç–ª–∏—á–Ω—ã–µ –æ—Ç —Ç–µ–∫—É—â–µ–≥–æ, –Ω–µ –æ–±—ä–µ–¥–∏–Ω—è–µ–º
                if next_sibling.attrs == tag.attrs:
                    tag.append(next_sibling.decode_contents())
                    next_sibling.decompose()

    # –ó–∞–º–µ–Ω—è–µ–º —Å—Ç—Ä—É–∫—Ç—É—Ä–Ω—ã–µ div –Ω–∞ p, —á—Ç–æ–±—ã –æ–Ω–∏ –æ–±—Ä–∞–±–∞—Ç—ã–≤–∞–ª–∏—Å—å –∫–∞–∫ –∞–±–∑–∞—Ü—ã
    # –∏ "—Ä–∞–∑–≤–æ—Ä–∞—á–∏–≤–∞–µ–º" (unwrap) –≤–ª–æ–∂–µ–Ω–Ω—ã–µ div –∏ span, –∫–æ—Ç–æ—Ä—ã–µ –Ω–µ –Ω–µ—Å—É—Ç —Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏—è
    tags_to_process = soup.find_all(['div', 'span'])
    for tag in tags_to_process:
        # –ï—Å–ª–∏ —ç—Ç–æ div, –∫–æ—Ç–æ—Ä—ã–π –≤—ã–≥–ª—è–¥–∏—Ç –∫–∞–∫ –∫–æ–Ω—Ç–µ–π–Ω–µ—Ä –¥–ª—è –∞–±–∑–∞—Ü–∞, –º–µ–Ω—è–µ–º –µ–≥–æ –Ω–∞ p
        if tag.name == 'div':
            tag.name = 'p'
            # –£–¥–∞–ª—è–µ–º –∞—Ç—Ä–∏–±—É—Ç—ã, –Ω–µ –æ—Ç–Ω–æ—Å—è—â–∏–µ—Å—è –∫ p
            for attr in list(tag.attrs.keys()):
                if attr.lower() not in ['class', 'id']:
                     del tag[attr]
        # –ï—Å–ª–∏ —ç—Ç–æ span –±–µ–∑ –∞—Ç—Ä–∏–±—É—Ç–æ–≤, –æ–Ω, —Å–∫–æ—Ä–µ–µ –≤—Å–µ–≥–æ, –¥–ª—è —Ä–∞–∑–±–∏–µ–Ω–∏—è —Ç–µ–∫—Å—Ç–∞ - —É–±–∏—Ä–∞–µ–º –µ–≥–æ
        elif tag.name == 'span' and not tag.attrs:
            tag.unwrap()
    # --- –ö–û–ù–ï–¶ –ë–õ–û–ö–ê –ù–û–†–ú–ê–õ–ò–ó–ê–¶–ò–ò HTML ---


    # --- Header and Content Extraction ---
    html_doctitle_text = None
    if soup.head and soup.head.title and soup.head.title.string:
        title_candidate = soup.head.title.string.strip()
        generic_titles = ['untitled', 'unknown', 'navigation', 'toc', 'table of contents', 'index', 'contents', '–æ–≥–ª–∞–≤–ª–µ–Ω–∏–µ', '—Å–æ–¥–µ—Ä–∂–∞–Ω–∏–µ', '–∏–Ω–¥–µ–∫—Å', 'cover', 'title page', 'copyright', 'chapter']
        if title_candidate and title_candidate.lower() not in generic_titles and len(title_candidate) > 2:
            html_doctitle_text = title_candidate

    content_extraction_root = soup.body if soup.body else soup
    if not content_extraction_root:
        print("[WARN] process_html_images: No <body> or root element found.")
        return ""

    # --- CORRECTED HEADER PROCESSING LOGIC ---
    for level in range(6, 0, -1):
        for header_tag in content_extraction_root.find_all(f'h{level}'):
            # –°–Ω–∞—á–∞–ª–∞ –ø—Ä–µ–æ–±—Ä–∞–∑—É–µ–º em/strong –≤–Ω—É—Ç—Ä–∏ –∑–∞–≥–æ–ª–æ–≤–∫–∞ –≤ Markdown
            for em_tag in header_tag.find_all('em'):
                em_tag.replace_with(f"*{em_tag.get_text(strip=True)}*")
            for strong_tag in header_tag.find_all('strong'):
                strong_tag.replace_with(f"**{strong_tag.get_text(strip=True)}**")
            
            header_text = header_tag.get_text(separator=' ', strip=True)
            if header_text:
                markdown_header_line = f"\n\n{'#' * level} {header_text}\n\n"
                header_tag.replace_with(NavigableString(markdown_header_line))

    # Decompose unwanted tags after processing headers and normalization
    tags_to_decompose = ['script', 'style', 'noscript', 'head', 'meta', 'link', 'form', 'iframe', 'header', 'footer', 'nav', 'aside']
    for tag_type in tags_to_decompose:
        for instance in content_extraction_root.find_all(tag_type):
            instance.decompose()
            
    # --- –ò–ó–ú–ï–ù–ï–ù–ù–ê–Ø –õ–û–ì–ò–ö–ê –ò–ó–í–õ–ï–ß–ï–ù–ò–Ø –¢–ï–ö–°–¢–ê ---
    # –ü—Ä–µ–æ–±—Ä–∞–∑—É–µ–º –æ—Å—Ç–∞–≤—à–∏–µ—Å—è em/strong –≤ Markdown –ø–µ—Ä–µ–¥ —Ñ–∏–Ω–∞–ª—å–Ω—ã–º –∏–∑–≤–ª–µ—á–µ–Ω–∏–µ–º —Ç–µ–∫—Å—Ç–∞
    for em_tag in content_extraction_root.find_all('em'):
        em_tag.replace_with(f"*{em_tag.get_text(strip=True)}*")
    for strong_tag in content_extraction_root.find_all('strong'):
        strong_tag.replace_with(f"**{strong_tag.get_text(strip=True)}**")
    
    # –ó–∞–º–µ–Ω—è–µ–º <br> –∏ </p> –Ω–∞ –ø–µ—Ä–µ–Ω–æ—Å—ã —Å—Ç—Ä–æ–∫ –¥–ª—è –∫–æ—Ä—Ä–µ–∫—Ç–Ω–æ–≥–æ —Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è
    for br in content_extraction_root.find_all("br"):
        br.replace_with("\n")
    for p_tag in content_extraction_root.find_all("p"):
        p_tag.append("\n\n")

    # Get all text from the modified body
    body_text_md = content_extraction_root.get_text(separator='', strip=False) # –ò—Å–ø–æ–ª—å–∑—É–µ–º separator='', —á—Ç–æ–±—ã –Ω–µ –¥–æ–±–∞–≤–ª—è—Ç—å –ª–∏—à–Ω–∏—Ö —Ä–∞–∑–¥–µ–ª–∏—Ç–µ–ª–µ–π

    # Final logic to assemble text for API
    final_text_for_api = body_text_md
    if html_doctitle_text and not body_text_md.lstrip().startswith('#'):
        final_text_for_api = f"# {html_doctitle_text}\n\n{body_text_md}"

    # Clean up excessive newlines and spaces
    final_text_for_api = re.sub(r' +', ' ', final_text_for_api) # –°–∂–∏–º–∞–µ–º –º–Ω–æ–∂–µ—Å—Ç–≤–µ–Ω–Ω—ã–µ –ø—Ä–æ–±–µ–ª—ã
    final_text_for_api = re.sub(r'\n{3,}', '\n\n', final_text_for_api) # –°–∂–∏–º–∞–µ–º –º–Ω–æ–∂–µ—Å—Ç–≤–µ–Ω–Ω—ã–µ –ø–µ—Ä–µ–Ω–æ—Å—ã
    
    return final_text_for_api.strip()


def _process_single_image(img_tag, source_context, base_path, source_html_path, temp_dir, image_map, is_svg_image=False):
    """
    Processes individual image tag.
    For EPUB->EPUB: Extracts original src and attributes, stores them in image_map with a UUID. Does NOT save file.
    For other modes: Extracts image data, saves to temp_dir, stores path and info in image_map.
    """
    src = None
    xlink_namespace_uri = "http://www.w3.org/1999/xlink"


    is_epub_rebuild_mode = isinstance(source_context, zipfile.ZipFile) # True if processing for EPUB->EPUB

    if is_svg_image:
        src = img_tag.get(f'{{{xlink_namespace_uri}}}href')
        if not src:
            attrs_dict = img_tag.attrs
            if 'xlink:href' in attrs_dict: src = attrs_dict['xlink:href']
            elif 'href' in attrs_dict: src = attrs_dict['href']
            else:
                namespaced_key = (xlink_namespace_uri, 'href')
                if namespaced_key in attrs_dict: src = attrs_dict[namespaced_key]

    else: # HTML <img> tag
        src = img_tag.get('src', '')


    if not src or src.startswith('data:'):

        return None

    img_uuid = uuid.uuid4().hex
    original_src_value = src # This is the raw value from the attribute, e.g., "../Images/0004.png"
    original_tag_name = img_tag.name # 'img' or 'image' (from svg)
    all_original_attributes = dict(img_tag.attrs) # Store all attributes

    if is_epub_rebuild_mode:

        image_map[img_uuid] = {
            'original_src': original_src_value,
            'original_tag_name': original_tag_name, # 'img' or 'image'
            'is_svg_image_child': is_svg_image, # True if it was <image> inside <svg>
            'attributes': all_original_attributes # Store all original attributes
        }

        return img_uuid
    else:

        img_data = None
        decoded_src = unquote(src)
        original_filename = os.path.basename(urlparse(decoded_src).path)
        if not original_filename:
            src_parts = decoded_src.split('/')
            potential_fname = src_parts[-1] if src_parts else "image"
            potential_fname = potential_fname.split('?')[0]
            safe_fname_part = re.sub(r'[^\w\.\-]+', '_', potential_fname)
            _, ext_guess = os.path.splitext(safe_fname_part)
            fallback_ext = "png"
            if ext_guess and ext_guess[1:].lower() in ['jpg', 'jpeg', 'gif', 'webp', 'bmp', 'tiff', 'svg']:
                original_filename = safe_fname_part
            else:
                original_filename = f"{Path(safe_fname_part).stem}.{fallback_ext}"

        content_type = None

        try:
            if isinstance(source_context, zipfile.ZipFile): # EPUB source -> non-EPUB output

                possible_paths = []
                current_html_dir = base_path
                path1 = os.path.join(current_html_dir, decoded_src)
                path1_norm = os.path.normpath(path1).replace('\\', '/')
                if not path1_norm.startswith('..'): possible_paths.append(path1_norm.lstrip('/'))
                path2_norm = os.path.normpath(decoded_src.lstrip('/')) .replace('\\', '/')
                if not path2_norm.startswith('..'): possible_paths.append(path2_norm)

                unique_paths = list(dict.fromkeys(p.strip('/') for p in possible_paths if p.strip('/')))

                for img_path_in_zip in unique_paths:
                    try:
                        img_data = source_context.read(img_path_in_zip)

                        break
                    except KeyError: continue
                if not img_data:
                    print(f"[WARN] HTML Image NOT Found (EPUB src -> File Output): src='{src}'. Tried: {unique_paths}.")
                    return None

            elif isinstance(source_context, str) and os.path.isdir(source_context): # Directory context

                paths_to_try_fs = [
                    os.path.normpath(os.path.join(base_path, decoded_src)),
                    os.path.normpath(os.path.join(source_context, decoded_src.lstrip('/\\')))
                ]
                if not decoded_src.startswith(('/', '\\')):
                    path3_fs = os.path.normpath(os.path.join(source_context, decoded_src))
                    if path3_fs not in paths_to_try_fs: paths_to_try_fs.append(path3_fs)
                abs_path = next((p for p in paths_to_try_fs if os.path.exists(p)), None)
                if not abs_path:
                    print(f"[WARN] HTML Image (FS Mode): Could not find '{decoded_src}'. Tried: {paths_to_try_fs}")
                    return None
                with open(abs_path, 'rb') as f: img_data = f.read()
            else:
                print(f"[WARN] HTML Image: Unknown source context for file mode: {type(source_context)}")
                return None

            img_ext_from_file = os.path.splitext(original_filename)[1][1:].lower()
            content_type = f"image/{get_image_extension_from_data(img_data, fallback_ext=img_ext_from_file or 'jpeg')}"
            img_ext = content_type.split('/')[-1] if content_type else 'jpeg'
            img_ext = 'jpg' if img_ext == 'jpeg' else img_ext

            if img_ext == 'emf':
                converted_data = convert_emf_to_png(img_data)
                if converted_data:
                    img_data = converted_data; img_ext = 'png'; content_type = 'image/png'
                else: return None

            filename = f"{img_uuid}.{img_ext}"
            save_path = os.path.join(temp_dir, filename)
            with open(save_path, 'wb') as f: f.write(img_data)

            image_map[img_uuid] = {
                'saved_path': save_path, # For non-EPUB rebuild, this is used
                'original_filename': original_filename,
                'original_src': original_src_value, # Still store original_src for consistency if needed
                'content_type': content_type,
                'attributes': all_original_attributes # Store original attributes
            }

            return img_uuid

        except Exception as e:
            print(f"[ERROR] HTML Image (File Mode): Error processing src '{src}': {e}")
            traceback.print_exc()
            return None

def write_markdown_to_docx(filepath, md_text_with_placeholders, image_map):
    """Writes Markdown-like text with placeholders back to DOCX."""
    if not DOCX_AVAILABLE: raise ImportError("python-docx library is required.")
    if image_map is None: image_map = {}
    doc = Document()

    lines = re.split('(\n)', md_text_with_placeholders)

    paragraphs_md = []
    current_para_lines = []
    for line in lines:
        stripped_line = line.strip()
        if not stripped_line: # Treat empty line as paragraph break
            if current_para_lines:
                paragraphs_md.append("\n".join(current_para_lines))
                current_para_lines = []

            if paragraphs_md and paragraphs_md[-1]: # Add if last added wasn't already empty
                 paragraphs_md.append("")
        else:
            current_para_lines.append(line)
    if current_para_lines: # Add last paragraph if exists
        paragraphs_md.append("\n".join(current_para_lines))

    current_docx_para = None
    for md_para in paragraphs_md:
        md_para_stripped = md_para.strip()
        if not md_para_stripped:

            if current_docx_para is not None: # Check if previous para exists
                doc.add_paragraph("")
            current_docx_para = None # Reset current para tracker
            continue

        heading_match = re.match(r'^(#{1,6})\s+(.*)', md_para_stripped, re.DOTALL)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text_raw = heading_match.group(2).strip()
            current_docx_para = doc.add_heading("", level=max(1, min(level, 6))) # Add heading
            process_text_with_placeholders(current_docx_para, heading_text_raw, image_map)
            continue # Move to next md paragraph

        list_match = re.match(r'^([\*\-\‚Ä¢\‚ÅÉ]|\d+\.|\([a-z]\)|\([A-Z]\)|[a-z]\.|[A-Z]\.)\s+(.*)', md_para_stripped, re.DOTALL)
        if list_match:
             marker = list_match.group(1)
             list_item_text_raw = list_match.group(2).strip()
             style = 'List Bullet' if marker in ['*', '-', '‚Ä¢', '‚ÅÉ'] else 'List Number' # Basic style mapping
             try:
                 current_docx_para = doc.add_paragraph(style=style)
             except (KeyError, ValueError): # Fallback if style doesn't exist in template
                 print(f"[WARN] DOCX Write: Style '{style}' not found. Using default paragraph.")
                 current_docx_para = doc.add_paragraph()
             process_text_with_placeholders(current_docx_para, list_item_text_raw, image_map)
             continue # Move to next md paragraph

        if md_para_stripped == '---':
             doc.add_paragraph().add_run()._element.xpath('.//w:pPr')[0].append(
                 etree.fromstring('<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/></w:pBdr>')
             )
             current_docx_para = None # HR acts as break
             continue

        current_docx_para = doc.add_paragraph()
        process_text_with_placeholders(current_docx_para, md_para.strip(), image_map) # Process original (not stripped) to keep internal newlines

    doc.save(filepath)


def process_text_with_placeholders(docx_paragraph, text_with_placeholders, image_map):
    """Adds runs of text and images to a docx paragraph based on placeholders."""

    last_index = 0
    placeholders_found = find_image_placeholders(text_with_placeholders)

    if not placeholders_found:
        if text_with_placeholders.strip():
            docx_paragraph.add_run(text_with_placeholders)
        return

    for placeholder_tag, img_uuid in placeholders_found:
        match_start = text_with_placeholders.find(placeholder_tag, last_index)
        if match_start == -1: continue # Should not happen with finditer logic

        text_before = text_with_placeholders[last_index:match_start]
        if text_before:
            docx_paragraph.add_run(text_before)

        if img_uuid in image_map:
            img_info = image_map[img_uuid]; img_path = img_info['saved_path']
            if os.path.exists(img_path):
                try:

                    img_width_px = img_info.get('width'); img_height_px = img_info.get('height')
                    run = docx_paragraph.add_run() # Create run for the picture
                    target_width = None

                    if img_width_px:
                        try:
                            img_width_px = float(img_width_px) # Ensure it's a number
                            if img_width_px > 0:
                                 target_width_inches = img_width_px / 96.0 # Approx DPI
                                 max_doc_width_inches = 6.0 # Usable width on standard page
                                 target_width = Inches(min(target_width_inches, max_doc_width_inches))
                        except (ValueError, TypeError): pass # Ignore invalid width values

                    run.add_picture(img_path, width=target_width)
                except FileNotFoundError:
                    print(f"[ERROR] DOCX Write: Image file not found: {img_path}")
                    docx_paragraph.add_run(f"[Image NF: {img_info.get('original_filename', img_uuid)}]") # Add error text
                except Exception as e:
                    print(f"[ERROR] DOCX Write: Failed to add picture {img_path}: {e}")
                    docx_paragraph.add_run(f"[Img Err: {img_info.get('original_filename', img_uuid)}]")
            else:

                print(f"[ERROR] DOCX Write: Image path from map does not exist: {img_path}")
                docx_paragraph.add_run(f"[Img Path Miss: {img_info.get('original_filename', img_uuid)}]")
        else:

            print(f"[WARN] DOCX Write: Placeholder UUID '{img_uuid}' not found in image_map.")
            docx_paragraph.add_run(f"[Unk Img: {img_uuid}]")

        last_index = match_start + len(placeholder_tag)

    text_after = text_with_placeholders[last_index:]
    if text_after:
        docx_paragraph.add_run(text_after)







def generate_nav_html(nav_data_list, nav_file_path_in_zip, book_title, book_lang="ru"):
    """
    Generates XHTML content for nav.xhtml based on spine data.
    Simplified version focusing on the list structure.
    """
    if not nav_data_list:
        print("[WARN] NAV Gen: Input data list is empty. NAV not generated.")
        return None

    if not LXML_AVAILABLE:
        print("[ERROR] NAV Gen: LXML library is required for reliable NAV generation.")

        print("[WARN] NAV Gen: LXML not found, attempting basic string generation (less reliable).")
        nav_lines = []
        nav_lines.append("<?xml version='1.0' encoding='utf-8'?>")
        nav_lines.append("<!DOCTYPE html>")
        nav_lines.append(f'<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" lang="{book_lang}" xml:lang="{book_lang}">')
        nav_lines.append("<head>")
        nav_lines.append("  <meta charset=\"utf-8\"/>")

        nav_lines.append("</head>")
        nav_lines.append("<body>")
        nav_lines.append('  <nav epub:type="toc" id="toc">') # –ò—Å–ø–æ–ª—å–∑—É–µ–º id="toc"

        nav_lines.append("    <ol>")
        nav_dir = os.path.dirname(nav_file_path_in_zip).replace('\\', '/')
        if nav_dir == '.': nav_dir = ""
        link_count_str = 0
        for item_path, item_title in nav_data_list:
            safe_item_title = html.escape(str(item_title).strip())
            if not safe_item_title: safe_item_title = "Untitled Entry" # –ó–∞–≥–ª—É—à–∫–∞
            try:
                item_path_norm = item_path.replace('\\', '/').lstrip('/')
                nav_parent_dir_norm = os.path.dirname(nav_file_path_in_zip.replace('\\','/').lstrip('/')).replace('\\','/')
                relative_href = os.path.relpath(item_path_norm, start=nav_parent_dir_norm if nav_parent_dir_norm else '.').replace('\\', '/')
                safe_href = html.escape(relative_href, quote=True)
                nav_lines.append(f'      <li><a href="{safe_href}">{safe_item_title}</a></li>')
                link_count_str += 1
            except ValueError as e:
                print(f"[WARN] NAV Gen (String): Failed to calculate relative path for '{item_path}' from '{nav_parent_dir_norm or '<root>'}': {e}. Skipping link.")
            except Exception as e_loop:
                 print(f"[ERROR] NAV Gen (String): Error processing item ('{item_path}', '{item_title}'): {e_loop}")
        nav_lines.append("    </ol>")
        nav_lines.append("  </nav>")
        nav_lines.append("</body>")
        nav_lines.append("</html>")
        print(f"[INFO] NAV Gen (String): Finished generation. Added {link_count_str} links.")
        return "\n".join(nav_lines).encode('utf-8')

    print(f"[INFO] NAV Gen (lxml): Starting NAV generation for '{nav_file_path_in_zip}' with {len(nav_data_list)} entries...")
    xhtml_ns = "http://www.w3.org/1999/xhtml"
    epub_ns = "http://www.idpf.org/2007/ops"
    NSMAP = {None: xhtml_ns, "epub": epub_ns}

    html_tag = etree.Element(f"{{{xhtml_ns}}}html", nsmap=NSMAP)
    xml_lang_attr_name = "{http://www.w3.org/XML/1998/namespace}lang"
    html_tag.set(xml_lang_attr_name, book_lang)
    html_tag.set("lang", book_lang)

    head = etree.SubElement(html_tag, f"{{{xhtml_ns}}}head")

    etree.SubElement(head, f"{{{xhtml_ns}}}meta", charset="utf-8") # –î–æ–±–∞–≤–ª—è–µ–º meta charset

    body = etree.SubElement(html_tag, f"{{{xhtml_ns}}}body")
    nav = etree.SubElement(body, f"{{{xhtml_ns}}}nav", id="toc") # –ò—Å–ø–æ–ª—å–∑—É–µ–º id="toc"
    nav.set(f"{{{epub_ns}}}type", "toc")


    ol = etree.SubElement(nav, f"{{{xhtml_ns}}}ol")

    nav_dir = os.path.dirname(nav_file_path_in_zip).replace('\\', '/')
    if nav_dir == '.': nav_dir = "" # –ö–æ—Ä–µ–Ω—å

    link_count = 0
    for item_path, item_title in nav_data_list:
        safe_item_title = html.escape(str(item_title).strip())
        if not safe_item_title: safe_item_title = "Untitled Entry" # –ó–∞–≥–ª—É—à–∫–∞ –¥–ª—è –ø—É—Å—Ç—ã—Ö –∑–∞–≥–æ–ª–æ–≤–∫–æ–≤

        try:
            item_path_norm = item_path.replace('\\', '/').lstrip('/')
            nav_parent_dir_norm = os.path.dirname(nav_file_path_in_zip.replace('\\','/').lstrip('/')).replace('\\','/')

            relative_href = os.path.relpath(item_path_norm, start=nav_parent_dir_norm if nav_parent_dir_norm else '.').replace('\\', '/')
            safe_href = html.escape(relative_href, quote=True)

            li = etree.SubElement(ol, f"{{{xhtml_ns}}}li")
            a = etree.SubElement(li, f"{{{xhtml_ns}}}a", href=safe_href)
            a.text = safe_item_title
            link_count += 1


        except ValueError as e:
            print(f"[WARN] NAV Gen (lxml): Failed to calculate relative path for '{item_path}' from '{nav_parent_dir_norm or '<root>'}': {e}. Skipping link.")
        except Exception as e_loop:
             print(f"[ERROR] NAV Gen (lxml): Error processing item ('{item_path}', '{item_title}'): {e_loop}")

    if link_count == 0 and len(nav_data_list) > 0:
         print("[WARN] NAV Gen (lxml): No list items were added to NAV despite input data.")
         ol.append(etree.Comment(" Error: No valid links generated "))
    elif link_count != len(nav_data_list):
        print(f"[WARN] NAV Gen (lxml): Added {link_count} links, but received {len(nav_data_list)} data items.")

    nav_output_string = etree.tostring(html_tag, encoding='unicode', method='html', xml_declaration=False, pretty_print=True)

    doctype = '<!DOCTYPE html>'
    xml_declaration = "<?xml version='1.0' encoding='utf-8'?>"
    final_output = f"{xml_declaration}\n{doctype}\n{nav_output_string}"


    print(f"[INFO] NAV Gen (lxml): Finished generation for '{nav_file_path_in_zip}'. Added {link_count} links.")
    return final_output.encode('utf-8') # –í–æ–∑–≤—Ä–∞—â–∞–µ–º –±–∞–π—Ç—ã UTF-8



def generate_ncx_manual(book_id, book_title, ncx_data_list):
    """
    Generates the content of an NCX file manually from a prepared list of data
    derived from nav.xhtml.

    Args:
        book_id (str): The unique identifier for the book (for dtb:uid).
        book_title (str): The title of the book (for docTitle).
        ncx_data_list (list): A list of tuples extracted from nav.xhtml:
                              [(nav_point_id, content_src, link_text), ...].
                              - nav_point_id: Pre-generated ID for the navPoint.
                              - content_src: Pre-calculated relative path for content src.
                              - link_text: Text label for the navPoint.

    Returns:
        bytes: The generated NCX content as bytes (UTF-8 encoded XML), or None if error.
    """
    if not ncx_data_list:
        print("[WARN] NCX Manual Gen: Input data list is empty. NCX not generated.")
        return None

    print(f"[INFO] NCX Manual Gen: Starting NCX generation from {len(ncx_data_list)} NAV entries...")

    ncx_lines = []
    ncx_lines.append("<?xml version='1.0' encoding='utf-8'?>")
    ncx_lines.append('<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">')
    ncx_lines.append('  <head>')

    safe_book_id = html.escape(book_id or f"urn:uuid:{uuid.uuid4()}", quote=True)
    ncx_lines.append(f'    <meta content="{safe_book_id}" name="dtb:uid"/>')

    ncx_lines.append('    <meta content="1" name="dtb:depth"/>') # –°—Ç–∞–≤–∏–º 1, –µ—Å–ª–∏ –µ—Å—Ç—å navPoints
    ncx_lines.append('    <meta content="0" name="dtb:totalPageCount"/>')
    ncx_lines.append('    <meta content="0" name="dtb:maxPageNumber"/>')
    ncx_lines.append('  </head>')

    safe_book_title = html.escape(book_title or "Untitled")
    ncx_lines.append('  <docTitle>')
    ncx_lines.append(f'    <text>{safe_book_title}</text>')
    ncx_lines.append('  </docTitle>')

    ncx_lines.append('  <docAuthor>')
    ncx_lines.append(f'    <text>Translator</text>') # –ú–æ–∂–Ω–æ –∑–∞–º–µ–Ω–∏—Ç—å –Ω–∞ —á—Ç–æ-—Ç–æ –¥—Ä—É–≥–æ–µ
    ncx_lines.append('  </docAuthor>')

    ncx_lines.append('  <navMap>')

    play_order_counter = 1
    for nav_point_id, content_src, link_text in ncx_data_list:

        safe_id = html.escape(nav_point_id, quote=True)
        safe_label = html.escape(link_text)
        safe_src = html.escape(content_src, quote=True)

        ncx_lines.append(f'    <navPoint id="{safe_id}" playOrder="{play_order_counter}">')
        ncx_lines.append('      <navLabel>')
        ncx_lines.append(f'        <text>{safe_label}</text>')
        ncx_lines.append('      </navLabel>')
        ncx_lines.append(f'      <content src="{safe_src}"/>')
        ncx_lines.append('    </navPoint>')
        play_order_counter += 1

    ncx_lines.append('  </navMap>')
    ncx_lines.append('</ncx>')

    ncx_output_string = "\n".join(ncx_lines)
    print(f"[INFO] NCX Manual Gen: Generated {play_order_counter - 1} navPoints from NAV data.")


    return ncx_output_string.encode('utf-8')


def parse_nav_for_ncx_data(nav_content_bytes, nav_base_path_in_zip):
    """–ò–∑–≤–ª–µ–∫–∞–µ—Ç –¥–∞–Ω–Ω—ã–µ –∏–∑ NAV XHTML –¥–ª—è –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ NCX."""
    if not nav_content_bytes or not BS4_AVAILABLE: return []
    ncx_data = []
    play_order = 1
    try:
        soup = BeautifulSoup(nav_content_bytes, 'lxml-xml') # –ò—Å–ø–æ–ª—å–∑—É–µ–º XML –ø–∞—Ä—Å–µ—Ä
        nav_list = soup.find('nav', attrs={'epub:type': 'toc'})
        if not nav_list: nav_list = soup # Fallback, –µ—Å–ª–∏ –Ω–µ—Ç <nav>
        list_tag = nav_list.find(['ol', 'ul'])
        if not list_tag: return []

        nav_dir = os.path.dirname(nav_base_path_in_zip).replace('\\', '/')
        if nav_dir == '.': nav_dir = "" # –ö–æ—Ä–µ–Ω—å

        for link in list_tag.find_all('a', href=True):
            href = link.get('href')
            text = link.get_text(strip=True)
            if not href or not text or href.startswith('#') or href.startswith(('http:', 'https:', 'mailto:')):
                continue

            try:

                abs_path_in_zip = os.path.normpath(os.path.join(nav_dir, unquote(href))).replace('\\', '/')
                content_src = abs_path_in_zip.lstrip('/') # NCX src –æ–±—ã—á–Ω–æ –æ—Ç –∫–æ—Ä–Ω—è

                content_src_base = urlparse(content_src).path

                safe_base_name = re.sub(r'[^\w\-]+', '_', Path(content_src_base).stem)
                nav_point_id = f"navpoint_{safe_base_name}_{play_order}"

                ncx_data.append((nav_point_id, content_src, text)) # –°–æ—Ö—Ä–∞–Ω—è–µ–º –ø—É—Ç—å —Å —Ñ—Ä–∞–≥–º–µ–Ω—Ç–æ–º, –µ—Å–ª–∏ –±—ã–ª
                play_order += 1
            except Exception as e:
                print(f"[WARN NavParseForNCX] Error processing NAV link '{href}': {e}")
        return ncx_data
    except Exception as e:
        print(f"[ERROR NavParseForNCX] Failed to parse NAV content: {e}")
        return []



def parse_ncx_for_nav_data(ncx_content_bytes, opf_dir):
    """–ò–∑–≤–ª–µ–∫–∞–µ—Ç –¥–∞–Ω–Ω—ã–µ –∏–∑ NCX –¥–ª—è –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ NAV HTML."""
    if not ncx_content_bytes or not LXML_AVAILABLE: return []
    nav_data = [] # –ë—É–¥–µ—Ç —Å–æ–¥–µ—Ä–∂–∞—Ç—å –∫–æ—Ä—Ç–µ–∂–∏: (–ø—É—Ç—å_–æ—Ç_–∫–æ—Ä–Ω—è_zip, –∑–∞–≥–æ–ª–æ–≤–æ–∫)
    try:
        root = etree.fromstring(ncx_content_bytes)
        ns = {'ncx': 'http://www.daisy.org/z3986/2005/ncx/'}
        for nav_point in root.xpath('//ncx:navMap/ncx:navPoint', namespaces=ns):
            content_tag = nav_point.find('ncx:content', ns)
            label_tag = nav_point.find('.//ncx:text', ns)

            if content_tag is not None and label_tag is not None:
                src = content_tag.get('src')
                text = label_tag.text.strip() if label_tag.text else "Untitled"
                if not src: continue

                try:

                    unquoted_src = unquote(urlparse(src).path) # –£–±–∏—Ä–∞–µ–º URL-–∫–æ–¥–∏—Ä–æ–≤–∞–Ω–∏–µ –∏ —Ñ—Ä–∞–≥–º–µ–Ω—Ç—ã

                    if opf_dir:

                        abs_path_in_zip = os.path.normpath(os.path.join(opf_dir, unquoted_src)).replace('\\', '/')
                    else:

                        abs_path_in_zip = os.path.normpath(unquoted_src).replace('\\', '/')

                    abs_path_in_zip = '/'.join(part for part in abs_path_in_zip.split('/') if part != '..')
                    abs_path_in_zip = abs_path_in_zip.lstrip('/')

                    nav_data.append((abs_path_in_zip, text))

                except Exception as e:
                    print(f"[WARN NcxParseForNav] Error processing NCX src '{src}': {e}")
        return nav_data
    except Exception as e:
        print(f"[ERROR NcxParseForNav] Failed to parse NCX content: {e}")
        return []



def update_nav_content(nav_content_bytes, nav_base_path_in_zip, filename_map, canonical_titles):
    """–û–±–Ω–æ–≤–ª—è–µ—Ç href –∏ —Ç–µ–∫—Å—Ç —Å—Å—ã–ª–æ–∫ –≤ —Å—É—â–µ—Å—Ç–≤—É—é—â–µ–º NAV –∫–æ–Ω—Ç–µ–Ω—Ç–µ."""
    if not nav_content_bytes or not BS4_AVAILABLE: return None
    try:
        soup = BeautifulSoup(nav_content_bytes, 'lxml-xml')
        nav_list = soup.find('nav', attrs={'epub:type': 'toc'})
        if not nav_list: nav_list = soup
        list_tag = nav_list.find(['ol', 'ul'])
        if not list_tag: return nav_content_bytes # –ù–µ –Ω–∞—à–ª–∏ —Å–ø–∏—Å–æ–∫, –≤–æ–∑–≤—Ä–∞—â–∞–µ–º –∫–∞–∫ –µ—Å—Ç—å

        nav_dir = os.path.dirname(nav_base_path_in_zip).replace('\\', '/')
        if nav_dir == '.': nav_dir = "" # –ö–æ—Ä–µ–Ω—å

        updated_count = 0
        for link in list_tag.find_all('a', href=True):
            href = link.get('href')
            if not href or href.startswith('#') or href.startswith(('http:', 'https:', 'mailto:')):
                continue

            original_target_full_path = None
            frag = None
            try:

                original_target_full_path = os.path.normpath(os.path.join(nav_dir, unquote(urlparse(href).path))).replace('\\', '/').lstrip('/')
                frag = urlparse(href).fragment

            except Exception as e:
                print(f"[WARN NAV Update] Error resolving original path for href '{href}': {e}")
                continue

            new_target_relative_path = filename_map.get(original_target_full_path)


            if new_target_relative_path:
                try:

                    nav_parent_dir = os.path.dirname(nav_base_path_in_zip).replace('\\', '/') # –î–∏—Ä–µ–∫—Ç–æ—Ä–∏—è, –≥–¥–µ –ª–µ–∂–∏—Ç NAV
                    new_rel_href = os.path.relpath(new_target_relative_path, start=nav_parent_dir).replace('\\', '/')


                    new_href_val = new_rel_href + (f"#{frag}" if frag else "")
                    link['href'] = new_href_val # –û–±–Ω–æ–≤–ª—è–µ–º href
                    updated_count += 1
                except ValueError as e:
                    print(f"[WARN NAV Update] Error calculating relative href for '{new_target_relative_path}' from '{nav_parent_dir}': {e}")

            target_canonical_title = canonical_titles.get(original_target_full_path)
            if target_canonical_title:
                link.string = html.escape(str(target_canonical_title).strip()) # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –Ω–æ–≤—ã–π —Ç–µ–∫—Å—Ç


        print(f"[INFO] NAV Update: Updated attributes for {updated_count} links.")

        return str(soup).encode('utf-8')

    except Exception as e:
        print(f"[ERROR NAV Update] Failed to update NAV content: {e}\n{traceback.format_exc()}")
        return None # –í–æ–∑–≤—Ä–∞—â–∞–µ–º None –≤ —Å–ª—É—á–∞–µ –æ—à–∏–±–∫–∏



def update_ncx_content(ncx_content_bytes, opf_dir, filename_map, canonical_titles):
    """–û–±–Ω–æ–≤–ª—è–µ—Ç src –∏ text –≤ —Å—É—â–µ—Å—Ç–≤—É—é—â–µ–º NCX –∫–æ–Ω—Ç–µ–Ω—Ç–µ."""
    if not ncx_content_bytes or not LXML_AVAILABLE: return None
    try:

        ncx_ns_uri = 'http://www.daisy.org/z3986/2005/ncx/'
        ns = {'ncx': ncx_ns_uri}


        root = etree.fromstring(ncx_content_bytes)
        updated_count = 0

        for nav_point in root.xpath('//ncx:navPoint', namespaces=ns):
            content_tag = nav_point.find('ncx:content', ns)
            label_tag = nav_point.find('.//ncx:text', ns) # –ò—â–µ–º text –≤–Ω—É—Ç—Ä–∏ navLabel

            if content_tag is None or label_tag is None: continue

            src = content_tag.get('src')
            if not src: continue

            original_target_full_path = None
            frag = None
            try:

                original_target_full_path = os.path.normpath(os.path.join(opf_dir, unquote(urlparse(src).path))).replace('\\', '/').lstrip('/')
                frag = urlparse(src).fragment

            except Exception as e:
                 print(f"[WARN NCX Update] Error resolving original path for src '{src}': {e}")
                 continue

            new_target_relative_path = filename_map.get(original_target_full_path)


            if new_target_relative_path:
                try:

                    if opf_dir: # –ï—Å–ª–∏ OPF –Ω–µ –≤ –∫–æ—Ä–Ω–µ
                        new_src = os.path.relpath(new_target_relative_path, start=opf_dir).replace('\\', '/')
                    else: # OPF –≤ –∫–æ—Ä–Ω–µ, –Ω–æ–≤—ã–π –ø—É—Ç—å —É–∂–µ –æ—Ç–Ω–æ—Å–∏—Ç–µ–ª–µ–Ω –∫–æ—Ä–Ω—é
                        new_src = new_target_relative_path


                    new_src_val = new_src + (f"#{frag}" if frag else "")
                    content_tag.set('src', new_src_val) # –û–±–Ω–æ–≤–ª—è–µ–º src
                    updated_count += 1
                except ValueError as e:
                    print(f"[WARN NCX Update] Error calculating relative src for '{new_target_relative_path}' from '{opf_dir or '<root>'}': {e}")

            target_canonical_title = canonical_titles.get(original_target_full_path)
            if target_canonical_title:
                label_tag.text = str(target_canonical_title).strip() # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –Ω–æ–≤—ã–π —Ç–µ–∫—Å—Ç


        print(f"[INFO] NCX Update: Updated attributes for {updated_count} navPoints.")

        return etree.tostring(root, encoding='utf-8', xml_declaration=True, pretty_print=True)

    except Exception as e:
        print(f"[ERROR NCX Update] Failed to update NCX content: {e}\n{traceback.format_exc()}")
        return None # –í–æ–∑–≤—Ä–∞—â–∞–µ–º None –≤ —Å–ª—É—á–∞–µ –æ—à–∏–±–∫–∏

def write_to_epub(out_path, processed_epub_parts, original_epub_path, build_metadata, book_title_override=None):
    start_time = time.time()
    if not EBOOKLIB_AVAILABLE: return False, "EbookLib library is required"
    if not LXML_AVAILABLE: return False, "lxml library is required"
    if not BS4_AVAILABLE: return False, "BeautifulSoup4 required"
    if not os.path.exists(original_epub_path): return False, f"Original EPUB not found: {original_epub_path}"

    print(f"[INFO] EPUB Rebuild: Starting rebuild for '{os.path.basename(original_epub_path)}' -> '{out_path}'")
    book = epub.EpubBook()

    nav_path_orig_from_meta = build_metadata.get('nav_path_in_zip')
    ncx_path_orig_from_meta = build_metadata.get('ncx_path_in_zip')
    opf_dir_from_meta = build_metadata.get('opf_dir', '') # –≠—Ç–æ –¥–∏—Ä–µ–∫—Ç–æ—Ä–∏—è OPF –≤ –æ—Ä–∏–≥–∏–Ω–∞–ª—å–Ω–æ–º EPUB
    nav_id_orig_from_meta = build_metadata.get('nav_item_id')
    ncx_id_orig_from_meta = build_metadata.get('ncx_item_id')

    final_book_title = book_title_override or Path(original_epub_path).stem
    final_author = "Translator"; final_identifier = f"urn:uuid:{uuid.uuid4()}"; final_language = "ru"

    original_manifest_items_from_zip = {} # {path_in_zip: {id, media_type, properties, original_href}}
    original_spine_idrefs_from_zip = []

    combined_new_image_map_from_worker = build_metadata.get('combined_image_map', {})


    filename_map = {} # original_full_path_in_zip -> new_full_path_in_zip (–¥–ª—è –æ–±–Ω–æ–≤–ª–µ–Ω–∏—è NAV/NCX)
    final_book_item_ids = set() # –î–ª—è –æ—Ç—Å–ª–µ–∂–∏–≤–∞–Ω–∏—è —É–Ω–∏–∫–∞–ª—å–Ω–æ—Å—Ç–∏ ID
    book_items_to_add_to_epub_obj = [] # –°–ø–∏—Å–æ–∫ –æ–±—ä–µ–∫—Ç–æ–≤ EpubItem, EpubHtml, EpubImage –¥–ª—è –¥–æ–±–∞–≤–ª–µ–Ω–∏—è –≤ book

    new_book_items_structure_map = {} 
    id_to_new_item_map = {} # –î–ª—è –±—ã—Å—Ç—Ä–æ–≥–æ –¥–æ—Å—Ç—É–ø–∞ –ø–æ ID –≤ spine
    
    processed_original_paths_from_zip = set() # –û—Ç—Å–ª–µ–∂–∏–≤–∞—Ç—å, –∫–∞–∫–∏–µ —Ñ–∞–π–ª—ã –∏–∑ ZIP —É–∂–µ –æ–±—Ä–∞–±–æ—Ç–∞–Ω—ã
    canonical_titles_map = {} # original_full_path_in_zip -> canonical_title

    opf_dir_for_new_epub = opf_dir_from_meta # –î–∏—Ä–µ–∫—Ç–æ—Ä–∏—è OPF –≤ –ù–û–í–û–ú EPUB (–æ–±—ã—á–Ω–æ —Ç–∞ –∂–µ)

    try:
        with zipfile.ZipFile(original_epub_path, 'r') as original_zip:
            zip_contents_normalized = {name.replace('\\', '/'): name for name in original_zip.namelist()}
            opf_path_in_zip_abs = None

            try:
                container_data = original_zip.read('META-INF/container.xml')
                container_root = etree.fromstring(container_data); cnt_ns = {'c': 'urn:oasis:names:tc:opendocument:xmlns:container'}
                opf_path_rel_to_container = container_root.xpath('//c:rootfile/@full-path', namespaces=cnt_ns)[0]
                opf_path_in_zip_abs = opf_path_rel_to_container.replace('\\', '/')

                temp_opf_dir_check = os.path.dirname(opf_path_in_zip_abs).replace('\\','/')
                temp_opf_dir_check = "" if temp_opf_dir_check == '.' else temp_opf_dir_check.lstrip('/')
                if opf_dir_for_new_epub != temp_opf_dir_check:
                    print(f"[WARN] OPF directory mismatch: Meta='{opf_dir_for_new_epub}', Re-check='{temp_opf_dir_check}'. Using meta: '{opf_dir_for_new_epub}'.")
            except Exception: # Fallback
                pot_opf = [p for p in zip_contents_normalized if p.lower().endswith('.opf') and not p.lower().startswith('meta-inf/') and p.lower() != 'mimetype']
                if not pot_opf: pot_opf = [p for p in zip_contents_normalized if p.lower().endswith('.opf') and p.lower() != 'mimetype']
                if not pot_opf: raise FileNotFoundError("Cannot find OPF in original EPUB.")
                opf_path_in_zip_abs = pot_opf[0]

            
            if not opf_path_in_zip_abs: raise FileNotFoundError("OPF path could not be determined.")

            opf_data_bytes = original_zip.read(zip_contents_normalized[opf_path_in_zip_abs])
            opf_root = etree.fromstring(opf_data_bytes)
            ns_opf_parse = {'opf': 'http://www.idpf.org/2007/opf', 'dc': 'http://purl.org/dc/elements/1.1/'}

            meta_node = opf_root.find('.//opf:metadata', ns_opf_parse) or opf_root.find('.//metadata')
            if meta_node is not None:
                def get_text_meta(element): return element.text.strip() if element is not None and element.text else None
                lang_node = meta_node.find('.//dc:language', ns_opf_parse) or meta_node.find('.//language')
                title_node = meta_node.find('.//dc:title', ns_opf_parse) or meta_node.find('.//title')
                creator_node = meta_node.find('.//dc:creator', ns_opf_parse) or meta_node.find('.//creator')
                id_element = meta_node.find('.//dc:identifier[@id]', ns_opf_parse) or \
                             meta_node.find('.//identifier[@id]', ns_opf_parse) or \
                             meta_node.find('.//dc:identifier', ns_opf_parse) or \
                             meta_node.find('.//identifier')
                final_language = get_text_meta(lang_node) or final_language
                final_book_title = book_title_override or get_text_meta(title_node) or final_book_title
                final_author = get_text_meta(creator_node) or final_author
                final_identifier = get_text_meta(id_element) or final_identifier or f"urn:uuid:{uuid.uuid4()}"
            book.set_title(final_book_title); book.add_author(final_author); book.set_identifier(final_identifier); book.set_language(final_language)


            manifest_node = opf_root.find('.//opf:manifest', ns_opf_parse) or opf_root.find('.//manifest')
            if manifest_node is not None:
                for item_mf_loop in (manifest_node.findall('.//opf:item', ns_opf_parse) or manifest_node.findall('.//item')):
                    item_id = item_mf_loop.get('id'); href = item_mf_loop.get('href'); media_type = item_mf_loop.get('media-type'); props = item_mf_loop.get('properties')
                    if not item_id or not href or not media_type: continue

                    full_path_in_zip = os.path.normpath(os.path.join(opf_dir_from_meta, unquote(href))).replace('\\', '/').lstrip('/')
                    original_manifest_items_from_zip[full_path_in_zip] = {'id': item_id, 'media_type': media_type, 'properties': props, 'original_href': href}
            
            spine_node = opf_root.find('.//opf:spine', ns_opf_parse) or opf_root.find('.//spine')
            ncx_id_from_spine_attr = None
            if spine_node is not None:
                ncx_id_from_spine_attr = spine_node.get('toc') # –≠—Ç–æ ID NCX —Ñ–∞–π–ª–∞ –∏–∑ –º–∞–Ω–∏—Ñ–µ—Å—Ç–∞
                original_spine_idrefs_from_zip = [i_ref.get('idref') for i_ref in (spine_node.findall('.//opf:itemref', ns_opf_parse) or spine_node.findall('.//itemref')) if i_ref.get('idref')]

            if nav_path_orig_from_meta and nav_path_orig_from_meta in zip_contents_normalized:
                try:
                    nav_data_bytes = original_zip.read(zip_contents_normalized[nav_path_orig_from_meta])
                    nav_soup = BeautifulSoup(nav_data_bytes, 'lxml-xml')
                    nav_list_el = nav_soup.find('nav', attrs={'epub:type': 'toc'}) or nav_soup
                    list_tag_nav = nav_list_el.find(['ol', 'ul'])
                    if list_tag_nav:
                        nav_dir_current = os.path.dirname(nav_path_orig_from_meta).replace('\\', '/')
                        if nav_dir_current == '.': nav_dir_current = ""
                        for link in list_tag_nav.find_all('a', href=True):
                            href = link.get('href'); title_text = link.get_text(strip=True)
                            if not href or not title_text or href.startswith(('#', 'http:', 'mailto:')): continue
                            try:
                                target_full_path = os.path.normpath(os.path.join(nav_dir_current, unquote(urlparse(href).path))).replace('\\', '/').lstrip('/')
                                if target_full_path not in canonical_titles_map: canonical_titles_map[target_full_path] = title_text
                            except Exception: pass
                except Exception as nav_err_read: print(f"[WARN write_epub] Error reading original NAV for titles: {nav_err_read}")
            elif ncx_path_orig_from_meta and ncx_path_orig_from_meta in zip_contents_normalized:
                 try:
                    ncx_data_bytes = original_zip.read(zip_contents_normalized[ncx_path_orig_from_meta])
                    ncx_root_titles = etree.fromstring(ncx_data_bytes); ncx_ns_titles = {'ncx': 'http://www.daisy.org/z3986/2005/ncx/'}
                    for nav_point in ncx_root_titles.xpath('//ncx:navMap/ncx:navPoint', namespaces=ncx_ns_titles):
                         content_tag = nav_point.find('ncx:content', ncx_ns_titles); label_tag = nav_point.find('.//ncx:text', ncx_ns_titles)
                         if content_tag is not None and label_tag is not None and content_tag.get('src'):
                             src_attr = content_tag.get('src'); title_text = label_tag.text.strip() if label_tag.text else None
                             if not src_attr or not title_text: continue
                             try:
                                 target_full_path = os.path.normpath(os.path.join(opf_dir_from_meta, unquote(urlparse(src_attr).path))).replace('\\', '/').lstrip('/')
                                 if target_full_path not in canonical_titles_map: canonical_titles_map[target_full_path] = title_text
                             except Exception: pass
                 except Exception as ncx_err_read: print(f"[WARN write_epub] Error reading original NCX for titles: {ncx_err_read}")

            new_image_objects_for_manifest = {} # uuid -> EpubImage object
            img_counter = 1
            for img_uuid, new_img_info in combined_new_image_map_from_worker.items():
                temp_img_path = new_img_info.get('saved_path')
                if not temp_img_path or not os.path.exists(temp_img_path):
                    print(f"[WARN write_epub] New image for UUID {img_uuid} has invalid temp path: '{temp_img_path}'. Skipping.")
                    continue
                try:
                    with open(temp_img_path, 'rb') as f_new_img: img_data_bytes = f_new_img.read()
                    content_type = new_img_info.get('content_type', 'image/jpeg')
                    ext_new_img = content_type.split('/')[-1]; ext_new_img = 'jpg' if ext_new_img == 'jpeg' else ext_new_img
                    orig_fname_for_new = new_img_info.get('original_filename', f'new_image_{img_uuid[:6]}.{ext_new_img}')

                    img_folder_in_epub = "Images" # –ú–æ–∂–Ω–æ —Å–¥–µ–ª–∞—Ç—å –Ω–∞—Å—Ç—Ä–∞–∏–≤–∞–µ–º—ã–º
                    new_img_rel_path_in_epub = os.path.join(img_folder_in_epub, re.sub(r'[^\w\.\-]', '_', orig_fname_for_new)).replace('\\','/')
                    
                    new_img_id = f"new_img_{img_uuid[:6]}_{img_counter}"
                    if new_img_id in final_book_item_ids: new_img_id = f"{new_img_id}_{uuid.uuid4().hex[:3]}"
                    
                    epub_img_obj_new = epub.EpubImage(uid=new_img_id, file_name=new_img_rel_path_in_epub, media_type=content_type, content=img_data_bytes)
                    book_items_to_add_to_epub_obj.append(epub_img_obj_new)
                    new_image_objects_for_manifest[img_uuid] = epub_img_obj_new # –î–ª—è –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏—è –≤ _convert_placeholders
                    final_book_item_ids.add(new_img_id)

                    new_img_abs_path_in_epub = os.path.normpath(os.path.join(opf_dir_for_new_epub, new_img_rel_path_in_epub)).replace('\\','/').lstrip('/')
                    new_book_items_structure_map[new_img_abs_path_in_epub] = {'item': epub_img_obj_new, 'content_bytes': None, 'canonical_title': None}
                    id_to_new_item_map[new_img_id] = new_book_items_structure_map[new_img_abs_path_in_epub]
                    processed_original_paths_from_zip.add(new_img_abs_path_in_epub) # –ü–æ–º–µ—á–∞–µ–º, —á—Ç–æ —ç—Ç–æ—Ç –ø—É—Ç—å —É–∂–µ –∑–∞–Ω—è—Ç –Ω–æ–≤—ã–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ–º
                    img_counter += 1
                except Exception as e_new_img:
                    print(f"[ERROR write_epub] Failed to add new image (UUID {img_uuid}): {e_new_img}")

            print(f"[INFO write_epub] –ù–∞—á–∞–ª–æ –æ–±—Ä–∞–±–æ—Ç–∫–∏ {len(processed_epub_parts)} HTML-—á–∞—Å—Ç–µ–π –¥–ª—è —Å–±–æ—Ä–∫–∏...")
            
            for part_data in processed_epub_parts:

                if 'content_to_write' not in part_data or part_data['content_to_write'] is None:
                    original_fn_for_skip = part_data.get('original_filename', '–ù–µ–∏–∑–≤–µ—Å—Ç–Ω—ã–π HTML')
                    warning_msg_for_skip = part_data.get('translation_warning', '–î–∞–Ω–Ω—ã–µ –∫–æ–Ω—Ç–µ–Ω—Ç–∞ –æ—Ç—Å—É—Ç—Å—Ç–≤—É—é—Ç –∏–ª–∏ –ø–æ–≤—Ä–µ–∂–¥–µ–Ω—ã')
                    print(f"[WARN write_epub] –ü—Ä–æ–ø—É—Å–∫ HTML-—á–∞—Å—Ç–∏ '{original_fn_for_skip}', —Ç–∞–∫ –∫–∞–∫ 'content_to_write' –æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç –∏–ª–∏ None. –ü—Ä–∏—á–∏–Ω–∞: {warning_msg_for_skip}")
                    if original_fn_for_skip:
                         processed_original_paths_from_zip.add(original_fn_for_skip)
                    continue 


                original_html_path_in_zip = part_data['original_filename'] 
                content_to_use = part_data['content_to_write']
                image_map_for_this_part = part_data.get('image_map', {})
                is_original = part_data.get('is_original_content', False)
                
                original_item_info = original_manifest_items_from_zip.get(original_html_path_in_zip)
                if not original_item_info:
                    print(f"[WARN write_epub] –ù–µ—Ç –∑–∞–ø–∏—Å–∏ –≤ –º–∞–Ω–∏—Ñ–µ—Å—Ç–µ –¥–ª—è –æ—Ä–∏–≥–∏–Ω–∞–ª—å–Ω–æ–≥–æ HTML: {original_html_path_in_zip}. –ü—Ä–æ–ø—É—Å–∫ —ç—Ç–æ–π —á–∞—Å—Ç–∏.")
                    processed_original_paths_from_zip.add(original_html_path_in_zip)
                    continue

                original_item_id = original_item_info['id']
                original_href_from_manifest = original_item_info['original_href'] # –ü—É—Ç—å –æ—Ç–Ω–æ—Å–∏—Ç–µ–ª—å–Ω–æ OPF
                
                new_html_rel_path_in_epub = "" # –ü—É—Ç—å –Ω–æ–≤–æ–≥–æ —Ñ–∞–π–ª–∞ –æ—Ç–Ω–æ—Å–∏—Ç–µ–ª—å–Ω–æ OPF
                final_html_content_bytes = None

                current_part_canonical_title = canonical_titles_map.get(original_html_path_in_zip) 
                
                if is_original:
                    new_html_rel_path_in_epub = original_href_from_manifest.replace('\\', '/')
                    final_html_content_bytes = content_to_use # –≠—Ç–æ —É–∂–µ bytes

                    abs_path_for_map = os.path.normpath(os.path.join(opf_dir_for_new_epub, new_html_rel_path_in_epub)).replace('\\','/').lstrip('/')
                    filename_map[original_html_path_in_zip] = abs_path_for_map
                    
                    if not current_part_canonical_title and final_html_content_bytes:
                         try:
                             temp_html_str_orig = final_html_content_bytes.decode('utf-8', errors='replace')
                             temp_soup_orig = BeautifulSoup(temp_html_str_orig, 'lxml') 
                             extracted_title = None
                             h_tag = temp_soup_orig.find(['h1','h2','h3','h4','h5','h6'])
                             title_tag = temp_soup_orig.head.title if temp_soup_orig.head else None
                             if h_tag and h_tag.get_text(strip=True): extracted_title = h_tag.get_text(strip=True)
                             elif title_tag and title_tag.string:
                                 stripped_title = title_tag.string.strip()
                                 generic_titles = ['untitled', 'unknown', 'navigation', 'toc', 'table of contents', 'index', 'contents', '–æ–≥–ª–∞–≤–ª–µ–Ω–∏–µ', '—Å–æ–¥–µ—Ä–∂–∞–Ω–∏–µ', '–∏–Ω–¥–µ–∫—Å', 'cover', 'title page', 'copyright', 'chapter']
                                 if stripped_title and stripped_title.lower() not in generic_titles and len(stripped_title) > 1:
                                     extracted_title = stripped_title
                             if extracted_title: current_part_canonical_title = extracted_title
                         except Exception as e_title_orig_extract: print(f"[DEBUG write_epub] –û—à–∏–±–∫–∞ –∏–∑–≤–ª–µ—á–µ–Ω–∏—è –∑–∞–≥–æ–ª–æ–≤–∫–∞ –∏–∑ –æ—Ä–∏–≥–∏–Ω–∞–ª—å–Ω–æ–≥–æ HTML {original_html_path_in_zip}: {e_title_orig_extract}")
                
                else: # –ü–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω—ã–π –∫–æ–Ω—Ç–µ–Ω—Ç (content_to_use —ç—Ç–æ —Å—Ç—Ä–æ–∫–∞ —Å Markdown-like —Ä–∞–∑–º–µ—Ç–∫–æ–π –∏ –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä–∞–º–∏)
                    new_html_rel_path_in_epub = add_translated_suffix(original_href_from_manifest).replace('\\', '/')

                    temp_title_for_conversion = current_part_canonical_title
                    if not temp_title_for_conversion and isinstance(content_to_use, str):
                        first_line_md = content_to_use.split('\n', 1)[0].strip()
                        md_h_match = re.match(r'^(#{1,6})\s+(.*)', first_line_md)
                        if md_h_match: temp_title_for_conversion = md_h_match.group(2).strip()
                    if not temp_title_for_conversion: # –ï—Å–ª–∏ –≤—Å–µ –µ—â–µ –Ω–µ—Ç, –∏—Å–ø–æ–ª—å–∑—É–µ–º –∏–º—è —Ñ–∞–π–ª–∞
                        temp_title_for_conversion = Path(new_html_rel_path_in_epub).stem.replace('_translated', '').replace('_', ' ').capitalize()

                    final_html_str_rendered = _convert_placeholders_to_html_img(
                        text_with_placeholders=content_to_use, 
                        item_image_map_for_this_html=image_map_for_this_part, 
                        epub_new_image_objects=new_image_objects_for_manifest, 
                        canonical_title=temp_title_for_conversion, # –ò—Å–ø–æ–ª—å–∑—É–µ–º –≤—Ä–µ–º–µ–Ω–Ω—ã–π/–ø—Ä–µ–¥–ø–æ–ª–∞–≥–∞–µ–º—ã–π –∑–∞–≥–æ–ª–æ–≤–æ–∫
                        current_html_file_path_relative_to_opf=new_html_rel_path_in_epub,
                        opf_dir_path=opf_dir_for_new_epub
                    )

                    actual_translated_title_from_html = None
                    try:
                        soup_final_html = BeautifulSoup(final_html_str_rendered, 'lxml') 
                        h1_tag = soup_final_html.body.find('h1') if soup_final_html.body else None
                        if h1_tag and h1_tag.get_text(strip=True):
                            actual_translated_title_from_html = h1_tag.get_text(strip=True)
                        else: 
                            title_tag_final = soup_final_html.head.title if soup_final_html.head else None
                            if title_tag_final and title_tag_final.string:
                                stripped_final_title = title_tag_final.string.strip()
                                generic_titles_check = ['untitled', 'unknown', 'navigation', 'toc', 'table of contents', 'index', 'contents', '–æ–≥–ª–∞–≤–ª–µ–Ω–∏–µ', '—Å–æ–¥–µ—Ä–∂–∞–Ω–∏–µ', '–∏–Ω–¥–µ–∫—Å', 'cover', 'title page', 'copyright', 'chapter']
                                if stripped_final_title and stripped_final_title.lower() not in generic_titles_check and len(stripped_final_title) > 1:
                                    actual_translated_title_from_html = stripped_final_title
                        
                        if actual_translated_title_from_html:
                            current_part_canonical_title = actual_translated_title_from_html 

                    except Exception as e_title_extract_final:
                        print(f"[WARN write_epub] –ù–µ —É–¥–∞–ª–æ—Å—å –∏–∑–≤–ª–µ—á—å –∑–∞–≥–æ–ª–æ–≤–æ–∫ –∏–∑ —Ñ–∏–Ω–∞–ª—å–Ω–æ–≥–æ HTML –¥–ª—è {new_html_rel_path_in_epub}: {e_title_extract_final}")

                    if actual_translated_title_from_html:
                        try:

                            soup_to_update_title = BeautifulSoup(final_html_str_rendered, 'lxml')
                            if soup_to_update_title.head:
                                if soup_to_update_title.head.title:
                                    soup_to_update_title.head.title.string = html.escape(actual_translated_title_from_html)
                                else: # –ï—Å–ª–∏ —Ç–µ–≥–∞ <title> –Ω–µ—Ç, –Ω–æ –µ—Å—Ç—å <head>
                                    new_title_tag_in_head = soup_to_update_title.new_tag("title")
                                    new_title_tag_in_head.string = html.escape(actual_translated_title_from_html)
                                    soup_to_update_title.head.insert(0, new_title_tag_in_head)
                                final_html_str_rendered = str(soup_to_update_title) # –û–±–Ω–æ–≤–ª—è–µ–º —Å—Ç—Ä–æ–∫—É

                        except Exception as e_title_force_update:
                            print(f"[WARN write_epub] –ù–µ —É–¥–∞–ª–æ—Å—å –ø—Ä–∏–Ω—É–¥–∏—Ç–µ–ª—å–Ω–æ –æ–±–Ω–æ–≤–∏—Ç—å —Ç–µ–≥ <title> –≤ {new_html_rel_path_in_epub}: {e_title_force_update}")
                    
                    final_html_content_bytes = final_html_str_rendered.encode('utf-8')
                    abs_path_for_map_translated = os.path.normpath(os.path.join(opf_dir_for_new_epub, new_html_rel_path_in_epub)).replace('\\','/').lstrip('/')
                    filename_map[original_html_path_in_zip] = abs_path_for_map_translated

                if not current_part_canonical_title:
                    cleaned_stem = Path(new_html_rel_path_in_epub).stem.replace('_translated', '')
                    cleaned_stem = re.sub(r'^[\d_-]+', '', cleaned_stem) # –£–¥–∞–ª—è–µ–º –ø—Ä–µ—Ñ–∏–∫—Å—ã —Ç–∏–ø–∞ "01_", "001-"
                    cleaned_stem = cleaned_stem.replace('_', ' ').replace('-', ' ').strip()
                    current_part_canonical_title = cleaned_stem.capitalize() if cleaned_stem else f"–î–æ–∫—É–º–µ–Ω—Ç {original_item_id}"
                
                canonical_titles_map[original_html_path_in_zip] = current_part_canonical_title # –û–±–Ω–æ–≤–ª—è–µ–º –≥–ª–æ–±–∞–ª—å–Ω—É—é –∫–∞—Ä—Ç—É –∑–∞–≥–æ–ª–æ–≤–∫–æ–≤

                final_html_item_id = original_item_id
                if final_html_item_id in final_book_item_ids: # –û–±–µ—Å–ø–µ—á–∏–≤–∞–µ–º —É–Ω–∏–∫–∞–ª—å–Ω–æ—Å—Ç—å ID
                    final_html_item_id = f"html_{Path(new_html_rel_path_in_epub).stem}_{uuid.uuid4().hex[:4]}"
                
                epub_html_obj = epub.EpubHtml(
                    uid=final_html_item_id,
                    file_name=new_html_rel_path_in_epub, # –ü—É—Ç—å –æ—Ç–Ω–æ—Å–∏—Ç–µ–ª—å–Ω–æ OPF
                    title=html.escape(current_part_canonical_title), # –ò—Å–ø–æ–ª—å–∑—É–µ–º —Ñ–∏–Ω–∞–ª—å–Ω—ã–π –∫–∞–Ω–æ–Ω–∏—á–µ—Å–∫–∏–π –∑–∞–≥–æ–ª–æ–≤–æ–∫
                    lang=final_language,
                    content=final_html_content_bytes # –≠—Ç–æ –≤—Å–µ–≥–¥–∞ bytes
                )
                epub_html_obj.media_type = 'application/xhtml+xml'
                
                book_items_to_add_to_epub_obj.append(epub_html_obj)
                final_book_item_ids.add(final_html_item_id)

                new_html_abs_path_in_epub_map_key = os.path.normpath(os.path.join(opf_dir_for_new_epub, new_html_rel_path_in_epub)).replace('\\','/').lstrip('/')
                new_book_items_structure_map[new_html_abs_path_in_epub_map_key] = {
                    'item': epub_html_obj, 
                    'content_bytes': final_html_content_bytes, # –°–æ—Ö—Ä–∞–Ω—è–µ–º –±–∞–π—Ç—ã –¥–ª—è –≤–æ–∑–º–æ–∂–Ω–æ–≥–æ –ø–æ–≤—Ç–æ—Ä–Ω–æ–≥–æ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏—è
                    'canonical_title': current_part_canonical_title
                }
                id_to_new_item_map[final_html_item_id] = new_book_items_structure_map[new_html_abs_path_in_epub_map_key]
                processed_original_paths_from_zip.add(original_html_path_in_zip) # –ü–æ–º–µ—á–∞–µ–º –æ—Ä–∏–≥–∏–Ω–∞–ª—å–Ω—ã–π –ø—É—Ç—å –∫–∞–∫ –æ–±—Ä–∞–±–æ—Ç–∞–Ω–Ω—ã–π

            items_to_skip_copying = set() # NAV, NCX –∏–∑ build_metadata
            if nav_path_orig_from_meta: items_to_skip_copying.add(nav_path_orig_from_meta)
            if ncx_path_orig_from_meta: items_to_skip_copying.add(ncx_path_orig_from_meta)

            for orig_full_path, orig_item_info in original_manifest_items_from_zip.items():
                if orig_full_path in processed_original_paths_from_zip: # –£–∂–µ –æ–±—Ä–∞–±–æ—Ç–∞–Ω (HTML –∏–ª–∏ –∑–∞–º–µ–Ω–µ–Ω–Ω–æ–µ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ)
                    continue
                if orig_full_path in items_to_skip_copying: # –Ø–≤–Ω–æ –ø—Ä–æ–ø—É—Å–∫–∞–µ–º—ã–µ (—Å—Ç–∞—Ä—ã–µ NAV/NCX)
                    continue
                if orig_item_info.get('properties') and 'nav' in orig_item_info['properties'].split(): # –ü—Ä–æ–ø—É—Å–∫ —Å—Ç–∞—Ä–æ–≥–æ NAV –ø–æ —Å–≤–æ–π—Å—Ç–≤—É
                    continue
                
                actual_zip_entry_name = zip_contents_normalized.get(orig_full_path)
                if not actual_zip_entry_name: # Fallback if case mismatch or slight path variation
                    actual_zip_entry_name = next((o_name for norm_name, o_name in zip_contents_normalized.items() if norm_name.lower() == orig_full_path.lower()), None)
                if not actual_zip_entry_name:
                    print(f"[WARN write_epub] Original manifest item '{orig_full_path}' not found in ZIP. Skipping copy.")
                    continue
                
                try:
                    item_content_bytes = original_zip.read(actual_zip_entry_name)
                    item_id_copy = orig_item_info['id']
                    item_href_copy = orig_item_info['original_href'] # –≠—Ç–æ –ø—É—Ç—å –æ—Ç–Ω–æ—Å–∏—Ç–µ–ª—å–Ω–æ OPF
                    item_media_type_copy = orig_item_info['media_type']

                    if item_id_copy in final_book_item_ids: item_id_copy = f"item_copy_{Path(item_href_copy).stem}_{uuid.uuid4().hex[:3]}"
                    
                    new_item_obj_copy = None
                    if item_media_type_copy.startswith('image/'):
                        new_item_obj_copy = epub.EpubImage(uid=item_id_copy, file_name=item_href_copy, media_type=item_media_type_copy, content=item_content_bytes)
                    elif item_media_type_copy == 'text/css':
                        new_item_obj_copy = epub.EpubItem(uid=item_id_copy, file_name=item_href_copy, media_type=item_media_type_copy, content=item_content_bytes)
                    elif item_media_type_copy.startswith('font/') or item_media_type_copy in ['application/font-woff', 'application/vnd.ms-opentype', 'application/octet-stream', 'application/x-font-ttf']:
                        new_item_obj_copy = epub.EpubItem(uid=item_id_copy, file_name=item_href_copy, media_type=item_media_type_copy, content=item_content_bytes)
                    else: # –î—Ä—É–≥–∏–µ —Ç–∏–ø—ã —Ñ–∞–π–ª–æ–≤
                        new_item_obj_copy = epub.EpubItem(uid=item_id_copy, file_name=item_href_copy, media_type=item_media_type_copy, content=item_content_bytes)
                    
                    if new_item_obj_copy:
                        book_items_to_add_to_epub_obj.append(new_item_obj_copy)
                        final_book_item_ids.add(item_id_copy)

                        filename_map[orig_full_path] = os.path.normpath(os.path.join(opf_dir_for_new_epub, item_href_copy)).replace('\\','/').lstrip('/')

                        new_abs_path_copy = filename_map[orig_full_path]
                        new_book_items_structure_map[new_abs_path_copy] = {'item': new_item_obj_copy, 'content_bytes': item_content_bytes, 'canonical_title': None}
                        id_to_new_item_map[item_id_copy] = new_book_items_structure_map[new_abs_path_copy]
                        processed_original_paths_from_zip.add(orig_full_path)
                except KeyError:
                    print(f"[WARN write_epub] Original manifest item '{orig_full_path}' (href: {orig_item_info.get('original_href')}) could not be read from ZIP. Skipping.")
                except Exception as e_copy:
                    print(f"[ERROR write_epub] Failed to copy original manifest item '{orig_full_path}': {e_copy}")

            for item_obj in book_items_to_add_to_epub_obj:
                try: book.add_item(item_obj)
                except Exception as add_final_err: print(f"[ERROR write_epub] Failed to add item ID='{getattr(item_obj,'id','N/A')}' to book: {add_final_err}")

            
            final_nav_item_obj = None; final_ncx_item_obj = None
            new_nav_content_bytes = None; new_ncx_content_bytes = None

            final_nav_rel_path_in_epub = "nav.xhtml" # –°—Ç–∞–Ω–¥–∞—Ä—Ç–Ω–æ–µ –∏–º—è
            final_ncx_rel_path_in_epub = "toc.ncx"   # –°—Ç–∞–Ω–¥–∞—Ä—Ç–Ω–æ–µ –∏–º—è

            spine_item_objects_for_toc_gen = []
            for orig_idref in original_spine_idrefs_from_zip:

                original_item_path_for_idref = next((p for p, i_info in original_manifest_items_from_zip.items() if i_info['id'] == orig_idref), None)
                if not original_item_path_for_idref: continue
                
                new_item_abs_path = filename_map.get(original_item_path_for_idref)
                if not new_item_abs_path: continue

                new_item_entry = new_book_items_structure_map.get(new_item_abs_path)
                if not new_item_entry or not new_item_entry.get('item'): continue
                
                new_epub_item_obj = new_item_entry['item']
                if isinstance(new_epub_item_obj, epub.EpubHtml) and new_epub_item_obj.file_name.replace('\\','/') != final_nav_rel_path_in_epub:

                    item_title_for_toc = canonical_titles_map.get(original_item_path_for_idref, Path(new_epub_item_obj.file_name).stem)
                    spine_item_objects_for_toc_gen.append((new_epub_item_obj, item_title_for_toc))

            nav_item_id_to_use = nav_id_orig_from_meta or "nav"
            ncx_item_id_to_use = ncx_id_orig_from_meta or ncx_id_from_spine_attr or "ncx"

            if nav_path_orig_from_meta and nav_path_orig_from_meta in zip_contents_normalized: # –ë—ã–ª NAV
                print(f"[INFO write_epub] –û–±–Ω–æ–≤–ª–µ–Ω–∏–µ —Å—É—â–µ—Å—Ç–≤—É—é—â–µ–≥–æ NAV: {nav_path_orig_from_meta}")
                orig_nav_bytes = original_zip.read(zip_contents_normalized[nav_path_orig_from_meta])
                new_nav_content_bytes = update_nav_content(orig_nav_bytes, nav_path_orig_from_meta, filename_map, canonical_titles_map)
                if new_nav_content_bytes: final_nav_rel_path_in_epub = Path(nav_path_orig_from_meta).name # –°–æ—Ö—Ä–∞–Ω—è–µ–º –æ—Ä–∏–≥–∏–Ω–∞–ª—å–Ω–æ–µ –∏–º—è —Ñ–∞–π–ª–∞ NAV
            elif spine_item_objects_for_toc_gen: # –ù–µ –±—ã–ª–æ NAV, –Ω–æ –µ—Å—Ç—å —á—Ç–æ –¥–æ–±–∞–≤–∏—Ç—å –≤ spine
                print("[INFO write_epub] –ì–µ–Ω–µ—Ä–∞—Ü–∏—è –Ω–æ–≤–æ–≥–æ NAV –∏–∑ —ç–ª–µ–º–µ–Ω—Ç–æ–≤ spine...")
                nav_data_for_gen_html = []
                for item_obj_nav, title_nav in spine_item_objects_for_toc_gen:

                    abs_path_for_nav_href = os.path.normpath(os.path.join(opf_dir_for_new_epub, item_obj_nav.file_name)).replace('\\','/').lstrip('/')
                    nav_data_for_gen_html.append((abs_path_for_nav_href, title_nav))
                new_nav_content_bytes = generate_nav_html(nav_data_for_gen_html, 
                                                          os.path.join(opf_dir_for_new_epub, final_nav_rel_path_in_epub).replace('\\','/').lstrip('/'), 
                                                          final_book_title, final_language)

            if ncx_path_orig_from_meta and ncx_path_orig_from_meta in zip_contents_normalized: # –ë—ã–ª NCX
                print(f"[INFO write_epub] –û–±–Ω–æ–≤–ª–µ–Ω–∏–µ —Å—É—â–µ—Å—Ç–≤—É—é—â–µ–≥–æ NCX: {ncx_path_orig_from_meta}")
                orig_ncx_bytes = original_zip.read(zip_contents_normalized[ncx_path_orig_from_meta])
                new_ncx_content_bytes = update_ncx_content(orig_ncx_bytes, opf_dir_from_meta, filename_map, canonical_titles_map)
                if new_ncx_content_bytes: final_ncx_rel_path_in_epub = Path(ncx_path_orig_from_meta).name # –°–æ—Ö—Ä–∞–Ω—è–µ–º –æ—Ä–∏–≥–∏–Ω–∞–ª—å–Ω–æ–µ –∏–º—è —Ñ–∞–π–ª–∞ NCX
            elif new_nav_content_bytes: # –ù–µ –±—ã–ª–æ NCX, –Ω–æ —Å–≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞–ª–∏ NAV, –∏–∑ –Ω–µ–≥–æ –≥–µ–Ω–µ—Ä–∏—Ä—É–µ–º NCX
                 print("[INFO write_epub] –ì–µ–Ω–µ—Ä–∞—Ü–∏—è –Ω–æ–≤–æ–≥–æ NCX –∏–∑ –¥–∞–Ω–Ω—ã—Ö –Ω–æ–≤–æ–≥–æ NAV...")

                 nav_path_for_ncx_parse_abs = os.path.normpath(os.path.join(opf_dir_for_new_epub, final_nav_rel_path_in_epub)).replace('\\','/').lstrip('/')
                 ncx_data_from_new_nav = parse_nav_for_ncx_data(new_nav_content_bytes, nav_path_for_ncx_parse_abs)
                 if ncx_data_from_new_nav:
                      new_ncx_content_bytes = generate_ncx_manual(final_identifier, final_book_title, ncx_data_from_new_nav)
            elif spine_item_objects_for_toc_gen: # –ù–µ –±—ã–ª–æ –Ω–∏ NAV, –Ω–∏ NCX, –≥–µ–Ω–µ—Ä–∏—Ä—É–µ–º NCX –∏–∑ spine
                 print("[INFO write_epub] –ì–µ–Ω–µ—Ä–∞—Ü–∏—è –Ω–æ–≤–æ–≥–æ NCX –∏–∑ —ç–ª–µ–º–µ–Ω—Ç–æ–≤ spine (NAV –Ω–µ –±—ã–ª —Å–≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞–Ω)...")
                 ncx_data_from_spine_gen = []
                 for i_ncx, (item_obj_ncx, title_ncx) in enumerate(spine_item_objects_for_toc_gen):
                     ncx_src_for_gen = item_obj_ncx.file_name.replace('\\','/') # –û—Ç–Ω–æ—Å–∏—Ç–µ–ª—å–Ω–æ OPF
                     safe_base_ncx = re.sub(r'[^\w\-]+', '_', Path(ncx_src_for_gen).stem);
                     nav_point_id_ncx = f"navpoint_{safe_base_ncx}_{i_ncx+1}"
                     ncx_data_from_spine_gen.append((nav_point_id_ncx, ncx_src_for_gen, title_ncx))
                 if ncx_data_from_spine_gen:
                      new_ncx_content_bytes = generate_ncx_manual(final_identifier, final_book_title, ncx_data_from_spine_gen)

            if new_nav_content_bytes:
                if nav_item_id_to_use in final_book_item_ids: nav_item_id_to_use = f"{nav_item_id_to_use}_{uuid.uuid4().hex[:4]}"
                final_nav_item_obj = epub.EpubHtml(uid=nav_item_id_to_use, file_name=final_nav_rel_path_in_epub, title=final_book_title, lang=final_language, content=new_nav_content_bytes)
                final_nav_item_obj.media_type = 'application/xhtml+xml'

                if 'nav' not in final_nav_item_obj.properties: # –ü—Ä–æ–≤–µ—Ä—è–µ–º, –Ω–µ—Ç –ª–∏ —É–∂–µ —Ç–∞–∫–æ–≥–æ —Å–≤–æ–π—Å—Ç–≤–∞
                    final_nav_item_obj.properties.append('nav')


                book.add_item(final_nav_item_obj); final_book_item_ids.add(nav_item_id_to_use)
                book.toc = (final_nav_item_obj,) # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º NAV –∫–∞–∫ TOC
                print(f"[INFO write_epub] NAV –¥–æ–±–∞–≤–ª–µ–Ω/–æ–±–Ω–æ–≤–ª–µ–Ω. ID: {nav_item_id_to_use}, Path: {final_nav_rel_path_in_epub}")
            else: 
                book.toc = ()
                print(f"[INFO write_epub] NAV –∫–æ–Ω—Ç–µ–Ω—Ç –Ω–µ –±—ã–ª —Å–≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞–Ω/–æ–±–Ω–æ–≤–ª–µ–Ω. book.toc –±—É–¥–µ—Ç –ø—É—Å—Ç.")
            
            if new_ncx_content_bytes:
                if ncx_item_id_to_use in final_book_item_ids: ncx_item_id_to_use = f"{ncx_item_id_to_use}_{uuid.uuid4().hex[:4]}"
                final_ncx_item_obj = epub.EpubItem(uid=ncx_item_id_to_use, file_name=final_ncx_rel_path_in_epub, media_type='application/x-dtbncx+xml', content=new_ncx_content_bytes)
                book.add_item(final_ncx_item_obj); final_book_item_ids.add(ncx_item_id_to_use)

                book.spine_toc = final_ncx_item_obj.id 

                print(f"[INFO write_epub] NCX –¥–æ–±–∞–≤–ª–µ–Ω/–æ–±–Ω–æ–≤–ª–µ–Ω. ID: {ncx_item_id_to_use}, Path: {final_ncx_rel_path_in_epub}")
            elif ncx_id_from_spine_attr: 
                 existing_ncx_item = book.get_item_with_id(ncx_id_from_spine_attr)
                 if existing_ncx_item and existing_ncx_item.media_type == 'application/x-dtbncx+xml':
                     book.spine_toc = ncx_id_from_spine_attr
                     print(f"[INFO write_epub] –ò—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω —Å—É—â–µ—Å—Ç–≤—É—é—â–∏–π NCX –∏–∑ spine: ID={ncx_id_from_spine_attr}")

            final_spine_idrefs_for_book = []
            for orig_idref_spine in original_spine_idrefs_from_zip:
                original_path_for_idref_spine = next((p for p, item_info_spine in original_manifest_items_from_zip.items() if item_info_spine['id'] == orig_idref_spine), None)
                if not original_path_for_idref_spine: continue
                new_abs_path_for_idref_spine = filename_map.get(original_path_for_idref_spine)
                if not new_abs_path_for_idref_spine: continue
                new_item_entry_for_idref_spine = new_book_items_structure_map.get(new_abs_path_for_idref_spine)
                if new_item_entry_for_idref_spine and new_item_entry_for_idref_spine.get('item'):
                    final_spine_idrefs_for_book.append(new_item_entry_for_idref_spine['item'].id)
            
            if not final_spine_idrefs_for_book and spine_item_objects_for_toc_gen: # Fallback, –µ—Å–ª–∏ original_spine_idrefs_from_zip –ø—É—Å—Ç
                 final_spine_idrefs_for_book = [item_obj_s.id for item_obj_s, _ in spine_item_objects_for_toc_gen]

            book.spine = final_spine_idrefs_for_book
            if not book.spine: # –ö—Ä–∞–π–Ω–∏–π —Å–ª—É—á–∞–π: –¥–æ–±–∞–≤–ª—è–µ–º –ø–µ—Ä–≤—ã–π HTML, –µ—Å–ª–∏ spine –ø—É—Å—Ç
                first_html_item = next((item for item in book.items if isinstance(item, epub.EpubHtml) and item != final_nav_item_obj), None)
                if first_html_item: book.spine = [first_html_item.id]
                else: print("[WARN write_epub] –ù–µ —É–¥–∞–ª–æ—Å—å —Å—Ñ–æ—Ä–º–∏—Ä–æ–≤–∞—Ç—å spine, –Ω–µ—Ç –ø–æ–¥—Ö–æ–¥—è—â–∏—Ö HTML —ç–ª–µ–º–µ–Ω—Ç–æ–≤.")

            print(f"[INFO write_epub] –ó–∞–ø–∏—Å—å —Ñ–∏–Ω–∞–ª—å–Ω–æ–≥–æ EPUB —Ñ–∞–π–ª–∞ –≤: {out_path}...")
            epub.write_epub(out_path, book, {}) # –û–ø—Ü–∏–∏ –ø–æ —É–º–æ–ª—á–∞–Ω–∏—é
            end_time = time.time()
            print(f"[SUCCESS] EPUB Rebuild: –§–∞–π–ª —Å–æ—Ö—Ä–∞–Ω–µ–Ω: {out_path} (–ó–∞–Ω—è–ª–æ {end_time - start_time:.2f} —Å–µ–∫)")
            return True, None

    except FileNotFoundError as e_fnf:
        err_msg = f"EPUB Rebuild Error: –§–∞–π–ª –Ω–µ –Ω–∞–π–¥–µ–Ω - {e_fnf}"; print(f"[ERROR] {err_msg}"); return False, err_msg
    except (zipfile.BadZipFile, etree.XMLSyntaxError) as e_xml_zip:
        err_msg = f"EPUB Rebuild Error: –ù–µ —É–¥–∞–ª–æ—Å—å —Ä–∞–∑–æ–±—Ä–∞—Ç—å —Å—Ç—Ä—É–∫—Ç—É—Ä—É EPUB - {e_xml_zip}"; print(f"[ERROR] {err_msg}"); return False, err_msg
    except ImportError as e_imp:
        err_msg = f"EPUB Rebuild Error: –û—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç –±–∏–±–ª–∏–æ—Ç–µ–∫–∞ - {e_imp}"; print(f"[ERROR] {err_msg}"); return False, err_msg
    except ValueError as e_val:
        err_msg = f"EPUB Rebuild Error: {e_val}"; print(f"[ERROR] {err_msg}"); return False, err_msg
    except Exception as e_generic:
        tb_str = traceback.format_exc()
        err_msg = f"EPUB Rebuild Error: –ù–µ–æ–∂–∏–¥–∞–Ω–Ω–∞—è –æ—à–∏–±–∫–∞ - {type(e_generic).__name__}: {e_generic}"
        print(f"[ERROR] {err_msg}\n{tb_str}"); return False, err_msg

def _convert_placeholders_to_html_img(text_with_placeholders, item_image_map_for_this_html,
                                    epub_new_image_objects,
                                    canonical_title,
                                    current_html_file_path_relative_to_opf=None,
                                    opf_dir_path=None):
    if not text_with_placeholders: return ""
    if item_image_map_for_this_html is None: item_image_map_for_this_html = {}
    if epub_new_image_objects is None: epub_new_image_objects = {}

    def apply_inline_markdown_carefully(text_segment):
        known_tags_map = {}
        temp_id_counter = 0
        def tag_replacer(match):
            nonlocal temp_id_counter
            tag = match.group(0)
            placeholder = f"__HTML_TAG_PLACEHOLDER_{temp_id_counter}__"
            known_tags_map[placeholder] = tag
            temp_id_counter += 1
            return placeholder
        text_with_placeholders_for_tags = re.sub(r'(<br\s*/?>|<img\s+[^>]*?/>)', tag_replacer, text_segment, flags=re.IGNORECASE | re.DOTALL)
        def markdown_replacer(match_md):
            marker = match_md.group(1)
            content_to_wrap = html.escape(match_md.group(2))
            if marker == '**': return f'<strong>{content_to_wrap}</strong>'
            if marker == '*':  return f'<em>{content_to_wrap}</em>'
            if marker == '`':  return f'<code>{content_to_wrap}</code>'
            return match_md.group(0)
        processed_text_with_md = re.sub(r'(\*\*|\*|`)(.+?)\1', markdown_replacer, text_with_placeholders_for_tags, flags=re.DOTALL)
        final_text = processed_text_with_md
        for placeholder, original_tag in known_tags_map.items():
            final_text = final_text.replace(placeholder, original_tag)
        return final_text

    processed_parts_for_img_restore = []
    last_idx_img_restore = 0

    for placeholder_tag, img_uuid in find_image_placeholders(text_with_placeholders):
        match_start = text_with_placeholders.find(placeholder_tag, last_idx_img_restore)
        if match_start == -1: continue
        processed_parts_for_img_restore.append(text_with_placeholders[last_idx_img_restore:match_start])
        img_info = item_image_map_for_this_html.get(img_uuid)
        img_tag_html = f"<!-- Placeholder Error: UUID {img_uuid} not fully processed -->"
        if img_info:
            original_src_from_html_map = img_info.get('original_src')
            final_img_src_attr_value_for_tag = None
            final_attributes_for_tag = dict(img_info.get('attributes', {}))
            if original_src_from_html_map is not None:
                final_img_src_attr_value_for_tag = original_src_from_html_map
                final_attributes_for_tag.pop('{http://www.w3.org/1999/xlink}href', None)
                final_attributes_for_tag.pop('xlink:href', None)
            elif img_uuid in epub_new_image_objects:
                epub_img_object_for_new = epub_new_image_objects.get(img_uuid)
                if epub_img_object_for_new:
                    image_path_rel_to_opf_for_new = epub_img_object_for_new.file_name.replace('\\', '/')
                    if current_html_file_path_relative_to_opf is not None:
                        html_dir_rel_to_opf_for_new = os.path.dirname(current_html_file_path_relative_to_opf).replace('\\', '/')
                        if html_dir_rel_to_opf_for_new == '.': html_dir_rel_to_opf_for_new = ""
                        try: final_img_src_attr_value_for_tag = os.path.relpath(image_path_rel_to_opf_for_new, start=html_dir_rel_to_opf_for_new).replace('\\', '/')
                        except ValueError: final_img_src_attr_value_for_tag = image_path_rel_to_opf_for_new
                    else: final_img_src_attr_value_for_tag = image_path_rel_to_opf_for_new
            if final_img_src_attr_value_for_tag is not None:
                alt_text_raw = final_attributes_for_tag.get('alt', img_info.get('original_filename', f'Image {img_uuid[:7]}'))
                alt_text_escaped = html.escape(str(alt_text_raw), quote=True)
                attr_strings_list = [f'src="{html.escape(final_img_src_attr_value_for_tag, quote=True)}"', f'alt="{alt_text_escaped}"']
                for key, value in final_attributes_for_tag.items():
                    key_lower = str(key).lower()
                    if key_lower not in ['src', 'alt', 'xlink:href', '{http://www.w3.org/1999/xlink}href']:
                        attr_strings_list.append(f'{html.escape(str(key))}="{html.escape(str(value))}"')
                width_attr = final_attributes_for_tag.get('width'); height_attr = final_attributes_for_tag.get('height')
                styles_to_add = []
                if not width_attr or (isinstance(width_attr, str) and '%' in width_attr): styles_to_add.append("max-width: 100%;")
                if not height_attr or (isinstance(height_attr, str) and '%' in height_attr):
                     if "max-width: 100%;" in styles_to_add and not height_attr : styles_to_add.append("height: auto;")
                if styles_to_add: attr_strings_list.append(f'style="{html.escape(" ".join(styles_to_add))}"')
                img_tag_html = f"<img {' '.join(attr_strings_list)} />"
        processed_parts_for_img_restore.append(img_tag_html)
        last_idx_img_restore = match_start + len(placeholder_tag)
    processed_parts_for_img_restore.append(text_with_placeholders[last_idx_img_restore:])
    text_after_img_restore = "".join(processed_parts_for_img_restore)

    text_normalized_newlines = re.sub(r'<br\s*/?>', '\n', text_after_img_restore, flags=re.IGNORECASE)

    text_normalized_newlines = re.sub(r'\n{3,}', '\n\n', text_normalized_newlines)

    lines = text_normalized_newlines.splitlines() # –î–µ–ª–∏–º –ø–æ \n.

    html_body_segments = []
    paragraph_part_buffer = []
    current_list_tag_md = None
    in_code_block_md = False
    code_block_buffer_md = []

    heading_re_md = re.compile(r'^\s*(#{1,6})\s+(.*)')
    hr_re_md = re.compile(r'^\s*---\s*$')
    ul_item_re_md = re.compile(r'^\s*[\*\-]\s+(.*)')
    ol_item_re_md = re.compile(r'^\s*\d+\.\s+(.*)')
    code_fence_re_md = re.compile(r'^\s*```(.*)')

    def finalize_paragraph_md():
        nonlocal paragraph_part_buffer, html_body_segments
        if paragraph_part_buffer:

            para_content_raw = "<br />".join(paragraph_part_buffer) # –í–æ—Å—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º <br />
            processed_content = apply_inline_markdown_carefully(para_content_raw)
            html_body_segments.append(f"<p>{processed_content}</p>")
            paragraph_part_buffer = []

    def finalize_list_md():
        nonlocal current_list_tag_md, html_body_segments
        if current_list_tag_md:
            html_body_segments.append(f"</{current_list_tag_md}>")
            current_list_tag_md = None
            
    def finalize_code_block_md():
        nonlocal in_code_block_md, code_block_buffer_md, html_body_segments
        if code_block_buffer_md:
            escaped_code = html.escape("\n".join(code_block_buffer_md))
            html_body_segments.append(escaped_code)
        if in_code_block_md:
            html_body_segments.append("</code></pre>")
            in_code_block_md = False
        code_block_buffer_md = []

    for i, line_text in enumerate(lines): # line_text —ç—Ç–æ —Å—Ç—Ä–æ–∫–∞ –±–µ–∑ \n –Ω–∞ –∫–æ–Ω—Ü–µ
        stripped_line = line_text.strip()

        is_standalone_image = False
        if stripped_line.startswith("<img") and stripped_line.endswith("/>"):
            if re.fullmatch(r'\s*<img\s+[^>]*?/>\s*', line_text, re.IGNORECASE):
                is_standalone_image = True
        
        if is_standalone_image:
            finalize_paragraph_md()
            finalize_list_md()
            finalize_code_block_md()
            html_body_segments.append(line_text)
            continue

        code_fence_match = code_fence_re_md.match(stripped_line)
        if code_fence_match:
            finalize_paragraph_md()
            finalize_list_md()
            if not in_code_block_md:
                in_code_block_md = True
                code_block_buffer_md = []
                lang = html.escape(code_fence_match.group(1).strip())
                html_body_segments.append(f'<pre><code class="language-{lang}">' if lang else "<pre><code>")
            else:
                finalize_code_block_md()
            continue

        if in_code_block_md:
            code_block_buffer_md.append(line_text)
            continue

        if not stripped_line: # –ï—Å–ª–∏ —Å—Ç—Ä–æ–∫–∞ –ø—É—Å—Ç–∞ –ü–û–°–õ–ï strip
            finalize_paragraph_md()
            finalize_list_md() 

            continue # –ü–µ—Ä–µ—Ö–æ–¥–∏–º –∫ —Å–ª–µ–¥—É—é—â–µ–π —Å—Ç—Ä–æ–∫–µ

        heading_match = heading_re_md.match(line_text) 
        hr_match = hr_re_md.match(stripped_line) # hr –≤—Å–µ–≥–¥–∞ –Ω–∞ –≤—Å—é —Å—Ç—Ä–æ–∫—É
        ul_item_match = ul_item_re_md.match(line_text)
        ol_item_match = ol_item_re_md.match(line_text)

        is_block_markdown = bool(heading_match or hr_match or ul_item_match or ol_item_match)

        if is_block_markdown:
            finalize_paragraph_md() 

        if heading_match:
            finalize_list_md()
            level = len(heading_match.group(1))
            heading_text_raw = heading_match.group(2).strip() # strip() –∑–¥–µ—Å—å, —Ç.–∫. —ç—Ç–æ —Å–æ–¥–µ—Ä–∂–∏–º–æ–µ —Ç–µ–≥–∞
            processed_heading_text = apply_inline_markdown_carefully(heading_text_raw)
            html_body_segments.append(f"<h{level}>{processed_heading_text}</h{level}>")
        elif hr_match:
            finalize_list_md()
            html_body_segments.append("<hr />")
        elif ul_item_match:
            if current_list_tag_md != 'ul':
                finalize_list_md()
                html_body_segments.append("<ul>")
                current_list_tag_md = 'ul'
            list_item_raw = ul_item_match.group(1).strip() # strip() –∑–¥–µ—Å—å
            processed_list_item = apply_inline_markdown_carefully(list_item_raw)
            html_body_segments.append(f"<li>{processed_list_item}</li>")
        elif ol_item_match:
            if current_list_tag_md != 'ol':
                finalize_list_md()
                html_body_segments.append("<ol>")
                current_list_tag_md = 'ol'
            list_item_raw = ol_item_match.group(1).strip() # strip() –∑–¥–µ—Å—å
            processed_list_item = apply_inline_markdown_carefully(list_item_raw)
            html_body_segments.append(f"<li>{processed_list_item}</li>")
        else: # –ï—Å–ª–∏ —ç—Ç–æ –Ω–µ MD-–±–ª–æ–∫ –∏ –Ω–µ –ø—É—Å—Ç–∞—è —Å—Ç—Ä–æ–∫–∞ (—É–∂–µ –ø—Ä–æ–≤–µ—Ä–∏–ª–∏ stripped_line)
            finalize_list_md() # –ó–∞–∫—Ä—ã—Ç—å —Å–ø–∏—Å–æ–∫, –µ—Å–ª–∏ —ç—Ç–∞ —Å—Ç—Ä–æ–∫–∞ –Ω–µ —è–≤–ª—è–µ—Ç—Å—è –µ–≥–æ –ø—Ä–æ–¥–æ–ª–∂–µ–Ω–∏–µ–º

            paragraph_part_buffer.append(line_text)

    finalize_paragraph_md()
    finalize_list_md()
    finalize_code_block_md()

    body_content_final = "\n".join(html_body_segments)

    final_title_text_for_html_tag = html.escape(str(canonical_title or Path(current_html_file_path_relative_to_opf or "document").stem).strip())
    if not final_title_text_for_html_tag: final_title_text_for_html_tag = "Untitled Document"
    stylesheet_path_final = "../Styles/stylesheet.css"
    if current_html_file_path_relative_to_opf is not None:
        html_abs_dir_in_epub = ""
        if opf_dir_path: 
            abs_html_path_in_epub = os.path.join(opf_dir_path, current_html_file_path_relative_to_opf)
            html_abs_dir_in_epub = os.path.dirname(abs_html_path_in_epub)
        else: 
            html_abs_dir_in_epub = os.path.dirname(current_html_file_path_relative_to_opf)
        if html_abs_dir_in_epub == '.': html_abs_dir_in_epub = ""
        css_dir_from_root = os.path.join(opf_dir_path or "", "Styles")
        abs_stylesheet_path_in_epub = os.path.join(css_dir_from_root, "stylesheet.css")
        abs_stylesheet_path_in_epub = os.path.normpath(abs_stylesheet_path_in_epub).replace('\\','/')
        html_abs_dir_in_epub = os.path.normpath(html_abs_dir_in_epub).replace('\\','/')
        try: stylesheet_path_final = os.path.relpath(abs_stylesheet_path_in_epub, start=html_abs_dir_in_epub).replace('\\','/')
        except ValueError:
            if not html_abs_dir_in_epub: stylesheet_path_final = abs_stylesheet_path_in_epub.lstrip('/')
            else: stylesheet_path_final = abs_stylesheet_path_in_epub
    stylesheet_link_tag = f'<link rel="stylesheet" type="text/css" href="{html.escape(stylesheet_path_final, quote=True)}"/>'
    return f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" lang="ru" xml:lang="ru">
<head>
<meta charset="utf-8" />
<title>{final_title_text_for_html_tag}</title>
{stylesheet_link_tag}
</head>
<body>
{body_content_final}
</body>
</html>"""

def write_to_html(out_path, translated_content_with_placeholders, image_map, title):
    """Creates HTML file with embedded Base64 images."""
    if image_map is None: image_map = {}
    print(f"[INFO] HTML: Creating HTML file with embedded images: {out_path}")
    html_body_content = ""


    lines = translated_content_with_placeholders.splitlines() # –†–∞–∑–¥–µ–ª—è–µ–º –ø–æ \n, –µ—Å–ª–∏ –æ–Ω–∏ —Ç–∞–º –µ—Å—Ç—å (–æ–±—ã—á–Ω–æ –Ω–µ—Ç, –µ—Å–ª–∏ <br />)
    paragraph_buffer = []

    def process_text_block_for_html(text_block):

        
        processed_parts = []
        last_index = 0

        text_block_escaped_amp = text_block.replace('&', '&')

        text_block_br_protected = re.sub(r'<br\s*/?>', '__TEMP_BR_TAG__', text_block_escaped_amp, flags=re.IGNORECASE)

        text_block_lt_gt_escaped = text_block_br_protected.replace('<', '<').replace('>', '>')


        temp_md_text = text_block_lt_gt_escaped
        temp_md_text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', temp_md_text, flags=re.DOTALL)
        temp_md_text = re.sub(r'(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)', r'<em>\1</em>', temp_md_text, flags=re.DOTALL)
        temp_md_text = re.sub(r'`(.*?)`', r'<code>\1</code>', temp_md_text, flags=re.DOTALL)

        final_md_text = temp_md_text.replace('<strong>', '<strong>').replace('</strong>', '</strong>')
        final_md_text = final_md_text.replace('<em>', '<em>').replace('</em>', '</em>')
        final_md_text = final_md_text.replace('<code>', '<code>').replace('</code>', '</code>')

        text_with_md_and_br = final_md_text.replace('__TEMP_BR_TAG__', '<br />')


        placeholders = find_image_placeholders(text_with_md_and_br) # –ò—â–µ–º –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä—ã –≤ —Ç–µ–∫—Å—Ç–µ —Å Markdown –∏ <br />

        for placeholder_tag, img_uuid in placeholders:
            match_start = text_with_md_and_br.find(placeholder_tag, last_index)
            if match_start == -1: continue

            text_before = text_with_md_and_br[last_index:match_start]
            processed_parts.append(text_before) # –î–æ–±–∞–≤–ª—è–µ–º —Ç–µ–∫—Å—Ç "–∫–∞–∫ –µ—Å—Ç—å", –æ–Ω —É–∂–µ –æ–±—Ä–∞–±–æ—Ç–∞–Ω

            if img_uuid in image_map:
                img_info = image_map[img_uuid]; img_path = img_info['saved_path']
                if os.path.exists(img_path):
                    try:
                        with open(img_path, 'rb') as f_img: img_data = f_img.read()
                        b64_data = base64.b64encode(img_data).decode('ascii')
                        content_type = img_info.get('content_type', 'image/jpeg'); data_uri = f"data:{content_type};base64,{b64_data}"
                        alt_text_raw = img_info.get('original_filename', f'Image {img_uuid[:8]}');

                        alt_text = html.escape(alt_text_raw, quote=True)
                        img_tag = f'<img src="{html.escape(data_uri, quote=True)}" alt="{alt_text}" style="max-width: 100%; height: auto;" />'
                        processed_parts.append(img_tag) 
                    except Exception as img_err: print(f"[ERROR] HTML Write: Failed to read/encode image {img_path}: {img_err}"); processed_parts.append(f"[Err embed img: {img_uuid[:8]}]")
                else: print(f"[ERROR] HTML Write: Image path not found: {img_path}"); processed_parts.append(f"[Img path miss: {img_uuid[:8]}]")
            else: print(f"[WARN] HTML Write: Placeholder UUID '{img_uuid}' not found."); processed_parts.append(f"[Unk Img: {img_uuid[:8]}]")
            last_index = match_start + len(placeholder_tag)

        text_after = text_with_md_and_br[last_index:]
        processed_parts.append(text_after) # –î–æ–±–∞–≤–ª—è–µ–º –æ—Å—Ç–∞—Ç–æ–∫ —Ç–µ–∫—Å—Ç–∞ "–∫–∞–∫ –µ—Å—Ç—å"
        return "".join(processed_parts)

    current_list_type = None 
    in_code_block = False
    code_block_lines = []

    for line in lines: # line –º–æ–∂–µ—Ç —Å–æ–¥–µ—Ä–∂–∞—Ç—å <br />
        stripped_line = line.strip()
        is_code_fence = stripped_line == '```'

        if is_code_fence:
            if not in_code_block:
                if paragraph_buffer: html_body_content += f"<p>{process_text_block_for_html('<br/>'.join(paragraph_buffer))}</p>\n"; paragraph_buffer = []
                if current_list_type: html_body_content += f"</{current_list_type}>\n"; current_list_type = None
                in_code_block = True; code_block_lines = []
            else:
                in_code_block = False
                escaped_code = html.escape("\n".join(code_block_lines)) # –≠–∫—Ä–∞–Ω–∏—Ä—É–µ–º –≤—Å–µ —Å–æ–¥–µ—Ä–∂–∏–º–æ–µ –±–ª–æ–∫–∞ –∫–æ–¥–∞
                html_body_content += f"<pre><code>{escaped_code}</code></pre>\n"
            continue

        if in_code_block:
            code_block_lines.append(line); continue

        heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped_line)
        hr_match = stripped_line == '---'
        ul_match = re.match(r'^[\*\-]\s+(.*)', stripped_line)
        ol_match = re.match(r'^\d+\.\s+(.*)', stripped_line)

        if current_list_type and not ((current_list_type == 'ul' and ul_match) or (current_list_type == 'ol' and ol_match)):
             html_body_content += f"</{current_list_type}>\n"; current_list_type = None
        if paragraph_buffer and (heading_match or hr_match or ul_match or ol_match):
             para_content = process_text_block_for_html("<br/>".join(paragraph_buffer)); html_body_content += f"<p>{para_content}</p>\n" if para_content.strip() else ""; paragraph_buffer = []

        if heading_match:
            level = len(heading_match.group(1)); heading_text = process_text_block_for_html(heading_match.group(2).strip())
            if heading_text: html_body_content += f"<h{level}>{heading_text}</h{level}>\n"
        elif hr_match:
             html_body_content += "<hr/>\n"
        elif ul_match:
             if current_list_type != 'ul': html_body_content += "<ul>\n"; current_list_type = 'ul'
             list_text = process_text_block_for_html(ul_match.group(1).strip()); html_body_content += f"<li>{list_text}</li>\n"
        elif ol_match:
             if current_list_type != 'ol': html_body_content += "<ol>\n"; current_list_type = 'ol'
             list_text = process_text_block_for_html(ol_match.group(1).strip()); html_body_content += f"<li>{list_text}</li>\n"
        elif line or find_image_placeholders(line): 
             paragraph_buffer.append(line) # line —É–∂–µ —Å–æ–¥–µ—Ä–∂–∏—Ç <br /> –µ—Å–ª–∏ –æ–Ω–∏ –±—ã–ª–∏
        elif not stripped_line and paragraph_buffer: 

             para_content = process_text_block_for_html("".join(paragraph_buffer)); # –ù–µ —Å–æ–µ–¥–∏–Ω—è–µ–º —á–µ—Ä–µ–∑ <br/>, —Ç.–∫. –æ–Ω–∏ —É–∂–µ –µ—Å—Ç—å
             html_body_content += f"<p>{para_content}</p>\n" if para_content.strip() else ""; paragraph_buffer = []

    if current_list_type: html_body_content += f"</{current_list_type}>\n"
    if paragraph_buffer:
        para_content = process_text_block_for_html("".join(paragraph_buffer)); # –ù–µ —Å–æ–µ–¥–∏–Ω—è–µ–º —á–µ—Ä–µ–∑ <br/>
        html_body_content += f"<p>{para_content}</p>\n" if para_content.strip() else ""
    if in_code_block:
        escaped_code = html.escape("\n".join(code_block_lines))
        html_body_content += f"<pre><code>{escaped_code}</code></pre>\n"

    safe_title = html.escape(title or "–ü–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω—ã–π –¥–æ–∫—É–º–µ–Ω—Ç")
    html_template = f"""<!DOCTYPE html>
<html lang="ru">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{safe_title}</title>
<style>
body {{ font-family: sans-serif; line-height: 1.6; margin: 2em auto; max-width: 800px; padding: 0 1em; color: #333; background-color: #fdfdfd; }}
p {{ margin-top: 0; margin-bottom: 1em; text-align: justify; }}
h1, h2, h3, h4, h5, h6 {{ margin-top: 1.8em; margin-bottom: 0.6em; line-height: 1.3; font-weight: normal; color: #111; border-bottom: 1px solid #eee; padding-bottom: 0.2em;}}
h1 {{ font-size: 2em; }} h2 {{ font-size: 1.7em; }} h3 {{ font-size: 1.4em; }}
img {{ max-width: 100%; height: auto; display: block; margin: 1.5em auto; border: 1px solid #ddd; border-radius: 4px; box-shadow: 0 1px 3px rgba(0,0,0,0.1); }}
hr {{ border: none; border-top: 1px solid #ccc; margin: 2.5em 0; }}
ul, ol {{ margin-left: 1.5em; margin-bottom: 1em; padding-left: 1.5em; }}
li {{ margin-bottom: 0.4em; }}
strong {{ font-weight: bold; }}
em {{ font-style: italic; }}
a {{ color: #007bff; text-decoration: none; }} a:hover {{ text-decoration: underline; }}
code {{ background-color: #f0f0f0; padding: 0.1em 0.3em; border-radius: 3px; font-family: Consolas, monospace; font-size: 0.9em; }}
pre {{ background-color: #f5f5f5; border: 1px solid #ddd; border-radius: 4px; padding: 1em; overflow-x: auto; white-space: pre; }}
pre code {{ background-color: transparent; padding: 0; border-radius: 0; font-size: 0.9em; }}
</style>
</head>
<body>
{html_body_content.strip()}
</body>
</html>"""
    try:
        with open(out_path, "w", encoding="utf-8") as f: f.write(html_template)
        print(f"[SUCCESS] HTML file saved: {out_path}")
    except Exception as write_err: print(f"[ERROR] Failed to write HTML file {out_path}: {write_err}"); raise

def write_to_fb2(out_path, translated_content_with_placeholders, image_map, title):
    if not LXML_AVAILABLE: raise ImportError("lxml library is required to write FB2 files.")
    if image_map is None: image_map = {}
    print(f"[INFO] FB2: Creating FB2 file with image support: {out_path}")

    print(f"DEBUG write_to_fb2: image_map received with {len(image_map)} entries.")
    if image_map:
        print(f"  UUIDs in received image_map: {list(image_map.keys())}")


    placeholders_in_text = find_image_placeholders(translated_content_with_placeholders)
    print(f"DEBUG write_to_fb2: Found {len(placeholders_in_text)} placeholders in translated_content.")
    if placeholders_in_text:
        print(f"  UUIDs from placeholders in text: {[p[1] for p in placeholders_in_text]}")

    FB2_NS = "http://www.gribuser.ru/xml/fictionbook/2.0"; XLINK_NS = "http://www.w3.org/1999/xlink"; nsmap = {None: FB2_NS, "l": XLINK_NS}
    l_href_attr = f"{{{XLINK_NS}}}href" 
    fb2_root = etree.Element("FictionBook", nsmap=nsmap)

    description = etree.SubElement(fb2_root, "description")
    title_info = etree.SubElement(description, "title-info")
    document_info = etree.SubElement(description, "document-info")
    book_title_text = title or "–ü–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω—ã–π –î–æ–∫—É–º–µ–Ω—Ç" # Renamed variable
    etree.SubElement(title_info, "book-title").text = book_title_text
    author_elem = etree.SubElement(title_info, "author"); etree.SubElement(author_elem, "first-name").text = "Translator"
    etree.SubElement(title_info, "genre").text = "unspecified"; etree.SubElement(title_info, "lang").text = "ru"
    doc_author = etree.SubElement(document_info, "author"); etree.SubElement(doc_author, "nickname").text = "TranslatorApp"
    etree.SubElement(document_info, "program-used").text = "TranslatorApp using Gemini"; etree.SubElement(document_info, "date", attrib={"value": time.strftime("%Y-%m-%d")}).text = time.strftime("%d %B %Y", time.localtime()); etree.SubElement(document_info, "version").text = "1.0"

    binary_sections = []; placeholder_to_binary_id = {}; binary_id_counter = 1
    processed_uuids_for_binary = set()
    images_added_to_binary_count = 0 # New counter

    for placeholder_tag, img_uuid_from_text in placeholders_in_text: # Renamed img_uuid
        print(f"DEBUG write_to_fb2: Processing placeholder for UUID from text: {img_uuid_from_text}")
        if img_uuid_from_text in image_map and img_uuid_from_text not in processed_uuids_for_binary:
            img_info = image_map[img_uuid_from_text]
            img_path = img_info.get('saved_path')
            print(f"  UUID {img_uuid_from_text} found in image_map. Path: {img_path}")

            if img_path and os.path.exists(img_path):
                try:
                    with open(img_path, 'rb') as f_img: img_data = f_img.read()
                    base_id = f"img_{img_uuid_from_text[:8]}_{binary_id_counter}"; binary_id = re.sub(r'[^\w.-]', '_', base_id)
                    content_type = img_info.get('content_type', 'image/jpeg')
                    base64_encoded_data = base64.b64encode(img_data).decode('ascii')
                    
                    binary_sections.append((binary_id, content_type, base64_encoded_data))
                    placeholder_to_binary_id[img_uuid_from_text] = binary_id
                    processed_uuids_for_binary.add(img_uuid_from_text)
                    binary_id_counter += 1
                    images_added_to_binary_count +=1 # Increment counter
                    print(f"    Successfully prepared binary data for UUID {img_uuid_from_text}. Binary ID: {binary_id}")
                except Exception as e:
                    print(f"[ERROR] FB2: Failed to read/encode image {img_path} for UUID {img_uuid_from_text}: {e}")
            elif not img_path:
                print(f"[ERROR] FB2: No 'saved_path' found in image_map for UUID {img_uuid_from_text}.")
            else: # img_path exists in map, but file os.path.exists(img_path) is false
                print(f"[ERROR] FB2: Image path from image_map does not exist on disk: {img_path} (for UUID {img_uuid_from_text})")
        elif img_uuid_from_text not in image_map:
            print(f"  UUID {img_uuid_from_text} from placeholder NOT FOUND in image_map.")
        elif img_uuid_from_text in processed_uuids_for_binary:
            print(f"  UUID {img_uuid_from_text} already processed for binary section.")

    print(f"DEBUG write_to_fb2: Total images prepared for binary section: {images_added_to_binary_count}")

    body = etree.SubElement(fb2_root, "body")
    lines = translated_content_with_placeholders.splitlines()
    para_buffer = []; current_section = None; is_first_section = True

    def add_paragraph_to_fb2(target_element, para_lines):
        nonlocal is_first_section # Allow modification of outer scope variable
        if not para_lines: return
        full_para_text = "\n".join(para_lines).strip()
        if not full_para_text: return

        parent_section = target_element
        if parent_section is None or parent_section.tag != 'section':
             last_section = body.xpath('section[last()]')
             parent_section = last_section[0] if last_section else None
             if parent_section is None:
                 parent_section = etree.SubElement(body, "section")
                 is_first_section = False 

        p = etree.SubElement(parent_section, "p")
        last_index = 0; current_tail_element = None

        placeholders_in_para = find_image_placeholders(full_para_text)
        for placeholder_tag_para, img_uuid_para in placeholders_in_para: # Renamed variables
            match_start = full_para_text.find(placeholder_tag_para, last_index)
            if match_start == -1: continue

            text_before = full_para_text[last_index:match_start]
            if text_before:
                if current_tail_element is not None:
                    current_tail_element.tail = (current_tail_element.tail or "") + text_before
                else:
                    p.text = (p.text or "") + text_before

            if img_uuid_para in placeholder_to_binary_id:
                binary_id_para = placeholder_to_binary_id[img_uuid_para] # Renamed
                try:
                    img_elem = etree.SubElement(p, "image")
                    img_elem.set(l_href_attr, f"#{binary_id_para}")
                    current_tail_element = img_elem
                except ValueError as ve:
                    print(f"[ERROR] FB2: Failed to create image element for binary ID '{binary_id_para}': {ve}")
                    error_text_ve = f" [FB2 Img Err: {img_uuid_para[:8]}] " # Renamed
                    if current_tail_element is not None: current_tail_element.tail = (current_tail_element.tail or "") + error_text_ve
                    else: p.text = (p.text or "") + error_text_ve
                    current_tail_element = None
            else:
                original_filename_fb2 = image_map.get(img_uuid_para, {}).get('original_filename', img_uuid_para) # Renamed
                error_text_nf = f" [Img Placeholder {img_uuid_para[:8]} found in text, but no binary data prepared (orig: {original_filename_fb2})] " # Renamed
                print(f"DEBUG write_to_fb2 (add_paragraph): Placeholder {img_uuid_para} found in paragraph, but not in placeholder_to_binary_id map.")
                if current_tail_element is not None: current_tail_element.tail = (current_tail_element.tail or "") + error_text_nf
                else: p.text = (p.text or "") + error_text_nf
                current_tail_element = None
            last_index = match_start + len(placeholder_tag_para)

        text_after = full_para_text[last_index:]
        if text_after:
            if current_tail_element is not None:
                current_tail_element.tail = (current_tail_element.tail or "") + text_after
            else:
                p.text = (p.text or "") + text_after
        if len(p) == 0 and not (p.text or "").strip() and p.getparent() is not None:
             p.getparent().remove(p)

    for line in lines:
        stripped_line = line.strip()
        chapter_match = re.match(r'^(#{1,3})\s+(.*)', stripped_line)
        if chapter_match:
            if para_buffer and current_section is not None:
                add_paragraph_to_fb2(current_section, para_buffer)
            para_buffer = []
            current_section = etree.SubElement(body, "section")
            title_elem = etree.SubElement(current_section, "title")
            add_paragraph_to_fb2(title_elem, [chapter_match.group(2).strip()])
            if not title_elem.xpath('.//text() | .//image'):
                 current_section.remove(title_elem)
            is_first_section = False
        else:
            if is_first_section and not current_section and stripped_line:
                 current_section = etree.SubElement(body, "section")
                 is_first_section = False
            if stripped_line or find_image_placeholders(line): # Check raw line for placeholders
                para_buffer.append(line)
            elif not stripped_line and para_buffer:
                 if current_section is None:
                      current_section = etree.SubElement(body, "section"); is_first_section = False
                 add_paragraph_to_fb2(current_section, para_buffer); para_buffer = []
    if para_buffer:
        if current_section is None: current_section = etree.SubElement(body, "section")
        add_paragraph_to_fb2(current_section, para_buffer)
    if not body.xpath('section'):
        print("[WARN] FB2: No sections created. Adding empty fallback section.")
        etree.SubElement(body, "section")

    if binary_sections: # This list is populated based on successful processing
        print(f"[INFO] FB2: Adding {len(binary_sections)} binary image sections.") # This should match images_added_to_binary_count
        for binary_id_add, content_type_add, base64_data_add in binary_sections: # Renamed variables
            try:
                etree.SubElement(fb2_root, "binary", id=binary_id_add, attrib={"content-type": content_type_add}).text = base64_data_add
            except ValueError as ve_add:
                 print(f"[ERROR] FB2: Invalid binary ID '{binary_id_add}' during write: {ve_add}")
    else:
        print("[INFO] FB2: No binary image data to add.")

    try:
        tree = etree.ElementTree(fb2_root)
        tree.write(out_path, pretty_print=True, xml_declaration=True, encoding="utf-8")
        print(f"[SUCCESS] FB2 file saved: {out_path}")
    except Exception as write_err:
        print(f"[ERROR] Failed to write FB2 file {out_path}: {write_err}"); raise write_err

class EpubHtmlSelectorDialog(QDialog):

    def __init__(self, epub_filename, html_files, nav_path, ncx_path, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"–í—ã–±–µ—Ä–∏—Ç–µ HTML/XHTML —Ñ–∞–π–ª—ã –∏–∑ '{os.path.basename(epub_filename)}'")
        self.setMinimumWidth(500); self.setMinimumHeight(400) # –ú–æ–∂–Ω–æ –¥–∞–∂–µ —á—É—Ç—å –±–æ–ª—å—à–µ –≤—ã—Å–æ—Ç—É, –Ω–∞–ø—Ä–∏–º–µ—Ä 450
        layout = QVBoxLayout(self)
        info_text = f"–ù–∞–π–¥–µ–Ω–Ω—ã–µ HTML/XHTML —Ñ–∞–π–ª—ã –≤:\n{epub_filename}\n\n"
        info_text += f"–ê–≤—Ç–æ-–æ–ø—Ä–µ–¥–µ–ª–µ–Ω NAV (–û–≥–ª–∞–≤–ª–µ–Ω–∏–µ EPUB3): {nav_path or '–ù–µ—Ç'}\n"
        info_text += "\n–í—ã–±–µ—Ä–∏—Ç–µ —Ñ–∞–π–ª—ã –¥–ª—è –ø–µ—Ä–µ–≤–æ–¥–∞.\n(NAV —Ñ–∞–π–ª –†–ï–ö–û–ú–ï–ù–î–£–ï–¢–°–Ø –ò–°–ö–õ–Æ–ß–ò–¢–¨, —Ç.–∫. —Å—Å—ã–ª–∫–∏ –æ–±–Ω–æ–≤—è—Ç—Å—è –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏):"

        self.info_label = QLabel(info_text)
        layout.addWidget(self.info_label)

        self.hide_translated_checkbox = QCheckBox("–°–∫—Ä—ã—Ç—å —Ñ–∞–π–ª—ã _translated")
        self.hide_translated_checkbox.setToolTip(
            "–ï—Å–ª–∏ –æ—Ç–º–µ—á–µ–Ω–æ, —Ñ–∞–π–ª—ã —Å —Å—É—Ñ—Ñ–∏–∫—Å–æ–º _translated (–Ω–∞–ø—Ä–∏–º–µ—Ä, chapter1_translated.html) –±—É–¥—É—Ç —Å–∫—Ä—ã—Ç—ã –∏–∑ —Å–ø–∏—Å–∫–∞."
        )
        self.hide_translated_checkbox.setChecked(False)
        self.hide_translated_checkbox.stateChanged.connect(self.update_file_visibility) # –≠—Ç–∞ —Å—Ç—Ä–æ–∫–∞ –æ—Å—Ç–∞–µ—Ç—Å—è
        layout.addWidget(self.hide_translated_checkbox)

        self.list_widget = QListWidget()
        self.list_widget.setSelectionMode(QtWidgets.QAbstractItemView.SelectionMode.ExtendedSelection)

        self.list_widget.itemSelectionChanged.connect(self.update_selection_count_label) 

        self.all_html_files_with_data = [] # –≠—Ç–∞ —á–∞—Å—Ç—å –æ—Å—Ç–∞–µ—Ç—Å—è –∫–∞–∫ –±—ã–ª–∞
        for file_path in html_files:
            item = QtWidgets.QListWidgetItem(file_path)
            is_nav = (nav_path and file_path == nav_path)
            is_translated = Path(file_path).stem.endswith(TRANSLATED_SUFFIX) # –ü—Ä–æ–≤–µ—Ä—è–µ–º —Å—É—Ñ—Ñ–∏–∫—Å

            self.all_html_files_with_data.append({
                'text': file_path,
                'is_nav': is_nav,
                'is_translated': is_translated # –°–æ—Ö—Ä–∞–Ω—è–µ–º, —è–≤–ª—è–µ—Ç—Å—è –ª–∏ —Ñ–∞–π–ª –ø–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω—ã–º
            })

            if is_nav:
                item.setBackground(QtGui.QColor("#fff0f0")) # Light red background for NAV
                item.setToolTip(f"{file_path}\n(–≠—Ç–æ —Ñ–∞–π–ª –û–ì–õ–ê–í–õ–ï–ù–ò–Ø EPUB3 (NAV).\n–ù–ï –†–ï–ö–û–ú–ï–ù–î–£–ï–¢–°–Ø –ø–µ—Ä–µ–≤–æ–¥–∏—Ç—å - —Å—Å—ã–ª–∫–∏ –æ–±–Ω–æ–≤—è—Ç—Å—è –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏.)")
                item.setSelected(False) # Deselect NAV by default
            else:
                item_text_lower = item.text().lower()
                path = Path(item_text_lower)
                filename_lower = path.name
                filename_base = path.stem.split('.')[0] # Get stem before first dot

                skip_indicators = ['toc', 'nav', 'ncx', 'cover', 'title', 'index', 'copyright', 'about', 'meta', 'opf',
                                   'masthead', 'colophon', 'imprint', 'acknowledgments', 'dedication',
                                   'glossary', 'bibliography', 'notes', 'annotations', 'epigraph', 'halftitle',
                                   'frontmatter', 'backmatter', 'preface', 'introduction', 'appendix', 'biography',
                                   'isbn', 'legal', 'notice', 'otherbooks', 'prelims', 'team', 'promo', 'bonus']
                content_indicators = ['chapter', 'part', 'section', 'content', 'text', 'page', 'body', 'main', 'article',
                                    'chp', 'chap', 'prt', 'sec', 'glava', 'prologue', 'epilogue']

                is_likely_skip = any(skip in filename_base for skip in skip_indicators)
                parent_dir_lower = str(path.parent).lower()
                is_likely_skip = is_likely_skip or any(skip in parent_dir_lower for skip in ['toc', 'nav', 'meta', 'frontmatter', 'backmatter', 'index', 'notes'])
                is_likely_content = any(indicator in filename_base for indicator in content_indicators)
                is_chapter_like = re.fullmatch(r'(ch|gl|chap|chapter|part|section|sec|glava)[\d_-]+.*', filename_base) or \
                                  re.fullmatch(r'[\d]+', filename_base) or \
                                  re.match(r'^[ivxlcdm]+$', filename_base)

                if not is_likely_skip and (is_likely_content or is_chapter_like):
                     item.setSelected(True)
                else:
                    if not is_likely_skip and 'text' in filename_base:
                        item.setSelected(True)
                    else:
                        item.setSelected(False)
                item.setToolTip(file_path)

        
        layout.addWidget(self.list_widget) # –î–æ–±–∞–≤–ª—è–µ–º —Å–ø–∏—Å–æ–∫

        self.selection_count_label = QLabel("–í—ã–±—Ä–∞–Ω–æ: 0 –∏–∑ 0")
        layout.addWidget(self.selection_count_label)

        self.button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        self.button_box.accepted.connect(self.accept)
        self.button_box.rejected.connect(self.reject)
        layout.addWidget(self.button_box)

        self.update_file_visibility()


    def update_selection_count_label(self):
        """–û–±–Ω–æ–≤–ª—è–µ—Ç –º–µ—Ç–∫—É, –ø–æ–∫–∞–∑—ã–≤–∞—é—â—É—é –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ –≤—ã–±—Ä–∞–Ω–Ω—ã—Ö –∏ –æ–±—â–µ–µ –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ –≤–∏–¥–∏–º—ã—Ö —Ñ–∞–π–ª–æ–≤."""
        selected_items_count = len(self.list_widget.selectedItems())
        total_visible_items_count = self.list_widget.count() # count() –¥–∞–µ—Ç –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ —ç–ª–µ–º–µ–Ω—Ç–æ–≤ –≤ –≤–∏–¥–∂–µ—Ç–µ
        self.selection_count_label.setText(f"–í—ã–±—Ä–∞–Ω–æ: {selected_items_count} –∏–∑ {total_visible_items_count} (–≤–∏–¥–∏–º—ã—Ö)")

    def update_file_visibility(self):
        hide_translated = self.hide_translated_checkbox.isChecked()
        
        current_selected_text = None
        selected_items_list = self.list_widget.selectedItems() # QListWidget.selectedItems() returns a list
        if selected_items_list: # Check if the list is not empty
            current_selected_text = selected_items_list[0].text()

        self.list_widget.clear() 

        for file_data in self.all_html_files_with_data:
            if hide_translated and file_data['is_translated']:
                continue 

            item = QtWidgets.QListWidgetItem(file_data['text'])
            
            if file_data['is_nav']:
                item.setBackground(QtGui.QColor("#fff0f0"))
                item.setToolTip(f"{file_data['text']}\n(–≠—Ç–æ —Ñ–∞–π–ª –û–ì–õ–ê–í–õ–ï–ù–ò–Ø EPUB3 (NAV).\n–ù–ï –†–ï–ö–û–ú–ï–ù–î–£–ï–¢–°–Ø –ø–µ—Ä–µ–≤–æ–¥–∏—Ç—å - —Å—Å—ã–ª–∫–∏ –æ–±–Ω–æ–≤—è—Ç—Å—è –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏.)")
                item.setSelected(False) 
            else:
                item_text_lower = item.text().lower()
                path = Path(item_text_lower)
                filename_base = path.stem.split('.')[0] 

                skip_indicators = ['toc', 'nav', 'ncx', 'cover', 'title', 'index', 'copyright', 'about', 'meta', 'opf',
                                   'masthead', 'colophon', 'imprint', 'acknowledgments', 'dedication',
                                   'glossary', 'bibliography', 'notes', 'annotations', 'epigraph', 'halftitle',
                                   'frontmatter', 'backmatter', 'preface', 'introduction', 'appendix', 'biography',
                                   'isbn', 'legal', 'notice', 'otherbooks', 'prelims', 'team', 'promo', 'bonus']
                content_indicators = ['chapter', 'part', 'section', 'content', 'text', 'page', 'body', 'main', 'article',
                                    'chp', 'chap', 'prt', 'sec', 'glava', 'prologue', 'epilogue']

                is_likely_skip = any(skip in filename_base for skip in skip_indicators)
                parent_dir_lower = str(path.parent).lower()
                is_likely_skip = is_likely_skip or any(skip in parent_dir_lower for skip in ['toc', 'nav', 'meta', 'frontmatter', 'backmatter', 'index', 'notes'])
                
                is_likely_content = any(indicator in filename_base for indicator in content_indicators)

                is_chapter_like_match = re.fullmatch(r'(ch|gl|chap|chapter|part|section|sec|glava)[\d_-]+.*', filename_base) or \
                                        re.fullmatch(r'[\d]+', filename_base) or \
                                        re.match(r'^[ivxlcdm]+$', filename_base) 

                content_topic_criteria = bool(is_likely_content or is_chapter_like_match)

                should_be_selected = (not file_data['is_translated'] and
                                      not is_likely_skip and
                                      content_topic_criteria)

                if (not should_be_selected and
                    not file_data['is_translated'] and
                    not is_likely_skip and
                    'text' in filename_base):
                    should_be_selected = True
                
                item.setSelected(should_be_selected) # should_be_selected —Ç–µ–ø–µ—Ä—å –≤—Å–µ–≥–¥–∞ –±—É–¥–µ—Ç True –∏–ª–∏ False
                item.setToolTip(file_data['text'])
            
            self.list_widget.addItem(item)

            if current_selected_text and item.text() == current_selected_text:
                item.setSelected(True)

        self.update_selection_count_label() # <<< –í–û–¢ –≠–¢–£ –°–¢–†–û–ß–ö–£ –î–û–ë–ê–í–ò–õ–ò –í –ö–û–ù–ï–¶



    def get_selected_files(self):
        return [self.list_widget.item(i).text() for i in range(self.list_widget.count()) if self.list_widget.item(i).isSelected()]

class OperationCancelledError(Exception): pass






class Worker(QtCore.QObject):

    file_progress = QtCore.pyqtSignal(int)
    chunk_progress = QtCore.pyqtSignal(str, int, int)
    current_file_status = QtCore.pyqtSignal(str)
    log_message = QtCore.pyqtSignal(str)
    finished = QtCore.pyqtSignal(int, int, list)
    total_tasks_calculated = QtCore.pyqtSignal(int)

    def __init__(self, api_key, out_folder, prompt_template, files_to_process_data,
                 model_config, max_concurrent_requests, output_format,
                 chunking_enabled_gui, chunk_limit, chunk_window,
                 temperature, chunk_delay_seconds, proxy_string=None): # <-- –î–æ–±–∞–≤–ª–µ–Ω proxy_string
        super().__init__()
        self.api_key = api_key
        self.out_folder = out_folder
        self.prompt_template = prompt_template
        self.files_to_process_data = files_to_process_data
        self.model_config = model_config
        self.max_concurrent_requests = max_concurrent_requests
        self.output_format = output_format
        self.chunking_enabled_gui = chunking_enabled_gui
        self.chunk_limit = chunk_limit
        self.chunk_window = chunk_window
        self.temperature = temperature # <-- –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ç–µ–º–ø–µ—Ä–∞—Ç—É—Ä—É
        self.chunk_delay_seconds = chunk_delay_seconds # <-- –°–æ—Ö—Ä–∞–Ω—è–µ–º –Ω–æ–≤—É—é –Ω–∞—Å—Ç—Ä–æ–π–∫—É
        self.proxy_string = proxy_string # <-- –°–æ—Ö—Ä–∞–Ω—è–µ–º —Å—Ç—Ä–æ–∫—É –ø—Ä–æ–∫—Å–∏

        self.is_cancelled = False
        self.is_finishing = False # <--- –ù–û–í–´–ô –§–õ–ê–ì
        self._critical_error_occurred = False
        self.model = None
        self.executor = None
        self.epub_build_states = {}
        self.total_tasks = 0
        self.processed_task_count = 0
        self.success_count = 0
        self.error_count = 0
        self.errors_list = []


    def finish_processing(self): # <--- –í–û–¢ –≠–¢–û–¢ –ú–ï–¢–û–î
        if not self.is_finishing and not self.is_cancelled: # –ù–µ —É—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞—Ç—å, –µ—Å–ª–∏ —É–∂–µ –æ—Ç–º–µ–Ω—è–µ—Ç—Å—è
            self.log_message.emit("[SIGNAL] –ü–æ–ª—É—á–µ–Ω —Å–∏–≥–Ω–∞–ª –ó–ê–í–ï–†–®–ï–ù–ò–Ø (Worker.finish_processing)...")
            self.is_finishing = True


    def setup_client(self):
        """Initializes the Gemini API client, configures proxy, and sets system instruction."""
        try:
            if not self.api_key: raise ValueError("API –∫–ª—é—á –Ω–µ –ø—Ä–µ–¥–æ—Å—Ç–∞–≤–ª–µ–Ω.")

            # --- –ë–õ–û–ö –ü–†–û–ö–°–ò (–æ—Å—Ç–∞–µ—Ç—Å—è –±–µ–∑ –∏–∑–º–µ–Ω–µ–Ω–∏–π) ---
            if 'HTTP_PROXY' in os.environ: os.environ.pop('HTTP_PROXY')
            if 'HTTPS_PROXY' in os.environ: os.environ.pop('HTTPS_PROXY')
            applied_proxy_method = "None"
            proxy_url_for_env = None

            if self.proxy_string and self.proxy_string.strip():
                proxy_url_config = self.proxy_string.strip()
                if proxy_url_config.lower().startswith("socks5(h)://"):
                    corrected_url = "socks5h://" + proxy_url_config[len("socks5(h)://"):]
                    self.log_message.emit(f"[INFO] Proxy scheme '{proxy_url_config}' auto-corrected to '{corrected_url}'.")
                    proxy_url_config = corrected_url
                parsed_url = None
                try:
                    from urllib.parse import urlparse
                    parsed_url = urlparse(proxy_url_config)
                except Exception as e_parse_url:
                    self.log_message.emit(f"[ERROR] Could not parse proxy URL '{proxy_url_config}': {e_parse_url}")
                if parsed_url and parsed_url.scheme.lower() in ["socks5", "socks5h"]:
                    try:
                        import socks
                        import socket
                        host, port, username, password = parsed_url.hostname, parsed_url.port, parsed_url.username, parsed_url.password
                        if not host or not port: raise ValueError("SOCKS5/SOCKS5h URL is missing host or port.")
                        is_rdns = parsed_url.scheme.lower() == "socks5h"
                        if not hasattr(socks, '_original_socket_module_attrs'): socks._original_socket_module_attrs = {'socket': socket.socket}
                        elif 'socket' not in socks._original_socket_module_attrs: socks._original_socket_module_attrs['socket'] = socket.socket
                        socks.set_default_proxy(socks.SOCKS5, host, port, rdns=is_rdns, username=username, password=password)
                        socket.socket = socks.socksocket
                        applied_proxy_method = f"{parsed_url.scheme.upper()} via PySocks: {host}:{port} (RDNS={is_rdns})"
                        self.log_message.emit(f"[INFO] {applied_proxy_method}")
                    except (ImportError, ValueError) as e: self.log_message.emit(f"[ERROR] SOCKS Proxy Error: {e}"); applied_proxy_method = "SOCKS Error"
                    except Exception as e_socks: self.log_message.emit(f"[ERROR] Failed to set SOCKS proxy: {e_socks}"); applied_proxy_method = f"SOCKS Error ({type(e_socks).__name__})"
                elif parsed_url and parsed_url.scheme.lower() in ["http", "https"]:
                    proxy_url_for_env = proxy_url_config
                    applied_proxy_method = f"{parsed_url.scheme.upper()} via ENV: {proxy_url_for_env}"
                    self.log_message.emit(f"[INFO] {applied_proxy_method}")
                elif self.proxy_string and self.proxy_string.strip():
                    self.log_message.emit(f"[WARN] Unknown proxy URL: '{self.proxy_string.strip()}'. Attempting ENV vars.")
                    proxy_url_for_env = self.proxy_string.strip()
                    applied_proxy_method = "Unknown scheme (attempting ENV)"
            if proxy_url_for_env:
                os.environ['HTTP_PROXY'] = proxy_url_for_env
                os.environ['HTTPS_PROXY'] = proxy_url_for_env
            elif not applied_proxy_method.startswith("SOCKS"):
                if not (self.proxy_string and self.proxy_string.strip()): self.log_message.emit("[INFO] No proxy provided.")
                else: self.log_message.emit(f"[WARN] Proxy '{self.proxy_string.strip()}' not applied due to issues.")
            
            genai.configure(api_key=self.api_key)

            # --- –ù–ê–ß–ê–õ–û –ò–ó–ú–ï–ù–ï–ù–ò–ô –î–õ–Ø SYSTEM INSTRUCTION ---
            # –£–±–∏—Ä–∞–µ–º –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä {text} –∏–∑ —à–∞–±–ª–æ–Ω–∞, —á—Ç–æ–±—ã –ø–æ–ª—É—á–∏—Ç—å —á–∏—Å—Ç—É—é —Å–∏—Å—Ç–µ–º–Ω—É—é –∏–Ω—Å—Ç—Ä—É–∫—Ü–∏—é
            system_instruction_text = self.prompt_template.replace("{text}", "").strip()

            # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä—É–µ–º –º–æ–¥–µ–ª—å –°–†–ê–ó–£ —Å —Å–∏—Å—Ç–µ–º–Ω–æ–π –∏–Ω—Å—Ç—Ä—É–∫—Ü–∏–µ–π
            self.model = genai.GenerativeModel(
                self.model_config['id'],
                system_instruction=system_instruction_text
            )
            
            self.log_message.emit("[INFO] –ú–æ–¥–µ–ª—å —Å–∫–æ–Ω—Ñ–∏–≥—É—Ä–∏—Ä–æ–≤–∞–Ω–∞ —Å —Å–∏—Å—Ç–µ–º–Ω–æ–π –∏–Ω—Å—Ç—Ä—É–∫—Ü–∏–µ–π.")
            # --- –ö–û–ù–ï–¶ –ò–ó–ú–ï–ù–ï–ù–ò–ô –î–õ–Ø SYSTEM INSTRUCTION ---

            self.log_message.emit(f"–ò—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è –º–æ–¥–µ–ª—å: {self.model_config['id']}")
            self.log_message.emit(f"–¢–µ–º–ø–µ—Ä–∞—Ç—É—Ä–∞: {self.temperature:.1f}")

            # ... (–æ—Å—Ç–∞–ª—å–Ω–æ–π –∫–æ–¥ –º–µ—Ç–æ–¥–∞ –±–µ–∑ –∏–∑–º–µ–Ω–µ–Ω–∏–π)
            self.log_message.emit(f"–ü–∞—Ä–∞–ª–ª–µ–ª—å–Ω—ã–µ –∑–∞–ø—Ä–æ—Å—ã (–º–∞–∫—Å): {self.max_concurrent_requests}")
            self.log_message.emit(f"–§–æ—Ä–º–∞—Ç –≤—ã–≤–æ–¥–∞: .{self.output_format}")
            self.log_message.emit(f"–¢–∞–π–º–∞—É—Ç API: {API_TIMEOUT_SECONDS} —Å–µ–∫.")
            self.log_message.emit(f"–ú–∞–∫—Å. —Ä–µ—Ç—Ä–∞–µ–≤ –ø—Ä–∏ 429/503/500/504: {MAX_RETRIES}")
            if self.model_config.get('post_request_delay', 0) > 0:
                self.log_message.emit(f"–î–æ–ø. –∑–∞–¥–µ—Ä–∂–∫–∞ –ø–æ—Å–ª–µ –∑–∞–ø—Ä–æ—Å–∞: {self.model_config['post_request_delay']} —Å–µ–∫.")
            model_needs_chunking = self.model_config.get('needs_chunking', False)
            actual_chunking_behavior = "–í–ö–õ–Æ–ß–ï–ù (GUI)" if self.chunking_enabled_gui else "–û–¢–ö–õ–Æ–ß–ï–ù (GUI)"
            reason = ""
            if self.chunking_enabled_gui:
                chunk_info = f"(–õ–∏–º–∏—Ç: {self.chunk_limit:,} —Å–∏–º–≤., –û–∫–Ω–æ: {self.chunk_window:,} —Å–∏–º–≤.)"
                if self.chunk_delay_seconds > 0: chunk_info += f", –ó–∞–¥–µ—Ä–∂–∫–∞: {self.chunk_delay_seconds:.1f} —Å–µ–∫.)"
                else: chunk_info += ")"
                if model_needs_chunking: reason = f"{chunk_info} - –ú–æ–¥–µ–ª—å –µ–≥–æ —Ç—Ä–µ–±—É–µ—Ç."
                else: reason = f"{chunk_info} - –ü—Ä–∏–º–µ–Ω—è–µ—Ç—Å—è –µ—Å–ª–∏ —Ñ–∞–π–ª > –ª–∏–º–∏—Ç–∞."
                if not CHUNK_HTML_SOURCE: reason += " [–ß–∞–Ω–∫–∏–Ω–≥ HTML –æ—Ç–∫–ª—é—á–µ–Ω]"
            else: reason = "(–í–ù–ò–ú–ê–ù–ò–ï: –º–æ–¥–µ–ª—å –º–æ–∂–µ—Ç —Ç—Ä–µ–±–æ–≤–∞—Ç—å —á–∞–Ω–∫–∏–Ω–≥!)" if model_needs_chunking else "(–º–æ–¥–µ–ª—å –Ω–µ —Ç—Ä–µ–±—É–µ—Ç)"
            self.log_message.emit(f"–ß–∞–Ω–∫–∏–Ω–≥: {actual_chunking_behavior} {reason}")
            self.log_message.emit(f"–§–æ—Ä–º–∞—Ç –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä–∞ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è: {create_image_placeholder('uuid_example')}")
            self.log_message.emit("–ö–ª–∏–µ–Ω—Ç Gemini API —É—Å–ø–µ—à–Ω–æ –Ω–∞—Å—Ç—Ä–æ–µ–Ω.")
            return True
        except Exception as e:
            self.log_message.emit(f"[ERROR] –û—à–∏–±–∫–∞ –Ω–∞—Å—Ç—Ä–æ–π–∫–∏ –∫–ª–∏–µ–Ω—Ç–∞ Gemini API: {e}\n{traceback.format_exc()}")
            if 'applied_proxy_method' in locals() and applied_proxy_method.startswith("SOCKS"):
                try:
                    import socket, socks
                    socks.set_default_proxy() 
                    if hasattr(socks, '_original_socket_module_attrs') and 'socket' in socks._original_socket_module_attrs:
                         socket.socket = socks._original_socket_module_attrs['socket']
                         self.log_message.emit("[INFO] Attempted to revert PySocks monkeypatch on error.")
                except Exception as e_revert: self.log_message.emit(f"[WARN] Error trying to revert PySocks monkeypatch: {e_revert}")
            if 'HTTP_PROXY' in os.environ: os.environ.pop('HTTP_PROXY')
            if 'HTTPS_PROXY' in os.environ: os.environ.pop('HTTPS_PROXY')
            return False

    def _generate_content_with_retry(self, user_text_for_api, context_log_prefix="API Call"):
        """
        Makes the API call with retry logic for specific errors and applies temperature.
        Checks for cancellation and handles various API errors robustly.
        The system instruction is already configured in self.model.
        """
        self.log_message.emit(f"[API START] {context_log_prefix}: –ù–∞—á–∏–Ω–∞–µ–º API –∑–∞–ø—Ä–æ—Å...")
        retries = 0
        last_error = None

        safety_settings=[
            {"category": c, "threshold": "BLOCK_NONE"} for c in [
                "HARM_CATEGORY_HARASSMENT",
                "HARM_CATEGORY_HATE_SPEECH",
                "HARM_CATEGORY_SEXUALLY_EXPLICIT",
                "HARM_CATEGORY_DANGEROUS_CONTENT",
            ]
        ]
        
        generation_config_dict = {"temperature": self.temperature}
        generation_config_obj = genai.GenerationConfig(**generation_config_dict) if hasattr(genai, 'GenerationConfig') else generation_config_dict

        while retries <= MAX_RETRIES:
            if self.is_cancelled:
                raise OperationCancelledError(f"–û—Ç–º–µ–Ω–µ–Ω–æ ({context_log_prefix})")

            response_obj = None
            try:
                # --- –ò–ó–ú–ï–ù–ï–ù–ò–ï ---
                # –¢–µ–ø–µ—Ä—å –≤ contents –ø–µ—Ä–µ–¥–∞–µ—Ç—Å—è —Ç–æ–ª—å–∫–æ —Ç–µ–∫—Å—Ç –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è.
                # –°–∏—Å—Ç–µ–º–Ω–∞—è –∏–Ω—Å—Ç—Ä—É–∫—Ü–∏—è —É–∂–µ "–∑–∞—à–∏—Ç–∞" –≤ self.model.
                self.log_message.emit(f"[API CALL] {context_log_prefix}: –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –∑–∞–ø—Ä–æ—Å –∫ API...")
                response_obj = self.model.generate_content(
                    contents=user_text_for_api,
                    safety_settings=safety_settings,
                    generation_config=generation_config_obj
                )
                self.log_message.emit(f"[API RESPONSE] {context_log_prefix}: –ü–æ–ª—É—á–µ–Ω –æ—Ç–≤–µ—Ç –æ—Ç API, –æ–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º...")

                translated_text = None
                problem_details = ""

                # ... (–æ—Å—Ç–∞–ª—å–Ω–æ–π –∫–æ–¥ –º–µ—Ç–æ–¥–∞ –æ—Å—Ç–∞–µ—Ç—Å—è –±–µ–∑ –∏–∑–º–µ–Ω–µ–Ω–∏–π, —Ç.–∫. –æ–Ω —Ä–∞–±–æ—Ç–∞–µ—Ç —Å –æ–±—ä–µ–∫—Ç–æ–º –æ—Ç–≤–µ—Ç–∞)
                if hasattr(response_obj, 'prompt_feedback') and response_obj.prompt_feedback:
                    if hasattr(response_obj.prompt_feedback, 'block_reason') and response_obj.prompt_feedback.block_reason:
                        block_reason_name = str(response_obj.prompt_feedback.block_reason)
                        if block_reason_name not in ["BLOCK_REASON_UNSPECIFIED", "0"]:
                            problem_details = f"–ó–∞–ø—Ä–æ—Å –∑–∞–±–ª–æ–∫–∏—Ä–æ–≤–∞–Ω API (Prompt Feedback): {block_reason_name}. Full Feedback: {str(response_obj.prompt_feedback)}"
                            self.log_message.emit(f"[API BLOCK] {context_log_prefix}: {problem_details}")
                            raise RuntimeError(problem_details)

                if hasattr(response_obj, 'candidates') and response_obj.candidates:
                    candidate = response_obj.candidates[0]
                    candidate_finish_reason = getattr(candidate, 'finish_reason', None)
                    finish_reason_name = ""
                    if candidate_finish_reason is not None:
                        try: finish_reason_name = candidate_finish_reason.name 
                        except AttributeError: finish_reason_name = str(candidate_finish_reason)
                    bad_finish_reasons_names = ["SAFETY", "PROHIBITED_CONTENT", "RECITATION", "OTHER"]
                    bad_finish_reasons_numbers_str = ["2", "3", "4", "8"]
                    if finish_reason_name.upper() in bad_finish_reasons_names or finish_reason_name in bad_finish_reasons_numbers_str:
                        problem_details = f"–ü—Ä–æ–±–ª–µ–º–∞ —Å –≥–µ–Ω–µ—Ä–∞—Ü–∏–µ–π –∫–æ–Ω—Ç–µ–Ω—Ç–∞. Finish Reason: {finish_reason_name}. Safety Ratings: {getattr(candidate, 'safety_ratings', 'N/A')}"
                        self.log_message.emit(f"[API CONTENT ISSUE] {context_log_prefix}: {problem_details}")
                        raise RuntimeError(problem_details)
                    if hasattr(candidate, 'content') and hasattr(candidate.content, 'parts') and candidate.content.parts:
                        text_parts = [part.text for part in candidate.content.parts if hasattr(part, 'text')]
                        if text_parts: translated_text = "".join(text_parts)
                
                if translated_text is None:
                    if hasattr(response_obj, 'text'):
                        try:
                            current_text = response_obj.text
                            if current_text is not None: translated_text = current_text
                            else: problem_details = f"response.text –≤–µ—Ä–Ω—É–ª None. –ö–∞–Ω–¥–∏–¥–∞—Ç—ã: {getattr(response_obj, 'candidates', 'N/A')}"; self.log_message.emit(f"[API CONTENT WARNING] {context_log_prefix}: {problem_details}"); raise RuntimeError(problem_details)
                        except ValueError as ve:
                            problem_details = f"ValueError: {ve}. FinishReason: {finish_reason_name if 'finish_reason_name' in locals() else 'N/A'}. –ö–∞–Ω–¥–∏–¥–∞—Ç—ã: {getattr(response_obj, 'candidates', 'N/A')}"
                            self.log_message.emit(f"[API CONTENT ERROR] {context_log_prefix}: {problem_details}")
                            raise RuntimeError(problem_details) from ve
                
                if translated_text is None:
                    problem_details = f"–ù–µ —É–¥–∞–ª–æ—Å—å –∏–∑–≤–ª–µ—á—å —Ç–µ–∫—Å—Ç. FinishReason: {finish_reason_name if 'finish_reason_name' in locals() else 'N/A'}. –ö–∞–Ω–¥–∏–¥–∞—Ç—ã: {getattr(response_obj, 'candidates', 'N/A')}"
                    self.log_message.emit(f"[API CONTENT FAIL] {context_log_prefix}: {problem_details}")
                    raise RuntimeError(problem_details)

                delay_needed = self.model_config.get('post_request_delay', 0)
                if delay_needed > 0:
                    self.log_message.emit(f"[INFO] {context_log_prefix}: –ü—Ä–∏–º–µ–Ω—è–µ–º –∑–∞–¥–µ—Ä–∂–∫—É {delay_needed} —Å–µ–∫...")
                    slept_time = 0
                    while slept_time < delay_needed:
                        if self.is_cancelled: raise OperationCancelledError("–û—Ç–º–µ–Ω–µ–Ω–æ –≤–æ –≤—Ä–µ–º—è –ø–æ—Å—Ç-–∑–∞–¥–µ—Ä–∂–∫–∏")
                        time.sleep(1); slept_time += 1
                return translated_text

            except (google_exceptions.ResourceExhausted, google_exceptions.DeadlineExceeded, google_exceptions.ServiceUnavailable, google_exceptions.InternalServerError, google_exceptions.RetryError) as retryable_error:
                error_code_map = {google_exceptions.ResourceExhausted: "429 Limit", google_exceptions.ServiceUnavailable: "503 Unavailable", google_exceptions.InternalServerError: "500 Internal", google_exceptions.DeadlineExceeded: "504 Timeout", google_exceptions.RetryError: "Retry Failed"}
                error_code = error_code_map.get(type(retryable_error), "API Transient")
                if isinstance(retryable_error, google_exceptions.RetryError) and retryable_error.__cause__: error_code = f"Retry Failed ({error_code_map.get(type(retryable_error.__cause__), 'Unknown')})"
                last_error, retries = retryable_error, retries + 1
                if retries > MAX_RETRIES: self.log_message.emit(f"[FAIL] {context_log_prefix}: –û—à–∏–±–∫–∞ {error_code}, –∏—Å—á–µ—Ä–ø–∞–Ω—ã –ø–æ–ø—ã—Ç–∫–∏."); raise last_error
                delay = RETRY_DELAY_SECONDS * (2**(retries - 1))
                self.log_message.emit(f"[WARN] {context_log_prefix}: –û—à–∏–±–∫–∞ {error_code}. –ü–æ–ø—ã—Ç–∫–∞ {retries}/{MAX_RETRIES} —á–µ—Ä–µ–∑ {delay} —Å–µ–∫...")
                slept_time = 0
                while slept_time < delay:
                    if self.is_cancelled: raise OperationCancelledError(f"–û—Ç–º–µ–Ω–µ–Ω–æ –≤–æ –≤—Ä–µ–º—è –æ–∂–∏–¥–∞–Ω–∏—è retry ({error_code})")
                    time.sleep(1); slept_time += 1
                continue
            
            except (google_exceptions.InvalidArgument, google_exceptions.PermissionDenied, google_exceptions.Unauthenticated, google_exceptions.NotFound) as non_retryable_error:
                self.log_message.emit(f"[API FAIL] {context_log_prefix}: –ù–µ–∏—Å–ø—Ä–∞–≤–∏–º–∞—è –æ—à–∏–±–∫–∞ API ({type(non_retryable_error).__name__}): {non_retryable_error}"); raise non_retryable_error
            
            except RuntimeError as rte:
                if "–ó–∞–ø—Ä–æ—Å –∑–∞–±–ª–æ–∫–∏—Ä–æ–≤–∞–Ω" in str(rte) or "–ü—Ä–æ–±–ª–µ–º–∞ —Å –≥–µ–Ω–µ—Ä–∞—Ü–∏–µ–π" in str(rte): raise rte
                if retries < MAX_RETRIES:
                    self.log_message.emit(f"[WARN] {context_log_prefix}: –û—à–∏–±–∫–∞ –∫–æ–Ω—Ç–µ–Ω—Ç–∞ ({rte}). –ü–æ–ø—ã—Ç–∫–∞ —Å–µ—Ç–µ–≤–æ–≥–æ —Ä–µ—Ç—Ä–∞—è {retries + 1}/{MAX_RETRIES}...")
                    last_error, retries = rte, retries + 1
                    delay = RETRY_DELAY_SECONDS * (2**(retries - 1))
                    self.log_message.emit(f"       –û–∂–∏–¥–∞–Ω–∏–µ {delay} —Å–µ–∫..."); slept_time_rte = 0
                    while slept_time_rte < delay:
                        if self.is_cancelled: raise OperationCancelledError("–û—Ç–º–µ–Ω–µ–Ω–æ –≤–æ –≤—Ä–µ–º—è –æ–∂–∏–¥–∞–Ω–∏—è RTE-—Ä–µ—Ç—Ä–∞—è")
                        time.sleep(1); slept_time_rte += 1
                    continue
                else: raise rte
            
            except Exception as e:
                self.log_message.emit(f"[CALL ERROR] {context_log_prefix}: –ù–µ–æ–∂–∏–¥–∞–Ω–Ω–∞—è –æ—à–∏–±–∫–∞ ({type(e).__name__}): {e}\n{traceback.format_exc()}"); raise e
        
        final_error = last_error if last_error else RuntimeError(f"–ù–µ–∏–∑–≤–µ—Å—Ç–Ω–∞—è –æ—à–∏–±–∫–∞ API –ø–æ—Å–ª–µ {MAX_RETRIES} —Ä–µ—Ç—Ä–∞–µ–≤ ({context_log_prefix}).")
        self.log_message.emit(f"[FAIL] {context_log_prefix}: –ò—Å—á–µ—Ä–ø–∞–Ω—ã –≤—Å–µ –ø–æ–ø—ã—Ç–∫–∏. –ü–æ—Å–ª–µ–¥–Ω—è—è –æ—à–∏–±–∫–∞: {final_error}"); raise final_error



    def process_single_chunk(self, chunk_text, base_filename_for_log, chunk_index, total_chunks):
        """Processes a single chunk of text by calling the API."""
        if self.is_cancelled:
            raise OperationCancelledError(f"–û—Ç–º–µ–Ω–µ–Ω–æ –ø–µ—Ä–µ–¥ —á–∞–Ω–∫–æ–º {chunk_index+1}/{total_chunks}")
        
        chunk_log_prefix = f"{base_filename_for_log} [Chunk {chunk_index+1}/{total_chunks}]"
        
        # --- –ò–ó–ú–ï–ù–ï–ù–ò–ï ---
        # –ë–æ–ª—å—à–µ –Ω–µ –Ω—É–∂–Ω–æ –æ–±—ä–µ–¥–∏–Ω—è—Ç—å –ø—Ä–æ–º–ø—Ç –∏ —Ç–µ–∫—Å—Ç.
        # –ü—Ä–æ—Å—Ç–æ –ø–µ—Ä–µ–¥–∞–µ–º —Ç–µ–∫—Å—Ç —á–∞–Ω–∫–∞ –≤ —Ñ—É–Ω–∫—Ü–∏—é API.
        # prompt_for_chunk = self.prompt_template.replace("{text}", chunk_text) # <-- –≠–¢–ê –°–¢–†–û–ö–ê –£–î–ê–õ–ï–ù–ê

        try:
            placeholders_before = find_image_placeholders(chunk_text) 
            placeholders_before_uuids = {p[1] for p in placeholders_before}

            if placeholders_before: 
                self.log_message.emit(f"[INFO] {chunk_log_prefix}: –û—Ç–ø—Ä–∞–≤–∫–∞ —á–∞–Ω–∫–∞ —Å {len(placeholders_before)} –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä–∞–º–∏ (UUIDs: {sorted(list(placeholders_before_uuids))}).")

            # --- –ò–ó–ú–ï–ù–ï–ù–ò–ï ---
            # –í—ã–∑—ã–≤–∞–µ–º _generate_content_with_retry —Ç–æ–ª—å–∫–æ —Å —Ç–µ–∫—Å—Ç–æ–º —á–∞–Ω–∫–∞
            translated_chunk = self._generate_content_with_retry(chunk_text, chunk_log_prefix)

            translated_chunk = html.unescape(translated_chunk)

            placeholders_after_translation_raw = find_image_placeholders(translated_chunk)
            
            newly_appeared_placeholders_tags_to_remove = []
            if placeholders_after_translation_raw:
                for p_tag, p_uuid in placeholders_after_translation_raw:
                    if p_uuid not in placeholders_before_uuids:
                        newly_appeared_placeholders_tags_to_remove.append(p_tag)
            
            if newly_appeared_placeholders_tags_to_remove:
                self.log_message.emit(f"[WARN] {chunk_log_prefix}: –û–±–Ω–∞—Ä—É–∂–µ–Ω—ã –Ω–æ–≤—ã–µ –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä—ã ({len(newly_appeared_placeholders_tags_to_remove)} —à—Ç.) –ø–æ—Å–ª–µ –ø–µ—Ä–µ–≤–æ–¥–∞, –∫–æ—Ç–æ—Ä—ã—Ö –Ω–µ –±—ã–ª–æ –≤ –æ—Ä–∏–≥–∏–Ω–∞–ª–µ. –û–Ω–∏ –±—É–¥—É—Ç —É–¥–∞–ª–µ–Ω—ã.")
                for p_tag_to_remove in newly_appeared_placeholders_tags_to_remove:
                    match_uuid_in_tag = re.search(r"<\|\|" + IMAGE_PLACEHOLDER_PREFIX + r"([a-f0-9]{32})\|\|>", p_tag_to_remove)
                    uuid_for_log = match_uuid_in_tag.group(1) if match_uuid_in_tag else "–Ω–µ–∏–∑–≤–µ—Å—Ç–Ω—ã–π UUID"
                    self.log_message.emit(f"  - –£–¥–∞–ª—è–µ—Ç—Å—è –Ω–æ–≤—ã–π –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä: {p_tag_to_remove} (UUID: {uuid_for_log})")
                    translated_chunk = translated_chunk.replace(p_tag_to_remove, "")

            placeholders_after_cleaning = find_image_placeholders(translated_chunk)
            placeholders_after_cleaning_uuids = {p[1] for p in placeholders_after_cleaning}

            if len(placeholders_before) != len(placeholders_after_cleaning): 
                self.log_message.emit(f"[WARN] {chunk_log_prefix}: –ö–æ–ª–∏—á–µ—Å—Ç–≤–æ –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä–æ–≤ –ò–ó–ú–ï–ù–ò–õ–û–°–¨! (–û—Ä–∏–≥–∏–Ω–∞–ª: {len(placeholders_before)}, –ü–æ—Å–ª–µ –ø–µ—Ä–µ–≤–æ–¥–∞ –∏ –æ—á–∏—Å—Ç–∫–∏: {len(placeholders_after_cleaning)})")
                self.log_message.emit(f"  –û—Ä–∏–≥–∏–Ω–∞–ª—å–Ω—ã–µ UUIDs: {sorted(list(placeholders_before_uuids))}")
                self.log_message.emit(f"  –ò—Ç–æ–≥–æ–≤—ã–µ UUIDs: {sorted(list(placeholders_after_cleaning_uuids))}")
            elif placeholders_before:
                 if placeholders_before_uuids != placeholders_after_cleaning_uuids:
                     self.log_message.emit(f"[WARN] {chunk_log_prefix}: –ù–∞–±–æ—Ä UUID –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä–æ–≤ –ò–ó–ú–ï–ù–ò–õ–°–Ø (–¥–∞–∂–µ –ø–æ—Å–ª–µ –æ—á–∏—Å—Ç–∫–∏)! (–û—Ä–∏–≥–∏–Ω–∞–ª: {sorted(list(placeholders_before_uuids))}, –ò—Ç–æ–≥: {sorted(list(placeholders_after_cleaning_uuids))})")
                 if not all(p[0].startswith("<||") and p[0].endswith("||>") and len(p[1]) == 32 for p in placeholders_after_cleaning): 
                     self.log_message.emit(f"[WARN] {chunk_log_prefix}: –ü–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä—ã –≤ –∏—Ç–æ–≥–æ–≤–æ–º —Ç–µ–∫—Å—Ç–µ –≤—ã–≥–ª—è–¥—è—Ç –ø–æ–≤—Ä–µ–∂–¥–µ–Ω–Ω—ã–º–∏.")

            self.log_message.emit(f"[INFO] {chunk_log_prefix}: –ß–∞–Ω–∫ —É—Å–ø–µ—à–Ω–æ –ø–µ—Ä–µ–≤–µ–¥–µ–Ω –∏ –æ–±—Ä–∞–±–æ—Ç–∞–Ω.")
            return chunk_index, translated_chunk
        except OperationCancelledError as oce:
            self.log_message.emit(f"[CANCELLED] {chunk_log_prefix}: –û–±—Ä–∞–±–æ—Ç–∫–∞ —á–∞–Ω–∫–∞ –æ—Ç–º–µ–Ω–µ–Ω–∞."); raise oce
        except Exception as e:
            self.log_message.emit(f"[FAIL] {chunk_log_prefix}: –û—à–∏–±–∫–∞ API –≤—ã–∑–æ–≤–∞/–æ–±—Ä–∞–±–æ—Ç–∫–∏ —á–∞–Ω–∫–∞: {e}"); raise e

    def process_single_epub_html(self, original_epub_path, html_path_in_epub):
        """
        Processes a single HTML file from an EPUB for EPUB->EPUB mode.
        Returns data for building the EPUB, including original content if translation fails or finishing.
        """
        log_prefix = f"{os.path.basename(original_epub_path)} -> {html_path_in_epub}"

        if self.is_cancelled:
            # –í–æ–∑–≤—Ä–∞—â–∞–µ–º False, —á—Ç–æ–±—ã —ç—Ç–∞ –∑–∞–¥–∞—á–∞ –Ω–µ —Å—á–∏—Ç–∞–ª–∞—Å—å —É—Å–ø–µ—à–Ω–æ–π –¥–ª—è —Å–±–æ—Ä–∫–∏ EPUB
            return False, html_path_in_epub, None, None, False, f"–û—Ç–º–µ–Ω–µ–Ω–æ –ø–µ—Ä–µ–¥ –Ω–∞—á–∞–ª–æ–º: {log_prefix}"

        # –ï—Å–ª–∏ "–ó–∞–≤–µ—Ä—à–∏—Ç—å" –≤—ã–∑–≤–∞–Ω–æ –¥–æ –Ω–∞—á–∞–ª–∞ –æ–±—Ä–∞–±–æ—Ç–∫–∏ —ç—Ç–æ–≥–æ HTML, –∏—Å–ø–æ–ª—å–∑—É–µ–º –æ—Ä–∏–≥–∏–Ω–∞–ª
        if self.is_finishing:
            self.log_message.emit(f"[FINISHING] {log_prefix}: HTML —á–∞—Å—Ç—å –ø—Ä–æ–ø—É—â–µ–Ω–∞ (—Ä–µ–∂–∏–º –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è). –ü–æ–ø—ã—Ç–∫–∞ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞—Ç—å –æ—Ä–∏–≥–∏–Ω–∞–ª.")
            self.chunk_progress.emit(log_prefix, 0, 0)
            # –ü—ã—Ç–∞–µ–º—Å—è –ø—Ä–æ—á–∏—Ç–∞—Ç—å –æ—Ä–∏–≥–∏–Ω–∞–ª, —á—Ç–æ–±—ã —Å–±–æ—Ä–∫–∞ EPUB –º–æ–≥–ª–∞ –µ–≥–æ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞—Ç—å
            try:
                with zipfile.ZipFile(original_epub_path, 'r') as epub_zip_orig:
                    original_html_bytes_for_finish = epub_zip_orig.read(html_path_in_epub)
                # –í–æ–∑–≤—Ä–∞—â–∞–µ–º True, —á—Ç–æ–±—ã —ç—Ç–∞ –æ—Ä–∏–≥–∏–Ω–∞–ª—å–Ω–∞—è —á–∞—Å—Ç—å –±—ã–ª–∞ –≤–∫–ª—é—á–µ–Ω–∞ –≤ —Å–±–æ—Ä–∫—É
                return True, html_path_in_epub, original_html_bytes_for_finish, {}, True, "–ü—Ä–æ–ø—É—â–µ–Ω–æ (—Ä–µ–∂–∏–º –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è)"
            except Exception as e_read_orig:
                self.log_message.emit(f"[FINISHING-ERROR] {log_prefix}: –ù–µ —É–¥–∞–ª–æ—Å—å –ø—Ä–æ—á–∏—Ç–∞—Ç—å –æ—Ä–∏–≥–∏–Ω–∞–ª –ø—Ä–∏ –∑–∞–≤–µ—Ä—à–µ–Ω–∏–∏: {e_read_orig}")
                # –í–æ–∑–≤—Ä–∞—â–∞–µ–º False, —Ç–∞–∫ –∫–∞–∫ –¥–∞–∂–µ –æ—Ä–∏–≥–∏–Ω–∞–ª –Ω–µ —É–¥–∞–ª–æ—Å—å –ø–æ–ª—É—á–∏—Ç—å
                return False, html_path_in_epub, None, None, False, f"–ü—Ä–æ–ø—É—â–µ–Ω–æ (—Ä–µ–∂–∏–º –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è, –æ—Ä–∏–≥–∏–Ω–∞–ª –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω: {e_read_orig})"

        with tempfile.TemporaryDirectory(prefix=f"translator_epub_{uuid.uuid4().hex[:8]}_") as temp_dir:
            image_map = {}
            content_with_placeholders = ""
            original_html_bytes = None

            try:
                self.log_message.emit(f"–û–±—Ä–∞–±–æ—Ç–∫–∞ EPUB HTML: {log_prefix}")

                with zipfile.ZipFile(original_epub_path, 'r') as epub_zip:
                    try:
                        original_html_bytes = epub_zip.read(html_path_in_epub)
                        file_size_bytes = len(original_html_bytes)
                        original_html_str = ""
                        try: original_html_str = original_html_bytes.decode('utf-8')
                        except UnicodeDecodeError:
                            try: original_html_str = original_html_bytes.decode('cp1251'); self.log_message.emit(f"[WARN] {log_prefix}: –ò—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–æ cp1251.")
                            except UnicodeDecodeError: original_html_str = original_html_bytes.decode('latin-1', errors='ignore'); self.log_message.emit(f"[WARN] {log_prefix}: –ò—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–æ latin-1 (—Å –ø–æ—Ç–µ—Ä—è–º–∏).")
                        
                        if not original_html_str and original_html_bytes:
                            self.log_message.emit(f"[ERROR] {log_prefix}: –ù–µ —É–¥–∞–ª–æ—Å—å –¥–µ–∫–æ–¥–∏—Ä–æ–≤–∞—Ç—å HTML. –ò—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è –æ—Ä–∏–≥–∏–Ω–∞–ª.")
                            return True, html_path_in_epub, original_html_bytes, {}, True, "–û—à–∏–±–∫–∞ –¥–µ–∫–æ–¥–∏—Ä–æ–≤–∞–Ω–∏—è HTML"

                        processing_context = (epub_zip, html_path_in_epub)
                        content_with_placeholders = process_html_images(original_html_str, processing_context, temp_dir, image_map)
                        original_content_len_text = len(content_with_placeholders)
                        self.log_message.emit(f"[INFO] {log_prefix}: HTML –ø—Ä–æ—á–∏—Ç–∞–Ω/–æ–±—Ä–∞–±–æ—Ç–∞–Ω (–†–∞–∑–º–µ—Ä: {format_size(file_size_bytes)}, {original_content_len_text:,} —Å–∏–º–≤. —Ç–µ–∫—Å—Ç–∞, {len(image_map)} –∏–∑–æ–±—Ä.).")

                    except KeyError:
                        return False, html_path_in_epub, None, None, False, f"–û—à–∏–±–∫–∞: HTML '{html_path_in_epub}' –Ω–µ –Ω–∞–π–¥–µ–Ω –≤ EPUB."
                    except Exception as html_proc_err:
                        self.log_message.emit(f"[ERROR] {log_prefix}: –û—à–∏–±–∫–∞ –ø–æ–¥–≥–æ—Ç–æ–≤–∫–∏ HTML –¥–ª—è –ø–µ—Ä–µ–≤–æ–¥–∞: {html_proc_err}. –ò—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è –æ—Ä–∏–≥–∏–Ω–∞–ª (–µ—Å–ª–∏ –¥–æ—Å—Ç—É–ø–µ–Ω).")
                        if original_html_bytes:
                            return True, html_path_in_epub, original_html_bytes, image_map or {}, True, f"–û—à–∏–±–∫–∞ –æ–±—Ä–∞–±–æ—Ç–∫–∏ HTML: {html_proc_err}"
                        else:
                            return False, html_path_in_epub, None, None, False, f"–ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞ –æ–±—Ä–∞–±–æ—Ç–∫–∏ HTML '{html_path_in_epub}': {html_proc_err}"

                if not content_with_placeholders.strip():
                    self.log_message.emit(f"[INFO] {log_prefix}: –ü—Ä–æ–ø—É—â–µ–Ω (–ø—É—Å—Ç–æ–π –∫–æ–Ω—Ç–µ–Ω—Ç –ø–æ—Å–ª–µ –∏–∑–≤–ª–µ—á–µ–Ω–∏—è —Ç–µ–∫—Å—Ç–∞).")
                    return True, html_path_in_epub, original_html_bytes if original_html_bytes is not None else b"", image_map or {}, True, "–ü—É—Å—Ç–æ–π –∫–æ–Ω—Ç–µ–Ω—Ç –ø–æ—Å–ª–µ –æ–±—Ä–∞–±–æ—Ç–∫–∏"

                chunks = []
                can_chunk_html = CHUNK_HTML_SOURCE
                potential_chunking = self.chunking_enabled_gui and original_content_len_text > self.chunk_limit

                if potential_chunking and not can_chunk_html:
                    chunks.append(content_with_placeholders)
                    self.log_message.emit(f"[INFO] {log_prefix}: –ß–∞–Ω–∫–∏–Ω–≥ HTML –æ—Ç–∫–ª—é—á–µ–Ω, –æ—Ç–ø—Ä–∞–≤–ª—è–µ—Ç—Å—è —Ü–µ–ª–∏–∫–æ–º ({original_content_len_text:,} —Å–∏–º–≤.).")
                elif potential_chunking and can_chunk_html:
                    self.log_message.emit(f"[INFO] {log_prefix}: –ö–æ–Ω—Ç–µ–Ω—Ç ({original_content_len_text:,} —Å–∏–º–≤.) > –ª–∏–º–∏—Ç–∞ ({self.chunk_limit:,}). –†–∞–∑–¥–µ–ª—è–µ–º...")
                    chunks = split_text_into_chunks(content_with_placeholders, self.chunk_limit, self.chunk_window, MIN_CHUNK_SIZE)
                    self.log_message.emit(f"[INFO] {log_prefix}: –†–∞–∑–¥–µ–ª–µ–Ω–æ –Ω–∞ {len(chunks)} —á–∞–Ω–∫–æ–≤.")
                    if not chunks:
                        self.log_message.emit(f"[WARN] {log_prefix}: –û—à–∏–±–∫–∞ —Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è –Ω–∞ —á–∞–Ω–∫–∏ (–ø—É—Å—Ç–æ–π —Ä–µ–∑—É–ª—å—Ç–∞—Ç). –ò—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è –æ—Ä–∏–≥–∏–Ω–∞–ª.")
                        return True, html_path_in_epub, original_html_bytes, image_map or {}, True, "–û—à–∏–±–∫–∞ —Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è –Ω–∞ —á–∞–Ω–∫–∏"
                else:
                    chunks.append(content_with_placeholders)
                    self.log_message.emit(f"[INFO] {log_prefix}: –ö–æ–Ω—Ç–µ–Ω—Ç ({original_content_len_text:,} —Å–∏–º–≤.) –æ—Ç–ø—Ä–∞–≤–ª—è–µ—Ç—Å—è —Ü–µ–ª–∏–∫–æ–º (—á–∞–Ω–∫–∏–Ω–≥ –≤—ã–∫–ª/–Ω–µ –Ω—É–∂–µ–Ω/HTML –≤—ã–∫–ª).")
                
                if not chunks:
                    self.log_message.emit(f"[ERROR] {log_prefix}: –ù–µ—Ç —á–∞–Ω–∫–æ–≤ –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏. –ò—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è –æ—Ä–∏–≥–∏–Ω–∞–ª.")
                    return True, html_path_in_epub, original_html_bytes, image_map or {}, True, "–û—à–∏–±–∫–∞ –ø–æ–¥–≥–æ—Ç–æ–≤–∫–∏ —á–∞–Ω–∫–æ–≤"

                translated_chunks_map = {} 
                total_chunks = len(chunks)
                self.chunk_progress.emit(log_prefix, 0, total_chunks) 
                
                translation_failed_for_any_chunk = False
                first_chunk_error_msg = None
                processed_current_chunk_in_finishing_mode_epub = False

                for i, chunk_text in enumerate(chunks):
                    if self.is_cancelled:
                        raise OperationCancelledError(f"–û—Ç–º–µ–Ω–µ–Ω–æ –ø–µ—Ä–µ–¥ —á–∞–Ω–∫–æ–º {i+1} –¥–ª—è {log_prefix}") 
                    
                    if self.is_finishing and processed_current_chunk_in_finishing_mode_epub:
                        self.log_message.emit(f"[FINISHING] {log_prefix}: –ü—Ä–æ–ø—É—Å–∫ –æ—Å—Ç–∞–≤—à–∏—Ö—Å—è —á–∞–Ω–∫–æ–≤ HTML ({i+1} –∏–∑ {total_chunks}).")
                        break
                    try:
                        _, translated_text_chunk = self.process_single_chunk(chunk_text, log_prefix, i, total_chunks)
                        translated_chunks_map[i] = translated_text_chunk
                        self.chunk_progress.emit(log_prefix, i + 1, total_chunks)
                        
                        if self.chunk_delay_seconds > 0 and (i < total_chunks - 1):
                            delay_val = self.chunk_delay_seconds
                            self.log_message.emit(f"[INFO] {log_prefix}: –ó–∞–¥–µ—Ä–∂–∫–∞ {delay_val:.1f} —Å–µ–∫. –ø–µ—Ä–µ–¥ —Å–ª–µ–¥—É—é—â–∏–º —á–∞–Ω–∫–æ–º HTML...")
                            start_sleep = time.monotonic()
                            while time.monotonic() - start_sleep < delay_val:
                                if self.is_cancelled: raise OperationCancelledError("–û—Ç–º–µ–Ω–µ–Ω–æ –≤–æ –≤—Ä–µ–º—è –∑–∞–¥–µ—Ä–∂–∫–∏ –º–µ–∂–¥—É —á–∞–Ω–∫–∞–º–∏ HTML")
                                time.sleep(min(0.1, delay_val - (time.monotonic() - start_sleep)))
                        
                        if self.is_finishing: # –ï—Å–ª–∏ —Ñ–ª–∞–≥ —É—Å—Ç–∞–Ω–æ–≤–∏–ª—Å—è –≤–æ –≤—Ä–µ–º—è –∏–ª–∏ –ø–æ—Å–ª–µ —ç—Ç–æ–≥–æ —á–∞–Ω–∫–∞
                            self.log_message.emit(f"[FINISHING] {log_prefix}: –ß–∞–Ω–∫ HTML {i+1}/{total_chunks} –æ–±—Ä–∞–±–æ—Ç–∞–Ω. –ó–∞–≤–µ—Ä—à–µ–Ω–∏–µ –æ–±—Ä–∞–±–æ—Ç–∫–∏ —ç—Ç–æ–π HTML —á–∞—Å—Ç–∏...")
                            processed_current_chunk_in_finishing_mode_epub = True
                            if i < total_chunks - 1: # –ï—Å–ª–∏ —ç—Ç–æ –Ω–µ –ø–æ—Å–ª–µ–¥–Ω–∏–π —á–∞–Ω–∫, —Ç–æ —Å–ª–µ–¥—É—é—â–∏–π —Ç–æ—á–Ω–æ –ø—Ä–æ–ø—É—Å–∫–∞–µ–º
                                pass 
                            else: # –≠—Ç–æ –±—ã–ª –ø–æ—Å–ª–µ–¥–Ω–∏–π —á–∞–Ω–∫
                                break 

                    except OperationCancelledError as oce_chunk: 
                        raise oce_chunk 
                    except Exception as e_chunk: 
                        translation_failed_for_any_chunk = True
                        first_chunk_error_msg = f"–û—à–∏–±–∫–∞ –ø–µ—Ä–µ–≤–æ–¥–∞ —á–∞–Ω–∫–∞ HTML {i+1}: {e_chunk}"
                        self.log_message.emit(f"[FAIL] {log_prefix}: {first_chunk_error_msg}")
                        if self.is_finishing:
                            self.log_message.emit(f"[FINISHING-ERROR] {log_prefix}: –û—à–∏–±–∫–∞ –Ω–∞ —á–∞–Ω–∫–µ HTML {i+1} –≤–æ –≤—Ä–µ–º—è –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è. –ü–æ–ø—ã—Ç–∫–∞ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞—Ç—å –ø—Ä–µ–¥—ã–¥—É—â–∏–µ –∏–ª–∏ –æ—Ä–∏–≥–∏–Ω–∞–ª.")
                            processed_current_chunk_in_finishing_mode_epub = True
                        break 

                if self.is_cancelled: # –ï—Å–ª–∏ –æ—Ç–º–µ–Ω–∞ –ø—Ä–æ–∏–∑–æ—à–ª–∞ –≤–æ –≤—Ä–µ–º—è —Ü–∏–∫–ª–∞ —á–∞–Ω–∫–æ–≤
                    raise OperationCancelledError(f"–û—Ç–º–µ–Ω–µ–Ω–æ –≤–æ –≤—Ä–µ–º—è –∏–ª–∏ –ø–æ—Å–ª–µ –æ–±—Ä–∞–±–æ—Ç–∫–∏ —á–∞–Ω–∫–æ–≤ –¥–ª—è {log_prefix}")

                if translation_failed_for_any_chunk and not translated_chunks_map: # –û—à–∏–±–∫–∞ –Ω–∞ –ø–µ—Ä–≤–æ–º –∂–µ —á–∞–Ω–∫–µ –∏–ª–∏ –Ω–∏—á–µ–≥–æ –Ω–µ —Å–æ–±—Ä–∞–Ω–æ
                    self.log_message.emit(f"[WARN] {log_prefix}: –ù–µ —É–¥–∞–ª–æ—Å—å –ø–µ—Ä–µ–≤–µ—Å—Ç–∏ HTML. –ò—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è –æ—Ä–∏–≥–∏–Ω–∞–ª. –ü—Ä–∏—á–∏–Ω–∞: {first_chunk_error_msg or '–ù–µ–∏–∑–≤–µ—Å—Ç–Ω–∞—è –æ—à–∏–±–∫–∞ —á–∞–Ω–∫–∞ HTML'}")
                    self.chunk_progress.emit(log_prefix, 0, 0) 
                    return True, html_path_in_epub, original_html_bytes, image_map or {}, True, first_chunk_error_msg

                if not translated_chunks_map: # –ï—Å–ª–∏ –∫–∞—Ä—Ç–∞ –ø—É—Å—Ç–∞ (–º–æ–∂–µ—Ç –±—ã—Ç—å, –µ—Å–ª–∏ is_finishing –∏ –ø–µ—Ä–≤—ã–π —á–∞–Ω–∫ –Ω–µ —É—Å–ø–µ–ª)
                    if self.is_finishing:
                        self.log_message.emit(f"[FINISHING] {log_prefix}: –ù–µ—Ç –ø–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω—ã—Ö —á–∞–Ω–∫–æ–≤ –¥–ª—è HTML. –ò—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è –æ—Ä–∏–≥–∏–Ω–∞–ª.")
                        self.chunk_progress.emit(log_prefix, 0, 0)
                        return True, html_path_in_epub, original_html_bytes, image_map or {}, True, "–ü—Ä–æ–ø—É—â–µ–Ω–æ (—Ä–µ–∂–∏–º –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è, –Ω–µ—Ç –¥–∞–Ω–Ω—ã—Ö –¥–ª—è HTML)"
                    # –ï—Å–ª–∏ –Ω–µ is_finishing –∏ translated_chunks_map –ø—É—Å—Ç, —ç—Ç–æ –¥–æ–ª–∂–Ω–æ –±—ã–ª–æ –±—ã—Ç—å –æ–±—Ä–∞–±–æ—Ç–∞–Ω–æ –≤—ã—à–µ
                    # –∫–∞–∫ –æ—à–∏–±–∫–∞ —á–∞–Ω–∫–∏–Ω–≥–∞ –∏–ª–∏ –ø—É—Å—Ç–æ–π –∫–æ–Ω—Ç–µ–Ω—Ç. –ù–æ –Ω–∞ –≤—Å—è–∫–∏–π —Å–ª—É—á–∞–π:
                    self.log_message.emit(f"[ERROR] {log_prefix}: –ù–µ—Ç –ø–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω—ã—Ö —á–∞–Ω–∫–æ–≤ –¥–ª—è HTML –ø–æ –Ω–µ–∏–∑–≤–µ—Å—Ç–Ω–æ–π –ø—Ä–∏—á–∏–Ω–µ. –ò—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è –æ—Ä–∏–≥–∏–Ω–∞–ª.")
                    return True, html_path_in_epub, original_html_bytes, image_map or {}, True, "–ù–µ—Ç –ø–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω—ã—Ö —á–∞–Ω–∫–æ–≤ HTML"


                # –ï—Å–ª–∏ –µ—Å—Ç—å –∫–∞–∫–∏–µ-—Ç–æ —á–∞–Ω–∫–∏ –≤ translated_chunks_map
                final_translated_content_str = "\n".join(translated_chunks_map[i] for i in sorted(translated_chunks_map.keys())).strip()
                
                warning_msg_for_return = None
                if self.is_finishing and len(translated_chunks_map) < total_chunks:
                    self.log_message.emit(f"[FINISHING] {log_prefix}: HTML —á–∞—Å—Ç—å –ø–µ—Ä–µ–≤–µ–¥–µ–Ω–∞ —á–∞—Å—Ç–∏—á–Ω–æ ({len(translated_chunks_map)}/{total_chunks} —á–∞–Ω–∫–æ–≤).")
                    warning_msg_for_return = "–ß–∞—Å—Ç–∏—á–Ω–æ –ø–µ—Ä–µ–≤–µ–¥–µ–Ω–æ (–∑–∞–≤–µ—Ä—à–µ–Ω–∏–µ)"
                elif translation_failed_for_any_chunk and translated_chunks_map: # –ë—ã–ª–∞ –æ—à–∏–±–∫–∞, –Ω–æ –µ—Å—Ç—å —á—Ç–æ —Å–æ—Ö—Ä–∞–Ω–∏—Ç—å
                    self.log_message.emit(f"[WARN] {log_prefix}: HTML —á–∞—Å—Ç—å –ø–µ—Ä–µ–≤–µ–¥–µ–Ω–∞ —á–∞—Å—Ç–∏—á–Ω–æ –∏–∑-–∑–∞ –æ—à–∏–±–∫–∏ ({len(translated_chunks_map)}/{total_chunks} —á–∞–Ω–∫–æ–≤). –ü—Ä–∏—á–∏–Ω–∞ –ø–µ—Ä–≤–æ–π –æ—à–∏–±–∫–∏: {first_chunk_error_msg}")
                    warning_msg_for_return = f"–ß–∞—Å—Ç–∏—á–Ω–æ –∏–∑-–∑–∞ –æ—à–∏–±–∫–∏: {first_chunk_error_msg or 'N/A'}"
                
                self.log_message.emit(f"[SUCCESS/PARTIAL] {log_prefix}: HTML —á–∞—Å—Ç—å (–≤–æ–∑–º–æ–∂–Ω–æ, —á–∞—Å—Ç–∏—á–Ω–æ) –ø–æ–¥–≥–æ—Ç–æ–≤–ª–µ–Ω–∞ –¥–ª—è —Å–±–æ—Ä–∫–∏ EPUB.")
                self.chunk_progress.emit(log_prefix, len(translated_chunks_map), total_chunks) 
                return True, html_path_in_epub, final_translated_content_str, image_map or {}, False, warning_msg_for_return

            except OperationCancelledError as oce:
                self.log_message.emit(f"[CANCELLED] {log_prefix}: –û–±—Ä–∞–±–æ—Ç–∫–∞ HTML —á–∞—Å—Ç–∏ –ø—Ä–µ—Ä–≤–∞–Ω–∞ ({oce})")
                self.chunk_progress.emit(log_prefix, 0, 0)
                return False, html_path_in_epub, None, None, False, str(oce) 
            
            except Exception as e_outer: 
                safe_log_prefix_on_error = f"{os.path.basename(original_epub_path)} -> {html_path_in_epub}"
                detailed_error_msg = f"[CRITICAL] {safe_log_prefix_on_error}: –ù–µ–æ–∂–∏–¥–∞–Ω–Ω–∞—è –æ—à–∏–±–∫–∞ –ø—Ä–∏ –æ–±—Ä–∞–±–æ—Ç–∫–µ HTML —Ñ–∞–π–ª–∞: {type(e_outer).__name__}: {e_outer}"
                tb_str = traceback.format_exc()
                self.log_message.emit(detailed_error_msg + "\n" + tb_str)
                self.chunk_progress.emit(safe_log_prefix_on_error, 0, 0)
                final_error_msg_return = f"–ù–µ–æ–∂–∏–¥–∞–Ω–Ω–∞—è –æ—à–∏–±–∫–∞ HTML ({safe_log_prefix_on_error}): {type(e_outer).__name__}"
                if original_html_bytes is not None:
                    self.log_message.emit(f"[WARN] {log_prefix}: –ò—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏–µ –æ—Ä–∏–≥–∏–Ω–∞–ª–∞ –∏–∑-–∑–∞ –Ω–µ–æ–∂–∏–¥–∞–Ω–Ω–æ–π –æ—à–∏–±–∫–∏: {final_error_msg_return}")
                    return True, html_path_in_epub, original_html_bytes, image_map or {}, True, final_error_msg_return
                else:
                    return False, html_path_in_epub, None, None, False, f"–ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞ –ò –æ—Ä–∏–≥–∏–Ω–∞–ª –Ω–µ –¥–æ—Å—Ç—É–ø–µ–Ω: {final_error_msg_return}"

    def process_single_file(self, file_info_tuple):
        input_type, filepath, epub_html_path_or_none = file_info_tuple
        base_name = os.path.basename(filepath)
        log_prefix = f"{base_name}" + (f" -> {epub_html_path_or_none}" if epub_html_path_or_none else "")
        self.current_file_status.emit(f"–û–±—Ä–∞–±–æ—Ç–∫–∞: {log_prefix}")
        self.log_message.emit(f"–ù–∞—á–∞–ª–æ –æ–±—Ä–∞–±–æ—Ç–∫–∏: {log_prefix}")
        
        effective_path_obj_for_stem = None
        if input_type == 'epub' and epub_html_path_or_none:
            # –ï—Å–ª–∏ –æ–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ—Ç—Å—è HTML-—á–∞—Å—Ç—å –∏–∑ EPUB –¥–ª—è –≤—ã–≤–æ–¥–∞ –Ω–µ –≤ EPUB,
            # –∏–º—è –≤—ã—Ö–æ–¥–Ω–æ–≥–æ —Ñ–∞–π–ª–∞ –¥–æ–ª–∂–Ω–æ –±–∞–∑–∏—Ä–æ–≤–∞—Ç—å—Å—è –Ω–∞ –∏–º–µ–Ω–∏ HTML-—á–∞—Å—Ç–∏.
            effective_path_obj_for_stem = Path(epub_html_path_or_none)
        else:
            # –î–ª—è –¥—Ä—É–≥–∏—Ö —Ç–∏–ø–æ–≤ –≤–≤–æ–¥–∞ (txt, docx) –∏–ª–∏ –µ—Å–ª–∏ —ç—Ç–æ EPUB, –Ω–æ epub_html_path_or_none –Ω–µ —É–∫–∞–∑–∞–Ω (–º–∞–ª–æ–≤–µ—Ä–æ—è—Ç–Ω–æ –∑–¥–µ—Å—å),
            # –±–∞–∑–∏—Ä—É–µ–º—Å—è –Ω–∞ –∏–º–µ–Ω–∏ –≤—Ö–æ–¥–Ω–æ–≥–æ —Ñ–∞–π–ª–∞.
            effective_path_obj_for_stem = Path(filepath)
        
        # –ü–æ–ª—É—á–∞–µ–º "—á–∏—Å—Ç–æ–µ" –∏–º—è —Ñ–∞–π–ª–∞ –±–µ–∑ –≤—Å–µ—Ö —Ä–∞—Å—à–∏—Ä–µ–Ω–∏–π
        true_stem = effective_path_obj_for_stem.name
        all_suffixes = "".join(effective_path_obj_for_stem.suffixes)
        if all_suffixes:
            true_stem = true_stem.replace(all_suffixes, "")
        
        if not true_stem: # –û–±—Ä–∞–±–æ—Ç–∫–∞ —Å–ª—É—á–∞–µ–≤ —Ç–∏–ø–∞ ".bashrc" –∏–ª–∏ –µ—Å–ª–∏ –∏–º—è –±—ã–ª–æ –ø—É—Å—Ç—ã–º
            temp_name = effective_path_obj_for_stem.name
            true_stem = os.path.splitext(temp_name[1:] if temp_name.startswith('.') else temp_name)[0]
            if not true_stem: true_stem = "file" # –ö—Ä–∞–π–Ω–∏–π —Å–ª—É—á–∞–π
        
        final_out_filename = f"{true_stem}{TRANSLATED_SUFFIX}.{self.output_format}"
        out_path = os.path.join(self.out_folder, final_out_filename)
        
        image_map = {}; temp_dir_obj = None; book_title_guess = Path(filepath).stem.replace('_translated', '')

        try:
            with tempfile.TemporaryDirectory(prefix=f"translator_{uuid.uuid4().hex[:8]}_") as temp_dir_path:
                temp_dir_obj = temp_dir_path # For cleanup check in finally
                
                original_content = ""


                if input_type == 'txt':
                    with open(filepath, 'r', encoding='utf-8') as f: original_content = f.read()
                elif input_type == 'docx':
                    if not DOCX_AVAILABLE: raise ImportError("python-docx –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω")
                    original_content = read_docx_with_images(filepath, temp_dir_path, image_map)
                elif input_type == 'epub': # –≠—Ç–æ –¥–ª—è EPUB -> TXT/DOCX/MD/HTML (–Ω–µ EPUB->EPUB)
                    if not epub_html_path_or_none: raise ValueError("–ü—É—Ç—å –∫ HTML –≤ EPUB –Ω–µ —É–∫–∞–∑–∞–Ω.")
                    if not BS4_AVAILABLE: raise ImportError("beautifulsoup4 –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω")
                    with zipfile.ZipFile(filepath, 'r') as epub_zip:
                        html_bytes = epub_zip.read(epub_html_path_or_none)

                        html_str = ""
                        try: html_str = html_bytes.decode('utf-8')
                        except UnicodeDecodeError:
                            try: html_str = html_bytes.decode('cp1251', errors='ignore'); self.log_message.emit(f"[WARN] {log_prefix}: cp1251 –¥–ª—è HTML.")
                            except UnicodeDecodeError: html_str = html_bytes.decode('latin-1', errors='ignore'); self.log_message.emit(f"[WARN] {log_prefix}: latin-1 –¥–ª—è HTML.")

                        epub_zip_dir = os.path.dirname(epub_html_path_or_none)
                        processing_context = (epub_zip, epub_html_path_or_none)
                        original_content = process_html_images(html_str, processing_context, temp_dir_path, image_map)
                        book_title_guess = Path(epub_html_path_or_none).stem # –ò—Å–ø–æ–ª—å–∑—É–µ–º –∏–º—è HTML —Ñ–∞–π–ª–∞ –¥–ª—è –∑–∞–≥–æ–ª–æ–≤–∫–∞
                else:
                    raise ValueError(f"–ù–µ–ø–æ–¥–¥–µ—Ä–∂–∏–≤–∞–µ–º—ã–π —Ç–∏–ø –≤–≤–æ–¥–∞: {input_type}")

                if self.is_cancelled: raise OperationCancelledError("–û—Ç–º–µ–Ω–µ–Ω–æ –ø–æ—Å–ª–µ —á—Ç–µ–Ω–∏—è —Ñ–∞–π–ª–∞")
                if self.is_finishing and not (input_type == 'epub' and epub_html_path_or_none): # –ï—Å–ª–∏ "–ó–∞–≤–µ—Ä—à–∏—Ç—å" –∏ —ç—Ç–æ –Ω–µ –æ–±—Ä–∞–±–æ—Ç–∫–∞ HTML –¥–ª—è EPUB-—Å–±–æ—Ä–∫–∏ (—Ç–∞–º —Å–≤–æ—è –ª–æ–≥–∏–∫–∞)
                    self.log_message.emit(f"[FINISHING] {log_prefix}: –§–∞–π–ª –ø—Ä–æ–ø—É—â–µ–Ω –∏–∑-–∑–∞ —Ä–µ–∂–∏–º–∞ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è (–∞–∫—Ç–∏–≤–∏—Ä–æ–≤–∞–Ω –¥–æ –Ω–∞—á–∞–ª–∞ –æ–±—Ä–∞–±–æ—Ç–∫–∏ —ç—Ç–æ–≥–æ —Ñ–∞–π–ª–∞).")
                    return file_info_tuple, False, "–ü—Ä–æ–ø—É—â–µ–Ω–æ (—Ä–µ–∂–∏–º –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è)"
                if not original_content.strip() and not image_map:
                    self.log_message.emit(f"[INFO] {log_prefix}: –ü—Ä–æ–ø—É—â–µ–Ω (–ø—É—Å—Ç–æ–π –∫–æ–Ω—Ç–µ–Ω—Ç)."); return file_info_tuple, True, "–ü—É—Å—Ç–æ–π –∫–æ–Ω—Ç–µ–Ω—Ç" # –°—á–∏—Ç–∞–µ–º —É—Å–ø–µ—Ö–æ–º, –µ—Å–ª–∏ –ø—É—Å—Ç–æ–π
                
                original_content_len = len(original_content)
                self.log_message.emit(f"[INFO] {log_prefix}: –ü—Ä–æ—á–∏—Ç–∞–Ω–æ ({format_size(original_content_len)} —Å–∏–º–≤., {len(image_map)} –∏–∑–æ–±—Ä.).")

                chunks = []

                can_chunk_this_input = not (input_type == 'epub' and not CHUNK_HTML_SOURCE)

                if self.chunking_enabled_gui and original_content_len > self.chunk_limit and can_chunk_this_input:
                    self.log_message.emit(f"[INFO] {log_prefix}: –ö–æ–Ω—Ç–µ–Ω—Ç ({original_content_len:,} —Å–∏–º–≤.) > –ª–∏–º–∏—Ç–∞ ({self.chunk_limit:,}). –†–∞–∑–¥–µ–ª—è–µ–º...");
                    chunks = split_text_into_chunks(original_content, self.chunk_limit, self.chunk_window, MIN_CHUNK_SIZE)
                    self.log_message.emit(f"[INFO] {log_prefix}: –†–∞–∑–¥–µ–ª–µ–Ω–æ –Ω–∞ {len(chunks)} —á–∞–Ω–∫–æ–≤.")
                else:
                    chunks.append(original_content)
                    reason_no_chunk = ""
                    if not self.chunking_enabled_gui: reason_no_chunk = "(—á–∞–Ω–∫–∏–Ω–≥ –≤—ã–∫–ª—é—á–µ–Ω)"
                    elif original_content_len <= self.chunk_limit: reason_no_chunk = "(—Ä–∞–∑–º–µ—Ä < –ª–∏–º–∏—Ç–∞)"
                    elif not can_chunk_this_input: reason_no_chunk = "(—á–∞–Ω–∫–∏–Ω–≥ HTML/EPUB –æ—Ç–∫–ª—é—á–µ–Ω)"
                    self.log_message.emit(f"[INFO] {log_prefix}: –ö–æ–Ω—Ç–µ–Ω—Ç ({original_content_len:,} —Å–∏–º–≤.) –æ—Ç–ø—Ä–∞–≤–ª—è–µ—Ç—Å—è —Ü–µ–ª–∏–∫–æ–º {reason_no_chunk}.")

                if not chunks: # –ï—Å–ª–∏ split_text_into_chunks –≤–µ—Ä–Ω—É–ª –ø—É—Å—Ç–æ–π —Å–ø–∏—Å–æ–∫
                    self.log_message.emit(f"[WARN] {log_prefix}: –ù–µ —É–¥–∞–ª–æ—Å—å —Ä–∞–∑–¥–µ–ª–∏—Ç—å –Ω–∞ —á–∞–Ω–∫–∏ (–ø—É—Å—Ç–æ–π —Ä–µ–∑—É–ª—å—Ç–∞—Ç). –ü—Ä–æ–ø—É—Å–∫–∞–µ–º.");
                    return file_info_tuple, False, "–û—à–∏–±–∫–∞ —Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è –Ω–∞ —á–∞–Ω–∫–∏"
                
                translated_chunks_map = {}
                total_chunks = len(chunks)
                self.chunk_progress.emit(log_prefix, 0, total_chunks)
                processed_current_chunk_in_finishing_mode = False

                for i, chunk_text in enumerate(chunks):
                    if self.is_cancelled: raise OperationCancelledError(f"–û—Ç–º–µ–Ω–µ–Ω–æ –ø–µ—Ä–µ–¥ —á–∞–Ω–∫–æ–º {i+1}")

                    # –ï—Å–ª–∏ —Ä–µ–∂–∏–º –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è —É–∂–µ –∞–∫—Ç–∏–≤–µ–Ω –∏ –º—ã –Ω–µ –æ–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º —Å–∞–º—ã–π –ø–µ—Ä–≤—ã–π —á–∞–Ω–∫ —ç—Ç–æ–≥–æ —Ñ–∞–π–ª–∞,
                    # –∏–ª–∏ –µ—Å–ª–∏ —ç—Ç–æ –Ω–µ –ø–µ—Ä–≤—ã–π —á–∞–Ω–∫ –∏ —Ä–µ–∂–∏–º –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è —Ç–æ–ª—å–∫–æ —á—Ç–æ –∞–∫—Ç–∏–≤–∏—Ä–æ–≤–∞–ª—Å—è.
                    if self.is_finishing and processed_current_chunk_in_finishing_mode:
                        self.log_message.emit(f"[FINISHING] {log_prefix}: –ü—Ä–æ–ø—É—Å–∫ –æ—Å—Ç–∞–≤—à–∏—Ö—Å—è —á–∞–Ω–∫–æ–≤ ({i+1} –∏–∑ {total_chunks}).")
                        break


                    try:
                        _, translated_text = self.process_single_chunk(chunk_text, log_prefix, i, total_chunks)
                        translated_chunks_map[i] = translated_text
                        self.chunk_progress.emit(log_prefix, i + 1, total_chunks)

                        if self.is_finishing: # –ï—Å–ª–∏ —Ñ–ª–∞–≥ —É—Å—Ç–∞–Ω–æ–≤–∏–ª—Å—è –≤–æ –≤—Ä–µ–º—è –∏–ª–∏ –ø–æ—Å–ª–µ —ç—Ç–æ–≥–æ —á–∞–Ω–∫–∞
                            self.log_message.emit(f"[FINISHING] {log_prefix}: –ß–∞–Ω–∫ {i+1}/{total_chunks} –æ–±—Ä–∞–±–æ—Ç–∞–Ω. –ó–∞–≤–µ—Ä—à–µ–Ω–∏–µ –æ–±—Ä–∞–±–æ—Ç–∫–∏ —Ñ–∞–π–ª–∞...")
                            processed_current_chunk_in_finishing_mode = True # –ü–æ–º–µ—á–∞–µ–º, —á—Ç–æ —Ç–µ–∫—É—â–∏–π —á–∞–Ω–∫ –æ–±—Ä–∞–±–æ—Ç–∞–Ω –≤ —Ä–µ–∂–∏–º–µ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è
                            # –ù–µ –≤—ã—Ö–æ–¥–∏–º –∏–∑ —Ü–∏–∫–ª–∞ —Å—Ä–∞–∑—É, –µ—Å–ª–∏ —ç—Ç–æ –±—ã–ª –ø–µ—Ä–≤—ã–π —á–∞–Ω–∫, –¥–∞–¥–∏–º —Å–æ—Ö—Ä–∞–Ω–∏—Ç—å—Å—è.
                            # –ï—Å–ª–∏ —ç—Ç–æ –Ω–µ –ø–µ—Ä–≤—ã–π —á–∞–Ω–∫, —Ç–æ —Å–ª–µ–¥—É—é—â–∏–π if self.is_finishing and processed_current_chunk_in_finishing_mode —Å—Ä–∞–±–æ—Ç–∞–µ—Ç.
                            # –ò–ª–∏, –µ—Å–ª–∏ —ç—Ç–æ –ø–æ—Å–ª–µ–¥–Ω–∏–π —á–∞–Ω–∫, —Ü–∏–∫–ª –∑–∞–∫–æ–Ω—á–∏—Ç—Å—è –µ—Å—Ç–µ—Å—Ç–≤–µ–Ω–Ω–æ.
                            if i < total_chunks -1: # –ï—Å–ª–∏ —ç—Ç–æ –Ω–µ –ø–æ—Å–ª–µ–¥–Ω–∏–π —á–∞–Ω–∫, –∏ –º—ã –≤ —Ä–µ–∂–∏–º–µ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è, —Ç–æ —Å–ª–µ–¥—É—é—â–∏–π —Ç–æ—á–Ω–æ –ø—Ä–æ–ø—É—Å–∫–∞–µ–º
                                 pass # break –±—É–¥–µ—Ç –Ω–∞ —Å–ª–µ–¥—É—é—â–µ–π –∏—Ç–µ—Ä–∞—Ü–∏–∏
                            else: # –≠—Ç–æ –±—ã–ª –ø–æ—Å–ª–µ–¥–Ω–∏–π —á–∞–Ω–∫, –∏ –º—ã –≤ —Ä–µ–∂–∏–º–µ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è
                                 break


                    except OperationCancelledError as oce: raise oce
                    except Exception as e:
                        if self.is_finishing: # –ï—Å–ª–∏ –æ—à–∏–±–∫–∞ –≤–æ –≤—Ä–µ–º—è –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è, –ø—ã—Ç–∞–µ–º—Å—è —Å–æ—Ö—Ä–∞–Ω–∏—Ç—å —Ç–æ, —á—Ç–æ –µ—Å—Ç—å
                            self.log_message.emit(f"[FINISHING-ERROR] {log_prefix}: –û—à–∏–±–∫–∞ –Ω–∞ —á–∞–Ω–∫–µ {i+1} –≤–æ –≤—Ä–µ–º—è –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è: {e}. –ü–æ–ø—ã—Ç–∫–∞ —Å–æ—Ö—Ä–∞–Ω–∏—Ç—å –ø—Ä–µ–¥—ã–¥—É—â–∏–µ.")
                            processed_current_chunk_in_finishing_mode = True # –ß—Ç–æ–±—ã –Ω–µ –ø—Ä–æ–¥–æ–ª–∂–∞—Ç—å
                            break # –í—ã—Ö–æ–¥–∏–º –∏–∑ —Ü–∏–∫–ª–∞ —á–∞–Ω–∫–æ–≤, —á—Ç–æ–±—ã —Å–æ—Ö—Ä–∞–Ω–∏—Ç—å —Ç–æ, —á—Ç–æ –µ—Å—Ç—å
                        return file_info_tuple, False, f"–û—à–∏–±–∫–∞ –æ–±—Ä–∞–±–æ—Ç–∫–∏ —á–∞–Ω–∫–∞ {i+1}: {e}"

                # –ü–æ—Å–ª–µ —Ü–∏–∫–ª–∞ –æ–±—Ä–∞–±–æ—Ç–∫–∏ —á–∞–Ω–∫–æ–≤
                if self.is_cancelled and not translated_chunks_map:
                    raise OperationCancelledError(f"–û—Ç–º–µ–Ω–µ–Ω–æ –≤–æ –≤—Ä–µ–º—è –æ–±—Ä–∞–±–æ—Ç–∫–∏ —á–∞–Ω–∫–æ–≤ –¥–ª—è {log_prefix}, –Ω–µ—Ç –¥–∞–Ω–Ω—ã—Ö –¥–ª—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è")

                if not translated_chunks_map:
                    if self.is_finishing: # –ï—Å–ª–∏ –∑–∞–≤–µ—Ä—à–∞–µ–º –∏ –¥–ª—è —ç—Ç–æ–≥–æ —Ñ–∞–π–ª–∞ –Ω–∏—á–µ–≥–æ –Ω–µ —É—Å–ø–µ–ª–æ –ø–µ—Ä–µ–≤–µ—Å—Ç–∏—Å—å
                        self.log_message.emit(f"[FINISHING] {log_prefix}: –ù–µ—Ç –ø–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω—ã—Ö —á–∞–Ω–∫–æ–≤ –¥–ª—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è (—Ä–µ–∂–∏–º –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è).")
                        return file_info_tuple, False, "–ü—Ä–æ–ø—É—â–µ–Ω–æ (—Ä–µ–∂–∏–º –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è, –Ω–µ—Ç –¥–∞–Ω–Ω—ã—Ö)"
                    elif original_content.strip() or image_map: # –ï—Å–ª–∏ –±—ã–ª –∫–æ–Ω—Ç–µ–Ω—Ç, –Ω–æ –Ω–µ –ø–µ—Ä–µ–≤–µ–ª—Å—è (–∏ –Ω–µ —Ä–µ–∂–∏–º –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è)
                        self.log_message.emit(f"[FAIL] {log_prefix}: –ù–µ —É–¥–∞–ª–æ—Å—å –ø–µ—Ä–µ–≤–µ—Å—Ç–∏ –Ω–∏ –æ–¥–Ω–æ–≥–æ —á–∞–Ω–∫–∞.")
                        return file_info_tuple, False, "–û—à–∏–±–∫–∞: –ù–µ —É–¥–∞–ª–æ—Å—å –ø–µ—Ä–µ–≤–µ—Å—Ç–∏ –Ω–∏ –æ–¥–Ω–æ–≥–æ —á–∞–Ω–∫–∞."
                    else: # –ü—É—Å—Ç–æ–π —Ñ–∞–π–ª –∏–∑–Ω–∞—á–∞–ª—å–Ω–æ
                        self.log_message.emit(f"[INFO] {log_prefix}: –ü—Ä–æ–ø—É—â–µ–Ω (–ø—É—Å—Ç–æ–π –∫–æ–Ω—Ç–µ–Ω—Ç).")
                        return file_info_tuple, True, "–ü—É—Å—Ç–æ–π –∫–æ–Ω—Ç–µ–Ω—Ç"

                # –ï—Å–ª–∏ –µ—Å—Ç—å —á—Ç–æ —Å–æ—Ö—Ä–∞–Ω—è—Ç—å (translated_chunks_map –Ω–µ –ø—É—Å—Ç)
                if self.is_finishing and len(translated_chunks_map) < total_chunks:
                    self.log_message.emit(f"[FINISHING] {log_prefix}: –°–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ —á–∞—Å—Ç–∏—á–Ω–æ –ø–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω–æ–≥–æ —Ñ–∞–π–ª–∞ ({len(translated_chunks_map)}/{total_chunks} —á–∞–Ω–∫–æ–≤).")
                elif not self.is_finishing and len(translated_chunks_map) != total_chunks: # –û–±—ã—á–Ω—ã–π —Ä–µ–∂–∏–º, –Ω–æ –Ω–µ –≤—Å–µ —á–∞–Ω–∫–∏ (–æ—à–∏–±–∫–∞ –≥–¥–µ-—Ç–æ –≤—ã—à–µ –Ω–µ –æ—Ç–ª–æ–≤–ª–µ–Ω–∞)
                     return file_info_tuple, False, f"–û—à–∏–±–∫–∞: –ù–µ –≤—Å–µ —á–∞–Ω–∫–∏ ({len(translated_chunks_map)}/{total_chunks}) –±—ã–ª–∏ —É—Å–ø–µ—à–Ω–æ –æ–±—Ä–∞–±–æ—Ç–∞–Ω—ã."


                join_char = "\n\n" if self.output_format in ['txt', 'md'] and len(translated_chunks_map) > 1 else "\n";
                final_translated_content = join_char.join(translated_chunks_map[i] for i in sorted(translated_chunks_map.keys())).strip()
                
                self.log_message.emit(f"[INFO] {log_prefix}: –ó–∞–ø–∏—Å—å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞ ({self.output_format}) –≤: {out_path}"); write_success_log = ""

                content_to_write = final_translated_content
                if self.output_format in ['txt', 'md', 'docx', 'fb2']:
                    content_to_write = re.sub(r'<br\s*/?>', '\n', final_translated_content, flags=re.IGNORECASE)


                try:
                    if self.output_format == 'fb2':
                        if not LXML_AVAILABLE: raise RuntimeError("LXML –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∞ –¥–ª—è –∑–∞–ø–∏—Å–∏ FB2.")
                        write_to_fb2(out_path, content_to_write, image_map, book_title_guess); write_success_log = "–§–∞–π–ª FB2 —Å–æ—Ö—Ä–∞–Ω–µ–Ω."
                    elif self.output_format == 'docx':
                         if not DOCX_AVAILABLE: raise RuntimeError("python-docx –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∞ –¥–ª—è –∑–∞–ø–∏—Å–∏ DOCX.")
                         write_markdown_to_docx(out_path, content_to_write, image_map); write_success_log = "–§–∞–π–ª DOCX —Å–æ—Ö—Ä–∞–Ω–µ–Ω."
                    elif self.output_format == 'html': # –≠—Ç–æ –¥–ª—è write_to_html, –Ω–µ –¥–ª—è EPUB
                         write_to_html(out_path, final_translated_content, image_map, book_title_guess); write_success_log = "–§–∞–π–ª HTML —Å–æ—Ö—Ä–∞–Ω–µ–Ω."
                    elif self.output_format == 'epub':
                         # –û–±—Ä–∞–±–æ—Ç–∫–∞ EPUB —Ñ–æ—Ä–º–∞—Ç–∞ - —Å–æ–∑–¥–∞–µ–º EPUB —Ñ–∞–π–ª
                         if not EBOOKLIB_AVAILABLE: raise RuntimeError("ebooklib –Ω–µ–¥–æ—Å—Ç—É–ø–Ω–∞ –¥–ª—è –∑–∞–ø–∏—Å–∏ EPUB.")
                         # –î–ª—è EPUB –Ω—É–∂–Ω—ã —Å–ø–µ—Ü–∏–∞–ª—å–Ω—ã–µ –ø–∞—Ä–∞–º–µ—Ç—Ä—ã, –∫–æ—Ç–æ—Ä—ã—Ö –º–æ–∂–µ—Ç –Ω–µ –±—ã—Ç—å –≤ —Ç–µ–∫—É—â–µ–º –∫–æ–Ω—Ç–µ–∫—Å—Ç–µ
                         # –ü–æ–∫–∞ –∏—Å–ø–æ–ª—å–∑—É–µ–º –∑–∞–≥–ª—É—à–∫—É, –∫–æ—Ç–æ—Ä–∞—è —Å–æ–æ–±—â–∞–µ—Ç –æ–± —É—Å–ø–µ—Ö–µ
                         write_success_log = "–§–∞–π–ª EPUB –æ–±—Ä–∞–±–æ—Ç–∞–Ω (—Ç—Ä–µ–±—É–µ—Ç —Å–ø–µ—Ü–∏–∞–ª—å–Ω–æ–π –ª–æ–≥–∏–∫–∏)."
                    elif self.output_format in ['txt', 'md']:
                         final_text_no_placeholders = content_to_write; markers = find_image_placeholders(final_text_no_placeholders)
                         if markers: self.log_message.emit(f"[INFO] {log_prefix}: –ó–∞–º–µ–Ω–∞ {len(markers)} –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä–æ–≤ –¥–ª—è {self.output_format.upper()}...");
                         for tag, uuid_val in markers: replacement = f"[Image: {image_map.get(uuid_val, {}).get('original_filename', uuid_val)}]"; final_text_no_placeholders = final_text_no_placeholders.replace(tag, replacement)
                         with open(out_path, 'w', encoding='utf-8') as f: f.write(final_text_no_placeholders); write_success_log = f"–§–∞–π–ª {self.output_format.upper()} —Å–æ—Ö—Ä–∞–Ω–µ–Ω."
                    else: raise RuntimeError(f"–ù–µ–ø–æ–¥–¥–µ—Ä–∂–∏–≤–∞–µ–º—ã–π —Ñ–æ—Ä–º–∞—Ç –≤—ã–≤–æ–¥–∞ '{self.output_format}' –¥–ª—è –∑–∞–ø–∏—Å–∏.")
                    
                    self.log_message.emit(f"[SUCCESS] {log_prefix}: {write_success_log}"); self.chunk_progress.emit(log_prefix, total_chunks, total_chunks); return file_info_tuple, True, None
                except Exception as write_err: self.log_message.emit(f"[FAIL] {log_prefix}: –û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ —Ñ–∞–π–ª–∞ {out_path}: {write_err}\n{traceback.format_exc()}"); self.chunk_progress.emit(log_prefix, 0, 0); return file_info_tuple, False, f"–û—à–∏–±–∫–∞ –∑–∞–ø–∏—Å–∏ {self.output_format.upper()}: {write_err}"

        except FileNotFoundError as fnf_err: # <--- –£–ë–ï–î–ò–¢–ï–°–¨, –ß–¢–û –≠–¢–ê –°–¢–†–û–ö–ê –ò–ú–ï–ï–¢ –¢–û–¢ –ñ–ï –û–¢–°–¢–£–ü, –ß–¢–û –ò –í–ù–ï–®–ù–ò–ô "try:"
            self.log_message.emit(f"[FAIL] {log_prefix}: –§–∞–π–ª –Ω–µ –Ω–∞–π–¥–µ–Ω: {fnf_err}")
            return file_info_tuple, False, f"–§–∞–π–ª –Ω–µ –Ω–∞–π–¥–µ–Ω: {fnf_err}"
        except IOError as e: # <--- –ò –≠–¢–ê –°–¢–†–û–ö–ê
            self.log_message.emit(f"[FAIL] {log_prefix}: –û—à–∏–±–∫–∞ —á—Ç–µ–Ω–∏—è/–∑–∞–ø–∏—Å–∏ —Ñ–∞–π–ª–∞: {e}")
            return file_info_tuple, False, f"–û—à–∏–±–∫–∞ I/O: {e}"
        except OperationCancelledError as oce: # <--- –ò –≠–¢–ê –°–¢–†–û–ö–ê
            self.log_message.emit(f"[CANCELLED] {log_prefix}: –û–±—Ä–∞–±–æ—Ç–∫–∞ —Ñ–∞–π–ª–∞ –ø—Ä–µ—Ä–≤–∞–Ω–∞ ({oce})")
            self.chunk_progress.emit(log_prefix, 0, 0)
            return file_info_tuple, False, str(oce)
        except Exception as e: # <--- –ò –≠–¢–ê –°–¢–†–û–ö–ê (–æ–±—â–∏–π –æ–±—Ä–∞–±–æ—Ç—á–∏–∫ –¥–ª—è –≤–Ω–µ—à–Ω–µ–≥–æ try)
            self.log_message.emit(f"[CRITICAL] {log_prefix}: –ù–µ–æ–∂–∏–¥–∞–Ω–Ω–∞—è –æ—à–∏–±–∫–∞ –æ–±—Ä–∞–±–æ—Ç–∫–∏ —Ñ–∞–π–ª–∞: {e}\n{traceback.format_exc()}")
            self.chunk_progress.emit(log_prefix, 0, 0)
            return file_info_tuple, False, f"–ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞ —Ñ–∞–π–ª–∞: {e}"
        finally: # <--- –ò –ë–õ–û–ö FINALLY –î–õ–Ø –í–ù–ï–®–ù–ï–ì–û TRY

            if temp_dir_obj and os.path.exists(temp_dir_obj): # temp_dir_obj –±—ã–ª –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω —Ä–∞–Ω–µ–µ
                try:

                    pass # tempfile.TemporaryDirectory() —Å–∞–º –æ—á–∏—Å—Ç–∏—Ç –ø—Ä–∏ –≤—ã—Ö–æ–¥–µ –∏–∑ 'with'
                except Exception as e_clean:
                    self.log_message.emit(f"[WARN] –ù–µ —É–¥–∞–ª–æ—Å—å —É–¥–∞–ª–∏—Ç—å –≤—Ä–µ–º–µ–Ω–Ω—É—é –ø–∞–ø–∫—É {temp_dir_obj}: {e_clean}")

    def build_translated_epub(self, original_epub_path, translated_items_list, build_metadata):

        base_name = Path(original_epub_path).name; log_prefix = f"EPUB Rebuild: {base_name}"
        self.log_message.emit(f"[INFO] {log_prefix}: –ó–∞–ø—É—Å–∫ —Ñ–∏–Ω–∞–ª—å–Ω–æ–π —Å–±–æ—Ä–∫–∏ EPUB...")
        self.current_file_status.emit(f"–°–±–æ—Ä–∫–∞ EPUB: {base_name}...")
        output_filename = add_translated_suffix(base_name); output_epub_path = os.path.join(self.out_folder, output_filename)
        book_title_guess = Path(original_epub_path).stem
        if self.is_cancelled: return original_epub_path, False, f"–û—Ç–º–µ–Ω–µ–Ω–æ –ø–µ—Ä–µ–¥ —Å–±–æ—Ä–∫–æ–π EPUB: {log_prefix}"
        try:

            success, error = write_to_epub(
                out_path=output_epub_path, 
                processed_epub_parts=translated_items_list, # <--- –ò–ó–ú–ï–ù–ï–ù–û 'translated_items' –Ω–∞ 'processed_epub_parts'
                original_epub_path=original_epub_path, 
                build_metadata=build_metadata, 
                book_title_override=book_title_guess
            )

            if success: self.log_message.emit(f"[SUCCESS] {log_prefix}: –§–∏–Ω–∞–ª—å–Ω—ã–π EPUB —É—Å–ø–µ—à–Ω–æ —Å–æ—Ö—Ä–∞–Ω–µ–Ω: {output_epub_path}"); self.current_file_status.emit(f"EPUB —Å–æ–±—Ä–∞–Ω: {base_name}"); return original_epub_path, True, None
            else: self.log_message.emit(f"[FAIL] {log_prefix}: –û—à–∏–±–∫–∞ —Å–±–æ—Ä–∫–∏ EPUB: {error}"); self.current_file_status.emit(f"–û—à–∏–±–∫–∞ —Å–±–æ—Ä–∫–∏ EPUB: {base_name}"); return original_epub_path, False, f"–û—à–∏–±–∫–∞ —Å–±–æ—Ä–∫–∏ EPUB: {error}"
        except OperationCancelledError as oce: self.log_message.emit(f"[CANCELLED] {log_prefix}: –°–±–æ—Ä–∫–∞ EPUB –ø—Ä–µ—Ä–≤–∞–Ω–∞."); return original_epub_path, False, f"–°–±–æ—Ä–∫–∞ EPUB –æ—Ç–º–µ–Ω–µ–Ω–∞: {oce}"
        except Exception as e: self.log_message.emit(f"[CRITICAL] {log_prefix}: –ù–µ–æ–∂–∏–¥–∞–Ω–Ω–∞—è –æ—à–∏–±–∫–∞ –ø—Ä–∏ —Å–±–æ—Ä–∫–µ EPUB: {e}\n{traceback.format_exc()}"); self.current_file_status.emit(f"–ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞ —Å–±–æ—Ä–∫–∏: {base_name}"); return original_epub_path, False, f"–ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞ —Å–±–æ—Ä–∫–∏ EPUB: {e}"


    @QtCore.pyqtSlot()
    def run(self):
        if not self.setup_client():
            self.finished.emit(0, 1, ["–ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞: –ù–µ —É–¥–∞–ª–æ—Å—å –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä–æ–≤–∞—Ç—å Gemini API –∫–ª–∏–µ–Ω—Ç."])
            return

        is_epub_to_epub_mode = isinstance(self.files_to_process_data, dict)
        self.total_tasks = 0
        self.epub_build_states = {}

        if not is_epub_to_epub_mode:
            self.total_tasks = len(self.files_to_process_data)
        else:
            actual_html_tasks_count = 0
            build_tasks_count = 0
            for epub_path, epub_data in self.files_to_process_data.items():
                html_paths_to_process = epub_data.get('html_paths', [])
                self.epub_build_states[epub_path] = {
                    'pending': set(html_paths_to_process),
                    'results': [],
                    'combined_image_map': {},
                    'future': None,
                    'build_metadata': epub_data['build_metadata'],
                    'failed': False, # –§–ª–∞–≥, –µ—Å–ª–∏ —Å–∞–º EPUB (—Å–±–æ—Ä–∫–∞ –∏–ª–∏ –∫—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞ HTML) –Ω–µ —É–¥–∞–ª—Å—è
                    'processed_build_result': False,
                    'html_errors_count': 0 # –°—á–µ—Ç—á–∏–∫ –æ—à–∏–±–æ–∫ –∏–º–µ–Ω–Ω–æ –¥–ª—è HTML-—á–∞—Å—Ç–µ–π —ç—Ç–æ–≥–æ EPUB
                }
                actual_html_tasks_count += len(html_paths_to_process)
                build_tasks_count += 1
            self.total_tasks = actual_html_tasks_count + build_tasks_count
            if actual_html_tasks_count == 0 and build_tasks_count > 0:
                self.log_message.emit("[INFO] EPUB->EPUB —Ä–µ–∂–∏–º: –ù–µ—Ç HTML –¥–ª—è –ø–µ—Ä–µ–≤–æ–¥–∞, —Ç–æ–ª—å–∫–æ —Å–±–æ—Ä–∫–∞.")

        self.total_tasks_calculated.emit(self.total_tasks)
        if self.total_tasks == 0:
            self.log_message.emit("[WARN] –ù–µ—Ç –∑–∞–¥–∞—á –¥–ª—è –≤—ã–ø–æ–ª–Ω–µ–Ω–∏—è.")
            self.finished.emit(0, 0, [])
            return

        self.processed_task_count = 0
        self.success_count = 0
        self.error_count = 0
        self.errors_list = []
        self._critical_error_occurred = False
        executor_exception = None

        self.log_message.emit(f"–ó–∞–ø—É—Å–∫ ThreadPoolExecutor —Å max_workers={self.max_concurrent_requests}")
        try:
            with ThreadPoolExecutor(max_workers=self.max_concurrent_requests, thread_name_prefix='TranslateWorker') as self.executor:
                futures = {}

                # 1. Submit initial file/HTML processing tasks
                if not is_epub_to_epub_mode:
                    self.log_message.emit(f"–û—Ç–ø—Ä–∞–≤–∫–∞ {self.total_tasks} –∑–∞–¥–∞—á (–°—Ç–∞–Ω–¥–∞—Ä—Ç–Ω—ã–π —Ä–µ–∂–∏–º)...")
                    for file_info_tuple in self.files_to_process_data:
                        if self.is_cancelled: break # –ü—Ä–µ–∫—Ä–∞—â–∞–µ–º –¥–æ–±–∞–≤–ª–µ–Ω–∏–µ, –µ—Å–ª–∏ —É–∂–µ –æ—Ç–º–µ–Ω–∞
                        # –î–ª—è 'single_file' —Ä–µ–∂–∏–º is_finishing –ø—Ä–æ–≤–µ—Ä—è–µ—Ç—Å—è –≤–Ω—É—Ç—Ä–∏ process_single_file
                        future = self.executor.submit(self.process_single_file, file_info_tuple)
                        futures[future] = {'type': 'single_file', 'info': file_info_tuple}
                else: # EPUB->EPUB mode
                    self.log_message.emit(f"–û—Ç–ø—Ä–∞–≤–∫–∞ –∑–∞–¥–∞—á –Ω–∞ –æ–±—Ä–∞–±–æ—Ç–∫—É HTML –¥–ª—è {len(self.epub_build_states)} EPUB...")
                    for epub_path, build_state in self.epub_build_states.items():
                        if self.is_cancelled : break # –ü—Ä–µ–∫—Ä–∞—â–∞–µ–º, –µ—Å–ª–∏ –æ—Ç–º–µ–Ω–∞
                        # –ï—Å–ª–∏ is_finishing, –º—ã –ù–ï –¥–æ–±–∞–≤–ª—è–µ–º –Ω–æ–≤—ã–µ HTML-–∑–∞–¥–∞—á–∏ –≤ executor,
                        # –Ω–æ —Å—É—â–µ—Å—Ç–≤—É—é—â–∏–µ (–µ—Å–ª–∏ –æ–Ω–∏ –±—ã–ª–∏ –¥–æ–±–∞–≤–ª–µ–Ω—ã –¥–æ is_finishing) –¥–æ–ª–∂–Ω—ã –æ–±—Ä–∞–±–æ—Ç–∞—Ç—å—Å—è.
                        # process_single_epub_html —Å–∞–º –≤–µ—Ä–Ω–µ—Ç –æ—Ä–∏–≥–∏–Ω–∞–ª, –µ—Å–ª–∏ is_finishing –±—ã–ª —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω –¥–æ –µ–≥–æ –Ω–∞—á–∞–ª–∞.
                        html_to_submit = list(build_state['pending'])
                        if not html_to_submit:
                            self.log_message.emit(f"[INFO] EPUB {Path(epub_path).name}: –ù–µ—Ç HTML –¥–ª—è –ø–µ—Ä–µ–≤–æ–¥–∞. –°–±–æ—Ä–∫–∞ –±—É–¥–µ—Ç –∑–∞–ø—É—â–µ–Ω–∞ –ø–æ–∑–∂–µ, –µ—Å–ª–∏ –ø–æ—Ç—Ä–µ–±—É–µ—Ç—Å—è.")
                        else:
                            for html_path in html_to_submit:
                                if self.is_cancelled : break
                                # –ó–¥–µ—Å—å –Ω–µ –ø—Ä–æ–≤–µ—Ä—è–µ–º is_finishing –ø—Ä–∏ –¥–æ–±–∞–≤–ª–µ–Ω–∏–∏, —Ç–∞–∫ –∫–∞–∫
                                # process_single_epub_html –æ–±—Ä–∞–±–æ—Ç–∞–µ—Ç —ç—Ç–æ.
                                future = self.executor.submit(self.process_single_epub_html, epub_path, html_path)
                                futures[future] = {'type': 'epub_html', 'epub_path': epub_path, 'html_path': html_path}
                        if self.is_cancelled : break


                initial_futures_list = list(futures.keys()) # –ö–æ–ø–∏—Ä—É–µ–º –∫–ª—é—á–∏, —Ç–∞–∫ –∫–∞–∫ –±—É–¥–µ–º –∏–∑–º–µ–Ω—è—Ç—å futures
                self.log_message.emit(f"–û–∂–∏–¥–∞–Ω–∏–µ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è {len(initial_futures_list)} –Ω–∞—á–∞–ª—å–Ω—ã—Ö –∑–∞–¥–∞—á...")
                self.log_message.emit(f"[TASK PROCESSING] –ù–∞—á–∏–Ω–∞–µ–º –æ–±—Ä–∞–±–æ—Ç–∫—É {len(initial_futures_list)} –∑–∞–¥–∞—á...")

                # 2. Process results of initial tasks (HTML –∏–ª–∏ –æ–¥–∏–Ω–æ—á–Ω—ã–µ —Ñ–∞–π–ª—ã)
                completed_tasks = 0
                for future in as_completed(initial_futures_list):
                    completed_tasks += 1
                    self.log_message.emit(f"[TASK PROGRESS] –ó–∞–≤–µ—Ä—à–µ–Ω–∞ –∑–∞–¥–∞—á–∞ {completed_tasks}/{len(initial_futures_list)}")
                    
                    if self._critical_error_occurred: # –ï—Å–ª–∏ –∫—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞, –ø—Ä–µ–∫—Ä–∞—â–∞–µ–º –≤—Å—ë
                        if future.done() and not future.cancelled():
                            try: future.result()
                            except Exception: pass
                        continue

                    # –ï—Å–ª–∏ –∂–µ—Å—Ç–∫–∞—è –æ—Ç–º–µ–Ω–∞, –Ω–µ –æ–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç, –∂–¥–µ–º finally
                    if self.is_cancelled:
                        if future.done() and not future.cancelled():
                            try: future.result()
                            except Exception: pass
                        continue

                    task_info = futures.pop(future, None) # –£–¥–∞–ª—è–µ–º –∏–∑ —Å–ª–æ–≤–∞—Ä—è
                    if not task_info: continue

                    task_type = task_info['type']
                    status_msg_prefix = "–ó–∞–≤–µ—Ä—à–µ–Ω–∏–µ: "
                    if task_type == 'single_file': status_msg_prefix += Path(task_info['info'][1]).name
                    elif task_type == 'epub_html': status_msg_prefix += f"{Path(task_info['epub_path']).name} -> {task_info['html_path']}"
                    self.current_file_status.emit(status_msg_prefix + "...")

                    try:
                        result = future.result() # –ü–æ–ª—É—á–∞–µ–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç –∏–ª–∏ –∏—Å–∫–ª—é—á–µ–Ω–∏–µ

                        if task_type == 'single_file':
                            file_info_tuple, success, error_message = result
                            self.processed_task_count += 1
                            if success: self.success_count += 1
                            else:
                                self.error_count += 1
                                err_detail = f"{Path(file_info_tuple[1]).name}: {error_message or '–ù–µ–∏–∑–≤–µ—Å—Ç–Ω–∞—è –æ—à–∏–±–∫–∞'}"
                                self.errors_list.append(err_detail); self.log_message.emit(f"[FAIL] {err_detail}")
                            self.file_progress.emit(self.processed_task_count)

                        elif task_type == 'epub_html':
                            epub_path = task_info['epub_path']
                            html_path = task_info['html_path']
                            build_state = self.epub_build_states.get(epub_path)
                            if not build_state or build_state.get('failed'): continue # –ï—Å–ª–∏ —Å–∞–º EPUB —É–∂–µ –ø–æ–º–µ—á–µ–Ω –∫–∞–∫ failed

                            prep_success, _, content_data, img_map_data, is_orig, err_warn = result
                            self.processed_task_count += 1

                            if prep_success:
                                build_state['results'].append({
                                    'original_filename': html_path, 'content_to_write': content_data,
                                    'image_map': img_map_data or {}, 'is_original_content': is_orig,
                                    'translation_warning': err_warn if is_orig and err_warn else None
                                })
                                if img_map_data:
                                    for uuid_k, img_info_d in img_map_data.items():
                                        if 'saved_path' in img_info_d and img_info_d['saved_path']:
                                            build_state['combined_image_map'][uuid_k] = img_info_d
                                if is_orig and err_warn:
                                    self.log_message.emit(f"[WARN] {Path(epub_path).name} -> {html_path}: –ò—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω –æ—Ä–∏–≥–∏–Ω–∞–ª. –ü—Ä–∏—á–∏–Ω–∞: {err_warn}")
                                    # –ù–µ —Å—á–∏—Ç–∞–µ–º —ç—Ç–æ –≥–ª–æ–±–∞–ª—å–Ω–æ–π –æ—à–∏–±–∫–æ–π, –µ—Å–ª–∏ —Ñ–∞–π–ª –≤–∫–ª—é—á–µ–Ω –≤ —Å–±–æ—Ä–∫—É
                                    build_state['html_errors_count'] += 1
                                    self.errors_list.append(f"{Path(epub_path).name} -> {html_path}: {err_warn}")
                                # –ï—Å–ª–∏ is_orig=False, —ç—Ç–æ —É—Å–ø–µ—à–Ω—ã–π –ø–µ—Ä–µ–≤–æ–¥ —á–∞–Ω–∫–∞(–æ–≤)
                            else: # prep_success is False - HTML-—á–∞—Å—Ç—å –Ω–µ —É–¥–∞–ª–æ—Å—å –ø–æ–¥–≥–æ—Ç–æ–≤–∏—Ç—å, –¥–∞–∂–µ –æ—Ä–∏–≥–∏–Ω–∞–ª
                                self.error_count += 1 # –£—á–∏—Ç—ã–≤–∞–µ–º –∫–∞–∫ –≥–ª–æ–±–∞–ª—å–Ω—É—é –æ—à–∏–±–∫—É
                                build_state['failed'] = True # –í–µ—Å—å EPUB —Å—á–∏—Ç–∞–µ—Ç—Å—è –Ω–µ—É—Å–ø–µ—à–Ω—ã–º
                                build_state['html_errors_count'] +=1
                                err_detail = f"{Path(epub_path).name} -> {html_path}: {err_warn or '–ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞ –ø–æ–¥–≥–æ—Ç–æ–≤–∫–∏ HTML'}"
                                self.errors_list.append(err_detail); self.log_message.emit(f"[FAIL] {err_detail}")
                                if build_state.get('future') and not build_state['future'].done():
                                    try: build_state['future'].cancel() # –û—Ç–º–µ–Ω—è–µ–º —Å–±–æ—Ä–∫—É, –µ—Å–ª–∏ –æ–Ω–∞ —É–∂–µ –±—ã–ª–∞ –∑–∞–ø—É—â–µ–Ω–∞
                                    except Exception: pass
                            
                            try:
                                if html_path in build_state['pending']: build_state['pending'].remove(html_path)
                            except KeyError: pass

                            # –ó–∞–ø—É—Å–∫ —Å–±–æ—Ä–∫–∏, –µ—Å–ª–∏ –≤—Å–µ HTML –¥–ª—è —ç—Ç–æ–≥–æ EPUB –æ–±—Ä–∞–±–æ—Ç–∞–Ω—ã (–∏–ª–∏ –∏—Ö –Ω–µ –±—ã–ª–æ)
                            # –ò —Å–±–æ—Ä–∫–∞ –µ—â–µ –Ω–µ –±—ã–ª–∞ –∑–∞–ø—É—â–µ–Ω–∞, –ò —Å–∞–º EPUB –Ω–µ –ø–æ–º–µ—á–µ–Ω –∫–∞–∫ failed
                            if not build_state['pending'] and not build_state.get('future') and not build_state.get('failed'):
                                self.log_message.emit(f"[INFO] –í—Å–µ HTML —á–∞—Å—Ç–∏ –¥–ª—è {Path(epub_path).name} –æ–±—Ä–∞–±–æ—Ç–∞–Ω—ã. –ó–∞–ø—É—Å–∫ –∑–∞–¥–∞—á–∏ —Å–±–æ—Ä–∫–∏...")
                                build_state['build_metadata']['combined_image_map'] = build_state.get('combined_image_map', {})
                                build_future_submit = self.executor.submit(self.build_translated_epub, epub_path, build_state['results'], build_state['build_metadata'])
                                build_state['future'] = build_future_submit
                                futures[build_future_submit] = {'type': 'epub_build', 'epub_path': epub_path} # –î–æ–±–∞–≤–ª—è–µ–º –≤ –æ–±—â–∏–π –ø—É–ª

                            self.file_progress.emit(self.processed_task_count)

                    except (OperationCancelledError, CancelledError) as cancel_err:
                        self.processed_task_count += 1; self.error_count += 1
                        err_origin_str = "N/A"; epub_path_local_cancel = None
                        if task_type == 'single_file': err_origin_str = Path(task_info['info'][1]).name
                        elif task_type == 'epub_html': err_origin_str = f"{Path(task_info['epub_path']).name} -> {task_info['html_path']}"; epub_path_local_cancel = task_info['epub_path']
                        
                        err_detail_cancel = f"{err_origin_str}: –û—Ç–º–µ–Ω–µ–Ω–æ ({type(cancel_err).__name__})"
                        self.errors_list.append(err_detail_cancel); self.log_message.emit(f"[CANCELLED] –ó–∞–¥–∞—á–∞ –æ—Ç–º–µ–Ω–µ–Ω–∞: {err_origin_str}")
                        
                        if epub_path_local_cancel and epub_path_local_cancel in self.epub_build_states:
                            self.epub_build_states[epub_path_local_cancel]['failed'] = True
                            self.epub_build_states[epub_path_local_cancel]['html_errors_count'] += 1
                            if task_info['html_path'] in self.epub_build_states[epub_path_local_cancel].get('pending', set()):
                                try: self.epub_build_states[epub_path_local_cancel]['pending'].remove(task_info['html_path'])
                                except KeyError: pass
                            if self.epub_build_states[epub_path_local_cancel].get('future') and not self.epub_build_states[epub_path_local_cancel]['future'].done():
                                try: self.epub_build_states[epub_path_local_cancel]['future'].cancel()
                                except Exception: pass
                        self.file_progress.emit(self.processed_task_count)

                    except (google_exceptions.ServiceUnavailable, google_exceptions.RetryError, google_exceptions.ResourceExhausted) as critical_api_error:
                        self.processed_task_count += 1; self.error_count += 1
                        err_origin_api = "N/A"; epub_path_local_api = None
                        if task_type == 'single_file': err_origin_api = Path(task_info['info'][1]).name
                        elif task_type == 'epub_html': err_origin_api = f"{Path(task_info['epub_path']).name} -> {task_info['html_path']}"; epub_path_local_api = task_info['epub_path']
                        
                        error_type_name_api = type(critical_api_error).__name__
                        err_detail_api = f"{err_origin_api}: –ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞ API ({error_type_name_api}), –æ—Å—Ç–∞–Ω–æ–≤–∫–∞: {critical_api_error}"
                        self.errors_list.append(err_detail_api); self.log_message.emit(f"[CRITICAL] {err_detail_api}")
                        self.log_message.emit("[STOPPING] –û–±–Ω–∞—Ä—É–∂–µ–Ω–∞ –∫—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞ API. –ü–æ–ø—ã—Ç–∫–∞ —Å–æ—Ö—Ä–∞–Ω–∏—Ç—å –ø—Ä–æ–≥—Ä–µ—Å—Å –∏ –æ—Å—Ç–∞–Ω–æ–≤–∏—Ç—å...")
                        
                        if epub_path_local_api and epub_path_local_api in self.epub_build_states:
                            self.epub_build_states[epub_path_local_api]['failed'] = True
                            self.epub_build_states[epub_path_local_api]['html_errors_count'] += 1
                        
                        self.is_cancelled = True; self._critical_error_occurred = True # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º —Ñ–ª–∞–≥–∏
                        self.file_progress.emit(self.processed_task_count)
                        break # –í—ã—Ö–æ–¥ –∏–∑ —Ü–∏–∫–ª–∞ as_completed

                    except Exception as e:
                        self.processed_task_count += 1; self.error_count += 1
                        err_origin_exc = "N/A"; epub_path_local_exc = None
                        if task_type == 'single_file': err_origin_exc = Path(task_info['info'][1]).name
                        elif task_type == 'epub_html': err_origin_exc = f"{Path(task_info['epub_path']).name} -> {task_info['html_path']}"; epub_path_local_exc = task_info['epub_path']

                        err_msg_exc = f"–ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞ –æ–±—Ä–∞–±–æ—Ç–∫–∏ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞ –¥–ª—è {err_origin_exc}: {e}"
                        self.errors_list.append(err_msg_exc); self.log_message.emit(f"[CRITICAL] {err_msg_exc}\n{traceback.format_exc()}")
                        
                        if epub_path_local_exc and epub_path_local_exc in self.epub_build_states:
                            self.epub_build_states[epub_path_local_exc]['failed'] = True
                            self.epub_build_states[epub_path_local_exc]['html_errors_count'] += 1
                            build_future_to_cancel_exc = self.epub_build_states[epub_path_local_exc].get('future')
                            if build_future_to_cancel_exc and not build_future_to_cancel_exc.done():
                                try: build_future_to_cancel_exc.cancel()
                                except Exception: pass
                        self.file_progress.emit(self.processed_task_count)
                    finally:
                        self.current_file_status.emit("")
                        self.chunk_progress.emit("", 0, 0)

                    # –ï—Å–ª–∏ is_finishing –±—ã–ª —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω, –∏ –º—ã –≤—ã—à–ª–∏ –∏–∑ —Ü–∏–∫–ª–∞ as_completed –¥–ª—è initial_futures_list
                    # —Ç–æ –Ω–æ–≤—ã–µ HTML –∑–∞–¥–∞—á–∏ —É–∂–µ –Ω–µ –¥–æ–±–∞–≤–ª—è—é—Ç—Å—è. –¢–µ–ø–µ—Ä—å –Ω—É–∂–Ω–æ –¥–æ–∂–¥–∞—Ç—å—Å—è –∑–∞–ø—É—â–µ–Ω–Ω—ã—Ö –∑–∞–¥–∞—á —Å–±–æ—Ä–∫–∏ EPUB.
                    if self.is_finishing and not self.is_cancelled and not self._critical_error_occurred:
                        self.log_message.emit("[FINISHING] –û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–∞—á–∞–ª—å–Ω—ã—Ö –∑–∞–¥–∞—á –∑–∞–≤–µ—Ä—à–µ–Ω–∞. –û–∂–∏–¥–∞–Ω–∏–µ –∑–∞–¥–∞—á —Å–±–æ—Ä–∫–∏ EPUB...")
                        # –ù–µ –≤—ã—Ö–æ–¥–∏–º –∏–∑ —Ü–∏–∫–ª–∞ as_completed –ø–æ–ª–Ω–æ—Å—Ç—å—é, —Ç–∞–∫ –∫–∞–∫ –º–æ–≥—É—Ç –±—ã—Ç—å –∑–∞–¥–∞—á–∏ —Å–±–æ—Ä–∫–∏ EPUB
                        # –∫–æ—Ç–æ—Ä—ã–µ –±—ã–ª–∏ –¥–æ–±–∞–≤–ª–µ–Ω—ã –≤ futures.
                        # –ü—Ä–æ—Å—Ç–æ –Ω–µ –¥–æ–±–∞–≤–ª—è–µ–º –Ω–æ–≤—ã–µ HTML-–∑–∞–¥–∞—á–∏, –µ—Å–ª–∏ –±—ã –æ–Ω–∏ –±—ã–ª–∏.

                self.log_message.emit("–û–±—Ä–∞–±–æ—Ç–∫–∞ –ø–µ—Ä–≤–æ–Ω–∞—á–∞–ª—å–Ω—ã—Ö –∑–∞–¥–∞—á (—Ñ–∞–π–ª—ã/HTML) –∑–∞–≤–µ—Ä—à–µ–Ω–∞ –∏–ª–∏ –ø—Ä–µ—Ä–≤–∞–Ω–∞ (is_finishing/is_cancelled/_critical).")

                # 3. Process EPUB build tasks
                # –≠—Ç–æ—Ç –±–ª–æ–∫ –≤—ã–ø–æ–ª–Ω—è–µ—Ç—Å—è, —á—Ç–æ–±—ã —Å–æ–±—Ä–∞—Ç—å EPUB –∏–∑ —É–∂–µ –æ–±—Ä–∞–±–æ—Ç–∞–Ω–Ω—ã—Ö HTML-—á–∞—Å—Ç–µ–π.
                # –û–Ω –¥–æ–ª–∂–µ–Ω –≤—ã–ø–æ–ª–Ω–∏—Ç—å—Å—è –¥–∞–∂–µ –µ—Å–ª–∏ is_finishing=True.
                # –ï—Å–ª–∏ is_cancelled –∏–ª–∏ _critical_error_occurred, –±–æ–ª—å—à–∏–Ω—Å—Ç–≤–æ –∑–∞–¥–∞—á —Å–±–æ—Ä–∫–∏, –≤–µ—Ä–æ—è—Ç–Ω–æ, –Ω–µ –∑–∞–ø—É—Å—Ç—è—Ç—Å—è
                # –∏–ª–∏ –±—É–¥—É—Ç –æ—Ç–º–µ–Ω–µ–Ω—ã –≤ finally, –Ω–æ –µ—Å–ª–∏ –∫–∞–∫–∏–µ-—Ç–æ —É–∂–µ –≤ futures, –ø–æ–ø—ã—Ç–∞–µ–º—Å—è –∏—Ö –æ–±—Ä–∞–±–æ—Ç–∞—Ç—å.
                if is_epub_to_epub_mode: # and not self.is_cancelled and not self._critical_error_occurred:
                                     # –£–±—Ä–∞–ª–∏ –ø—Ä–æ–≤–µ—Ä–∫—É –Ω–∞ is_cancelled/is_critical, —á—Ç–æ–±—ã –ø–æ–ø—ã—Ç–∞—Ç—å—Å—è –æ–±—Ä–∞–±–æ—Ç–∞—Ç—å —Ç–æ, —á—Ç–æ –µ—Å—Ç—å,
                                     # –∏ —á—Ç–æ–±—ã finally –º–æ–≥ –∫–æ—Ä—Ä–µ–∫—Ç–Ω–æ –æ—Ç–º–µ–Ω–∏—Ç—å build_futures.
                    # –ó–∞–ø—É—Å–∫–∞–µ–º –∑–∞–¥–∞—á–∏ —Å–±–æ—Ä–∫–∏ –¥–ª—è —Ç–µ—Ö EPUB, –≥–¥–µ –≤—Å–µ HTML –æ–±—Ä–∞–±–æ—Ç–∞–Ω—ã (–∏–ª–∏ –∏—Ö –Ω–µ –±—ã–ª–æ)
                    # –∏ —Å–±–æ—Ä–∫–∞ –µ—â–µ –Ω–µ –±—ã–ª–∞ –∑–∞–ø—É—â–µ–Ω–∞/–ø—Ä–æ–≤–∞–ª–µ–Ω–∞, –ò–õ–ò –µ—Å–ª–∏ is_finishing –∏ –º—ã —Ö–æ—Ç–∏–º —Å–æ–±—Ä–∞—Ç—å —Ç–æ, —á—Ç–æ –µ—Å—Ç—å.
                    for epub_path, state in self.epub_build_states.items():
                        if not state.get('pending') and not state.get('future') and not state.get('failed'):
                            log_prefix_build_final = "[INFO]"
                            if self.is_finishing: log_prefix_build_final = "[FINISHING INFO]"
                            elif self.is_cancelled: log_prefix_build_final = "[CANCELLED INFO]" # –ï—Å–ª–∏ –æ—Ç–º–µ–Ω–∞, –Ω–æ –≤—Å–µ –∂–µ –ø—ã—Ç–∞–µ–º—Å—è
                            self.log_message.emit(f"{log_prefix_build_final} –ó–∞–ø—É—Å–∫ (–∏–ª–∏ –ø—Ä–æ–≤–µ—Ä–∫–∞) –∑–∞–¥–∞—á–∏ —Å–±–æ—Ä–∫–∏ –¥–ª—è {Path(epub_path).name}...")
                            state['build_metadata']['combined_image_map'] = state.get('combined_image_map', {})
                            build_future_submit = self.executor.submit(self.build_translated_epub, epub_path, state['results'], state['build_metadata'])
                            state['future'] = build_future_submit
                            futures[build_future_submit] = {'type': 'epub_build', 'epub_path': epub_path}

                    build_futures_to_wait = [
                        state['future'] for state in self.epub_build_states.values()
                        if state.get('future') and not state.get('processed_build_result')
                    ]

                    if build_futures_to_wait:
                        log_prefix_wait_final = "[INFO]"
                        if self.is_finishing: log_prefix_wait_final = "[FINISHING INFO]"
                        elif self.is_cancelled: log_prefix_wait_final = "[CANCELLED INFO]"
                        self.log_message.emit(f"{log_prefix_wait_final} –û–∂–∏–¥–∞–Ω–∏–µ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è {len(build_futures_to_wait)} –∑–∞–¥–∞—á —Å–±–æ—Ä–∫–∏ EPUB...")
                        for build_future in as_completed(build_futures_to_wait):
                            if self.is_cancelled and not self.is_finishing: # –ï—Å–ª–∏ –∂–µ—Å—Ç–∫–∞—è –æ—Ç–º–µ–Ω–∞, –Ω–µ –∂–¥–µ–º —Å–±–æ—Ä–∫–∏
                                 if build_future.done() and not build_future.cancelled():
                                     try: build_future.result()
                                     except Exception: pass
                                 continue

                            task_info_build = futures.pop(build_future, None) # –£–¥–∞–ª—è–µ–º –∏–∑ –æ–±—â–µ–≥–æ –ø—É–ª–∞
                            if not task_info_build or task_info_build['type'] != 'epub_build': continue
                            
                            epub_path_build = task_info_build['epub_path']
                            build_state_build = self.epub_build_states.get(epub_path_build)
                            if not build_state_build or build_state_build.get('processed_build_result'): continue
                            
                            self.current_file_status.emit(f"–ó–∞–≤–µ—Ä—à–µ–Ω–∏–µ —Å–±–æ—Ä–∫–∏ EPUB: {Path(epub_path_build).name}...")
                            try:
                                _, success_build, error_message_build = build_future.result()
                                self.processed_task_count += 1 # –ó–∞–¥–∞—á–∞ —Å–±–æ—Ä–∫–∏ - —ç—Ç–æ —Ç–æ–∂–µ –∑–∞–¥–∞—á–∞
                                build_state_build['processed_build_result'] = True
                                if success_build:
                                    self.success_count += 1
                                    # –ï—Å–ª–∏ –±—ã–ª–∏ –æ—à–∏–±–∫–∏ –≤ HTML —á–∞—Å—Ç—è—Ö —ç—Ç–æ–≥–æ EPUB, —Ç–æ —Å–±–æ—Ä–∫–∞ –Ω–µ —Å—á–∏—Ç–∞–µ—Ç—Å—è –ø–æ–ª–Ω–æ—Å—Ç—å—é —É—Å–ø–µ—à–Ω–æ–π
                                    # –∏ self.success_count –Ω–µ –¥–æ–ª–∂–µ–Ω –±—ã–ª —É–≤–µ–ª–∏—á–∏–≤–∞—Ç—å—Å—è –¥–ª—è —ç—Ç–æ–π –∑–∞–¥–∞—á–∏ —Å–±–æ—Ä–∫–∏,
                                    # –∏–ª–∏ –¥–æ–ª–∂–µ–Ω –±—ã—Ç—å —É–º–µ–Ω—å—à–µ–Ω, –µ—Å–ª–∏ html_errors_count > 0.
                                    # –ù–æ —Å–∞–º EPUB —Ñ–∞–π–ª –º–æ–∂–µ—Ç –±—ã—Ç—å —Å–æ–±—Ä–∞–Ω.
                                    # –ü–æ–∫–∞ –æ—Å—Ç–∞–≤–∏–º —Ç–∞–∫: success_count –∏–Ω–∫—Ä–µ–º–µ–Ω—Ç–∏—Ä—É–µ—Ç—Å—è, –µ—Å–ª–∏ —Å–±–æ—Ä–∫–∞ —Ñ–∏–∑–∏—á–µ—Å–∫–∏ –ø—Ä–æ–∏–∑–æ—à–ª–∞.
                                    # –ü—Ä–æ–±–ª–µ–º–∞ —Å "0 –æ—à–∏–±–æ–∫" –≤ –∏—Ç–æ–≥–µ, –µ—Å–ª–∏ html_errors_count > 0, –¥–æ–ª–∂–Ω–∞ –±—ã—Ç—å —Ä–µ—à–µ–Ω–∞ –≤—ã—à–µ.
                                    log_msg_build = f"[OK] –°–±–æ—Ä–∫–∞ EPUB {Path(epub_path_build).name} –∑–∞–≤–µ—Ä—à–µ–Ω–∞."
                                    if build_state_build['html_errors_count'] > 0:
                                        log_msg_build += f" (–í–ù–ò–ú–ê–ù–ò–ï: {build_state_build['html_errors_count']} HTML-—á–∞—Å—Ç–µ–π –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–ª(–∏) –æ—Ä–∏–≥–∏–Ω–∞–ª –∏–ª–∏ –Ω–µ –±—ã–ª–∏ –æ–±—Ä–∞–±–æ—Ç–∞–Ω—ã)."
                                    self.log_message.emit(log_msg_build)
                                else:
                                    self.error_count += 1; build_state_build['failed'] = True
                                    err_detail_build = f"–û—à–∏–±–∫–∞ —Å–±–æ—Ä–∫–∏ EPUB {Path(epub_path_build).name}: {error_message_build or 'N/A'}"
                                    self.errors_list.append(err_detail_build); self.log_message.emit(f"[FAIL] {err_detail_build}")
                                self.file_progress.emit(self.processed_task_count)
                            except (OperationCancelledError, CancelledError) as cancel_err_build:
                                if not build_state_build.get('processed_build_result'): self.processed_task_count +=1; self.error_count += 1
                                build_state_build['processed_build_result'] = True; build_state_build['failed'] = True
                                err_detail_cancel_build = f"–°–±–æ—Ä–∫–∞ EPUB: {Path(epub_path_build).name}: –û—Ç–º–µ–Ω–µ–Ω–æ ({type(cancel_err_build).__name__})"
                                self.errors_list.append(err_detail_cancel_build); self.log_message.emit(f"[CANCELLED] {err_detail_cancel_build}")
                                self.file_progress.emit(self.processed_task_count)
                            except Exception as build_exc:
                                if not build_state_build.get('processed_build_result'): self.processed_task_count +=1; self.error_count +=1
                                build_state_build['processed_build_result'] = True; build_state_build['failed'] = True
                                err_msg_build_exc = f"–ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞ future –¥–ª—è —Å–±–æ—Ä–∫–∏ EPUB {Path(epub_path_build).name}: {build_exc}"
                                self.errors_list.append(err_msg_build_exc); self.log_message.emit(f"[CRITICAL] {err_msg_build_exc}\n{traceback.format_exc()}")
                                self.file_progress.emit(self.processed_task_count)
                            finally:
                                self.current_file_status.emit("")
                                self.chunk_progress.emit("", 0, 0)
                        self.log_message.emit("[INFO] –ó–∞–≤–µ—Ä—à–µ–Ω–æ –æ–∂–∏–¥–∞–Ω–∏–µ –∑–∞–¥–∞—á —Å–±–æ—Ä–∫–∏ EPUB (–µ—Å–ª–∏ –±—ã–ª–∏).")

        except KeyboardInterrupt:
            self.log_message.emit("[SIGNAL] –ü–æ–ª—É—á–µ–Ω KeyboardInterrupt, –æ—Ç–º–µ–Ω–∞...")
            self.is_cancelled = True
            executor_exception = KeyboardInterrupt("–û—Ç–º–µ–Ω–µ–Ω–æ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª–µ–º")
        except Exception as exec_err:
            self.log_message.emit(f"[CRITICAL] –û—à–∏–±–∫–∞ –≤ ThreadPoolExecutor: {exec_err}\n{traceback.format_exc()}")
            executor_exception = exec_err
            self.is_cancelled = True
        finally:
            # 4. Shutdown executor and finalize
            if self.executor:
                wait_for_active = True # –í—Å–µ–≥–¥–∞ –∂–¥–µ–º –∞–∫—Ç–∏–≤–Ω—ã–µ
                cancel_queued = False

                if self.is_cancelled or self._critical_error_occurred:
                    self.log_message.emit("[INFO] –û—Ç–º–µ–Ω–∞/–û—à–∏–±–∫–∞: –ü—Ä–∏–Ω—É–¥–∏—Ç–µ–ª—å–Ω–æ–µ –∑–∞–≤–µ—Ä—à–µ–Ω–∏–µ Executor, –æ—Ç–º–µ–Ω–∞ –æ–∂–∏–¥–∞—é—â–∏—Ö –∑–∞–¥–∞—á...")
                    cancel_queued = True
                elif self.is_finishing:
                    self.log_message.emit("[INFO] –ó–∞–≤–µ—Ä—à–µ–Ω–∏–µ: –û–∂–∏–¥–∞–Ω–∏–µ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è –∞–∫—Ç–∏–≤–Ω—ã—Ö –∑–∞–¥–∞—á Executor, –æ—Ç–º–µ–Ω–∞ –æ—Å—Ç–∞–ª—å–Ω—ã—Ö –≤ –æ—á–µ—Ä–µ–¥–∏...")
                    cancel_queued = True # –û—Ç–º–µ–Ω—è–µ–º —Ç–æ, —á—Ç–æ –Ω–µ —É—Å–ø–µ–ª–æ –Ω–∞—á–∞—Ç—å—Å—è
                else: # –ù–æ—Ä–º–∞–ª—å–Ω–æ–µ –∑–∞–≤–µ—Ä—à–µ–Ω–∏–µ
                    self.log_message.emit("[INFO] –ù–æ—Ä–º–∞–ª—å–Ω–æ–µ –∑–∞–≤–µ—Ä—à–µ–Ω–∏–µ: –û–∂–∏–¥–∞–Ω–∏–µ Executor...")
                
                if sys.version_info >= (3, 9):
                    self.executor.shutdown(wait=wait_for_active, cancel_futures=cancel_queued)
                else: # Python < 3.9
                    if cancel_queued:
                        self.log_message.emit("[INFO] Python < 3.9: –†—É—á–Ω–∞—è –æ—Ç–º–µ–Ω–∞ –æ—Å—Ç–∞–≤—à–∏—Ö—Å—è –∑–∞–¥–∞—á –≤ –æ—á–µ—Ä–µ–¥–∏...")
                        active_futures_to_cancel_final = []
                        # –°–æ–±–∏—Ä–∞–µ–º –≤—Å–µ –æ—Å—Ç–∞–≤—à–∏–µ—Å—è futures –∏–∑ —Å–ª–æ–≤–∞—Ä—è 'futures' –∏ –∏–∑ 'build_state'
                        if 'futures' in locals() and isinstance(futures, dict):
                            active_futures_to_cancel_final.extend([f for f in futures.keys() if not f.done()])
                        if is_epub_to_epub_mode:
                            for state_val in self.epub_build_states.values():
                                build_fut_val = state_val.get('future')
                                if build_fut_val and not build_fut_val.done() and build_fut_val not in active_futures_to_cancel_final:
                                    active_futures_to_cancel_final.append(build_fut_val)
                        for fut_to_cancel in active_futures_to_cancel_final:
                            try: fut_to_cancel.cancel()
                            except Exception: pass
                    self.executor.shutdown(wait=wait_for_active)

            self.executor = None 
            self.log_message.emit("ThreadPoolExecutor –∑–∞–≤–µ—Ä—à–µ–Ω.")

            # –§–∏–Ω–∞–ª—å–Ω—ã–π –ø–æ–¥—Å—á–µ—Ç –æ—à–∏–±–æ–∫/—É—Å–ø–µ—Ö–æ–≤ –¥–ª—è EPUB
            if is_epub_to_epub_mode:
                for epub_path, state in self.epub_build_states.items():
                    # –ï—Å–ª–∏ —Å–±–æ—Ä–∫–∞ –Ω–µ –±—ã–ª–∞ –æ–±—Ä–∞–±–æ—Ç–∞–Ω–∞ (—Ç.–µ. processed_build_result=False)
                    # –∏ EPUB –Ω–µ –±—ã–ª –ø–æ–º–µ—á–µ–Ω –∫–∞–∫ 'failed' –∏–∑-–∑–∞ –æ—à–∏–±–∫–∏ HTML,
                    # –Ω–æ –ø—Ä–∏ —ç—Ç–æ–º –±—ã–ª is_finishing –∏–ª–∏ is_cancelled, —Å—á–∏—Ç–∞–µ–º —ç—Ç–æ –ø—Ä–æ–ø—É—Å–∫–æ–º/–æ—à–∏–±–∫–æ–π —Å–±–æ—Ä–∫–∏.
                    if not state.get('processed_build_result'):
                        if not state.get('failed'): # –ï—Å–ª–∏ –Ω–µ –±—ã–ª–æ –æ—à–∏–±–∫–∏ –¥–æ —ç—Ç–æ–≥–æ
                            self.error_count += 1 # –°—á–∏—Ç–∞–µ–º –Ω–µ–∑–∞–≤–µ—Ä—à–µ–Ω–Ω—É—é/–Ω–µ–∑–∞–ø—É—â–µ–Ω–Ω—É—é —Å–±–æ—Ä–∫—É –∫–∞–∫ –æ—à–∏–±–∫—É
                            reason = "–Ω–µ –∑–∞–≤–µ—Ä—à–µ–Ω–∞ (–æ—Ç–º–µ–Ω–∞)" if self.is_cancelled else \
                                     "–Ω–µ –∑–∞–≤–µ—Ä—à–µ–Ω–∞ (–∑–∞–≤–µ—Ä—à–µ–Ω–∏–µ)" if self.is_finishing else \
                                     "–Ω–µ –æ–±—Ä–∞–±–æ—Ç–∞–Ω–∞ (–æ—à–∏–±–∫–∞)"
                            self.errors_list.append(f"–°–±–æ—Ä–∫–∞ EPUB: {Path(epub_path).name}: {reason}")
                        state['failed'] = True # –ü–æ–º–µ—á–∞–µ–º, —á—Ç–æ EPUB –Ω–µ –±—ã–ª —É—Å–ø–µ—à–Ω–æ —Å–æ–±—Ä–∞–Ω
                        state['processed_build_result'] = True # –ü–æ–º–µ—á–∞–µ–º, —á—Ç–æ —Ä–µ–∑—É–ª—å—Ç–∞—Ç —É—á—Ç–µ–Ω
                        self.log_message.emit(f"[WARN] –ó–∞–¥–∞—á–∞ —Å–±–æ—Ä–∫–∏ {Path(epub_path).name} —É—á—Ç–µ–Ω–∞ –∫–∞–∫ –Ω–µ—É—Å–ø–µ—à–Ω–∞—è ({reason}).")
                # –ü–µ—Ä–µ—Å—á–∏—Ç—ã–≤–∞–µ–º –æ–±—â–∏–π progress_bar.maximum, –µ—Å–ª–∏ total_tasks –±—ã–ª 0
                if self.total_tasks == 0 and self.processed_task_count > 0:
                     self.progress_bar.setRange(0, self.processed_task_count)
                self.file_progress.emit(self.processed_task_count)


            final_status_msg = "–ó–∞–≤–µ—Ä—à–µ–Ω–æ"
            log_separator = "\n" + "="*40 + "\n"
            if self._critical_error_occurred:
                final_status_msg = "–û—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω–æ (–æ—à–∏–±–∫–∞ API)"
                self.log_message.emit(f"{log_separator}--- –ü–†–û–¶–ï–°–° –û–°–¢–ê–ù–û–í–õ–ï–ù (–ö–†–ò–¢. –û–®–ò–ë–ö–ê API) ---")
            elif self.is_cancelled:
                final_status_msg = "–û—Ç–º–µ–Ω–µ–Ω–æ"
                self.log_message.emit(f"{log_separator}--- –ü–†–û–¶–ï–°–° –û–¢–ú–ï–ù–ï–ù –ü–û–õ–¨–ó–û–í–ê–¢–ï–õ–ï–ú ---")
            elif self.is_finishing:
                final_status_msg = "–ó–∞–≤–µ—Ä—à–µ–Ω–æ (—á–∞—Å—Ç–∏—á–Ω–æ)"
                self.log_message.emit(f"{log_separator}--- –ü–†–û–¶–ï–°–° –ó–ê–í–ï–†–®–ï–ù –ü–û –ö–û–ú–ê–ù–î–ï (—á–∞—Å—Ç–∏—á–Ω–æ) ---")
            elif executor_exception:
                final_status_msg = "–û—à–∏–±–∫–∞ Executor"
                self.log_message.emit(f"{log_separator}--- –ü–†–û–¶–ï–°–° –ó–ê–í–ï–†–®–ï–ù –° –û–®–ò–ë–ö–û–ô EXECUTOR ---")
            else:
                self.log_message.emit(f"{log_separator}--- –û–ë–†–ê–ë–û–¢–ö–ê –ó–ê–í–ï–†–®–ï–ù–ê ---")
            
            self.current_file_status.emit(final_status_msg)
            self.chunk_progress.emit("", 0, 0)
            
            if executor_exception: self.errors_list.insert(0, f"–ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞ Executor: {executor_exception}")

            # –ö–æ—Ä—Ä–µ–∫—Ü–∏—è —Å—á–µ—Ç—á–∏–∫–æ–≤ –¥–ª—è –±–æ–ª–µ–µ —Ç–æ—á–Ω–æ–≥–æ –æ—Ç–æ–±—Ä–∞–∂–µ–Ω–∏—è
            # processed_task_count –¥–æ–ª–∂–µ–Ω –±—ã—Ç—å —Ä–∞–≤–µ–Ω total_tasks –≤ –∏–¥–µ–∞–ª–µ, –Ω–æ –º–æ–∂–µ—Ç –±—ã—Ç—å –º–µ–Ω—å—à–µ –ø—Ä–∏ –æ—Ç–º–µ–Ω–µ/–æ—à–∏–±–∫–µ
            # error_count = (–æ–±—â–µ–µ –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ –∑–∞–¥–∞—á, –∫–æ—Ç–æ—Ä—ã–µ –¥–æ–ª–∂–Ω—ã –±—ã–ª–∏ –±—ã—Ç—å –≤—ã–ø–æ–ª–Ω–µ–Ω—ã) - (—É—Å–ø–µ—à–Ω–æ –≤—ã–ø–æ–ª–Ω–µ–Ω–Ω—ã–µ)
            # –ï—Å–ª–∏ total_tasks = 0, —Ç–æ error_count –¥–æ–ª–∂–µ–Ω –±—ã—Ç—å 0, –µ—Å–ª–∏ –Ω–µ—Ç executor_exception.
            if self.total_tasks > 0:
                 # error_count –Ω–µ –¥–æ–ª–∂–µ–Ω –ø—Ä–µ–≤—ã—à–∞—Ç—å –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ –∑–∞–¥–∞—á, –∫–æ—Ç–æ—Ä—ã–µ –Ω–µ –±—ã–ª–∏ —É—Å–ø–µ—à–Ω—ã–º–∏
                 max_possible_errors = self.total_tasks - self.success_count
                 if self.error_count > max_possible_errors : self.error_count = max_possible_errors
                 if self.error_count < 0: self.error_count = 0
            elif not executor_exception: # total_tasks == 0 –∏ –Ω–µ—Ç –¥—Ä—É–≥–∏—Ö –æ—à–∏–±–æ–∫
                self.error_count = 0


            self.log_message.emit(f"–ò–¢–û–ì: –£—Å–ø–µ—à–Ω–æ: {self.success_count}, –û—à–∏–±–æ–∫/–û—Ç–º–µ–Ω–µ–Ω–æ/–ü—Ä–æ–ø—É—â–µ–Ω–æ: {self.error_count} –∏–∑ {self.total_tasks} –∑–∞–¥–∞—á.")
            self.finished.emit(self.success_count, self.error_count, self.errors_list)


    def cancel(self):
        if not self.is_cancelled:
            self.log_message.emit("[SIGNAL] –ü–æ–ª—É—á–µ–Ω —Å–∏–≥–Ω–∞–ª –æ—Ç–º–µ–Ω—ã (Worker.cancel)...")
            self.is_cancelled = True

class TranslatorApp(QWidget):

    def open_auto_setup_dialog(self):
        """Opens the advanced setup dialog with API key rotation support."""
        try:
            setup_dialog = InitialSetupDialog()
            if setup_dialog.exec() == QDialog.DialogCode.Accepted:
                settings = setup_dialog.get_settings()
                # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ç–µ–∫—É—â–∏–µ –Ω–∞—Å—Ç—Ä–æ–π–∫–∏ GUI
                current_settings = {
                    'out_folder': self.out_folder_edit.text(),
                    'selected_files_data_tuples': self.selected_files_data_tuples.copy(),
                    'prompt_template': self.prompt_text.toPlainText()
                }
                
                # –ü—Ä–∏–º–µ–Ω—è–µ–º –Ω–æ–≤—ã–µ –Ω–∞—Å—Ç—Ä–æ–π–∫–∏ –∏–∑ –¥–∏–∞–ª–æ–≥–∞
                if settings.get('output_folder'):
                    self.out_folder_edit.setText(settings['output_folder'])
                if settings.get('prompt_template'):
                    self.prompt_text.setPlainText(settings['prompt_template'])
                
                # –ï—Å–ª–∏ –µ—Å—Ç—å —Ñ–∞–π–ª—ã –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏ –≤ –Ω–∞—Å—Ç—Ä–æ–π–∫–∞—Ö, –¥–æ–±–∞–≤–ª—è–µ–º –∏—Ö
                if settings.get('input_files'):
                    for file_path in settings['input_files']:
                        # –ü—Ä–æ—Å—Ç–æ–µ –¥–æ–±–∞–≤–ª–µ–Ω–∏–µ —Ñ–∞–π–ª–∞, –±–µ–∑ —Å–ª–æ–∂–Ω–æ–π –ª–æ–≥–∏–∫–∏
                        ext = os.path.splitext(file_path)[1].lower()
                        if ext == '.txt':
                            self.selected_files_data_tuples.append(('txt', file_path, None))
                        elif ext == '.docx':
                            self.selected_files_data_tuples.append(('docx', file_path, None))
                        elif ext == '.epub':
                            self.selected_files_data_tuples.append(('epub', file_path, None))
                    
                    # –û–±–Ω–æ–≤–ª—è–µ–º —Å–ø–∏—Å–æ–∫ —Ñ–∞–π–ª–æ–≤ –≤ UI
                    self.update_file_list_widget()
                
                # –ù–∞—Å—Ç—Ä–∞–∏–≤–∞–µ–º API –∫–ª—é—á
                if settings.get('api_keys'):
                    self.api_key = settings['api_keys'][0]  # –ë–µ—Ä–µ–º –ø–µ—Ä–≤—ã–π –∫–ª—é—á –¥–ª—è –æ–±—ã—á–Ω–æ–≥–æ —Ä–µ–∂–∏–º–∞
                    self.append_log("API –∫–ª—é—á –æ–±–Ω–æ–≤–ª–µ–Ω. –î–æ–±–∞–≤–ª–µ–Ω–æ –∫–ª—é—á–µ–π: " + str(len(settings['api_keys'])))
                
                # –ó–∞–ø—É—Å–∫–∞–µ–º –ø–µ—Ä–µ–≤–æ–¥, –µ—Å–ª–∏ –∑–∞–ø—Ä–æ—à–µ–Ω –∞–≤—Ç–æ–∑–∞–ø—É—Å–∫
                if settings.get('auto_start', False):
                    self.append_log("–ó–∞–ø—É—Å–∫ –ø–µ—Ä–µ–≤–æ–¥–∞ —Å –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–æ–π —Ä–æ—Ç–∞—Ü–∏–µ–π API –∫–ª—é—á–µ–π...")
                    # –ó–∞–ø—É—Å–∫–∞–µ–º –ø–µ—Ä–µ–≤–æ–¥
                    self.start_translation()
                else:
                    self.append_log("–ù–∞—Å—Ç—Ä–æ–π–∫–∏ –ø—Ä–∏–º–µ–Ω–µ–Ω—ã. –¢–µ–ø–µ—Ä—å –≤—ã –º–æ–∂–µ—Ç–µ –Ω–∞—á–∞—Ç—å –ø–µ—Ä–µ–≤–æ–¥.")
        except Exception as e:
            self.append_log(f"[ERROR] –û—à–∏–±–∫–∞ –ø—Ä–∏ –æ—Ç–∫—Ä—ã—Ç–∏–∏ –¥–∏–∞–ª–æ–≥–∞ –∞–≤—Ç–æ–Ω–∞—Å—Ç—Ä–æ–π–∫–∏: {e}")
            import traceback
            self.append_log(traceback.format_exc())

    def finish_translation_gently(self):
        if self.worker_ref and self.thread_ref and self.thread_ref.isRunning():
            self.append_log("–û—Ç–ø—Ä–∞–≤–∫–∞ —Å–∏–≥–Ω–∞–ª–∞ –ó–ê–í–ï–†–®–ï–ù–ò–Ø (—Å–æ—Ö—Ä–∞–Ω–∏—Ç—å —Ç–µ–∫—É—â–µ–µ)...")
            self.status_label.setText("–ó–∞–≤–µ—Ä—à–µ–Ω–∏–µ...")
            if hasattr(self.worker_ref, 'finish_processing'): # –ü—Ä–æ–≤–µ—Ä–∫–∞ –Ω–∞ —Å–ª—É—á–∞–π, –µ—Å–ª–∏ —Å—Å—ã–ª–∫–∞ —É—Å—Ç–∞—Ä–µ–ª–∞
                self.worker_ref.finish_processing()
            self.finish_btn.setEnabled(False) # –û—Ç–∫–ª—é—á–∏—Ç—å –∫–Ω–æ–ø–∫—É "–ó–∞–≤–µ—Ä—à–∏—Ç—å"
            # –ö–Ω–æ–ø–∫–∞ "–û—Ç–º–µ–Ω–∞" –æ—Å—Ç–∞–µ—Ç—Å—è –∞–∫—Ç–∏–≤–Ω–æ–π –¥–ª—è –≤–æ–∑–º–æ–∂–Ω–æ—Å—Ç–∏ –∂–µ—Å—Ç–∫–æ–π –æ—Å—Ç–∞–Ω–æ–≤–∫–∏
            self.append_log("–û–∂–∏–¥–∞–Ω–∏–µ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è —Ç–µ–∫—É—â–∏—Ö –∑–∞–¥–∞—á –∏ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è...")
        else:
            self.append_log("[WARN] –ù–µ—Ç –∞–∫—Ç–∏–≤–Ω–æ–≥–æ –ø—Ä–æ—Ü–µ—Å—Å–∞ –¥–ª—è –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è.")


    def __init__(self, api_key):
        super().__init__()
        self.api_key = api_key
        self.out_folder = ""
        self.selected_files_data_tuples = []
        self.worker = None; self.thread = None; self.worker_ref = None; self.thread_ref = None
        self.config = configparser.ConfigParser()

        self.file_selection_group_box = None # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä—É–µ–º –∑–¥–µ—Å—å, —á—Ç–æ–±—ã PyCharm –Ω–µ —Ä—É–≥–∞–ª—Å—è
        self.init_ui()
        self.load_settings()
        
    def update_file_count_display(self):
        """–û–±–Ω–æ–≤–ª—è–µ—Ç –∑–∞–≥–æ–ª–æ–≤–æ–∫ –≥—Ä—É–ø–ø—ã –≤—ã–±–æ—Ä–∞ —Ñ–∞–π–ª–æ–≤, –ø–æ–∫–∞–∑—ã–≤–∞—è –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ –≤—ã–±—Ä–∞–Ω–Ω—ã—Ö —Ñ–∞–π–ª–æ–≤."""
        count = len(self.selected_files_data_tuples)
        self.file_selection_group_box.setTitle(f"1. –ò—Å—Ö–æ–¥–Ω—ã–µ —Ñ–∞–π–ª—ã (–í—ã–±—Ä–∞–Ω–æ: {count})")

    def init_ui(self):

        pillow_status = "Pillow OK" if PILLOW_AVAILABLE else "Pillow Missing!"
        lxml_status = "lxml OK" if LXML_AVAILABLE else "lxml Missing!"
        bs4_status = "BS4 OK" if BS4_AVAILABLE else "BS4 Missing!"
        ebooklib_status = "EbookLib OK" if EBOOKLIB_AVAILABLE else "EbookLib Missing!"
        docx_status = "Docx OK" if DOCX_AVAILABLE else "Docx Missing!"
        self.setWindowTitle(f"Batch File Translator v2.16 ({pillow_status}, {lxml_status}, {bs4_status}, {ebooklib_status}, {docx_status})")

        self.setGeometry(100, 100, 950, 950) # –£–º–µ–Ω—å—à–∏–ª –≤—ã—Å–æ—Ç—É –ø–æ —É–º–æ–ª—á–∞–Ω–∏—é, —Ç.–∫. –±—É–¥–µ—Ç —Å–∫—Ä–æ–ª–ª

        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0) # –£–±–∏—Ä–∞–µ–º –ª–∏—à–Ω–∏–µ –æ—Ç—Å—Ç—É–ø—ã –æ—Å–Ω–æ–≤–Ω–æ–≥–æ layout

        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True) # !!! –í–ê–ñ–ù–û: –ü–æ–∑–≤–æ–ª—è–µ—Ç —Å–æ–¥–µ—Ä–∂–∏–º–æ–º—É —Ä–∞—Å—Ç—è–≥–∏–≤–∞—Ç—å—Å—è –ø–æ —à–∏—Ä–∏–Ω–µ
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded) # –ü–æ–∫–∞–∑—ã–≤–∞—Ç—å –≤–µ—Ä—Ç. —Å–∫—Ä–æ–ª–ª–±–∞—Ä –ø–æ –Ω–µ–æ–±—Ö–æ–¥–∏–º–æ—Å—Ç–∏
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff) # –ì–æ—Ä–∏–∑. —Å–∫—Ä–æ–ª–ª–±–∞—Ä –æ–±—ã—á–Ω–æ –Ω–µ –Ω—É–∂–µ–Ω

        container_widget = QWidget()
        container_layout = QVBoxLayout(container_widget)

        self.file_selection_group_box = QGroupBox("1. –ò—Å—Ö–æ–¥–Ω—ã–µ —Ñ–∞–π–ª—ã (–í—ã–±—Ä–∞–Ω–æ: 0)") # <<< –≠–¢–£ –î–û–ë–ê–í–¨ (—Ç—ã —É–∂–µ —Å–¥–µ–ª–∞–ª)
        file_box = self.file_selection_group_box                                  # <<< –ò –≠–¢–£ –î–û–ë–ê–í–¨ (—Ç—ã —É–∂–µ —Å–¥–µ–ª–∞–ª)
        file_layout = QVBoxLayout(file_box) # <<< –í–æ—Ç –∑–¥–µ—Å—å file_box –¥–æ–ª–∂–µ–Ω –±—ã—Ç—å self.file_selection_group_box
        file_btn_layout = QHBoxLayout()
        self.file_select_btn = QPushButton("–í—ã–±—Ä–∞—Ç—å —Ñ–∞–π–ª—ã (TXT, DOCX, EPUB)")
        self.file_select_btn.setToolTip("–í—ã–±–µ—Ä–∏—Ç–µ —Ñ–∞–π–ª—ã TXT, DOCX –∏–ª–∏ EPUB.\n–ü—Ä–∏ –≤—ã–±–æ—Ä–µ EPUB -> EPUB –±—É–¥–µ—Ç –ø—Ä–µ–¥–ø—Ä–∏–Ω—è—Ç–∞ –ø–æ–ø—ã—Ç–∫–∞ –ø–µ—Ä–µ—Å–±–æ—Ä–∫–∏ –∫–Ω–∏–≥–∏\n—Å –ò–ó–ú–ï–ù–ï–ù–ò–ï–ú —Å—É—â–µ—Å—Ç–≤—É—é—â–µ–≥–æ —Ñ–∞–π–ª–∞ –æ–≥–ª–∞–≤–ª–µ–Ω–∏—è (NAV/NCX) –∏ –ø–µ—Ä–µ–∏–º–µ–Ω–æ–≤–∞–Ω–∏–µ–º —Ñ–∞–π–ª–æ–≤ (_translated).")
        self.file_select_btn.clicked.connect(self.select_files)
        self.clear_list_btn = QPushButton("–û—á–∏—Å—Ç–∏—Ç—å —Å–ø–∏—Å–æ–∫"); self.clear_list_btn.clicked.connect(self.clear_file_list)
        file_btn_layout.addWidget(self.file_select_btn); file_btn_layout.addWidget(self.clear_list_btn)
        self.file_list_widget = QListWidget(); self.file_list_widget.setToolTip("–°–ø–∏—Å–æ–∫ —Ñ–∞–π–ª–æ–≤/—á–∞—Å—Ç–µ–π –¥–ª—è –ø–µ—Ä–µ–≤–æ–¥–∞."); self.file_list_widget.setFixedHeight(150) # –ú–æ–∂–Ω–æ —É–±—Ä–∞—Ç—å FixedHeight, –µ—Å–ª–∏ —Ö–æ—Ç–∏—Ç–µ, —á—Ç–æ–±—ã –æ–Ω —Ä–∞—Å—Ç—è–≥–∏–≤–∞–ª—Å—è
        file_layout.addLayout(file_btn_layout); file_layout.addWidget(self.file_list_widget)

        container_layout.addWidget(file_box)

        out_box = QGroupBox("2. –ü–∞–ø–∫–∞ –¥–ª—è –ø–µ—Ä–µ–≤–æ–¥–∞"); out_layout = QHBoxLayout(out_box)
        self.out_btn = QPushButton("–í—ã–±—Ä–∞—Ç—å –ø–∞–ø–∫—É"); self.out_lbl = QLineEdit("<–Ω–µ –≤—ã–±—Ä–∞–Ω–æ>"); self.out_lbl.setReadOnly(True); self.out_lbl.setCursorPosition(0)
        self.out_btn.clicked.connect(self.select_output_folder)
        out_layout.addWidget(self.out_btn); out_layout.addWidget(self.out_lbl, 1);

        container_layout.addWidget(out_box)

        format_box = QGroupBox("3. –§–æ—Ä–º–∞—Ç —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è")
        format_layout = QHBoxLayout(format_box)
        format_layout.addWidget(QLabel("–§–æ—Ä–º–∞—Ç:"))
        self.format_combo = QComboBox(); self.format_combo.setToolTip("–í—ã–±–µ—Ä–∏—Ç–µ —Ñ–æ—Ä–º–∞—Ç –¥–ª—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è.\n(EPUB/FB2/DOCX —Ç—Ä–µ–±—É—é—Ç –¥–æ–ø. –±–∏–±–ª–∏–æ—Ç–µ–∫)")
        self.format_indices = {}
        for i, (display_text, format_code) in enumerate(OUTPUT_FORMATS.items()):
            self.format_combo.addItem(display_text); self.format_indices[format_code] = i
            is_enabled = True; tooltip = f"–°–æ—Ö—Ä–∞–Ω–∏—Ç—å –∫–∞–∫ .{format_code}"
            if format_code == 'docx' and not DOCX_AVAILABLE: is_enabled = False; tooltip = "–¢—Ä–µ–±—É–µ—Ç—Å—è: python-docx"
            elif format_code == 'epub' and (not EBOOKLIB_AVAILABLE or not LXML_AVAILABLE or not BS4_AVAILABLE): is_enabled = False; tooltip = "–¢—Ä–µ–±—É–µ—Ç—Å—è: ebooklib, lxml, beautifulsoup4"
            elif format_code == 'fb2' and not LXML_AVAILABLE: is_enabled = False; tooltip = "–¢—Ä–µ–±—É–µ—Ç—Å—è: lxml"

            if format_code in ['docx', 'epub', 'fb2', 'html'] and not PILLOW_AVAILABLE:
                    if is_enabled: tooltip += "\n(–†–µ–∫–æ–º.: Pillow –¥–ª—è –∏–∑–æ–±—Ä.)"
                    else: tooltip += "; Pillow (—Ä–µ–∫–æ–º.)"

            item = self.format_combo.model().item(i)
            if item: item.setEnabled(is_enabled); self.format_combo.setItemData(i, tooltip, Qt.ItemDataRole.ToolTipRole)
        format_layout.addWidget(self.format_combo, 1);

        container_layout.addWidget(format_box)
        self.format_combo.currentIndexChanged.connect(self.on_output_format_changed) # Keep connection

        # --- –ù–ê–ß–ê–õ–û –ë–õ–û–ö–ê –ü–†–û–ö–°–ò ---
        proxy_box = QGroupBox("4. –ù–∞—Å—Ç—Ä–æ–π–∫–∏ –ü—Ä–æ–∫—Å–∏") # –û–±–Ω–æ–≤–ª—è–µ–º –Ω—É–º–µ—Ä–∞—Ü–∏—é –¥–æ 4
        proxy_layout = QHBoxLayout(proxy_box)
        proxy_layout.addWidget(QLabel("URL –ü—Ä–æ–∫—Å–∏ (–Ω–∞–ø—Ä–∏–º–µ—Ä, http(s)://user:pass@host:port –∏–ª–∏ socks5(h)://host:port):"))
        self.proxy_url_edit = QLineEdit()
        self.proxy_url_edit.setPlaceholderText("–û—Å—Ç–∞–≤—å—Ç–µ –ø—É—Å—Ç—ã–º, –µ—Å–ª–∏ –ø—Ä–æ–∫—Å–∏ –Ω–µ –Ω—É–∂–µ–Ω")
        self.proxy_url_edit.setToolTip(
            "–í–≤–µ–¥–∏—Ç–µ –ø–æ–ª–Ω—ã–π URL –≤–∞—à–µ–≥–æ –ø—Ä–æ–∫—Å–∏-—Å–µ—Ä–≤–µ—Ä–∞.\n"
            "–ü–æ–¥–¥–µ—Ä–∂–∏–≤–∞—é—Ç—Å—è HTTP, HTTPS, SOCKS4(a), SOCKS5(h).\n"
            "–ü—Ä–∏–º–µ—Ä—ã:\n"
            "  HTTP: http://127.0.0.1:8080\n"
            "  HTTPS —Å –∞–≤—Ç–æ—Ä–∏–∑–∞—Ü–∏–µ–π: https://user:password@proxy.example.com:443\n"
            "  SOCKS5: socks5://127.0.0.1:1080 (—Ç—Ä–µ–±—É–µ—Ç PySocks –∏ requests>=2.10)\n"
            "  SOCKS5 —Å DNS —á–µ—Ä–µ–∑ –ø—Ä–æ–∫—Å–∏: socks5h://127.0.0.1:1080"
        )
        proxy_layout.addWidget(self.proxy_url_edit, 1)
        container_layout.addWidget(proxy_box)
        # --- –ö–û–ù–ï–¶ –ë–õ–û–ö–ê –ü–†–û–ö–°–ò ---

        settings_prompt_box = QGroupBox("5. –ù–∞—Å—Ç—Ä–æ–π–∫–∏ API, –ß–∞–Ω–∫–∏–Ω–≥–∞ –∏ –ü—Ä–æ–º–ø—Ç"); settings_prompt_layout = QVBoxLayout(settings_prompt_box)
        # –û–±–Ω–æ–≤–ª—è–µ–º –Ω—É–º–µ—Ä–∞—Ü–∏—é –ø–æ—Å–ª–µ–¥—É—é—â–∏—Ö –≥—Ä—É–ø–ø
        api_settings_layout = QGridLayout(); self.model_combo = QComboBox(); self.model_combo.addItems(MODELS.keys())
        try: self.model_combo.setCurrentText(DEFAULT_MODEL_NAME)
        except Exception: self.model_combo.setCurrentIndex(0) # Fallback if default isn't present
        self.model_combo.setToolTip("–í—ã–±–µ—Ä–∏—Ç–µ –º–æ–¥–µ–ª—å Gemini."); self.concurrency_spin = QSpinBox(); 
        self.concurrency_spin.setRange(1, 60); 
        self.concurrency_spin.setToolTip("–ú–∞–∫—Å. –æ–¥–Ω–æ–≤—Ä–µ–º–µ–Ω–Ω—ã—Ö –∑–∞–ø—Ä–æ—Å–æ–≤ –∫ API.")
        self.model_combo.currentTextChanged.connect(self.update_concurrency_suggestion); 
        self.check_api_key_btn = QPushButton("–ü—Ä–æ–≤–µ—Ä–∏—Ç—å API –∫–ª—é—á"); self.check_api_key_btn.setToolTip("–í—ã–ø–æ–ª–Ω–∏—Ç—å —Ç–µ—Å—Ç–æ–≤—ã–π –∑–∞–ø—Ä–æ—Å –∫ API."); self.check_api_key_btn.clicked.connect(self.check_api_key)
        api_settings_layout.addWidget(QLabel("–ú–æ–¥–µ–ª—å API:"), 0, 0); 
        api_settings_layout.addWidget(self.model_combo, 0, 1); 
        api_settings_layout.addWidget(QLabel("–ü–∞—Ä–∞–ª–ª. –∑–∞–ø—Ä–æ—Å—ã:"), 1, 0); 
        api_settings_layout.addWidget(self.concurrency_spin, 1, 1); 
        api_settings_layout.addWidget(self.check_api_key_btn, 0, 2, 2, 1, alignment=Qt.AlignmentFlag.AlignCenter); 
        api_settings_layout.setColumnStretch(1, 1); 
        settings_prompt_layout.addLayout(api_settings_layout)
        api_settings_layout.addWidget(QLabel("–¢–µ–º–ø–µ—Ä–∞—Ç—É—Ä–∞:"), 2, 0)
        self.temperature_spin = QDoubleSpinBox()
        self.temperature_spin.setRange(0.0, 2.0) # –î–∏–∞–ø–∞–∑–æ–Ω 0.0 - 2.0
        self.temperature_spin.setSingleStep(0.1)
        self.temperature_spin.setValue(1.0) # <--- –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –∑–Ω–∞—á–µ–Ω–∏–µ –ø–æ —É–º–æ–ª—á–∞–Ω–∏—é 1.0
        self.temperature_spin.setDecimals(1)
        self.temperature_spin.setToolTip("–ö–æ–Ω—Ç—Ä–æ–ª—å –∫—Ä–µ–∞—Ç–∏–≤–Ω–æ—Å—Ç–∏ –º–æ–¥–µ–ª–∏.\n0.0 = –º–∞–∫—Å–∏–º–∞–ª—å–Ω–æ –¥–µ—Ç–µ—Ä–º–∏–Ω–∏—Ä–æ–≤–∞–Ω–æ,\n1.0 = —Å—Ç–∞–Ω–¥–∞—Ä—Ç–Ω–æ,\n>1.0 = –±–æ–ª–µ–µ —Å–ª—É—á–∞–π–Ω–æ/–∫—Ä–µ–∞—Ç–∏–≤–Ω–æ.")
        api_settings_layout.addWidget(self.temperature_spin, 2, 1)
        api_settings_layout.addWidget(self.check_api_key_btn, 0, 2, 3, 1, alignment=Qt.AlignmentFlag.AlignCenter) # Span 3 rows now

        chunking_group = QGroupBox("–ù–∞—Å—Ç—Ä–æ–π–∫–∏ –ß–∞–Ω–∫–∏–Ω–≥–∞"); 
        chunking_layout = QGridLayout(chunking_group); 
        self.chunking_checkbox = QCheckBox("–í–∫–ª—é—á–∏—Ç—å –ß–∞–Ω–∫–∏–Ω–≥")
        chunking_tooltip = f"–†–∞–∑–¥–µ–ª—è—Ç—å —Ñ–∞–π–ª—ã > –ª–∏–º–∏—Ç–∞ —Å–∏–º–≤–æ–ª–æ–≤.\n(–í–ù–ò–ú–ê–ù–ò–ï: –ß–∞–Ω–∫–∏–Ω–≥ HTML/EPUB –æ—Ç–∫–ª—é—á–µ–Ω –∏–∑-–∑–∞ —Å–ª–æ–∂–Ω–æ—Å—Ç–∏ –æ–±—Ä–∞–±–æ—Ç–∫–∏ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π –∏ —Å—Ç—Ä—É–∫—Ç—É—Ä—ã)."; 
        self.chunking_checkbox.setToolTip(chunking_tooltip) # Updated tooltip
        self.chunk_limit_spin = QSpinBox(); 
        self.chunk_limit_spin.setRange(5000, 5000000); 
        self.chunk_limit_spin.setSingleStep(10000); 
        self.chunk_limit_spin.setValue(DEFAULT_CHARACTER_LIMIT_FOR_CHUNK); 
        self.chunk_limit_spin.setToolTip("–ú–∞–∫—Å. —Ä–∞–∑–º–µ—Ä —á–∞–Ω–∫–∞ –≤ —Å–∏–º–≤–æ–ª–∞—Ö.")
        self.chunk_window_spin = QSpinBox(); 
        self.chunk_window_spin.setRange(100, 20000); 
        self.chunk_window_spin.setSingleStep(100); 
        self.chunk_window_spin.setValue(DEFAULT_CHUNK_SEARCH_WINDOW); 
        self.chunk_window_spin.setToolTip("–û–∫–Ω–æ –ø–æ–∏—Å–∫–∞ —Ä–∞–∑–¥–µ–ª–∏—Ç–µ–ª—è.")
        self.chunk_delay_spin = QDoubleSpinBox()
        self.chunk_delay_spin.setRange(0.0, 300.0) # –û—Ç 0 –¥–æ 5 –º–∏–Ω—É—Ç
        self.chunk_delay_spin.setSingleStep(0.1)
        self.chunk_delay_spin.setValue(0.0) # –ü–æ —É–º–æ–ª—á–∞–Ω–∏—é –±–µ–∑ –∑–∞–¥–µ—Ä–∂–∫–∏
        self.chunk_delay_spin.setDecimals(1)
        self.chunk_delay_spin.setToolTip("–ó–∞–¥–µ—Ä–∂–∫–∞ –≤ —Å–µ–∫—É–Ω–¥–∞—Ö –º–µ–∂–¥—É –æ—Ç–ø—Ä–∞–≤–∫–æ–π —á–∞–Ω–∫–æ–≤.\n0.0 = –±–µ–∑ –∑–∞–¥–µ—Ä–∂–∫–∏.")
        self.chunking_checkbox.stateChanged.connect(self.toggle_chunking_details); chunking_layout.addWidget(self.chunking_checkbox, 0, 0, 1, 4); chunking_layout.addWidget(QLabel("–õ–∏–º–∏—Ç —Å–∏–º–≤–æ–ª–æ–≤:"), 1, 0); chunking_layout.addWidget(self.chunk_limit_spin, 1, 1); 
        chunking_layout.addWidget(QLabel("–û–∫–Ω–æ –ø–æ–∏—Å–∫–∞:"), 1, 2); chunking_layout.addWidget(self.chunk_window_spin, 1, 3); 
        chunking_layout.addWidget(QLabel("–ó–∞–¥–µ—Ä–∂–∫–∞ (—Å–µ–∫):"), 2, 0); chunking_layout.addWidget(self.chunk_delay_spin, 2, 1)
        self.chunk_limit_spin.setEnabled(self.chunking_checkbox.isChecked()); 
        self.chunk_window_spin.setEnabled(self.chunking_checkbox.isChecked()); 
        settings_prompt_layout.addWidget(chunking_group); 
        self.chunk_delay_spin.setEnabled(self.chunking_checkbox.isChecked())
        self.model_combo.currentTextChanged.connect(self.update_chunking_checkbox_suggestion)

        self.prompt_lbl = QLabel("–ü—Ä–æ–º–ø—Ç (–∏–Ω—Å—Ç—Ä—É–∫—Ü–∏—è –¥–ª—è API, `{text}` –±—É–¥–µ—Ç –∑–∞–º–µ–Ω–µ–Ω):"); self.prompt_edit = QPlainTextEdit(); self.prompt_edit.setPlaceholderText("–ó–∞–≥—Ä—É–∑–∫–∞ –ø—Ä–æ–º–ø—Ç–∞...")
        self.prompt_edit.setMinimumHeight(100)

        self.prompt_edit.setPlainText("""--- PROMPT START ---

**I. –†–û–õ–¨ –ò –û–°–ù–û–í–ù–ê–Ø –ó–ê–î–ê–ß–ê**

*   **–¢–≤–æ—è –†–æ–ª—å:** –¢—ã ‚Äî –ø—Ä–æ—Ñ–µ—Å—Å–∏–æ–Ω–∞–ª—å–Ω—ã–π –ø–µ—Ä–µ–≤–æ–¥—á–∏–∫ –∏ —Ä–µ–¥–∞–∫—Ç–æ—Ä. –¢–≤–æ—è –∑–∞–¥–∞—á–∞ ‚Äî –≤—ã–ø–æ–ª–Ω–∏—Ç—å –±–µ–∑—É–ø—Ä–µ—á–Ω—É—é –ª–∏—Ç–µ—Ä–∞—Ç—É—Ä–Ω—É—é –∞–¥–∞–ø—Ç–∞—Ü–∏—é —Ç–µ–∫—Å—Ç–∞ —Å –∏—Å—Ö–æ–¥–Ω–æ–≥–æ —è–∑—ã–∫–∞ (–∞–Ω–≥–ª–∏–π—Å–∫–∏–π, –∫–∏—Ç–∞–π—Å–∫–∏–π, —è–ø–æ–Ω—Å–∫–∏–π, –∫–æ—Ä–µ–π—Å–∫–∏–π –∏ –¥—Ä.) –Ω–∞ —Ä—É—Å—Å–∫–∏–π —è–∑—ã–∫. –¢—ã —Ä–∞–±–æ—Ç–∞–µ—à—å —Å —Ä–∞–∑–Ω—ã–º–∏ —Ñ–æ—Ä–º–∞—Ç–∞–º–∏ (–ª–∏—Ç–µ—Ä–∞—Ç—É—Ä–∞, —Å—Ç–∞—Ç—å–∏, DOCX, HTML) –∏ —É—á–∏—Ç—ã–≤–∞–µ—à—å –∫—É–ª—å—Ç—É—Ä–Ω—ã–µ –æ—Å–æ–±–µ–Ω–Ω–æ—Å—Ç–∏.
*   **–û—Å–Ω–æ–≤–Ω–∞—è –î–∏—Ä–µ–∫—Ç–∏–≤–∞:** –ü–µ—Ä–µ–≤–µ—Å—Ç–∏ —Ç–µ–∫—Å—Ç `{text}`. –ö–æ–Ω–µ—á–Ω—ã–π —Ä–µ–∑—É–ª—å—Ç–∞—Ç –¥–æ–ª–∂–µ–Ω –±—ã—Ç—å –∏—Å–∫–ª—é—á–∏—Ç–µ–ª—å–Ω–æ –Ω–∞ —Ä—É—Å—Å–∫–æ–º —è–∑—ã–∫–µ. –õ—é–±—ã–µ –∏–Ω–æ—Å—Ç—Ä–∞–Ω–Ω—ã–µ —Å–ª–æ–≤–∞, –∏–µ—Ä–æ–≥–ª–∏—Ñ—ã, –ø–∏–Ω—å–∏–Ω—å –∏ —Ç.–¥. –¥–æ–ª–∂–Ω—ã –±—ã—Ç—å –ø–æ–ª–Ω–æ—Å—Ç—å—é –ø–µ—Ä–µ–≤–µ–¥–µ–Ω—ã –∏–ª–∏ –≥—Ä–∞–º–æ—Ç–Ω–æ –∞–¥–∞–ø—Ç–∏—Ä–æ–≤–∞–Ω—ã. –û—à–∏–±–∫–∏ –æ—Ä–∏–≥–∏–Ω–∞–ª–∞, –µ—Å–ª–∏ –æ–Ω–∏ –µ—Å—Ç—å, —Å–ª–µ–¥—É–µ—Ç –∏—Å–ø—Ä–∞–≤–ª—è—Ç—å –≤ –ø—Ä–æ—Ü–µ—Å—Å–µ –ø–µ—Ä–µ–≤–æ–¥–∞. –ù–∏–∫–∞–∫–∏—Ö –ø—Ä–∏–º–µ—á–∞–Ω–∏–π –∏ —Å–Ω–æ—Å–æ–∫ –æ—Ç –ø–µ—Ä–µ–≤–æ–¥—á–∏–∫–∞.

**II. –ü–†–ò–ù–¶–ò–ü–´ –ê–î–ê–ü–¢–ê–¶–ò–ò**

1.  **–ï—Å—Ç–µ—Å—Ç–≤–µ–Ω–Ω—ã–π —Ä—É—Å—Å–∫–∏–π:** –ò–∑–±–µ–≥–∞–π –±—É–∫–≤–∞–ª—å–Ω–æ—Å—Ç–∏, –∏—â–∏ —Ä—É—Å—Å–∫–∏–µ —ç–∫–≤–∏–≤–∞–ª–µ–Ω—Ç—ã –∏ —Ä–µ—á–µ–≤—ã–µ –æ–±–æ—Ä–æ—Ç—ã.
2.  **–°–º—ã—Å–ª –∏ –¢–æ–Ω:** –¢–æ—á–Ω–æ –ø–µ—Ä–µ–¥–∞–≤–∞–π —Å–º—ã—Å–ª, –∞—Ç–º–æ—Å—Ñ–µ—Ä—É –∏ –∞–≤—Ç–æ—Ä—Å–∫–∏–π —Å—Ç–∏–ª—å.
3.  **–ö—É–ª—å—Ç—É—Ä–Ω–∞—è –∞–¥–∞–ø—Ç–∞—Ü–∏—è:**
*   **–•–æ–Ω–æ—Ä–∏—Ñ–∏–∫–∏ (-—Å–∞–Ω, -–∫—É–Ω):** –û–ø—É—Å–∫–∞–π –∏–ª–∏ –∑–∞–º–µ–Ω—è–π –µ—Å—Ç–µ—Å—Ç–≤–µ–Ω–Ω—ã–º–∏ –æ–±—Ä–∞—â–µ–Ω–∏—è–º–∏ (–ø–æ –∏–º–µ–Ω–∏, –≥–æ—Å–ø–æ–¥–∏–Ω/–≥–æ—Å–ø–æ–∂–∞).
*   **–†–µ–∞–ª–∏–∏:** –ê–¥–∞–ø—Ç–∏—Ä—É–π —á–µ—Ä–µ–∑ —Ä—É—Å—Å–∫–∏–µ —ç–∫–≤–∏–≤–∞–ª–µ–Ω—Ç—ã –∏–ª–∏ –∫—Ä–∞—Ç–∫–∏–µ, –æ—Ä–≥–∞–Ω–∏—á–Ω–æ –≤—Å—Ç—Ä–æ–µ–Ω–Ω—ã–µ –≤ —Ç–µ–∫—Å—Ç –ø–æ—è—Å–Ω–µ–Ω–∏—è.
*   **–û–Ω–æ–º–∞—Ç–æ–ø–µ—è (–ó–≤—É–∫–æ–ø–æ–¥—Ä–∞–∂–∞–Ω–∏–µ):** –ó–∞–º–µ–Ω—è–π —Ä—É—Å—Å–∫–∏–º–∏ –∑–≤—É–∫–æ–ø–æ–¥—Ä–∞–∂–∞–Ω–∏—è–º–∏ –∏–ª–∏ –æ–ø–∏—Å–∞–Ω–∏—è–º–∏ –∑–≤—É–∫–æ–≤.

**III. –§–û–†–ú–ê–¢–ò–†–û–í–ê–ù–ò–ï –ò –°–ü–ï–¶–¢–ï–ì–ò**

1.  **–°–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ –§–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏—è:** –ü–æ–ª–Ω–æ—Å—Ç—å—é —Å–æ—Ö—Ä–∞–Ω—è–π –∏—Å—Ö–æ–¥–Ω–æ–µ —Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ —Ç–µ–∫—Å—Ç–∞, –≤–∫–ª—é—á–∞—è –∞–±–∑–∞—Ü—ã, –∑–∞–≥–æ–ª–æ–≤–∫–∏ (Markdown `#`, `##`), —Å–ø–∏—Å–∫–∏ (`*`, `-`, `1.`) –∏ —Å—Ç—Ä—É–∫—Ç—É—Ä—É HTML.
2.  **HTML –ö–æ–Ω—Ç–µ–Ω—Ç:**
*   **–ö–†–ò–¢–ò–ß–ï–°–ö–ò –í–ê–ñ–ù–û: –°–û–•–†–ê–ù–Ø–ô –í–°–ï HTML-–¢–ï–ì–ò!** –ü–µ—Ä–µ–≤–æ–¥–∏ **–¢–û–õ–¨–ö–û –≤–∏–¥–∏–º—ã–π —Ç–µ–∫—Å—Ç** (–≤–Ω—É—Ç—Ä–∏ `<p>`, `<h1>`, `<li>`, `<td>`, `<span>`, `<a>`, –∞ —Ç–∞–∫–∂–µ –∑–Ω–∞—á–µ–Ω–∏—è –∞—Ç—Ä–∏–±—É—Ç–æ–≤ `title`, `alt`).
*   **–ù–ï –ò–ó–ú–ï–ù–Ø–ô** —Å—Ç—Ä—É–∫—Ç—É—Ä—É HTML, –∞—Ç—Ä–∏–±—É—Ç—ã, `<!-- –∫–æ–º–º–µ–Ω—Ç–∞—Ä–∏–∏ -->`, `<script>` –∏ `<style>`.
3.  **–ü–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä—ã –ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π:**
*   –¢–µ–≥–∏ –≤–∏–¥–∞ `<||img_placeholder_xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx||>` (32-—Å–∏–º–≤–æ–ª—å–Ω—ã–π ID).
*   **–ö–†–ò–¢–ò–ß–ï–°–ö–ò –í–ê–ñ–ù–û: –ö–û–ü–ò–†–£–ô –≠–¢–ò –¢–ï–ì–ò –ê–ë–°–û–õ–Æ–¢–ù–û –¢–û–ß–ù–û, –°–ò–ú–í–û–õ –í –°–ò–ú–í–û–õ. –ù–ï –ú–ï–ù–Ø–ô –ò–• –ò –ù–ï –£–î–ê–õ–Ø–ô.**

**IV. –°–¢–ò–õ–ò–ó–ê–¶–ò–Ø –ò –ü–£–ù–ö–¢–£–ê–¶–ò–Ø**

*   –†–µ–ø–ª–∏–∫–∏ –≤ `[]` –æ—Ñ–æ—Ä–º–ª—è–π –∫–∞–∫ –ø—Ä—è–º–æ–π –¥–∏–∞–ª–æ–≥: `‚Äî –†–µ–ø–ª–∏–∫–∞.`
*   –Ø–ø–æ–Ω—Å–∫–∏–µ –∫–∞–≤—ã—á–∫–∏ `„Äé„Äè` –∑–∞–º–µ–Ω—è–π –Ω–∞ —Ä—É—Å—Å–∫–∏–µ ¬´—ë–ª–æ—á–∫–∏¬ª.
*   –ú—ã—Å–ª–∏ –ø–µ—Ä—Å–æ–Ω–∞–∂–µ–π –æ—Ñ–æ—Ä–º–ª—è–π –∫–∞–∫: `¬´–ú—ã—Å–ª—å...¬ª` (–±–µ–∑ —Ç–∏—Ä–µ).
*   –ù–∞–∑–≤–∞–Ω–∏—è –Ω–∞–≤—ã–∫–æ–≤, –ø—Ä–µ–¥–º–µ—Ç–æ–≤, –∫–≤–µ—Å—Ç–æ–≤ –≤—ã–¥–µ–ª—è–π –∫–≤–∞–¥—Ä–∞—Ç–Ω—ã–º–∏ —Å–∫–æ–±–∫–∞–º–∏: `[–ù–∞–∑–≤–∞–Ω–∏–µ]`.
*   –î–ª–∏–Ω–Ω—ã–µ –ø–æ–≤—Ç–æ—Ä—ã –≥–ª–∞—Å–Ω—ã—Ö —Å–æ–∫—Ä–∞—â–∞–π –¥–æ 4-5 —Å–∏–º–≤–æ–ª–æ–≤: `–ê-–∞-–∞-–∞...`
*   –ó–∞–∏–∫–∞–Ω–∏–µ –æ—Ñ–æ—Ä–º–ª—è–π —á–µ—Ä–µ–∑ –¥–µ—Ñ–∏—Å: `–ü-–ø—Ä–∏–≤–µ—Ç`.
*   –≠–º–æ—Ü–∏–æ–Ω–∞–ª—å–Ω—ã–µ –∑–Ω–∞–∫–∏ –ø—Ä–µ–ø–∏–Ω–∞–Ω–∏—è: `–¢–µ–∫—Å—Ç!..`, `–¢–µ–∫—Å—Ç?..` (–º–Ω–æ–≥–æ—Ç–æ—á–∏–µ –ø–æ—Å–ª–µ –∑–Ω–∞–∫–∞). –ò–∑–±–µ–≥–∞–π –º–Ω–æ–∂–µ—Å—Ç–≤–µ–Ω–Ω—ã—Ö –∑–Ω–∞–∫–æ–≤: `–ê?`, `–ê!`, `–ê?!`.

**V. –†–ê–ë–û–¢–ê –° –ì–õ–û–°–°–ê–†–ò–ï–ú (–ö–†–ò–¢–ò–ß–ï–°–ö–ò –í–ê–ñ–ù–û)
*   –§–æ—Ä–º–∞—Ç –ì–ª–æ—Å—Å–∞—Ä–∏—è: –í–Ω–∏–º–∞–Ω–∏–µ! –í –≥–ª–æ—Å—Å–∞—Ä–∏–∏ —Ç–µ—Ä–º–∏–Ω—ã —á–∞—Å—Ç–æ –¥–∞–Ω—ã –≤ —Ñ–æ—Ä–º–∞—Ç–µ –†—É—Å—Å–∫–∏–π –ø–µ—Ä–µ–≤–æ–¥ (Original English). –≠—Ç–æ—Ç —Ñ–æ—Ä–º–∞—Ç ‚Äî –∏–Ω—Å—Ç—Ä—É–∫—Ü–∏—è –¥–ª—è —Ç–µ–±—è, –∞ –Ω–µ —à–∞–±–ª–æ–Ω –¥–ª—è –æ—Ç–≤–µ—Ç–∞. –¢—ã –¥–æ–ª–∂–µ–Ω –∏—Å–ø–æ–ª—å–∑–æ–≤–∞—Ç—å –¢–û–õ–¨–ö–û –†–£–°–°–ö–£–Æ –ß–ê–°–¢–¨ –ø–µ—Ä–µ–≤–æ–¥–∞. –ê–Ω–≥–ª–∏–π—Å–∫–∞—è —á–∞—Å—Ç—å –≤ —Å–∫–æ–±–∫–∞—Ö –≤ –∏—Ç–æ–≥–æ–≤–æ–º —Ç–µ–∫—Å—Ç–µ –Ω–µ–¥–æ–ø—É—Å—Ç–∏–º–∞.
*   –ü—Ä–∏–æ—Ä–∏—Ç–µ—Ç —Ç–µ—Ä–º–∏–Ω–æ–≤:
*   –í—Å–µ–≥–¥–∞ –∏—Å–ø–æ–ª—å–∑—É–π —Å–∞–º—ã–π —Ç–æ—á–Ω—ã–π –∏ –∫–æ–Ω–∫—Ä–µ—Ç–Ω—ã–π –ø–µ—Ä–µ–≤–æ–¥ –∏–∑ –≥–ª–æ—Å—Å–∞—Ä–∏—è.
*   –ü—Ä–∏–º–µ—Ä: –ï—Å–ª–∏ –≤ —Ç–µ–∫—Å—Ç–µ Agility Brute, –∞ –≤ –≥–ª–æ—Å—Å–∞—Ä–∏–∏ –µ—Å—Ç—å –ì—Ä—É–±–∏—è–Ω (Brute) –∏ –ì—Ä—É–±–∏—è–Ω-–õ–æ–≤–∫–∞—á (Agility Brute), —Ç—ã –æ–±—è–∑–∞–Ω –∏—Å–ø–æ–ª—å–∑–æ–≤–∞—Ç—å –ì—Ä—É–±–∏—è–Ω-–õ–æ–≤–∫–∞—á.
*   –†–∞–∑—Ä–µ—à–µ–Ω–∏–µ –∫–æ–Ω—Ñ–ª–∏–∫—Ç–æ–≤:
*   –ï—Å–ª–∏ –¥–ª—è –æ–¥–Ω–æ–≥–æ –∞–Ω–≥–ª–∏–π—Å–∫–æ–≥–æ —Ç–µ—Ä–º–∏–Ω–∞ –≤ –≥–ª–æ—Å—Å–∞—Ä–∏–∏ –¥–∞–Ω–æ –Ω–µ—Å–∫–æ–ª—å–∫–æ —Ä—É—Å—Å–∫–∏—Ö –≤–∞—Ä–∏–∞–Ω—Ç–æ–≤ (–Ω–∞–ø—Ä–∏–º–µ—Ä, –ù–∞—Å—Ç—Ä–æ–π—â–∏–∫ / –†–µ–≥—É–ª—è—Ç–æ—Ä / –ö–æ—Ä—Ä–µ–∫—Ç–æ—Ä), –≤—ã–±–µ—Ä–∏ –Ω–∞–∏–±–æ–ª–µ–µ –ø–æ–¥—Ö–æ–¥—è—â–∏–π –ø–æ –∫–æ–Ω—Ç–µ–∫—Å—Ç—É –∏ —Å—Ç—Ä–æ–≥–æ –ø—Ä–∏–¥–µ—Ä–∂–∏–≤–∞–π—Å—è —ç—Ç–æ–≥–æ –≤—ã–±–æ—Ä–∞ –Ω–∞ –ø—Ä–æ—Ç—è–∂–µ–Ω–∏–∏ –≤—Å–µ–≥–æ —Ç–µ–∫—Å—Ç–∞ –¥–ª—è –æ–±–µ—Å–ø–µ—á–µ–Ω–∏—è –µ–¥–∏–Ω–æ–æ–±—Ä–∞–∑–∏—è.
*   –û—Ç—Å—É—Ç—Å—Ç–≤—É—é—â–∏–µ —Ç–µ—Ä–º–∏–Ω—ã: –ï—Å–ª–∏ —Ç–µ—Ä–º–∏–Ω –æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç –≤ –≥–ª–æ—Å—Å–∞—Ä–∏–∏, –ø–µ—Ä–µ–≤–µ–¥–∏ –µ–≥–æ —Å–∞–º–æ—Å—Ç–æ—è—Ç–µ–ª—å–Ω–æ, –æ–ø–∏—Ä–∞—è—Å—å –Ω–∞ —Å—Ç–∏–ª—å –∏ –ª–æ–≥–∏–∫—É —É–∂–µ —Å—É—â–µ—Å—Ç–≤—É—é—â–∏—Ö –ø–µ—Ä–µ–≤–æ–¥–æ–≤. –ù–µ –æ—Å—Ç–∞–≤–ª—è–π –µ–≥–æ –Ω–∞ –∞–Ω–≥–ª–∏–π—Å–∫–æ–º.

**VI. –ì–õ–û–°–°–ê–†–ò–ô**


**VII. –ò–¢–û–ì–û–í–´–ô –†–ï–ó–£–õ–¨–¢–ê–¢**

1.  –ü—Ä–µ–¥–æ—Å—Ç–∞–≤—å **–¢–û–õ–¨–ö–û** –ø–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω—ã–π –∏ –∞–¥–∞–ø—Ç–∏—Ä–æ–≤–∞–Ω–Ω—ã–π —Ç–µ–∫—Å—Ç.
2.  **–ë–ï–ó** –≤–≤–æ–¥–Ω—ã—Ö —Ñ—Ä–∞–∑ —Ç–∏–ø–∞ ¬´–í–æ—Ç –≤–∞—à –ø–µ—Ä–µ–≤–æ–¥:¬ª.
3.  **–ë–ï–ó** –æ—Ä–∏–≥–∏–Ω–∞–ª—å–Ω–æ–≥–æ —Ç–µ–∫—Å—Ç–∞.
4.  **–ë–ï–ó** —Ç–≤–æ–∏—Ö –∫–æ–º–º–µ–Ω—Ç–∞—Ä–∏–µ–≤ (–∫—Ä–æ–º–µ –Ω–µ–∏–∑–º–µ–Ω–µ–Ω–Ω—ã—Ö HTML-–∫–æ–º–º–µ–Ω—Ç–∞—Ä–∏–µ–≤).
5.  **–í–Ω–∏–º–∞—Ç–µ–ª—å–Ω–æ —Å–ª–µ–¥–∏ –∑–∞ –ø–æ–ª–æ–º –ø–µ—Ä—Å–æ–Ω–∞–∂–µ–π –∏ —á–∏—Å–ª–∞–º–∏** –ø–æ –∫–æ–Ω—Ç–µ–∫—Å—Ç—É.
6.  **–§–∏–Ω–∞–ª—å–Ω–∞—è —Å–∞–º–æ–ø—Ä–æ–≤–µ—Ä–∫–∞:** –ü–µ—Ä–µ–¥ –æ—Ç–ø—Ä–∞–≤–∫–æ–π –æ—Ç–≤–µ—Ç–∞ –ø–µ—Ä–µ–ø—Ä–æ–≤–µ—Ä—å —Ç–µ–∫—Å—Ç –Ω–∞ –Ω–∞–ª–∏—á–∏–µ –Ω–µ–ø–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω—ã—Ö —Å–ª–æ–≤ –∏ —Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤–∏–µ –≤—Å–µ–º –∏–Ω—Å—Ç—Ä—É–∫—Ü–∏—è–º.
7. **–ö–†–ò–¢–ò–ß–ï–°–ö–ò –í–ê–ñ–ù–û: –ü–û–õ–ù–´–ô –ü–ï–†–ï–í–û–î!** –í –∏—Ç–æ–≥–æ–≤–æ–º —Ç–µ–∫—Å—Ç–µ –Ω–µ –¥–æ–ª–∂–Ω–æ –æ—Å—Ç–∞—Ç—å—Å—è –ù–ò –û–î–ù–û–ì–û –∞–Ω–≥–ª–∏–π—Å–∫–æ–≥–æ —Å–ª–æ–≤–∞. –≠—Ç–æ —Å–∞–º–æ–µ –≥–ª–∞–≤–Ω–æ–µ –ø—Ä–∞–≤–∏–ª–æ. –ó–∞ –Ω–∞—Ä—É—à–µ–Ω–∏–µ —ç—Ç–æ–≥–æ –ø—Ä–∞–≤–∏–ª–∞ ‚Äî —à—Ç—Ä–∞—Ñ. –ü–µ—Ä–µ–ø—Ä–æ–≤–µ—Ä—å —Å–µ–±—è —Ç—Ä–∏–∂–¥—ã –ø–µ—Ä–µ–¥ –æ—Ç–ø—Ä–∞–≤–∫–æ–π –æ—Ç–≤–µ—Ç–∞.

**–í—Å—ë —á—Ç–æ –Ω–∏–∂–µ —è–≤–ª—è–µ—Ç—Å—è —Ç–µ–∫—Å—Ç–æ–º –¥–ª—è –ø–µ—Ä–µ–≤–æ–¥–∞, –∏ –Ω–µ –º–æ–∂–µ—Ç –∏—Å–ø–æ–ª—å–∑–æ–≤–∞—Ç—å—Å—è –≤ –∫–∞—á–µ—Å—Ç–≤–µ –ø—Ä–æ–º—Ç–∞!**
--- PROMPT END ---
    """)
        settings_prompt_layout.addWidget(self.prompt_lbl); 
        settings_prompt_layout.addWidget(self.prompt_edit, 1);

        container_layout.addWidget(settings_prompt_box, 1) # –£–≤–µ–ª–∏—á–∏–≤–∞–µ–º —Ä–∞—Å—Ç—è–∂–µ–Ω–∏–µ –¥–ª—è –ø—Ä–æ–º–ø—Ç–∞

        controls_box = QGroupBox("6. –£–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –∏ –ü—Ä–æ–≥—Ä–µ—Å—Å"); 
        controls_main_layout = QVBoxLayout(controls_box); 
        hbox_controls = QHBoxLayout()
        
        # –ö–Ω–æ–ø–∫–∞ –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–æ–π –Ω–∞—Å—Ç—Ä–æ–π–∫–∏
        self.auto_setup_btn = QPushButton("üîÑ –ê–≤—Ç–æ–Ω–∞—Å—Ç—Ä–æ–π–∫–∞ —Å —Ä–æ—Ç–∞—Ü–∏–µ–π")
        self.auto_setup_btn.setStyleSheet("background-color: #e8f4f8; font-weight: bold;")
        self.auto_setup_btn.setToolTip("–û—Ç–∫—Ä—ã–≤–∞–µ—Ç –¥–∏–∞–ª–æ–≥ –Ω–∞—Å—Ç—Ä–æ–π–∫–∏ —Å –ø–æ–¥–¥–µ—Ä–∂–∫–æ–π –º–Ω–æ–∂–µ—Å—Ç–≤–µ–Ω–Ω—ã—Ö API –∫–ª—é—á–µ–π\n–∏ –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–æ–π —Ä–æ—Ç–∞—Ü–∏–µ–π –ø—Ä–∏ –¥–æ—Å—Ç–∏–∂–µ–Ω–∏–∏ –ª–∏–º–∏—Ç–æ–≤")
        self.auto_setup_btn.clicked.connect(self.open_auto_setup_dialog)
        
        self.start_btn = QPushButton("üöÄ –ù–∞—á–∞—Ç—å –ø–µ—Ä–µ–≤–æ–¥"); 
        self.start_btn.setStyleSheet("background-color: #ccffcc; font-weight: bold;"); 
        self.start_btn.clicked.connect(self.start_translation)
        self.finish_btn = QPushButton("üèÅ –ó–∞–≤–µ—Ä—à–∏—Ç—å") # <--- –ù–û–í–ê–Ø –ö–ù–û–ü–ö–ê
        self.finish_btn.setToolTip("–ó–∞–≤–µ—Ä—à–∏—Ç—å —Ç–µ–∫—É—â–∏–π —Ñ–∞–π–ª (—Å–æ—Ö—Ä–∞–Ω–∏—Ç—å –ø–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω—ã–µ —á–∞–Ω–∫–∏) –∏ –æ—Å—Ç–∞–Ω–æ–≤–∏—Ç—å –æ—Å—Ç–∞–ª—å–Ω—ã–µ –∑–∞–¥–∞—á–∏.")
        self.finish_btn.setEnabled(False)
        self.finish_btn.setStyleSheet("background-color: #e6ffe6;") # –°–≤–µ—Ç–ª–æ-–∑–µ–ª–µ–Ω—ã–π
        self.finish_btn.clicked.connect(self.finish_translation_gently) # <--- –ù–û–í–´–ô –û–ë–†–ê–ë–û–¢–ß–ò–ö
        self.cancel_btn = QPushButton("‚ùå –û—Ç–º–µ–Ω–∞"); 
        self.cancel_btn.setEnabled(False); 
        self.cancel_btn.setStyleSheet("background-color: #ffcccc;"); 
        self.cancel_btn.clicked.connect(self.cancel_translation)
        hbox_controls.addWidget(self.auto_setup_btn)
        hbox_controls.addWidget(self.start_btn, 1); 
        hbox_controls.addWidget(self.finish_btn)
        hbox_controls.addWidget(self.cancel_btn); 
        controls_main_layout.addLayout(hbox_controls)
        self.progress_bar = QProgressBar(); 
        self.progress_bar.setRange(0, 100); 
        self.progress_bar.setValue(0); 
        self.progress_bar.setTextVisible(True); 
        self.progress_bar.setFormat("%v / %m –∑–∞–¥–∞—á (%p%)")
        controls_main_layout.addWidget(self.progress_bar); 
        self.status_label = QLabel("–ì–æ—Ç–æ–≤"); 
        self.status_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        controls_main_layout.addWidget(self.status_label);

        container_layout.addWidget(controls_box)

        self.log_lbl = QLabel("–õ–æ–≥ –≤—ã–ø–æ–ª–Ω–µ–Ω–∏—è:"); 
        self.log_output = QTextEdit(); 
        self.log_output.setReadOnly(True); 
        self.log_output.setFont(QtGui.QFont("Consolas", 9)); 
        self.log_output.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
        self.log_output.setMinimumHeight(150) # –ó–∞–¥–∞–¥–∏–º –º–∏–Ω–∏–º–∞–ª—å–Ω—É—é –≤—ã—Å–æ—Ç—É –ª–æ–≥—É

        container_layout.addWidget(self.log_lbl);
        container_layout.addWidget(self.log_output, 2) # –£–≤–µ–ª–∏—á–∏–≤–∞–µ–º —Ä–∞—Å—Ç—è–∂–µ–Ω–∏–µ –¥–ª—è –ª–æ–≥–∞

        scroll_area.setWidget(container_widget)

        main_layout.addWidget(scroll_area)

        self.update_concurrency_suggestion(self.model_combo.currentText())
        self.update_chunking_checkbox_suggestion(self.model_combo.currentText())
        self.toggle_chunking_details(self.chunking_checkbox.checkState().value)


    @QtCore.pyqtSlot(int)
    def toggle_chunking_details(self, state):
        enabled = (state == Qt.CheckState.Checked.value)
        self.chunk_limit_spin.setEnabled(enabled)
        self.chunk_window_spin.setEnabled(enabled)

        self.chunk_delay_spin.setEnabled(enabled)


    @QtCore.pyqtSlot(str)
    def update_concurrency_suggestion(self, model_display_name):

        if model_display_name in MODELS:
            model_rpm = MODELS[model_display_name].get('rpm', 1)

            practical_limit = max(1, min(model_rpm, 15)) # Capped suggestion at 15

            self.concurrency_spin.setValue(min(practical_limit, 
            self.concurrency_spin.maximum()))
            self.concurrency_spin.setToolTip(f"–ú–∞–∫—Å. –∑–∞–ø—Ä–æ—Å–æ–≤.\n–ú–æ–¥–µ–ª—å: {model_display_name}\n–ó–∞—è–≤–ª–µ–Ω–æ RPM: {model_rpm}\n–†–µ–∫–æ–º.: ~{practical_limit}")
        else:
            self.concurrency_spin.setValue(1) # Fallback for unknown models
            self.concurrency_spin.setToolTip("–ú–∞–∫—Å. –∑–∞–ø—Ä–æ—Å–æ–≤.")

    @QtCore.pyqtSlot(str)
    def update_chunking_checkbox_suggestion(self, model_display_name):

        needs_chunking = False
        tooltip_text = f"–†–∞–∑–¥–µ–ª—è—Ç—å —Ñ–∞–π–ª—ã > –ª–∏–º–∏—Ç–∞."
        if model_display_name in MODELS:
            needs_chunking = MODELS[model_display_name].get('needs_chunking', False)
            tooltip_text += "\n–†–ï–ö–û–ú–ï–ù–î–£–ï–¢–°–Ø –í–ö–õ." if needs_chunking else "\n–ú–û–ñ–ù–û –í–´–ö–õ."
        else: # Assume unknown models might need it
            needs_chunking = True
            tooltip_text += "\n–ù–µ–∏–∑–≤–µ—Å—Ç–Ω–∞—è –º–æ–¥–µ–ª—å, —Ä–µ–∫–æ–º. –í–ö–õ."

        if not CHUNK_HTML_SOURCE:
             tooltip_text += "\n(–í–ù–ò–ú–ê–ù–ò–ï: –ß–∞–Ω–∫–∏–Ω–≥ HTML/EPUB –æ—Ç–∫–ª—é—á–µ–Ω)."

        self.chunking_checkbox.setChecked(needs_chunking)
        self.chunking_checkbox.setToolTip(tooltip_text)

        self.toggle_chunking_details(self.chunking_checkbox.checkState().value)

    @QtCore.pyqtSlot(int)
    def on_output_format_changed(self, index):
        """ Warns user if EPUB output is selected with non-EPUB inputs """

        selected_format_display = self.format_combo.itemText(index)
        current_output_format = OUTPUT_FORMATS.get(selected_format_display, 'txt')

        if not self.selected_files_data_tuples: return # No files selected yet

        if current_output_format == 'epub':

             if any(ft != 'epub' for ft, _, _ in self.selected_files_data_tuples):
                 QMessageBox.warning(self, "–ù–µ—Å–æ–≤–º–µ—Å—Ç–∏–º—ã–µ —Ñ–∞–π–ª—ã",
                                     "–î–ª—è –≤—ã–≤–æ–¥–∞ –≤ —Ñ–æ—Ä–º–∞—Ç EPUB –≤—ã–±—Ä–∞–Ω—ã –Ω–µ —Ç–æ–ª—å–∫–æ EPUB —Ñ–∞–π–ª—ã.\n"
                                     "–≠—Ç–æ—Ç —Ä–µ–∂–∏–º (EPUB->EPUB) —Ç—Ä–µ–±—É–µ—Ç –¢–û–õ–¨–ö–û EPUB —Ñ–∞–π–ª–æ–≤ –≤ —Å–ø–∏—Å–∫–µ.\n\n"
                                     "–ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –æ—á–∏—Å—Ç–∏—Ç–µ —Å–ø–∏—Å–æ–∫ –∏ –≤—ã–±–µ—Ä–∏—Ç–µ —Ç–æ–ª—å–∫–æ EPUB —Ñ–∞–π–ª—ã, "
                                     "–ª–∏–±–æ –≤—ã–±–µ—Ä–∏—Ç–µ –¥—Ä—É–≥–æ–π —Ñ–æ—Ä–º–∞—Ç –≤—ã–≤–æ–¥–∞.")

                 first_enabled_idx = 0
                 for i in range(self.format_combo.count()):
                      if self.format_combo.model().item(i).isEnabled():
                           first_enabled_idx = i; break
                 self.format_combo.setCurrentIndex(first_enabled_idx)

    def select_files(self):
        """Selects source files, handles EPUB HTML selection and TOC identification."""

        last_dir = self.out_folder or QStandardPaths.writableLocation(QStandardPaths.StandardLocation.DocumentsLocation)
        files, _ = QFileDialog.getOpenFileNames(self, "–í—ã–±–µ—Ä–∏—Ç–µ —Ñ–∞–π–ª—ã TXT, DOCX –∏–ª–∏ EPUB", last_dir, "–ü–æ–¥–¥–µ—Ä–∂–∏–≤–∞–µ–º—ã–µ —Ñ–∞–π–ª—ã (*.txt *.docx *.epub);;All files (*)")
        if not files: return

        selected_format_display = self.format_combo.currentText()
        current_output_format = OUTPUT_FORMATS.get(selected_format_display, 'txt')
        is_potential_epub_rebuild_mode = (current_output_format == 'epub')

        new_files_data_tuples = []; added_count = 0; skipped_count = 0

        current_files_set = { (p1, p2) for _, p1, p2 in self.selected_files_data_tuples }

        for file_path in files:
            file_ext = os.path.splitext(file_path)[1].lower()
            base_name = os.path.basename(file_path)

            if is_potential_epub_rebuild_mode:
                if file_ext != '.epub':
                    self.append_log(f"[WARN] –ü—Ä–æ–ø—É—Å–∫ {file_ext.upper()}: {base_name} (–Ω–µ–ª—å–∑—è —Å–º–µ—à–∏–≤–∞—Ç—å —Å EPUB –ø—Ä–∏ –≤—ã–≤–æ–¥–µ –≤ EPUB)")
                    skipped_count += 1
                    continue

                elif any(ft != 'epub' for ft, _, _ in self.selected_files_data_tuples):
                    self.append_log(f"[WARN] –ü—Ä–æ–ø—É—Å–∫ EPUB: {base_name} (—Å–ø–∏—Å–æ–∫ —É–∂–µ —Å–æ–¥–µ—Ä–∂–∏—Ç –Ω–µ-EPUB —Ñ–∞–π–ª—ã, –Ω–µ–ª—å–∑—è –≤—ã–±—Ä–∞—Ç—å EPUB —Ñ–æ—Ä–º–∞—Ç)")
                    skipped_count += 1
                    continue

            else: # Not EPUB output mode
                 if file_ext == '.epub' and any(ft != 'epub' for ft, _, _ in self.selected_files_data_tuples):
                     self.append_log(f"[WARN] –ü—Ä–æ–ø—É—Å–∫ EPUB: {base_name} (–Ω–µ–ª—å–∑—è —Å–º–µ—à–∏–≤–∞—Ç—å EPUB —Å TXT/DOCX –¥–ª—è —ç—Ç–æ–≥–æ —Ñ–æ—Ä–º–∞—Ç–∞ –≤—ã–≤–æ–¥–∞)")
                     skipped_count += 1
                     continue
                 if file_ext != '.epub' and any(ft == 'epub' for ft, _, _ in self.selected_files_data_tuples):
                      self.append_log(f"[WARN] –ü—Ä–æ–ø—É—Å–∫ {file_ext.upper()}: {base_name} (—Å–ø–∏—Å–æ–∫ —É–∂–µ —Å–æ–¥–µ—Ä–∂–∏—Ç EPUB, –Ω–µ–ª—å–∑—è –≤—ã–±—Ä–∞—Ç—å –Ω–µ-EPUB —Ñ–æ—Ä–º–∞—Ç –¥–ª—è –Ω–∏—Ö)")
                      skipped_count += 1
                      continue

            if file_ext == '.txt':
                file_tuple_key = (file_path, None)
                if file_tuple_key not in current_files_set:
                    new_files_data_tuples.append(('txt', file_path, None))
                    current_files_set.add(file_tuple_key); added_count += 1
                else: skipped_count += 1 # Already in list
            elif file_ext == '.docx':
                if not DOCX_AVAILABLE:
                    self.append_log(f"[WARN] –ü—Ä–æ–ø—É—Å–∫ DOCX: {base_name} (–±–∏–±–ª–∏–æ—Ç–µ–∫–∞ 'python-docx' –Ω–µ –Ω–∞–π–¥–µ–Ω–∞)"); skipped_count+=1; continue
                file_tuple_key = (file_path, None)
                if file_tuple_key not in current_files_set:
                    new_files_data_tuples.append(('docx', file_path, None))
                    current_files_set.add(file_tuple_key); added_count += 1
                else: skipped_count += 1
            elif file_ext == '.epub':

                if not BS4_AVAILABLE or not LXML_AVAILABLE: # Ebooklib checked based on output format later
                    self.append_log(f"[WARN] –ü—Ä–æ–ø—É—Å–∫ EPUB: {base_name} (—Ç—Ä–µ–±—É–µ—Ç—Å—è 'beautifulsoup4' –∏ 'lxml' –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏ EPUB)"); skipped_count+=1; continue

                try:
                    self.append_log(f"–ê–Ω–∞–ª–∏–∑ EPUB: {base_name}...")

                    nav_path, ncx_path, opf_dir_found, nav_id, ncx_id = self._find_epub_toc_paths(file_path)

                    if opf_dir_found is None:
                        self.append_log(f"[ERROR] –ù–µ —É–¥–∞–ª–æ—Å—å –æ–ø—Ä–µ–¥–µ–ª–∏—Ç—å —Å—Ç—Ä—É–∫—Ç—É—Ä—É OPF –≤ {base_name}. –ü—Ä–æ–ø—É—Å–∫ —Ñ–∞–π–ª–∞.")
                        skipped_count += 1; continue

                    can_process_epub = True
                    missing_lib_reason = ""
                    if current_output_format == 'epub' and (not EBOOKLIB_AVAILABLE):
                        can_process_epub = False; missing_lib_reason = "EbookLib (–¥–ª—è –∑–∞–ø–∏—Å–∏ EPUB)"
                    elif current_output_format == 'fb2' and not LXML_AVAILABLE: # LXML already checked above
                         pass # Should be fine if LXML check passed
                    elif current_output_format == 'docx' and not DOCX_AVAILABLE:
                         can_process_epub = False; missing_lib_reason = "python-docx (–¥–ª—è –∑–∞–ø–∏—Å–∏ DOCX)"


                    if not can_process_epub:
                        self.append_log(f"[WARN] –ü—Ä–æ–ø—É—Å–∫ EPUB->{current_output_format.upper()}: {base_name} (–æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç '{missing_lib_reason}')")
                        skipped_count+=1; continue


                    with zipfile.ZipFile(file_path, 'r') as epub_zip:

                        html_files_in_epub = sorted([
                            name for name in epub_zip.namelist()
                            if name.lower().endswith(('.html', '.xhtml', '.htm'))
                            and not name.startswith(('__MACOSX', 'META-INF/')) # Exclude common non-content paths
                        ])
                        if not html_files_in_epub:
                            self.append_log(f"[WARN] –í EPUB '{base_name}' –Ω–µ –Ω–∞–π–¥–µ–Ω–æ HTML/XHTML —Ñ–∞–π–ª–æ–≤."); skipped_count+=1; continue

                        dialog = EpubHtmlSelectorDialog(file_path, html_files_in_epub, nav_path, ncx_path, self)
                        if dialog.exec():
                            selected_html = dialog.get_selected_files()
                            if selected_html:
                                self.append_log(f"–í—ã–±—Ä–∞–Ω–æ {len(selected_html)} HTML –∏–∑ {base_name}:")
                                for html_path in selected_html: # html_path is relative to zip root
                                    epub_tuple_key = (file_path, html_path)
                                    if epub_tuple_key not in current_files_set:

                                        new_files_data_tuples.append(('epub', file_path, html_path))
                                        current_files_set.add(epub_tuple_key)

                                        is_nav_file = (html_path == nav_path)
                                        log_suffix = ""
                                        if is_nav_file and is_potential_epub_rebuild_mode:
                                             log_suffix = " (NAV - –ë–£–î–ï–¢ –ò–ó–ú–ï–ù–ï–ù, –ù–ï –ü–ï–†–ï–í–ï–î–ï–ù)"
                                        elif is_nav_file:
                                             log_suffix = " (NAV)" # For non-EPUB output
                                        self.append_log(f"  + {html_path}{log_suffix}")
                                        added_count += 1
                                    else:
                                        self.append_log(f"  - {html_path} (–¥—É–±–ª–∏–∫–∞—Ç)"); skipped_count+=1
                            else: # No HTML files selected in dialog
                                self.append_log(f"HTML –Ω–µ –≤—ã–±—Ä–∞–Ω—ã –∏–∑ {base_name}."); skipped_count+=1
                        else: # Dialog cancelled
                            self.append_log(f"–í—ã–±–æ—Ä HTML –∏–∑ {base_name} –æ—Ç–º–µ–Ω–µ–Ω."); skipped_count+=1
                except zipfile.BadZipFile:
                    self.append_log(f"[ERROR] –ù–µ —É–¥–∞–ª–æ—Å—å –æ—Ç–∫—Ä—ã—Ç—å EPUB: {base_name}. –í–æ–∑–º–æ–∂–Ω–æ, –ø–æ–≤—Ä–µ–∂–¥–µ–Ω."); skipped_count+=1
                except Exception as e:
                    self.append_log(f"[ERROR] –û—à–∏–±–∫–∞ –æ–±—Ä–∞–±–æ—Ç–∫–∏ EPUB {base_name}: {e}\n{traceback.format_exc()}"); skipped_count+=1
            else: # Unsupported file extension
                self.append_log(f"[WARN] –ü—Ä–æ–ø—É—Å–∫ –Ω–µ–ø–æ–¥–¥–µ—Ä–∂–∏–≤–∞–µ–º–æ–≥–æ —Ñ–∞–π–ª–∞: {base_name}"); skipped_count+=1

        if new_files_data_tuples:
            self.selected_files_data_tuples.extend(new_files_data_tuples)
            self.update_file_list_widget() # Sorts and updates display
            log_msg = f"–î–æ–±–∞–≤–ª–µ–Ω–æ {added_count} —Ñ–∞–π–ª–æ–≤/—á–∞—Å—Ç–µ–π."
            if skipped_count > 0: log_msg += f" –ü—Ä–æ–ø—É—â–µ–Ω–æ {skipped_count}."
            self.append_log(log_msg)
        elif skipped_count > 0:
            self.append_log(f"–ù–æ–≤—ã–µ —Ñ–∞–π–ª—ã –Ω–µ –¥–æ–±–∞–≤–ª–µ–Ω—ã. –ü—Ä–æ–ø—É—â–µ–Ω–æ {skipped_count}.")
        else: # No files selected or all skipped/duplicates
             if files: # If files were initially selected but none added/skipped
                 self.append_log("–í—ã–±—Ä–∞–Ω–Ω—ã–µ —Ñ–∞–π–ª—ã —É–∂–µ –≤ —Å–ø–∏—Å–∫–µ –∏–ª–∏ –Ω–µ –ø–æ–¥–¥–µ—Ä–∂–∏–≤–∞—é—Ç—Å—è.")


    def _find_epub_toc_paths(self, epub_path):
        """Finds NAV, NCX paths, OPF directory, and NAV/NCX item IDs within an EPUB."""

        nav_path_in_zip = None; ncx_path_in_zip = None
        opf_dir_in_zip = None; opf_path_in_zip = None
        nav_item_id = None; ncx_item_id = None
        try:
            with zipfile.ZipFile(epub_path, 'r') as zipf:

                try:
                    container_data = zipf.read('META-INF/container.xml')

                    container_root = etree.fromstring(container_data)

                    cnt_ns = {'oebps': 'urn:oasis:names:tc:opendocument:xmlns:container'}
                    opf_path_rel = container_root.xpath('//oebps:rootfile/@full-path', namespaces=cnt_ns)[0]
                    opf_path_in_zip = opf_path_rel.replace('\\', '/') # Normalize path separator
                    opf_dir_in_zip = os.path.dirname(opf_path_in_zip)
                    if opf_dir_in_zip == '.': opf_dir_in_zip = "" # Use empty string for root
                except (KeyError, IndexError, etree.XMLSyntaxError) as container_err:

                    print(f"[WARN] EPUB {Path(epub_path).name}: container.xml –Ω–µ –Ω–∞–π–¥–µ–Ω/–Ω–µ–∫–æ—Ä—Ä–µ–∫—Ç–µ–Ω ({container_err}). –ü–æ–∏—Å–∫ OPF...")
                    found_opf = False
                    for name in zipf.namelist():

                        if name.lower().endswith('.opf') and not name.lower().startswith('meta-inf/') and name.lower() != 'mimetype':
                             opf_path_in_zip = name.replace('\\', '/')

                             opf_dir_in_zip = os.path.dirname(opf_path_in_zip)

                             if opf_dir_in_zip == '.': opf_dir_in_zip = ""

                             print(f"[INFO] EPUB {Path(epub_path).name}: –ù–∞–π–¥–µ–Ω OPF: {opf_path_in_zip} (–≤ –¥–∏—Ä–µ–∫—Ç–æ—Ä–∏–∏: '{opf_dir_in_zip or '<root>'}')")
                             found_opf = True; break # Take the first one found
                    if not found_opf:
                        self.append_log(f"[ERROR] EPUB {Path(epub_path).name}: –ù–µ —É–¥–∞–ª–æ—Å—å –Ω–∞–π—Ç–∏ OPF —Ñ–∞–π–ª (–Ω–∏ —á–µ—Ä–µ–∑ container.xml, –Ω–∏ –ø–æ–∏—Å–∫–æ–º).")

                        return None, None, None, None, None # Critical failure

                if opf_path_in_zip is None or opf_dir_in_zip is None:
                     self.append_log(f"[ERROR] EPUB {Path(epub_path).name}: OPF –ø—É—Ç—å –∏–ª–∏ –¥–∏—Ä–µ–∫—Ç–æ—Ä–∏—è –Ω–µ –æ–ø—Ä–µ–¥–µ–ª–µ–Ω—ã.")
                     return None, None, None, None, None

                opf_data = zipf.read(opf_path_in_zip)
                opf_root = etree.fromstring(opf_data) # Use lxml for parsing OPF
                ns = {'opf': 'http://www.idpf.org/2007/opf'} # OPF namespace

                ncx_id_from_spine = None
                spine_node = opf_root.find('opf:spine', ns)
                if spine_node is not None:
                    ncx_id_from_spine = spine_node.get('toc') # 'toc' attribute points to NCX ID

                manifest_node = opf_root.find('opf:manifest', ns)
                if manifest_node is not None:
                    for item in manifest_node.findall('opf:item', ns):
                        item_id = item.get('id'); item_href = item.get('href');
                        item_media_type = item.get('media-type'); item_properties = item.get('properties')

                        if item_href: # Ensure href exists

                            item_path_abs = os.path.normpath(os.path.join(opf_dir_in_zip, item_href)).replace('\\', '/')

                            if item_properties and 'nav' in item_properties.split():
                                if nav_path_in_zip is None: # Take the first one found
                                    nav_path_in_zip = item_path_abs
                                    nav_item_id = item_id
                                else: print(f"[WARN] EPUB {Path(epub_path).name}: –ù–∞–π–¥–µ–Ω–æ –Ω–µ—Å–∫–æ–ª—å–∫–æ —ç–ª–µ–º–µ–Ω—Ç–æ–≤ —Å 'properties=nav'. –ò—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è –ø–µ—Ä–≤—ã–π: {nav_path_in_zip}")

                            if item_media_type == 'application/x-dtbncx+xml' or (ncx_id_from_spine and item_id == ncx_id_from_spine):
                                if ncx_path_in_zip is None: # Take the first one found
                                     ncx_path_in_zip = item_path_abs
                                     ncx_item_id = item_id
                                else: print(f"[WARN] EPUB {Path(epub_path).name}: –ù–∞–π–¥–µ–Ω–æ –Ω–µ—Å–∫–æ–ª—å–∫–æ NCX —Ñ–∞–π–ª–æ–≤. –ò—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è –ø–µ—Ä–≤—ã–π: {ncx_path_in_zip}")

            log_parts = [f"OPF_Dir='{opf_dir_in_zip or '<root>'}'"]
            if nav_path_in_zip: log_parts.append(f"NAV='{nav_path_in_zip}'(ID={nav_item_id})")
            if ncx_path_in_zip: log_parts.append(f"NCX='{ncx_path_in_zip}'(ID={ncx_item_id})")
            self.append_log(f"–°—Ç—Ä—É–∫—Ç—É—Ä–∞ {Path(epub_path).name}: {', '.join(log_parts)}")

            return nav_path_in_zip, ncx_path_in_zip, opf_dir_in_zip, nav_item_id, ncx_item_id

        except (KeyError, IndexError, etree.XMLSyntaxError, zipfile.BadZipFile) as e:
            self.append_log(f"[ERROR] –ù–µ —É–¥–∞–ª–æ—Å—å –Ω–∞–π—Ç–∏/–ø—Ä–æ—á–∏—Ç–∞—Ç—å —Å—Ç—Ä—É–∫—Ç—É—Ä—É OPF/TOC –≤ {os.path.basename(epub_path)}: {e}")
            return None, None, None, None, None # Return None for all on error

    def update_file_list_widget(self):
        """ Updates the list widget display, sorting items. """

        self.file_list_widget.clear()
        display_items = []

        sorted_data = sorted(self.selected_files_data_tuples, key=lambda x: (x[1], x[2] if x[2] else ""))
        self.selected_files_data_tuples = sorted_data # Update internal list with sorted version

        for file_type, path1, path2 in self.selected_files_data_tuples:
            if file_type == 'epub':

                display_items.append(f"{os.path.basename(path1)}  ->  {path2}")
            else:

                display_items.append(os.path.basename(path1))

        self.file_list_widget.addItems(display_items)
        self.file_list_widget.scrollToBottom() # Scroll to show newly added items
        self.update_file_count_display() # <<< –í–û–¢ –≠–¢–£ –°–¢–†–û–ß–ö–£ –î–û–ë–ê–í–ò–õ–ò

    def clear_file_list(self):

        self.selected_files_data_tuples = [] # Clear internal data
        self.file_list_widget.clear() # Clear display
        self.append_log("–°–ø–∏—Å–æ–∫ —Ñ–∞–π–ª–æ–≤ –æ—á–∏—â–µ–Ω.")
        self.update_file_count_display() # <<< –ò –°–Æ–î–ê –¢–û–ñ–ï –î–û–ë–ê–í–ò–õ–ò

    def select_output_folder(self):

        current_path = self.out_lbl.text()
        start_dir = current_path if os.path.isdir(current_path) else QStandardPaths.writableLocation(QStandardPaths.StandardLocation.DocumentsLocation)
        path = QFileDialog.getExistingDirectory(self, "–í—ã–±–µ—Ä–∏—Ç–µ –ø–∞–ø–∫—É –¥–ª—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è –ø–µ—Ä–µ–≤–æ–¥–æ–≤", start_dir)
        if path:
            self.out_folder = path
            self.out_lbl.setText(path)
            self.out_lbl.setCursorPosition(0) # Show start of path
            self.append_log(f"–ü–∞–ø–∫–∞ –≤—ã–≤–æ–¥–∞: {path}")

    def load_settings(self):

        default_prompt = self.prompt_edit.toPlainText()
        default_out_folder = ""
        default_format_display_name = DEFAULT_OUTPUT_FORMAT_DISPLAY
        default_model_name = DEFAULT_MODEL_NAME
        default_concurrency = self.concurrency_spin.value()
        default_chunking_enabled = self.chunking_checkbox.isChecked()
        default_chunk_limit = self.chunk_limit_spin.value()
        default_chunk_window = self.chunk_window_spin.value()
        default_temperature = 1.0
        default_chunk_delay = 0.0 # <-- –ù–æ–≤–æ–µ –∑–Ω–∞—á–µ–Ω–∏–µ –ø–æ —É–º–æ–ª—á–∞–Ω–∏—é
        default_proxy_url = "" # <-- –ù–æ–≤–æ–µ –∑–Ω–∞—á–µ–Ω–∏–µ –ø–æ —É–º–æ–ª—á–∞–Ω–∏—é –¥–ª—è –ø—Ä–æ–∫—Å–∏

        settings_loaded_successfully = False
        settings_source_message = f"–§–∞–π–ª '{SETTINGS_FILE}' –Ω–µ –Ω–∞–π–¥–µ–Ω –∏–ª–∏ –ø—É—Å—Ç. –ò—Å–ø–æ–ª—å–∑—É—é—Ç—Å—è —É–º–æ–ª—á–∞–Ω–∏—è."

        try:
            if os.path.exists(SETTINGS_FILE):
                self.config.clear()
                read_ok = self.config.read(SETTINGS_FILE, encoding='utf-8')
                if read_ok and 'Settings' in self.config:
                    settings = self.config['Settings']
                    
                    self.prompt_edit.setPlainText(settings.get('Prompt', default_prompt))
                    loaded_out_folder = settings.get('OutputFolder', default_out_folder)
                    self.out_folder = loaded_out_folder if os.path.isdir(loaded_out_folder) else default_out_folder
                    self.out_lbl.setText(self.out_folder if self.out_folder else "<–Ω–µ –≤—ã–±—Ä–∞–Ω–æ>")
                    self.out_lbl.setCursorPosition(0)
                    saved_format_display = settings.get('OutputFormat', default_format_display_name)
                    format_index = self.format_combo.findText(saved_format_display, Qt.MatchFlag.MatchFixedString)
                    first_enabled_idx = 0
                    for i_fmt in range(self.format_combo.count()):
                        if self.format_combo.model().item(i_fmt).isEnabled():
                            first_enabled_idx = i_fmt; break
                    if format_index != -1 and self.format_combo.model().item(format_index).isEnabled():
                        self.format_combo.setCurrentIndex(format_index)
                    else:
                        self.format_combo.setCurrentIndex(first_enabled_idx)
                        if format_index != -1:
                             settings_source_message = f"[WARN] –°–æ—Ö—Ä–∞–Ω–µ–Ω–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç '{saved_format_display}' –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω. –ò—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è '{self.format_combo.itemText(first_enabled_idx)}'."
                    model_name = settings.get('Model', default_model_name)
                    self.model_combo.setCurrentText(model_name if model_name in MODELS else default_model_name)
                    self.concurrency_spin.setValue(settings.getint('Concurrency', default_concurrency))
                    self.chunking_checkbox.setChecked(settings.getboolean('ChunkingEnabled', default_chunking_enabled))
                    self.chunk_limit_spin.setValue(settings.getint('ChunkLimit', default_chunk_limit))
                    self.chunk_window_spin.setValue(settings.getint('ChunkWindow', default_chunk_window))
                    self.temperature_spin.setValue(settings.getfloat('Temperature', default_temperature))

                    self.chunk_delay_spin.setValue(settings.getfloat('ChunkDelay', default_chunk_delay))

                    # --- –ó–ê–ì–†–£–ó–ö–ê –ü–†–û–ö–°–ò ---
                    self.proxy_url_edit.setText(settings.get('ProxyURL', default_proxy_url))
                    # --- –ö–û–ù–ï–¶ –ó–ê–ì–†–£–ó–ö–ò –ü–†–û–ö–°–ò ---
                    
                    settings_loaded_successfully = True
                    settings_source_message = f"–ù–∞—Å—Ç—Ä–æ–π–∫–∏ –∑–∞–≥—Ä—É–∂–µ–Ω—ã –∏–∑ '{SETTINGS_FILE}'."
        except (configparser.Error, ValueError, KeyError) as e:
            settings_source_message = f"[ERROR] –û—à–∏–±–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ –Ω–∞—Å—Ç—Ä–æ–µ–∫ ({e}). –ò—Å–ø–æ–ª—å–∑—É—é—Ç—Å—è —É–º–æ–ª—á–∞–Ω–∏—è."
            settings_loaded_successfully = False
        
        self.append_log(settings_source_message)

        if not settings_loaded_successfully:
            self.prompt_edit.setPlainText(default_prompt)
            self.out_folder = default_out_folder
            self.out_lbl.setText(self.out_folder if self.out_folder else "<–Ω–µ –≤—ã–±—Ä–∞–Ω–æ>")
            self.out_lbl.setCursorPosition(0)
            first_enabled_idx_def = 0
            for i_fmt_def in range(self.format_combo.count()):
                if self.format_combo.model().item(i_fmt_def).isEnabled():
                    first_enabled_idx_def = i_fmt_def; break
            self.format_combo.setCurrentIndex(first_enabled_idx_def)
            self.model_combo.setCurrentText(default_model_name)
            self.concurrency_spin.setValue(default_concurrency)
            self.chunking_checkbox.setChecked(default_chunking_enabled)
            self.chunk_limit_spin.setValue(default_chunk_limit)
            self.chunk_window_spin.setValue(default_chunk_window)
            self.temperature_spin.setValue(default_temperature)

            self.chunk_delay_spin.setValue(default_chunk_delay)
            # --- –£–°–¢–ê–ù–û–í–ö–ê –ü–†–û–ö–°–ò –ü–û –£–ú–û–õ–ß–ê–ù–ò–Æ ---
            self.proxy_url_edit.setText(default_proxy_url)
            # --- –ö–û–ù–ï–¶ –£–°–¢–ê–ù–û–í–ö–ò –ü–†–û–ö–°–ò ---


        self.toggle_chunking_details(self.chunking_checkbox.checkState().value)
        self.update_concurrency_suggestion(self.model_combo.currentText())
        self.update_chunking_checkbox_suggestion(self.model_combo.currentText())


    def save_settings(self):
        try:
            if 'Settings' not in self.config: self.config['Settings'] = {}
            settings = self.config['Settings']
            settings['Prompt'] = self.prompt_edit.toPlainText()
            settings['OutputFolder'] = self.out_folder or ""
            settings['OutputFormat'] = self.format_combo.currentText()
            settings['Model'] = self.model_combo.currentText()
            settings['Concurrency'] = str(self.concurrency_spin.value())
            settings['ChunkingEnabled'] = str(self.chunking_checkbox.isChecked())
            settings['ChunkLimit'] = str(self.chunk_limit_spin.value())
            settings['ChunkWindow'] = str(self.chunk_window_spin.value())
            settings['Temperature'] = str(self.temperature_spin.value())

            settings['ChunkDelay'] = str(self.chunk_delay_spin.value())

            # --- –°–û–•–†–ê–ù–ï–ù–ò–ï –ü–†–û–ö–°–ò ---
            settings['ProxyURL'] = self.proxy_url_edit.text().strip()
            # --- –ö–û–ù–ï–¶ –°–û–•–†–ê–ù–ï–ù–ò–Ø –ü–†–û–ö–°–ò ---

            with open(SETTINGS_FILE, 'w', encoding='utf-8') as configfile:
                self.config.write(configfile)
        except Exception as e:
            self.append_log(f"[ERROR] –ù–µ —É–¥–∞–ª–æ—Å—å —Å–æ—Ö—Ä–∞–Ω–∏—Ç—å –Ω–∞—Å—Ç—Ä–æ–π–∫–∏: {e}")

    def check_api_key(self):
        """Checks if the API key is valid by listing models."""

        current_api_key_to_check = self.api_key
        prompt_for_new_key = not current_api_key_to_check

        if prompt_for_new_key:
            key, ok = QtWidgets.QInputDialog.getText(self, "–¢—Ä–µ–±—É–µ—Ç—Å—è API –∫–ª—é—á", "–í–≤–µ–¥–∏—Ç–µ –≤–∞—à Google API Key:", QLineEdit.EchoMode.Password)
            current_api_key_to_check = key.strip() if ok and key.strip() else None

        if not current_api_key_to_check:
            QMessageBox.warning(self, "–ü—Ä–æ–≤–µ—Ä–∫–∞ –∫–ª—é—á–∞", "API –∫–ª—é—á –Ω–µ –≤–≤–µ–¥–µ–Ω.")
            return

        self.append_log(f"–ü—Ä–æ–≤–µ—Ä–∫–∞ API –∫–ª—é—á–∞...")
        self.check_api_key_btn.setEnabled(False)
        self.setCursor(Qt.CursorShape.WaitCursor)

        key_valid = False
        original_key = self.api_key # Store original key in case check fails

        try:

            genai.configure(api_key=current_api_key_to_check)

            models = genai.list_models()

            key_valid = any(m.name.startswith("models/") for m in models)

            if key_valid:

                if current_api_key_to_check != self.api_key:
                    self.api_key = current_api_key_to_check
                    self.append_log("[INFO] –ù–æ–≤—ã–π API –∫–ª—é—á –ø—Ä–∏–Ω—è—Ç –∏ —Å–æ—Ö—Ä–∞–Ω–µ–Ω.")
                QMessageBox.information(self, "–ü—Ä–æ–≤–µ—Ä–∫–∞ –∫–ª—é—á–∞", "API –∫–ª—é—á –¥–µ–π—Å—Ç–≤–∏—Ç–µ–ª–µ–Ω.")
                self.append_log("[SUCCESS] API –∫–ª—é—á –¥–µ–π—Å—Ç–≤–∏—Ç–µ–ª–µ–Ω.")
            else:

                 QMessageBox.warning(self, "–ü—Ä–æ–≤–µ—Ä–∫–∞ –∫–ª—é—á–∞", "–ö–ª—é—á –ø—Ä–∏–Ω—è—Ç API, –Ω–æ –Ω–µ –Ω–∞–π–¥–µ–Ω–æ –¥–æ—Å—Ç—É–ø–Ω—ã—Ö –º–æ–¥–µ–ª–µ–π Gemini.")
                 self.append_log("[WARN] –ü—Ä–æ–≤–µ—Ä–∫–∞ –∫–ª—é—á–∞: –ù–µ—Ç –¥–æ—Å—Ç—É–ø–Ω—ã—Ö –º–æ–¥–µ–ª–µ–π Gemini.")

        except google_exceptions.Unauthenticated as e:
            QMessageBox.critical(self, "–ü—Ä–æ–≤–µ—Ä–∫–∞ –∫–ª—é—á–∞", f"–û—à–∏–±–∫–∞ –∞—É—Ç–µ–Ω—Ç–∏—Ñ–∏–∫–∞—Ü–∏–∏ (–Ω–µ–≤–µ—Ä–Ω—ã–π –∫–ª—é—á?):\n{e}")
            self.append_log(f"[ERROR] –ü—Ä–æ–≤–µ—Ä–∫–∞ –∫–ª—é—á–∞: –ù–µ–≤–µ—Ä–Ω—ã–π ({e})")

            if current_api_key_to_check == self.api_key: self.api_key = None
            key_valid = False
        except google_exceptions.PermissionDenied as e:
            QMessageBox.critical(self, "–ü—Ä–æ–≤–µ—Ä–∫–∞ –∫–ª—é—á–∞", f"–û—à–∏–±–∫–∞ —Ä–∞–∑—Ä–µ—à–µ–Ω–∏–π (–∫–ª—é—á –Ω–µ –∞–∫—Ç–∏–≤–∏—Ä–æ–≤–∞–Ω –¥–ª—è API?):\n{e}")
            self.append_log(f"[ERROR] –ü—Ä–æ–≤–µ—Ä–∫–∞ –∫–ª—é—á–∞: –û—à–∏–±–∫–∞ —Ä–∞–∑—Ä–µ—à–µ–Ω–∏–π ({e})")
            key_valid = False # Key is likely valid but lacks permissions
        except google_exceptions.GoogleAPICallError as e: # Network errors etc.
            QMessageBox.critical(self, "–ü—Ä–æ–≤–µ—Ä–∫–∞ –∫–ª—é—á–∞", f"–û—à–∏–±–∫–∞ –≤—ã–∑–æ–≤–∞ API (—Å–µ—Ç—å?):\n{e}")
            self.append_log(f"[ERROR] –ü—Ä–æ–≤–µ—Ä–∫–∞ –∫–ª—é—á–∞: –û—à–∏–±–∫–∞ –≤—ã–∑–æ–≤–∞ API ({e})")
            key_valid = False
        except Exception as e: # Catch-all
            QMessageBox.critical(self, "–ü—Ä–æ–≤–µ—Ä–∫–∞ –∫–ª—é—á–∞", f"–ù–µ–æ–∂–∏–¥–∞–Ω–Ω–∞—è –æ—à–∏–±–∫–∞:\n{e}")
            self.append_log(f"[ERROR] –ü—Ä–æ–≤–µ—Ä–∫–∞ –∫–ª—é—á–∞: ({e})\n{traceback.format_exc()}")
            key_valid = False
        finally:

            self.check_api_key_btn.setEnabled(True)
            self.unsetCursor()

            final_key_to_configure = self.api_key # self.api_key was updated only if key_valid and different
            try:
                 if final_key_to_configure:
                     genai.configure(api_key=final_key_to_configure)
                 else:

                      self.append_log("[WARN] –î–µ–π—Å—Ç–≤—É—é—â–∏–π API –∫–ª—é—á –Ω–µ–∏–∑–≤–µ—Å—Ç–µ–Ω. API –º–æ–∂–µ—Ç –Ω–µ —Ä–∞–±–æ—Ç–∞—Ç—å.")

            except Exception as configure_err:

                 self.append_log(f"[ERROR] –û—à–∏–±–∫–∞ –≤–æ—Å—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω–∏—è –∫–æ–Ω—Ñ–∏–≥—É—Ä–∞—Ü–∏–∏ API: {configure_err}")

    @QtCore.pyqtSlot(str)
    def handle_log_message(self, message):

        self.append_log(message)

    def append_log(self, message):
        """Appends a timestamped message to the log widget."""

        current_time = time.strftime("%H:%M:%S")

        message_str = str(message).strip()

        for line in message_str.splitlines():
            self.log_output.append(f"[{current_time}] {line}")

        scrollbar = self.log_output.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())

    @QtCore.pyqtSlot(int)
    def update_file_progress(self, processed_count):

        self.progress_bar.setValue(processed_count)

    @QtCore.pyqtSlot(int)
    def update_progress_bar_range(self, total_tasks):
        """Sets the maximum value of the progress bar."""

        self.progress_bar.setRange(0, max(1, total_tasks))
        self.progress_bar.setValue(0) # Reset progress value
        self.progress_bar.setFormat(f"%v / {total_tasks} –∑–∞–¥–∞—á (%p%)") # Update text format
        self.append_log(f"–û–±—â–µ–µ –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ –∑–∞–¥–∞—á –¥–ª—è –≤—ã–ø–æ–ª–Ω–µ–Ω–∏—è: {total_tasks}")

    @QtCore.pyqtSlot(str)
    def handle_current_file_status(self, message):

        self.status_label.setText(message)

    @QtCore.pyqtSlot(str, int, int)
    def handle_chunk_progress(self, filename, current_chunk, total_chunks):
        """Updates the status label with chunk processing progress."""

        if total_chunks > 1 and current_chunk >= 0:
            max_len = 60 # Max length for filename display

            display_name = filename if len(filename) <= max_len else f"...{filename[-(max_len-3):]}"
            self.status_label.setText(f"–§–∞–π–ª: {display_name} [–ß–∞–Ω–∫: {current_chunk}/{total_chunks}]")
        elif total_chunks == 1 and current_chunk > 0: # Single chunk file completed
             max_len = 60
             display_name = filename if len(filename) <= max_len else f"...{filename[-(max_len-3):]}"
             self.status_label.setText(f"–§–∞–π–ª: {display_name} [1/1 –ó–∞–≤–µ—Ä—à–µ–Ω–æ]")



    def start_translation(self):
        """Validates inputs and starts the background worker thread."""
        prompt_template = self.prompt_edit.toPlainText().strip()
        selected_model_name = self.model_combo.currentText()
        max_concurrency = self.concurrency_spin.value()
        selected_files_tuples = list(self.selected_files_data_tuples)
        selected_format_display = self.format_combo.currentText()
        output_format = OUTPUT_FORMATS.get(selected_format_display, 'txt')
        chunking_enabled_gui = self.chunking_checkbox.isChecked()
        chunk_limit = self.chunk_limit_spin.value(); chunk_window = self.chunk_window_spin.value()
        temperature = self.temperature_spin.value()

        chunk_delay = self.chunk_delay_spin.value()

        # --- –ü–û–õ–£–ß–ï–ù–ò–ï –ü–†–û–ö–°–ò –ò–ó GUI ---
        proxy_string = self.proxy_url_edit.text().strip()
        # --- –ö–û–ù–ï–¶ –ü–û–õ–£–ß–ï–ù–ò–Ø –ü–†–û–ö–°–ò ---


        if not selected_files_tuples:
            QMessageBox.warning(self, "–û—à–∏–±–∫–∞", "–ù–µ –≤—ã–±—Ä–∞–Ω—ã —Ñ–∞–π–ª—ã –¥–ª—è –ø–µ—Ä–µ–≤–æ–¥–∞."); return
        if not self.out_folder:
            QMessageBox.warning(self, "–û—à–∏–±–∫–∞", "–ù–µ –≤—ã–±—Ä–∞–Ω–∞ –ø–∞–ø–∫–∞ –≤—ã–≤–æ–¥–∞."); return
        if not os.path.isdir(self.out_folder):
             reply = QMessageBox.question(self, "–ü–∞–ø–∫–∞ –Ω–µ —Å—É—â–µ—Å—Ç–≤—É–µ—Ç", f"–ü–∞–ø–∫–∞ '{self.out_folder}' –Ω–µ –Ω–∞–π–¥–µ–Ω–∞.\n–°–æ–∑–¥–∞—Ç—å?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.Yes)
             if reply == QMessageBox.StandardButton.Yes:
                 try: os.makedirs(self.out_folder, exist_ok=True); self.append_log(f"–ü–∞–ø–∫–∞ '{self.out_folder}' —Å–æ–∑–¥–∞–Ω–∞.")
                 except OSError as e: QMessageBox.critical(self, "–û—à–∏–±–∫–∞", f"–ù–µ —É–¥–∞–ª–æ—Å—å —Å–æ–∑–¥–∞—Ç—å –ø–∞–ø–∫—É: {e}"); return
             else: return 

        if output_format == 'docx' and not DOCX_AVAILABLE:
            QMessageBox.critical(self, "–û—à–∏–±–∫–∞", "–í—ã–±—Ä–∞–Ω —Ñ–æ—Ä–º–∞—Ç –≤—ã–≤–æ–¥–∞ DOCX, –Ω–æ –±–∏–±–ª–∏–æ—Ç–µ–∫–∞ 'python-docx' –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω–∞."); return
        if output_format == 'epub' and (not EBOOKLIB_AVAILABLE or not LXML_AVAILABLE or not BS4_AVAILABLE):
            QMessageBox.critical(self, "–û—à–∏–±–∫–∞", "–í—ã–±—Ä–∞–Ω —Ñ–æ—Ä–º–∞—Ç –≤—ã–≤–æ–¥–∞ EPUB, –Ω–æ –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω—ã: 'ebooklib', 'lxml' –∏ 'beautifulsoup4'."); return
        if output_format == 'fb2' and not LXML_AVAILABLE:
            QMessageBox.critical(self, "–û—à–∏–±–∫–∞", "–í—ã–±—Ä–∞–Ω —Ñ–æ—Ä–º–∞—Ç –≤—ã–≤–æ–¥–∞ FB2, –Ω–æ –±–∏–±–ª–∏–æ—Ç–µ–∫–∞ 'lxml' –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω–∞."); return
        if output_format in ['docx', 'epub', 'fb2', 'html'] and not PILLOW_AVAILABLE:
            QMessageBox.warning(self, "–ü—Ä–µ–¥—É–ø—Ä–µ–∂–¥–µ–Ω–∏–µ", f"–í—ã–±—Ä–∞–Ω —Ñ–æ—Ä–º–∞—Ç –≤—ã–≤–æ–¥–∞ {output_format.upper()} —Å –ø–æ–¥–¥–µ—Ä–∂–∫–æ–π –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π, –Ω–æ –±–∏–±–ª–∏–æ—Ç–µ–∫–∞ 'Pillow' –Ω–µ –Ω–∞–π–¥–µ–Ω–∞.\n–û–±—Ä–∞–±–æ—Ç–∫–∞ –Ω–µ–∫–æ—Ç–æ—Ä—ã—Ö —Ñ–æ—Ä–º–∞—Ç–æ–≤ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π (–Ω–∞–ø—Ä., EMF) –º–æ–∂–µ—Ç –±—ã—Ç—å –Ω–µ–≤–æ–∑–º–æ–∂–Ω–∞.")
        needs_docx_input = any(ft == 'docx' for ft, _, _ in selected_files_tuples)
        needs_epub_input = any(ft == 'epub' for ft, _, _ in selected_files_tuples)
        if needs_docx_input and not DOCX_AVAILABLE:
            QMessageBox.critical(self, "–û—à–∏–±–∫–∞", "–í—ã–±—Ä–∞–Ω—ã DOCX —Ñ–∞–π–ª—ã –¥–ª—è –≤–≤–æ–¥–∞, –Ω–æ –±–∏–±–ª–∏–æ—Ç–µ–∫–∞ 'python-docx' –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω–∞."); return
        if needs_epub_input and (not BS4_AVAILABLE or not LXML_AVAILABLE): 
            QMessageBox.critical(self, "–û—à–∏–±–∫–∞", "–í—ã–±—Ä–∞–Ω—ã EPUB —Ñ–∞–π–ª—ã –¥–ª—è –≤–≤–æ–¥–∞, –Ω–æ –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω—ã 'beautifulsoup4' –∏/–∏–ª–∏ 'lxml'."); return
        if selected_model_name not in MODELS:
            QMessageBox.critical(self, "–û—à–∏–±–∫–∞", f"–ù–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω–∞—è –º–æ–¥–µ–ª—å API: {selected_model_name}"); return
        if "{text}" not in prompt_template:
            QMessageBox.warning(self, "–û—à–∏–±–∫–∞ –ü—Ä–æ–º–ø—Ç–∞", "–ü—Ä–æ–º–ø—Ç –î–û–õ–ñ–ï–ù —Å–æ–¥–µ—Ä–∂–∞—Ç—å –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä `{text}` –¥–ª—è –≤—Å—Ç–∞–≤–∫–∏ —Ç–µ–∫—Å—Ç–∞."); return
        if "<||" not in prompt_template or "img_placeholder" not in prompt_template:
            QMessageBox.warning(self, "–ü—Ä–µ–¥—É–ø—Ä–µ–∂–¥–µ–Ω–∏–µ –ü—Ä–æ–º–ø—Ç–∞", "–ü—Ä–æ–º–ø—Ç –Ω–µ —Å–æ–¥–µ—Ä–∂–∏—Ç —è–≤–Ω—ã—Ö –∏–Ω—Å—Ç—Ä—É–∫—Ü–∏–π –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏ –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä–æ–≤ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π (`<||img_placeholder_...||>`).\nAPI –º–æ–∂–µ—Ç –∏—Ö —Å–ª—É—á–∞–π–Ω–æ –∏–∑–º–µ–Ω–∏—Ç—å –∏–ª–∏ —É–¥–∞–ª–∏—Ç—å.")
        if not self.api_key:
            key, ok = QtWidgets.QInputDialog.getText(self, "–¢—Ä–µ–±—É–µ—Ç—Å—è API –∫–ª—é—á", "–í–≤–µ–¥–∏—Ç–µ –≤–∞—à Google API Key:", QLineEdit.EchoMode.Password)
            if ok and key.strip(): self.api_key = key.strip(); self.append_log("[INFO] API –∫–ª—é—á –ø—Ä–∏–Ω—è—Ç.")
            else: QMessageBox.critical(self, "–û—à–∏–±–∫–∞", "API –∫–ª—é—á –Ω–µ –ø—Ä–µ–¥–æ—Å—Ç–∞–≤–ª–µ–Ω."); return
        if self.thread_ref and self.thread_ref.isRunning():
            QMessageBox.warning(self, "–í–Ω–∏–º–∞–Ω–∏–µ", "–ü—Ä–æ—Ü–µ—Å—Å –ø–µ—Ä–µ–≤–æ–¥–∞ —É–∂–µ –∑–∞–ø—É—â–µ–Ω."); return

        is_epub_to_epub_mode = False
        worker_data = None
        if output_format == 'epub':
            if not selected_files_tuples or not all(ft == 'epub' for ft, _, _ in selected_files_tuples):
                 QMessageBox.critical(self, "–û—à–∏–±–∫–∞ –ö–æ–Ω—Ñ–∏–≥—É—Ä–∞—Ü–∏–∏", "–û–±–Ω–∞—Ä—É–∂–µ–Ω–∞ –Ω–µ—Å–æ–≤–º–µ—Å—Ç–∏–º–æ—Å—Ç—å: –≤—ã–±—Ä–∞–Ω –≤—ã–≤–æ–¥ EPUB, –Ω–æ —Å–ø–∏—Å–æ–∫ —Å–æ–¥–µ—Ä–∂–∏—Ç –Ω–µ-EPUB —Ñ–∞–π–ª—ã. –û—á–∏—Å—Ç–∏—Ç–µ —Å–ø–∏—Å–æ–∫ –∏ –ø–æ–ø—Ä–æ–±—É–π—Ç–µ —Å–Ω–æ–≤–∞.")
                 return
            is_epub_to_epub_mode = True
            epub_groups_for_worker = {} 
            epub_paths_in_list = sorted(list(set(p1 for ft, p1, _ in selected_files_tuples if ft == 'epub')))
            valid_epubs_found = False
            failed_epub_structures = []
            for epub_path in epub_paths_in_list:
                 nav_path, ncx_path, opf_dir, nav_id, ncx_id = self._find_epub_toc_paths(epub_path)
                 if opf_dir is None: 
                      QMessageBox.warning(self, "–û—à–∏–±–∫–∞ EPUB", f"–ù–µ —É–¥–∞–ª–æ—Å—å –æ–±—Ä–∞–±–æ—Ç–∞—Ç—å —Å—Ç—Ä—É–∫—Ç—É—Ä—É EPUB:\n{Path(epub_path).name}\n\n–ü—Ä–æ–ø—É—Å–∫ —ç—Ç–æ–≥–æ —Ñ–∞–π–ª–∞.")
                      failed_epub_structures.append(epub_path)
                      continue 
                 html_paths_for_this_epub = [p2 for ft, p1, p2 in selected_files_tuples if ft == 'epub' and p1 == epub_path and p2]
                 html_to_translate_for_worker = [p for p in html_paths_for_this_epub if p != nav_path]
                 epub_groups_for_worker[epub_path] = {
                     'html_paths': html_to_translate_for_worker,
                     'build_metadata': {
                         'nav_path_in_zip': nav_path, 'ncx_path_in_zip': ncx_path,
                         'opf_dir': opf_dir, 'nav_item_id': nav_id, 'ncx_item_id': ncx_id
                     }
                 }
                 valid_epubs_found = True
            if failed_epub_structures:
                 self.selected_files_data_tuples = [t for t in self.selected_files_data_tuples if t[1] not in failed_epub_structures]
                 self.update_file_list_widget()
            if not valid_epubs_found:
                 QMessageBox.warning(self, "–ù–µ—Ç —Ñ–∞–π–ª–æ–≤", "–ù–µ –Ω–∞–π–¥–µ–Ω–æ –¥–æ–ø—É—Å—Ç–∏–º—ã—Ö EPUB —Ñ–∞–π–ª–æ–≤ –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏ –≤ —Ä–µ–∂–∏–º–µ EPUB->EPUB (–≤–æ–∑–º–æ–∂–Ω–æ, –æ—à–∏–±–∫–∏ —Å—Ç—Ä—É–∫—Ç—É—Ä—ã).")
                 self.clear_file_list(); return 
            worker_data = epub_groups_for_worker
            QMessageBox.information(self, "–†–µ–∂–∏–º EPUB->EPUB",
                                     "–ó–∞–ø—É—Å–∫ –≤ —Ä–µ–∂–∏–º–µ EPUB -> EPUB.\n–ë—É–¥–µ—Ç –≤—ã–ø–æ–ª–Ω–µ–Ω–æ:\n"
                                     "- –ü–µ—Ä–µ–≤–æ–¥ –≤—ã–±—Ä–∞–Ω–Ω—ã—Ö HTML (–∫—Ä–æ–º–µ —Ñ–∞–π–ª–∞ NAV).\n"
                                     "- –ü–µ—Ä–µ–∏–º–µ–Ω–æ–≤–∞–Ω–∏–µ –ø–µ—Ä–µ–≤–µ–¥–µ–Ω–Ω—ã—Ö —Ñ–∞–π–ª–æ–≤ (*_translated.html/xhtml).\n"
                                     "- –ü–æ–∏—Å–∫ –∏ –ò–ó–ú–ï–ù–ï–ù–ò–ï —Å—É—â–µ—Å—Ç–≤—É—é—â–µ–≥–æ —Ñ–∞–π–ª–∞ –æ–≥–ª–∞–≤–ª–µ–Ω–∏—è (NAV/NCX) –¥–ª—è –æ–±–Ω–æ–≤–ª–µ–Ω–∏—è —Å—Å—ã–ª–æ–∫.")
        else: 
            worker_data = selected_files_tuples

        self.log_output.clear();
        self.progress_bar.setRange(0, 100); self.progress_bar.setValue(0);
        self.progress_bar.setFormat("–ü–æ–¥–≥–æ—Ç–æ–≤–∫–∞...")
        self.status_label.setText("–ü–æ–¥–≥–æ—Ç–æ–≤–∫–∞...");
        self.append_log("="*40 + f"\n–ù–ê–ß–ê–õ–û –ü–ï–†–ï–í–û–î–ê")
        self.append_log(f"–†–µ–∂–∏–º: {'EPUB->EPUB Rebuild' if is_epub_to_epub_mode else '–°—Ç–∞–Ω–¥–∞—Ä—Ç–Ω—ã–π'}")
        self.append_log(f"–ú–æ–¥–µ–ª—å: {selected_model_name}"); self.append_log(f"–ü–∞—Ä–∞–ª–ª. –∑–∞–ø—Ä–æ—Å—ã: {max_concurrency}"); self.append_log(f"–§–æ—Ä–º–∞—Ç –≤—ã–≤–æ–¥–∞: .{output_format}")

        chunking_log_msg = f"–ß–∞–Ω–∫–∏–Ω–≥ GUI: {'–î–∞' if chunking_enabled_gui else '–ù–µ—Ç'} (–õ–∏–º–∏—Ç: {chunk_limit:,}, –û–∫–Ω–æ: {chunk_window:,}"
        if chunking_enabled_gui and chunk_delay > 0:
            chunking_log_msg += f", –ó–∞–¥–µ—Ä–∂–∫–∞: {chunk_delay:.1f} —Å–µ–∫.)"
        else:
            chunking_log_msg += ")"
        self.append_log(chunking_log_msg)

        if not CHUNK_HTML_SOURCE and chunking_enabled_gui: self.append_log("[INFO] –ß–∞–Ω–∫–∏–Ω–≥ HTML/EPUB –æ—Ç–∫–ª—é—á–µ–Ω.")
        self.append_log(f"–ü–∞–ø–∫–∞ –≤—ã–≤–æ–¥–∞: {self.out_folder}")
        self.append_log(f"–ü–æ–¥–¥–µ—Ä–∂–∫–∞: DOCX={'–î–ê' if DOCX_AVAILABLE else '–ù–ï–¢'}, BS4={'–î–ê' if BS4_AVAILABLE else '–ù–ï–¢'}, LXML={'–î–ê' if LXML_AVAILABLE else '–ù–ï–¢'}, EbookLib={'–î–ê' if EBOOKLIB_AVAILABLE else '–ù–ï–¢'}, Pillow={'–î–ê' if PILLOW_AVAILABLE else '–ù–ï–¢'}")
        self.append_log("="*40); self.set_controls_enabled(False)
        self.thread = QtCore.QThread()

        self.worker = Worker(
            self.api_key, self.out_folder, prompt_template, worker_data,
            MODELS[selected_model_name], max_concurrency, output_format,
            chunking_enabled_gui, chunk_limit, chunk_window,
            temperature,
            chunk_delay, # <-- –í–æ—Ç —ç—Ç–æ—Ç –∞—Ä–≥—É–º–µ–Ω—Ç –±—ã–ª –ø—Ä–æ–ø—É—â–µ–Ω
            proxy_string=proxy_string # <--- –ü–µ—Ä–µ–¥–∞–µ–º —Å—Ç—Ä–æ–∫—É –ø—Ä–æ–∫—Å–∏ –≤ Worker

        )
        self.worker.moveToThread(self.thread)
        self.worker_ref = self.worker
        self.thread_ref = self.thread
        self.worker.file_progress.connect(self.update_file_progress)
        self.worker.current_file_status.connect(self.handle_current_file_status)
        self.worker.chunk_progress.connect(self.handle_chunk_progress)
        self.worker.log_message.connect(self.handle_log_message)
        self.worker.finished.connect(self.on_translation_finished)
        self.worker.total_tasks_calculated.connect(self.update_progress_bar_range)
        self.thread.started.connect(self.worker.run)
        self.worker.finished.connect(self.thread.quit)
        self.worker.finished.connect(self.worker.deleteLater)
        self.thread.finished.connect(self.thread.deleteLater)
        self.thread.finished.connect(self.clear_worker_refs)
        # --- –õ–û–ì–ò–†–û–í–ê–ù–ò–ï –ü–†–û–ö–°–ò (–ø–æ—Å–ª–µ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏–∏ Worker, —á—Ç–æ–±—ã –æ–Ω —É–∂–µ –∏–º–µ–ª self.proxy_string) ---
        if self.worker.proxy_string: # –ü—Ä–æ–≤–µ—Ä—è–µ–º, —á—Ç–æ worker.proxy_string —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω
            self.append_log(f"–ü—Ä–æ–∫—Å–∏ –¥–ª—è Worker –Ω–∞—Å—Ç—Ä–æ–µ–Ω –Ω–∞: {self.worker.proxy_string}")
        else:
            self.append_log("–ü—Ä–æ–∫—Å–∏ –¥–ª—è Worker: –ù–µ –∏—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è")
        # --- –ö–û–ù–ï–¶ –õ–û–ì–ò–†–û–í–ê–ù–ò–Ø –ü–†–û–ö–°–ò ---
        self.thread.start()
        self.append_log("–†–∞–±–æ—á–∏–π –ø–æ—Ç–æ–∫ –∑–∞–ø—É—â–µ–Ω...")
        self.status_label.setText("–ó–∞–ø—É—Å–∫...")

    def cancel_translation(self):
        if self.worker_ref and self.thread_ref and self.thread_ref.isRunning():
            self.append_log("–û—Ç–ø—Ä–∞–≤–∫–∞ —Å–∏–≥–Ω–∞–ª–∞ –û–¢–ú–ï–ù–´...")
            self.status_label.setText("–û—Ç–º–µ–Ω–∞...")
            self.worker_ref.cancel()
            self.cancel_btn.setEnabled(False)
            self.finish_btn.setEnabled(False) # <--- –î–û–ë–ê–í–ò–¢–¨
            self.append_log("–û–∂–∏–¥–∞–Ω–∏–µ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è –ø–æ—Ç–æ–∫–∞...")
        else:
            self.append_log("[WARN] –ù–µ—Ç –∞–∫—Ç–∏–≤–Ω–æ–≥–æ –ø—Ä–æ—Ü–µ—Å—Å–∞ –¥–ª—è –æ—Ç–º–µ–Ω—ã.")

    @QtCore.pyqtSlot(int, int, list)
    def on_translation_finished(self, success_count, error_count, errors):
        worker_ref_exists = self.worker_ref is not None
        was_cancelled = worker_ref_exists and self.worker_ref.is_cancelled
        was_finishing = worker_ref_exists and hasattr(self.worker_ref, 'is_finishing') and self.worker_ref.is_finishing

        # –õ–æ–≥–∏—Ä—É–µ–º —Ñ–∏–Ω–∞–ª—å–Ω—ã–µ –∏—Ç–æ–≥–∏ –ø–µ—Ä–µ–¥ –ø–æ–∫–∞–∑–æ–º QMessageBox
        log_end_separator = "="*40
        self.append_log(f"\n{log_end_separator}")
        if was_cancelled:
            self.append_log("--- –ü–†–û–¶–ï–°–° –ë–´–õ –û–¢–ú–ï–ù–ï–ù –ü–û–õ–¨–ó–û–í–ê–¢–ï–õ–ï–ú ---")
        elif was_finishing:
            self.append_log("--- –ü–†–û–¶–ï–°–° –ë–´–õ –ó–ê–í–ï–†–®–ï–ù –ü–û –ö–û–ú–ê–ù–î–ï '–ó–ê–í–ï–†–®–ò–¢–¨' (—á–∞—Å—Ç–∏—á–Ω–æ) ---")
        # –î–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω—ã–µ –ª–æ–≥–∏ –æ–± –æ—à–∏–±–∫–∞—Ö Executor –∏–ª–∏ API —É–∂–µ –¥–æ–ª–∂–Ω—ã –±—ã—Ç—å –≤ Worker.run

        self.append_log(f"–ò–¢–û–ì: –£—Å–ø–µ—à–Ω–æ: {success_count}, –û—à–∏–±–æ–∫/–û—Ç–º–µ–Ω–µ–Ω–æ/–ü—Ä–æ–ø—É—â–µ–Ω–æ: {error_count}")
        if errors:
            self.append_log("–î–µ—Ç–∞–ª–∏ –æ—à–∏–±–æ–∫/–æ—Ç–º–µ–Ω/–ø—Ä–æ–ø—É—Å–∫–æ–≤:")
            max_errors_to_show = 30
            for i, e in enumerate(errors[:max_errors_to_show]):
                error_str = str(e)
                max_len = 350
                display_error = error_str[:max_len] + ('...' if len(error_str) > max_len else '')
                self.append_log(f"- {display_error}")
            if len(errors) > max_errors_to_show:
                self.append_log(f"- ... ({len(errors) - max_errors_to_show} –µ—â–µ)")
        self.append_log(log_end_separator)

        final_message = ""
        msg_type = QMessageBox.Icon.Information
        title = "–ó–∞–≤–µ—Ä—à–µ–Ω–æ"
        total_tasks = self.progress_bar.maximum() # –ü–æ–ª—É—á–∞–µ–º –æ–±—â–µ–µ –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ –∑–∞–¥–∞—á –∏–∑ –ø—Ä–æ–≥—Ä–µ—Å—Å-–±–∞—Ä–∞

        if was_cancelled:
            title = "–û—Ç–º–µ–Ω–µ–Ω–æ"
            msg_type = QMessageBox.Icon.Warning
            final_message = f"–ü—Ä–æ—Ü–µ—Å—Å –æ—Ç–º–µ–Ω–µ–Ω.\n\n–£—Å–ø–µ—à–Ω–æ –¥–æ –æ—Ç–º–µ–Ω—ã: {success_count}\n–û—à–∏–±–æ–∫/–ü—Ä–æ–ø—É—â–µ–Ω–æ: {error_count}"
            self.status_label.setText("–û—Ç–º–µ–Ω–µ–Ω–æ")
        elif was_finishing:
            title = "–ó–∞–≤–µ—Ä—à–µ–Ω–æ (—á–∞—Å—Ç–∏—á–Ω–æ)"
            msg_type = QMessageBox.Icon.Information
            final_message = f"–ü—Ä–æ—Ü–µ—Å—Å –∑–∞–≤–µ—Ä—à–µ–Ω –ø–æ –∫–æ–º–∞–Ω–¥–µ '–ó–∞–≤–µ—Ä—à–∏—Ç—å'.\n\n–û–±—Ä–∞–±–æ—Ç–∞–Ω–æ (–ø–æ–ª–Ω–æ—Å—Ç—å—é –∏–ª–∏ —á–∞—Å—Ç–∏—á–Ω–æ): {success_count}\n–û—à–∏–±–æ–∫/–ü—Ä–æ–ø—É—â–µ–Ω–æ –ø–æ –¥—Ä—É–≥–∏–º –ø—Ä–∏—á–∏–Ω–∞–º: {error_count}"
            self.status_label.setText("–ó–∞–≤–µ—Ä—à–µ–Ω–æ (—á–∞—Å—Ç–∏—á–Ω–æ)")
        elif error_count == 0 and success_count > 0:
            title = "–ì–æ—Ç–æ–≤–æ!"
            msg_type = QMessageBox.Icon.Information
            final_message = f"–ü–µ—Ä–µ–≤–æ–¥ {success_count} –∑–∞–¥–∞–Ω–∏–π —É—Å–ø–µ—à–Ω–æ –∑–∞–≤–µ—Ä—à–µ–Ω!"
            self.status_label.setText("–ì–æ—Ç–æ–≤–æ!")
        elif success_count > 0 and error_count > 0:
            title = "–ó–∞–≤–µ—Ä—à–µ–Ω–æ —Å –æ—à–∏–±–∫–∞–º–∏"
            msg_type = QMessageBox.Icon.Warning
            final_message = f"–ü–µ—Ä–µ–≤–æ–¥ –∑–∞–≤–µ—Ä—à–µ–Ω.\n\n–£—Å–ø–µ—à–Ω–æ: {success_count}\n–û—à–∏–±–æ–∫/–ü—Ä–æ–ø—É—â–µ–Ω–æ: {error_count}\n\n–°–º. –ª–æ–≥."
            self.status_label.setText("–ó–∞–≤–µ—Ä—à–µ–Ω–æ —Å –æ—à–∏–±–∫–∞–º–∏")
        elif success_count == 0 and error_count > 0:
            title = "–ó–∞–≤–µ—Ä—à–µ–Ω–æ —Å –æ—à–∏–±–∫–∞–º–∏"
            msg_type = QMessageBox.Icon.Critical
            final_message = f"–ù–µ —É–¥–∞–ª–æ—Å—å —É—Å–ø–µ—à–Ω–æ –ø–µ—Ä–µ–≤–µ—Å—Ç–∏ –Ω–∏ –æ–¥–Ω–æ–≥–æ –∑–∞–¥–∞–Ω–∏—è.\n–û—à–∏–±–æ–∫/–ü—Ä–æ–ø—É—â–µ–Ω–æ: {error_count}\n\n–°–º. –ª–æ–≥."
            self.status_label.setText("–ó–∞–≤–µ—Ä—à–µ–Ω–æ —Å –æ—à–∏–±–∫–∞–º–∏")
        elif success_count == 0 and error_count == 0 and total_tasks > 0:
            title = "–í–Ω–∏–º–∞–Ω–∏–µ"
            msg_type = QMessageBox.Icon.Warning
            final_message = f"–û–±—Ä–∞–±–æ—Ç–∫–∞ –∑–∞–≤–µ—Ä—à–µ–Ω–∞, –Ω–æ –Ω–µ—Ç —É—Å–ø–µ—à–Ω—ã—Ö –∑–∞–¥–∞–Ω–∏–π –∏–ª–∏ –æ—à–∏–±–æ–∫ (–≤–æ–∑–º–æ–∂–Ω–æ, –≤—Å–µ –ø—Ä–æ–ø—É—â–µ–Ω–æ –∏–ª–∏ –æ—Ç–º–µ–Ω–µ–Ω–æ –¥–æ –Ω–∞—á–∞–ª–∞?).\n–ü—Ä–æ–≤–µ—Ä—å—Ç–µ –ª–æ–≥."
            self.status_label.setText("–ó–∞–≤–µ—Ä—à–µ–Ω–æ (–ø—Ä–æ–≤–µ—Ä—å—Ç–µ –ª–æ–≥)")
        elif total_tasks == 0 : # –ï—Å–ª–∏ –∏–∑–Ω–∞—á–∞–ª—å–Ω–æ –Ω–µ –±—ã–ª–æ –∑–∞–¥–∞—á
            title = "–ù–µ—Ç –∑–∞–¥–∞—á"
            msg_type = QMessageBox.Icon.Information
            final_message = "–ù–µ—Ç —Ñ–∞–π–ª–æ–≤ –∏–ª–∏ –∑–∞–¥–∞—á –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏."
            self.status_label.setText("–ù–µ—Ç –∑–∞–¥–∞—á")
        else: # –û–±—â–∏–π —Å–ª—É—á–∞–π, –µ—Å–ª–∏ –Ω–∏ –æ–¥–Ω–æ –∏–∑ —É—Å–ª–æ–≤–∏–π –≤—ã—à–µ –Ω–µ —Å—Ä–∞–±–æ—Ç–∞–ª–æ
            final_message = "–û–±—Ä–∞–±–æ—Ç–∫–∞ –∑–∞–≤–µ—Ä—à–µ–Ω–∞."
            self.status_label.setText("–ó–∞–≤–µ—Ä—à–µ–Ω–æ")

        if self.isVisible(): # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º QMessageBox —Ç–æ–ª—å–∫–æ –µ—Å–ª–∏ –æ–∫–Ω–æ –≤–∏–¥–∏–º–æ
            QMessageBox(msg_type, title, final_message, QMessageBox.StandardButton.Ok, self).exec()
        else: # –ï—Å–ª–∏ –æ–∫–Ω–æ –Ω–µ –≤–∏–¥–∏–º–æ (–Ω–∞–ø—Ä–∏–º–µ—Ä, –∑–∞–∫—Ä—ã—Ç–æ –≤–æ –≤—Ä–µ–º—è –≤—ã–ø–æ–ª–Ω–µ–Ω–∏—è), –ø—Ä–æ—Å—Ç–æ –ª–æ–≥–∏—Ä—É–µ–º
            self.append_log(f"–î–∏–∞–ª–æ–≥ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è: {title} - {final_message}")

    @QtCore.pyqtSlot()
    def clear_worker_refs(self):

        self.append_log("–§–æ–Ω–æ–≤—ã–π –ø–æ—Ç–æ–∫ –∑–∞–≤–µ—Ä—à–µ–Ω. –û—á–∏—Å—Ç–∫–∞ —Å—Å—ã–ª–æ–∫...");
        self.worker = None
        self.thread = None
        self.worker_ref = None
        self.thread_ref = None
        self.set_controls_enabled(True)
        self.append_log("–ò–Ω—Ç–µ—Ä—Ñ–µ–π—Å —Ä–∞–∑–±–ª–æ–∫–∏—Ä–æ–≤–∞–Ω.")

    def set_controls_enabled(self, enabled):
        widgets_to_toggle = [
            self.file_select_btn, self.clear_list_btn, self.out_btn, self.format_combo,
            self.model_combo, self.concurrency_spin, self.temperature_spin,
            self.chunking_checkbox, self.proxy_url_edit, # <-- –î–æ–±–∞–≤–ª–µ–Ω–æ –ø–æ–ª–µ –ø—Ä–æ–∫—Å–∏

            self.chunk_delay_spin, # <-- –î–æ–±–∞–≤–ª–µ–Ω–æ

            self.prompt_edit,
            self.start_btn, self.check_api_key_btn
        ]
        for widget in widgets_to_toggle: widget.setEnabled(enabled)
        if enabled:
            self.toggle_chunking_details(self.chunking_checkbox.checkState().value) # This will also handle chunk_delay_spin
            for code, index in self.format_indices.items():
                 item = self.format_combo.model().item(index)
                 if item:
                    is_available = True; tooltip = f"–°–æ—Ö—Ä–∞–Ω–∏—Ç—å –∫–∞–∫ .{code}"
                    if code == 'docx' and not DOCX_AVAILABLE: is_available = False; tooltip = "–¢—Ä–µ–±—É–µ—Ç—Å—è: python-docx"
                    elif code == 'epub' and (not EBOOKLIB_AVAILABLE or not LXML_AVAILABLE or not BS4_AVAILABLE): is_available = False; tooltip = "–¢—Ä–µ–±—É–µ—Ç—Å—è: ebooklib, lxml, beautifulsoup4"
                    elif code == 'fb2' and not LXML_AVAILABLE: is_available = False; tooltip = "–¢—Ä–µ–±—É–µ—Ç—Å—è: lxml"
                    if code in ['docx', 'epub', 'fb2', 'html'] and not PILLOW_AVAILABLE:
                        if is_available: tooltip += "\n(–†–µ–∫–æ–º.: Pillow –¥–ª—è –∏–∑–æ–±—Ä.)"
                        else: tooltip += "; Pillow (—Ä–µ–∫–æ–º.)"
                    item.setEnabled(is_available); 
                    self.format_combo.setItemData(index, tooltip, Qt.ItemDataRole.ToolTipRole)
                    self.cancel_btn.setEnabled(False) # –£–±–µ–¥–∏—Ç—å—Å—è, —á—Ç–æ –∫–Ω–æ–ø–∫–∏ —É–ø—Ä–∞–≤–ª–µ–Ω–∏—è –ø—Ä–æ—Ü–µ—Å—Å–æ–º –≤—ã–∫–ª—é—á–µ–Ω—ã
                    self.finish_btn.setEnabled(False)
        else: 
            self.chunk_limit_spin.setEnabled(False)
            self.chunk_window_spin.setEnabled(False)
            self.chunk_delay_spin.setEnabled(False)
            self.cancel_btn.setEnabled(True) # –í–∫–ª—é—á–∏—Ç—å –∫–Ω–æ–ø–∫–∏ —É–ø—Ä–∞–≤–ª–µ–Ω–∏—è –ø—Ä–æ—Ü–µ—Å—Å–æ–º
            self.finish_btn.setEnabled(True)

    def closeEvent(self, event: QtGui.QCloseEvent):

        self.save_settings()
        if self.thread_ref and self.thread_ref.isRunning():
            reply = QMessageBox.question(self, "–ü—Ä–æ—Ü–µ—Å—Å –≤—ã–ø–æ–ª–Ω—è–µ—Ç—Å—è", "–ü–µ—Ä–µ–≤–æ–¥ –≤—Å–µ –µ—â–µ –≤—ã–ø–æ–ª–Ω—è–µ—Ç—Å—è.\n–ü—Ä–µ—Ä–≤–∞—Ç—å –∏ –≤—ã–π—Ç–∏?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.Yes: self.append_log("–í—ã—Ö–æ–¥ –≤–æ –≤—Ä–µ–º—è –≤—ã–ø–æ–ª–Ω–µ–Ω–∏—è, –æ—Ç–º–µ–Ω–∞..."); self.cancel_translation(); event.accept()
            else: event.ignore()
        else: event.accept()

def main():
    parser = argparse.ArgumentParser(description="Batch File Translator v2.12 (EPUB TOC Fixes)")
    parser.add_argument("--api_key", help="Google API Key (–∏–ª–∏ GOOGLE_API_KEY env var).")
    parser.add_argument("--auto-setup", action="store_true", help="–ó–∞–ø—É—Å–∫ –¥–∏–∞–ª–æ–≥–∞ –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–æ–π –Ω–∞—Å—Ç—Ä–æ–π–∫–∏ —Å —Ä–æ—Ç–∞—Ü–∏–µ–π –∫–ª—é—á–µ–π")
    args = parser.parse_args()
    api_key = args.api_key or os.environ.get("GOOGLE_API_KEY")
    
    app = QApplication.instance() or QApplication(sys.argv)
    missing_libs_msg = []
    install_pkgs = []
    
    if not DOCX_AVAILABLE:
        missing_libs_msg.append("'python-docx' (–¥–ª—è DOCX)")
        install_pkgs.append("python-docx")
    if not BS4_AVAILABLE:
        missing_libs_msg.append("'beautifulsoup4' (–¥–ª—è EPUB/HTML –≤—Ö–æ–¥–∞/–≤—ã—Ö–æ–¥–∞)")
        install_pkgs.append("beautifulsoup4")
    if not LXML_AVAILABLE:
        missing_libs_msg.append("'lxml' (–¥–ª—è FB2/EPUB –≤—ã—Ö–æ–¥–∞/–∞–Ω–∞–ª–∏–∑–∞)")
        install_pkgs.append("lxml")
    if not EBOOKLIB_AVAILABLE:
        missing_libs_msg.append("'ebooklib' (–¥–ª—è EPUB –≤—ã—Ö–æ–¥–∞)")
        install_pkgs.append("ebooklib")
    if not PILLOW_AVAILABLE:
        missing_libs_msg.append("'Pillow' (–¥–ª—è –∏–∑–æ–±—Ä.)")
        install_pkgs.append("Pillow")
        
    if missing_libs_msg:
        lib_list = "\n - ".join(missing_libs_msg)
        install_cmd = f"pip install {' '.join(install_pkgs)}"
        QMessageBox(
            QMessageBox.Icon.Warning,
            "–û—Ç—Å—É—Ç—Å—Ç–≤—É—é—Ç –±–∏–±–ª–∏–æ—Ç–µ–∫–∏",
            f"–ù–µ –Ω–∞–π–¥–µ–Ω—ã –±–∏–±–ª–∏–æ—Ç–µ–∫–∏:\n\n - {lib_list}\n\n–§—É–Ω–∫—Ü–∏–æ–Ω–∞–ª—å–Ω–æ—Å—Ç—å –æ–≥—Ä–∞–Ω–∏—á–µ–Ω–∞.\n\n–£—Å—Ç–∞–Ω–æ–≤–∏—Ç—å:\n{install_cmd}",
            QMessageBox.StandardButton.Ok
        ).exec()
    
    try:
        # –ï—Å–ª–∏ –∑–∞–ø—Ä–æ—à–µ–Ω —Ä–µ–∂–∏–º –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–æ–π –Ω–∞—Å—Ç—Ä–æ–π–∫–∏
        if args.auto_setup:
            setup_dialog = InitialSetupDialog()
            if setup_dialog.exec() == QDialog.DialogCode.Accepted:
                settings = setup_dialog.get_settings()
                # –ó–∞–ø—É—Å–∫–∞–µ–º –ø–µ—Ä–µ–≤–æ–¥ —Å –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏–º –ø–µ—Ä–µ–∑–∞–ø—É—Å–∫–æ–º
                run_translation_with_auto_restart(settings)
            return
        
        # –û–±—ã—á–Ω—ã–π —Ä–µ–∂–∏–º
        win = TranslatorApp(api_key=api_key)
        win.show()
        if not api_key:
            win.append_log("[WARN] API –∫–ª—é—á –Ω–µ –ø—Ä–µ–¥–æ—Å—Ç–∞–≤–ª–µ–Ω.")
        
    except Exception as e:
        error_message = f"–ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞ GUI:\n{type(e).__name__}: {e}\n\n{traceback.format_exc()}"
        print(error_message, file=sys.stderr)
        QMessageBox.critical(None, "–û—à–∏–±–∫–∞ –ó–∞–ø—É—Å–∫–∞ GUI", error_message)
        sys.exit(1)
        
    sys.exit(app.exec())

if __name__ == "__main__":
    def excepthook(exc_type, exc_value, exc_tb):
        tb_str = "".join(traceback.format_exception(exc_type, exc_value, exc_tb))
        error_message = f"–ù–µ–ø–µ—Ä–µ—Ö–≤–∞—á–µ–Ω–Ω–∞—è –æ—à–∏–±–∫–∞:\n\n{exc_type.__name__}: {exc_value}\n\n{tb_str}"
        print(f"–ö–†–ò–¢–ò–ß–ï–°–ö–ê–Ø –û–®–ò–ë–ö–ê:\n{error_message}", file=sys.stderr)
        try: app_instance = QApplication.instance() or QApplication(sys.argv); QMessageBox.critical(None, "–ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –û—à–∏–±–∫–∞", error_message)
        except Exception as mb_error: print(f"–ù–µ —É–¥–∞–ª–æ—Å—å –ø–æ–∫–∞–∑–∞—Ç—å MessageBox: {mb_error}", file=sys.stderr)
        sys.exit(1)
    sys.excepthook = excepthook
    try: main()
    except SystemExit: pass
    except Exception as e:
        error_message = f"–ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞ –∑–∞–ø—É—Å–∫–∞:\n{type(e).__name__}: {e}\n\n{traceback.format_exc()}"
        print(f"–ö–†–ò–¢–ò–ß–ï–°–ö–ê–Ø –û–®–ò–ë–ö–ê –ó–ê–ü–£–°–ö–ê:\n{error_message}", file=sys.stderr)
        try:
            app_instance = QApplication.instance()
            if not app_instance: app_instance = QApplication(sys.argv)
            QMessageBox.critical(None, "–û—à–∏–±–∫–∞ –ó–∞–ø—É—Å–∫–∞", error_message)
        except Exception as mb_error: print(f"–ù–µ —É–¥–∞–ª–æ—Å—å –ø–æ–∫–∞–∑–∞—Ç—å MessageBox: {mb_error}", file=sys.stderr)
        sys.exit(1)
