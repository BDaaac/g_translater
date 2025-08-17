#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Telegram бот для перевода файлов с использованием TransGemini.py
"""

import os
import sys
import re
import tempfile
import asyncio
import logging
import subprocess
import zipfile
import uuid
import shutil
import datetime
import json
import time
import threading
from pathlib import Path
from typing import Dict, Any, Optional, List, Union

# Импортируем исключение для обработки ошибок Telegram
from telegram.error import BadRequest

# Устанавливаем зависимости если не установлены
def ensure_package(package_name, import_name=None):
    import_name = import_name or package_name
    try:
        __import__(import_name.split('.')[0])
    except ImportError:
        print(f"Устанавливаю {package_name}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package_name])

# Устанавливаем необходимые пакеты
ensure_package("python-telegram-bot==20.7")
ensure_package("google-generativeai", "google.generativeai")
ensure_package("python-docx", "docx")
ensure_package("beautifulsoup4", "bs4")
ensure_package("lxml")
ensure_package("PyQt6", "PyQt6")
ensure_package("ebooklib")

# Импортируем библиотеки Telegram
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackQueryHandler, ContextTypes, filters
from telegram.constants import ParseMode

# Импортируем функции из TransGemini.py
from TransGemini import (
    MODELS,
    OUTPUT_FORMATS,
    Worker,
    write_to_epub,
    ApiKeyManager,
    RateLimitTracker,
    InitialSetupDialog,
    TranslationSessionManager,
    EpubCreator,
    TranslatedChaptersManagerDialog,
    ContextManager,
    DynamicGlossaryFilter,
    run_translation_with_auto_restart
)

# Импортируем Google API исключения для проверки ключа
try:
    import google.generativeai as genai
    from google.api_core import exceptions as google_exceptions
except ImportError:
    genai = None
    google_exceptions = None

# Настройка логирования
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# Поддерживаемые форматы файлов (соответствуют TransGemini.py)
SUPPORTED_FORMATS = {
    'txt': ['.txt'],
    'docx': ['.docx'], 
    'html': ['.html', '.htm'],
    'epub': ['.epub'],
    'xml': ['.xml'],
    'fb2': ['.fb2']
}

def get_possible_output_formats(input_format: str) -> list:
    """Возвращает возможные выходные форматы для данного входного формата"""
    # Используем OUTPUT_FORMATS из TransGemini.py
    available_formats = []
    for display_name, format_code in OUTPUT_FORMATS.items():
        # Включаем все доступные форматы из TransGemini
        available_formats.append((display_name, format_code))
    return available_formats

def process_text_block_for_chapter_html(text_block: str) -> str:
    """Обрабатывает блок текста для HTML, сохраняя структуру как в TransGemini"""
    from html import escape
    import re
    
    # Защищаем амперсанды
    text_block_escaped_amp = text_block.replace('&', '&amp;')
    
    # Защищаем существующие <br/> теги
    text_block_br_protected = re.sub(r'<br\s*/?>', '__TEMP_BR_TAG__', text_block_escaped_amp, flags=re.IGNORECASE)
    
    # Экранируем < и >
    text_block_lt_gt_escaped = text_block_br_protected.replace('<', '&lt;').replace('>', '&gt;')
    
    # Обрабатываем простое markdown форматирование
    temp_md_text = text_block_lt_gt_escaped
    temp_md_text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', temp_md_text, flags=re.DOTALL)
    temp_md_text = re.sub(r'(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)', r'<em>\1</em>', temp_md_text, flags=re.DOTALL)
    temp_md_text = re.sub(r'`(.*?)`', r'<code>\1</code>', temp_md_text, flags=re.DOTALL)
    
    # Возвращаем защищенные теги
    final_text = temp_md_text.replace('__TEMP_BR_TAG__', '<br/>')
    
    return final_text

def create_chapter_html(chapter_title: str, content: str, chapter_num: int) -> str:
    """Создает HTML контент для главы в стиле TransGemini с сохранением структуры"""
    from html import escape
    import re
    
    # Экранируем заголовок
    title_escaped = escape(chapter_title)
    
    # Очищаем контент от мусорных фраз AI
    content = clean_ai_response(content)
    
    # Разделяем контент на строки (сохраняем оригинальную структуру)
    lines = content.splitlines()
    
    html_body_content = ""
    paragraph_buffer = []
    current_list_type = None
    in_code_block = False
    code_block_lines = []
    
    def flush_paragraph_buffer():
        """Очищает буфер абзаца и добавляет его в HTML"""
        nonlocal html_body_content, paragraph_buffer
        if paragraph_buffer:
            # Соединяем строки абзаца через <br/>
            para_content = process_text_block_for_chapter_html('<br/>'.join(paragraph_buffer))
            if para_content.strip():
                html_body_content += f"    <p>{para_content}</p>\n"
            paragraph_buffer = []
    
    def close_current_list():
        """Закрывает текущий список"""
        nonlocal html_body_content, current_list_type
        if current_list_type:
            html_body_content += f"    </{current_list_type}>\n"
            current_list_type = None
    
    for line in lines:
        stripped_line = line.strip()
        is_code_fence = stripped_line == '```'
        
        # Обработка блоков кода
        if is_code_fence:
            if not in_code_block:
                flush_paragraph_buffer()
                close_current_list()
                in_code_block = True
                code_block_lines = []
            else:
                in_code_block = False
                escaped_code = escape("\n".join(code_block_lines))
                html_body_content += f"    <pre><code>{escaped_code}</code></pre>\n"
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            continue
        
        # Проверка на заголовки, списки и разделители
        heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped_line)
        hr_match = stripped_line == '---'
        ul_match = re.match(r'^[\*\-]\s+(.*)', stripped_line)
        ol_match = re.match(r'^\d+\.\s+(.*)', stripped_line)
        
        # Закрываем список если нужно
        if current_list_type and not ((current_list_type == 'ul' and ul_match) or (current_list_type == 'ol' and ol_match)):
            close_current_list()
        
        # Очищаем буфер перед специальными элементами
        if paragraph_buffer and (heading_match or hr_match or ul_match or ol_match):
            flush_paragraph_buffer()
        
        if heading_match:
            # Заголовок
            level = len(heading_match.group(1))
            heading_text = process_text_block_for_chapter_html(heading_match.group(2).strip())
            if heading_text:
                html_body_content += f"    <h{level}>{heading_text}</h{level}>\n"
        elif hr_match:
            # Горизонтальная линия
            html_body_content += "    <hr/>\n"
        elif ul_match:
            # Неупорядоченный список
            if current_list_type != 'ul':
                html_body_content += "    <ul>\n"
                current_list_type = 'ul'
            list_text = process_text_block_for_chapter_html(ul_match.group(1).strip())
            html_body_content += f"      <li>{list_text}</li>\n"
        elif ol_match:
            # Упорядоченный список
            if current_list_type != 'ol':
                html_body_content += "    <ol>\n"
                current_list_type = 'ol'
            list_text = process_text_block_for_chapter_html(ol_match.group(1).strip())
            html_body_content += f"      <li>{list_text}</li>\n"
        elif line.strip():
            # Обычная строка - добавляем в буфер абзаца
            paragraph_buffer.append(line)
        elif not stripped_line and paragraph_buffer:
            # Пустая строка - завершаем текущий абзац
            flush_paragraph_buffer()
    
    # Завершаем оставшиеся элементы
    close_current_list()
    flush_paragraph_buffer()
    
    if in_code_block:
        escaped_code = escape("\n".join(code_block_lines))
        html_body_content += f"    <pre><code>{escaped_code}</code></pre>\n"
    
    # Если контент пустой, создаем минимальный абзац
    if not html_body_content.strip():
        processed_content = process_text_block_for_chapter_html(content.strip())
        html_body_content = f"    <p>{processed_content}</p>\n"
    
    # Создаем полный HTML в стиле TransGemini
    html_content = f'''<?xml version='1.0' encoding='utf-8'?>
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
  <title>{title_escaped}</title>
  <link rel="stylesheet" type="text/css" href="../style/default.css"/>
  <meta http-equiv="Content-Type" content="text/html; charset=utf-8"/>
</head>
<body>
  <h1>{title_escaped}</h1>
{html_body_content.rstrip()}
</body>
</html>'''
    
    return html_content

def clean_ai_response(content: str) -> str:
    """Очищает переведенный контент от служебных фраз AI"""
    try:
        # Список фраз, которые нужно удалить из начала текста
        ai_garbage_patterns = [
            r'^(?:конечно[,!]?\s*)?вот\s+перевод[:\s]*',
            r'^вот\s+переведенный\s+текст[:\s]*',
            r'^перевод[:\s]*',
            r'^переведенный\s+текст[:\s]*',
            r'^переведено[:\s]*',
            r'^результат\s+перевода[:\s]*',
            r'^переведенная\s+версия[:\s]*',
            r'^вот\s+результат[:\s]*',
            r'^конечно[,!]?\s*',
            r'^да[,!]?\s*вот\s*',
            r'^хорошо[,!]?\s*',
            r'^отлично[,!]?\s*',
            r'^готово[,!]?\s*',
            r'^вот\s+он[:\s]*',
            r'^смотри[,!]?\s*',
            r'^держи[,!]?\s*',
            r'^пожалуйста[,!]?\s*',
            r'^\*\*перевод\*\*[:\s]*',
            r'^\*\*переведенный\s+текст\*\*[:\s]*',
            r'^here\s+is\s+the\s+translation[:\s]*',
            r'^translation[:\s]*',
            r'^translated\s+text[:\s]*',
            r'^of\s+course[,!]?\s*here\s*',
            r'^sure[,!]?\s*here\s*',
            r'^here\s+you\s+go[:\s]*',
            r'^вот\s+и\s+всё[,!]?\s*',
            r'^готово[,!]?\s*вот\s*'
        ]
        
        # Применяем каждый паттерн для очистки начала
        cleaned_content = content.strip()
        for pattern in ai_garbage_patterns:
            cleaned_content = re.sub(pattern, '', cleaned_content, flags=re.IGNORECASE | re.MULTILINE)
            cleaned_content = cleaned_content.strip()
        
        # Удаляем пустые строки в начале
        while cleaned_content.startswith('\n'):
            cleaned_content = cleaned_content[1:]
        
        # Удаляем markdown форматирование в начале, если есть
        cleaned_content = re.sub(r'^```[a-z]*\n?', '', cleaned_content, flags=re.MULTILINE)
        cleaned_content = re.sub(r'\n?```$', '', cleaned_content, flags=re.MULTILINE)
        
        # Удаляем повторяющиеся переносы строк
        cleaned_content = re.sub(r'\n{3,}', '\n\n', cleaned_content)
        
        # Удаляем лишние пробелы в начале строк
        cleaned_content = re.sub(r'^\s+', '', cleaned_content, flags=re.MULTILINE)
        
        if len(content) != len(cleaned_content):
            logger.info(f"🧹 Очищен контент: было {len(content)} символов, стало {len(cleaned_content)} символов")
        
        return cleaned_content.strip()
        
    except Exception as e:
        logger.error(f"❌ Ошибка при очистке AI ответа: {e}")
        return content

def smart_split_content(content: str, target_chapters: int) -> list:
    """Умно разделяет контент на главы, сохраняя структуру диалогов и абзацев"""
    try:
        logger.info(f"📝 Умное разделение контента ({len(content)} символов) на {target_chapters} глав")
        
        # Сначала пробуем разделить по явным маркерам глав
        chapter_patterns = [
            r'\n\s*(Глава|ГЛАВА|Chapter|CHAPTER)\s+\d+',
            r'\n\s*(Часть|ЧАСТЬ|Part|PART)\s+\d+',
            r'\n\s*\d+\.\s*[А-ЯA-Z]',
            r'\n\s*[IVX]+\.\s*[А-ЯA-Z]',
        ]
        
        for pattern in chapter_patterns:
            splits = re.split(pattern, content, flags=re.MULTILINE | re.IGNORECASE)
            if len(splits) > 1 and len(splits) <= target_chapters * 2:
                # Убираем пустые части и объединяем с заголовками
                chapters = []
                for i, part in enumerate(splits):
                    if part.strip():
                        chapters.append(part.strip())
                if len(chapters) >= 2:
                    logger.info(f"📖 Разделили контент по паттерну на {len(chapters)} частей")
                    return chapters
        
        # Если не получилось найти явные маркеры, разделяем по структурным границам
        # Сохраняем оригинальную структуру (переносы строк)
        
        # Ищем естественные границы (двойные переносы строк)
        sections = content.split('\n\n')
        
        if len(sections) <= target_chapters:
            # Слишком мало секций, возвращаем как есть
            return [content]
        
        # Группируем секции по размеру
        sections_per_chapter = max(1, len(sections) // target_chapters)
        chapters = []
        
        for i in range(0, len(sections), sections_per_chapter):
            chapter_sections = sections[i:i + sections_per_chapter]
            if chapter_sections:
                # Соединяем секции обратно двойными переносами для сохранения структуры
                chapter_content = '\n\n'.join(sec.strip() for sec in chapter_sections if sec.strip())
                if chapter_content:
                    chapters.append(chapter_content)
        
        # Если последняя глава получилась слишком короткой, объединяем с предыдущей
        if len(chapters) > 1 and len(chapters[-1]) < 200:  # Увеличили минимум
            chapters[-2] = chapters[-2] + '\n\n' + chapters[-1]
            chapters.pop()
        
        logger.info(f"📝 Разделили контент по структурным границам на {len(chapters)} глав")
        
        # Логируем структуру каждой главы
        for i, chapter in enumerate(chapters[:3]):  # Показываем первые 3
            lines_count = len(chapter.split('\n'))
            paragraphs_count = len([p for p in chapter.split('\n\n') if p.strip()])
            logger.info(f"  Глава {i+1}: {len(chapter)} символов, {lines_count} строк, {paragraphs_count} абзацев")
        
        return chapters
        
    except Exception as e:
        logger.error(f"❌ Ошибка при умном разделении контента: {e}")
        return [content]

def create_epub_from_original(original_epub_path: str, translated_content: str, output_path: str, title_override: str = None) -> bool:
    """
    Заглушка для создания EPUB - теперь TransGemini.py делает это сам
    """
    logger.info("create_epub_from_original: TransGemini.py теперь создает EPUB файлы напрямую")
    return False  # Не используется, так как TransGemini создает файлы сам


def create_epub_from_text(content: str, title: str, author: str, output_path: str, chapters_info: dict = None) -> bool:
    """
    Заглушка для создания EPUB - теперь TransGemini.py делает это сам
    """
    logger.info("create_epub_from_text: TransGemini.py теперь создает EPUB файлы напрямую")
    return False  # Не используется, так как TransGemini создает файлы сам

def extract_epub_metadata(epub_path: str) -> dict:
    """
    Извлекает необходимые метаданные из оригинального EPUB для функции write_to_epub
    """
    try:
        with zipfile.ZipFile(epub_path, 'r') as epub_zip:
            # Читаем container.xml для получения пути к OPF
            try:
                container_data = epub_zip.read('META-INF/container.xml')
                from xml.etree import ElementTree as ET
                container_root = ET.fromstring(container_data)
                
                # Находим путь к OPF файлу
                opf_path = None
                for rootfile in container_root.iter():
                    if rootfile.tag.endswith('rootfile'):
                        opf_path = rootfile.get('full-path')
                        break
                
                if not opf_path:
                    # Fallback - ищем .opf файлы
                    opf_files = [name for name in epub_zip.namelist() if name.endswith('.opf')]
                    opf_path = opf_files[0] if opf_files else None
                
                if not opf_path:
                    logger.warning("Не найден OPF файл в EPUB")
                    return {
                        'opf_dir': '',
                        'nav_path_in_zip': None,
                        'ncx_path_in_zip': None,
                        'nav_item_id': None,
                        'ncx_item_id': None,
                        'combined_image_map': {}
                    }
                
                # Определяем директорию OPF
                opf_dir = os.path.dirname(opf_path).replace('\\', '/')
                if opf_dir == '.':
                    opf_dir = ''
                
                # Читаем OPF файл для поиска NAV и NCX
                opf_data = epub_zip.read(opf_path)
                opf_root = ET.fromstring(opf_data)
                
                nav_path = None
                ncx_path = None
                nav_id = None
                ncx_id = None
                
                # Ищем элементы manifest для NAV и NCX
                for item in opf_root.iter():
                    if item.tag.endswith('item'):
                        href = item.get('href', '')
                        media_type = item.get('media-type', '')
                        properties = item.get('properties', '')
                        item_id = item.get('id', '')
                        
                        # NAV файл
                        if 'nav' in properties or 'nav' in href.lower():
                            nav_path = os.path.join(opf_dir, href).replace('\\', '/') if opf_dir else href
                            nav_id = item_id
                            
                        # NCX файл
                        elif media_type == 'application/x-dtbncx+xml' or href.endswith('.ncx'):
                            ncx_path = os.path.join(opf_dir, href).replace('\\', '/') if opf_dir else href
                            ncx_id = item_id
                
                logger.info(f"📋 Извлечены метаданные EPUB:")
                logger.info(f"   OPF dir: '{opf_dir}'")
                logger.info(f"   NAV path: '{nav_path}' (ID: {nav_id})")
                logger.info(f"   NCX path: '{ncx_path}' (ID: {ncx_id})")
                
                return {
                    'opf_dir': opf_dir,
                    'nav_path_in_zip': nav_path,
                    'ncx_path_in_zip': ncx_path,
                    'nav_item_id': nav_id,
                    'ncx_item_id': ncx_id,
                    'combined_image_map': {}
                }
                
            except Exception as e:
                logger.error(f"❌ Ошибка чтения container.xml: {e}")
                return {
                    'opf_dir': '',
                    'nav_path_in_zip': None,
                    'ncx_path_in_zip': None,
                    'nav_item_id': None,
                    'ncx_item_id': None,
                    'combined_image_map': {}
                }
                
    except Exception as e:
        logger.error(f"❌ Ошибка извлечения метаданных EPUB: {e}")
        return {
            'opf_dir': '',
            'nav_path_in_zip': None,
            'ncx_path_in_zip': None,
            'nav_item_id': None,
            'ncx_item_id': None,
            'combined_image_map': {}
        }


class UserState:
    def __init__(self):
        self.step = "waiting_file"  # waiting_file -> format_selection -> api_key -> chapter_selection -> translating
        self.action_type: str = "translate"  # "translate" или "glossary"
        self.file_path: Optional[str] = None
        self.file_name: Optional[str] = None
        self.file_format: Optional[str] = None
        self.output_format: Optional[str] = None
        self.api_key: Optional[str] = None
        self.api_keys: List[str] = []  # Список API ключей для ротации
        self.use_key_rotation: bool = False  # Использовать ротацию ключей
        self.target_language: str = "русский"
        self.model: str = list(MODELS.keys())[0] if MODELS else "Gemini 2.0 Flash"  # Используем первую доступную модель
        self.start_chapter: int = 1
        self.chapter_count: int = 0  # 0 = все главы
        self.total_chapters: int = 0  # Определяется при анализе файла
        self.chapters_info: Optional[Dict[str, Any]] = None  # Детальная информация о главах
        self.custom_prompt: Optional[str] = None  # Кастомный промпт для перевода
        self.temperature: float = 1.0  # Температура для генерации
        self.glossary_path: Optional[str] = None  # Путь к файлу глоссария
        self.glossary_data: Dict[str, Any] = {}  # Данные глоссария
        self.session_data: Dict[str, Any] = {}  # Данные сессии для восстановления
        self.proxy_string: Optional[str] = None  # Строка прокси
        
    def get_settings_dict(self) -> Dict[str, Any]:
        """Возвращает словарь с настройками для передачи в run_translation_with_auto_restart"""
        return {
            'api_keys': self.api_keys if self.use_key_rotation else [self.api_key] if self.api_key else [],
            'output_folder': str(Path(self.file_path).parent) if self.file_path else "",
            'prompt_template': self.custom_prompt,
            'input_files': [self.file_path] if self.file_path else [],
            'model_name': self.model,
            'output_format': self.output_format,
            'temperature': self.temperature,
            'glossary_data': self.glossary_data,
            'auto_start': True,
            'proxy_string': self.proxy_string
        }

# Состояния пользователя
USER_STATES = {}

def get_user_state(user_id: int) -> UserState:
    if user_id not in USER_STATES:
        USER_STATES[user_id] = UserState()
    return USER_STATES[user_id]

def reset_user_state(user_id: int):
    if user_id in USER_STATES:
        del USER_STATES[user_id]

async def handle_apikeys_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик команды /apikeys для управления множественными API ключами"""
    user_id = update.effective_user.id
    user_state = get_user_state(user_id)
    
    # Получаем текущие ключи
    current_keys = user_state.api_keys
    
    # Создаем сообщение
    message = "🔑 **Управление API ключами**\n\n"
    
    if current_keys:
        message += f"📋 У вас настроено {len(current_keys)} ключей:\n"
        for i, key in enumerate(current_keys, 1):
            # Показываем только первые и последние символы ключа для безопасности
            masked_key = key[:5] + "..." + key[-3:] if len(key) > 10 else "***"
            message += f"{i}. `{masked_key}`\n"
    else:
        message += "⚠️ У вас пока нет настроенных API ключей.\n"
    
    message += "\nДля управления ключами используйте команды:\n"
    message += "• `/addkey ВАШ_КЛЮЧ` - добавить новый ключ\n"
    message += "• `/removekey НОМЕР` - удалить ключ по номеру\n"
    message += "• `/clearkeys` - удалить все ключи\n"
    message += "• `/rotation on/off` - включить/выключить автоматическую ротацию ключей\n\n"
    message += f"🔄 Ротация ключей: **{'Включена' if user_state.use_key_rotation else 'Выключена'}**"
    
    await update.message.reply_text(message, parse_mode=ParseMode.MARKDOWN)

async def handle_addkey_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик команды /addkey для добавления API ключа"""
    user_id = update.effective_user.id
    user_state = get_user_state(user_id)
    
    # Удаляем сообщение с ключом для безопасности
    try:
        await update.message.delete()
    except Exception as e:
        logger.warning(f"Не удалось удалить сообщение с ключом: {e}")
    
    # Проверяем аргументы команды
    if not context.args or not context.args[0].strip():
        await update.message.reply_text("⚠️ Пожалуйста, укажите API ключ: `/addkey ВАШ_КЛЮЧ`", parse_mode=ParseMode.MARKDOWN)
        return
    
    new_key = context.args[0].strip()
    
    # Проверяем формат ключа
    if not re.match(r'^[A-Za-z0-9_-]+$', new_key):
        await update.message.reply_text("⚠️ API ключ имеет некорректный формат. Ключ должен содержать только буквы, цифры, дефисы и подчеркивания.")
        return
    
    # Добавляем ключ
    if new_key not in user_state.api_keys:
        user_state.api_keys.append(new_key)
        
        # Если это первый ключ, также устанавливаем его как основной
        if not user_state.api_key:
            user_state.api_key = new_key
        
        await update.message.reply_text(f"✅ API ключ добавлен. Всего ключей: {len(user_state.api_keys)}")
    else:
        await update.message.reply_text("ℹ️ Этот ключ уже добавлен в список.")

async def handle_removekey_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик команды /removekey для удаления API ключа по номеру"""
    user_id = update.effective_user.id
    user_state = get_user_state(user_id)
    
    # Проверяем аргументы команды
    if not context.args or not context.args[0].strip():
        await update.message.reply_text("⚠️ Пожалуйста, укажите номер ключа для удаления: `/removekey НОМЕР`", parse_mode=ParseMode.MARKDOWN)
        return
    
    try:
        key_index = int(context.args[0].strip()) - 1
        if key_index < 0 or key_index >= len(user_state.api_keys):
            await update.message.reply_text(f"⚠️ Некорректный номер ключа. Доступны номера от 1 до {len(user_state.api_keys)}.")
            return
        
        removed_key = user_state.api_keys.pop(key_index)
        
        # Если удаляем ключ, который был установлен как основной, обновляем основной ключ
        if user_state.api_key == removed_key:
            user_state.api_key = user_state.api_keys[0] if user_state.api_keys else None
        
        await update.message.reply_text(f"✅ API ключ #{key_index+1} удален. Осталось ключей: {len(user_state.api_keys)}")
        
    except ValueError:
        await update.message.reply_text("⚠️ Пожалуйста, укажите корректный номер ключа.")

async def handle_clearkeys_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик команды /clearkeys для удаления всех API ключей"""
    user_id = update.effective_user.id
    user_state = get_user_state(user_id)
    
    # Создаем клавиатуру для подтверждения
    keyboard = [
        [
            InlineKeyboardButton("Да, удалить все", callback_data="confirm_clear_keys"),
            InlineKeyboardButton("Отмена", callback_data="cancel_clear_keys")
        ]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.message.reply_text(
        f"⚠️ Вы уверены, что хотите удалить все {len(user_state.api_keys)} API ключей?",
        reply_markup=reply_markup
    )

async def handle_rotation_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик команды /rotation для включения/выключения ротации ключей"""
    user_id = update.effective_user.id
    user_state = get_user_state(user_id)
    
    # Проверяем аргументы команды
    if not context.args or context.args[0].strip().lower() not in ["on", "off"]:
        current_status = "включена" if user_state.use_key_rotation else "выключена"
        await update.message.reply_text(
            f"🔄 Текущий статус ротации ключей: **{current_status}**\n\n"
            "Для изменения укажите:\n"
            "• `/rotation on` - включить ротацию\n"
            "• `/rotation off` - выключить ротацию",
            parse_mode=ParseMode.MARKDOWN
        )
        return
    
    # Меняем статус ротации
    new_status = context.args[0].strip().lower() == "on"
    user_state.use_key_rotation = new_status
    
    # Проверяем наличие достаточного количества ключей для ротации
    if new_status and len(user_state.api_keys) < 2:
        await update.message.reply_text(
            "⚠️ Ротация ключей включена, но у вас меньше 2 ключей.\n"
            "Добавьте больше ключей с помощью `/addkey` для эффективной ротации.",
            parse_mode=ParseMode.MARKDOWN
        )
    else:
        status_text = "включена" if new_status else "выключена"
        await update.message.reply_text(f"✅ Ротация API ключей {status_text}.")

async def handle_keys_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик колбэков для управления ключами"""
    query = update.callback_query
    user_id = query.from_user.id
    user_state = get_user_state(user_id)
    
    await query.answer()  # Отвечаем на колбэк
    
    if query.data == "confirm_clear_keys":
        # Очищаем все ключи
        keys_count = len(user_state.api_keys)
        user_state.api_keys = []
        user_state.api_key = None
        await query.message.edit_text(f"🗑️ Все {keys_count} API ключей удалены.")
    elif query.data == "cancel_clear_keys":
        # Отменяем удаление
        await query.message.edit_text("❌ Удаление API ключей отменено.")

async def handle_settings_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик команды /settings для управления настройками"""
    user_id = update.effective_user.id
    user_state = get_user_state(user_id)
    
    # Создаем клавиатуру с настройками
    keyboard = [
        [InlineKeyboardButton("🔑 API ключи", callback_data="settings_apikeys")],
        [InlineKeyboardButton("🧠 Модель перевода", callback_data="settings_model")],
        [InlineKeyboardButton("🌡️ Температура", callback_data="settings_temperature")],
        [InlineKeyboardButton("📝 Промпт", callback_data="settings_prompt")],
        [InlineKeyboardButton("🔄 Ротация ключей", callback_data="settings_rotation")],
        [InlineKeyboardButton("🔍 Глоссарий", callback_data="settings_glossary")],
        [InlineKeyboardButton("🌐 Прокси", callback_data="settings_proxy")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    # Формируем текущие настройки
    model_name = user_state.model
    temperature = user_state.temperature
    rotation_status = "Включена" if user_state.use_key_rotation else "Выключена"
    api_keys_count = len(user_state.api_keys)
    has_custom_prompt = "Да" if user_state.custom_prompt else "Нет (используется стандартный)"
    has_glossary = "Да" if user_state.glossary_data else "Нет"
    has_proxy = "Настроен" if user_state.proxy_string else "Не используется"
    
    message = (
        "⚙️ **Настройки перевода**\n\n"
        f"🔑 API ключи: {api_keys_count} шт.\n"
        f"🧠 Модель: {model_name}\n"
        f"🌡️ Температура: {temperature}\n"
        f"🔄 Ротация ключей: {rotation_status}\n"
        f"📝 Кастомный промпт: {has_custom_prompt}\n"
        f"🔍 Глоссарий: {has_glossary}\n"
        f"🌐 Прокси: {has_proxy}\n\n"
        "Выберите настройку для изменения:"
    )
    
    await update.message.reply_text(message, reply_markup=reply_markup, parse_mode=ParseMode.MARKDOWN)

async def handle_settings_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик колбэков для настроек"""
    query = update.callback_query
    user_id = query.from_user.id
    user_state = get_user_state(user_id)
    
    await query.answer()  # Отвечаем на колбэк
    
    if query.data == "settings_apikeys":
        # Отображаем управление API ключами
        await handle_settings_apikeys(query, user_state)
    elif query.data == "settings_model":
        # Отображаем выбор модели
        await handle_settings_model(query, user_state)
    elif query.data == "settings_temperature":
        # Отображаем настройку температуры
        await handle_settings_temperature(query, user_state)
    elif query.data == "settings_prompt":
        # Отображаем настройку промпта
        await handle_settings_prompt(query, user_state)
    elif query.data == "settings_rotation":
        # Отображаем настройку ротации ключей
        await handle_settings_rotation(query, user_state)
    elif query.data == "settings_glossary":
        # Отображаем управление глоссарием
        await handle_settings_glossary(query, user_state)
    elif query.data == "settings_proxy":
        # Отображаем настройку прокси
        await handle_settings_proxy(query, user_state)
    elif query.data.startswith("set_model_"):
        # Обрабатываем выбор модели
        model_key = query.data[10:]
        if model_key in MODELS:
            user_state.model = model_key
            await query.message.edit_text(f"✅ Выбрана модель: {model_key}", reply_markup=None)
        else:
            await query.message.edit_text(f"⚠️ Неизвестная модель: {model_key}", reply_markup=None)
    elif query.data.startswith("set_temp_"):
        # Обрабатываем выбор температуры
        try:
            temp_value = float(query.data[9:])
            user_state.temperature = temp_value
            await query.message.edit_text(f"✅ Установлена температура: {temp_value}", reply_markup=None)
        except ValueError:
            await query.message.edit_text("⚠️ Некорректное значение температуры", reply_markup=None)
    elif query.data == "toggle_rotation":
        # Переключаем ротацию ключей
        user_state.use_key_rotation = not user_state.use_key_rotation
        status = "включена" if user_state.use_key_rotation else "выключена"
        
        if user_state.use_key_rotation and len(user_state.api_keys) < 2:
            await query.message.edit_text(
                f"⚠️ Ротация ключей {status}, но у вас меньше 2 ключей.\n"
                "Добавьте больше ключей с помощью `/addkey` для эффективной ротации.",
                parse_mode=ParseMode.MARKDOWN
            )
        else:
            await query.message.edit_text(f"✅ Ротация API ключей {status}.", reply_markup=None)
    elif query.data == "set_custom_prompt":
        # Запрашиваем новый промпт
        await query.message.edit_text(
            "📝 Пожалуйста, отправьте новый промпт-шаблон для перевода.\n\n"
            "Ваш промпт должен содержать `{text}` для указания места вставки переводимого текста.\n\n"
            "Для отмены отправьте /cancel.",
            parse_mode=ParseMode.MARKDOWN
        )
        user_state.step = "waiting_custom_prompt"
    elif query.data == "reset_prompt":
        # Сбрасываем промпт к стандартному
        user_state.custom_prompt = None
        await query.message.edit_text("✅ Промпт сброшен к стандартному.", reply_markup=None)
    elif query.data == "set_proxy":
        # Запрашиваем новый прокси
        await query.message.edit_text(
            "🌐 Пожалуйста, отправьте URL прокси-сервера.\n\n"
            "Формат: `http(s)://user:pass@host:port` или `socks5(h)://host:port`\n\n"
            "Для отключения прокси отправьте `none`.\n"
            "Для отмены отправьте /cancel.",
            parse_mode=ParseMode.MARKDOWN
        )
        user_state.step = "waiting_proxy"
    elif query.data == "reset_proxy":
        # Отключаем прокси
        user_state.proxy_string = None
        await query.message.edit_text("✅ Прокси отключен.", reply_markup=None)
    elif query.data == "upload_glossary":
        # Запрашиваем файл глоссария
        # Сохраняем текущий шаг, чтобы вернуться после загрузки
        if user_state.step != "waiting_glossary":
            user_state.session_data["previous_step"] = user_state.step
        
        user_state.step = "waiting_glossary"
        
        await query.message.edit_text(
            "📚 Пожалуйста, отправьте файл глоссария в формате JSON.\n\n"
            "Файл должен содержать словарь в формате:\n"
            "```\n{\n  \"term1\": \"перевод1\",\n  \"term2\": \"перевод2\"\n}\n```\n\n"
            "Для отмены отправьте /cancel.",
            parse_mode=ParseMode.MARKDOWN
        )
    elif query.data == "remove_glossary":
        # Очищаем данные глоссария
        terms_count = len(user_state.glossary_data) if user_state.glossary_data else 0
        user_state.glossary_data = {}
        
        success_message = f"✅ Глоссарий успешно удален ({terms_count} терминов)."
        
        # Проверяем, был ли предыдущий шаг выбором глав
        if user_state.session_data.get("previous_step") == "chapter_selection":
            # Создаем кнопку возврата к выбору глав
            keyboard = [
                [InlineKeyboardButton("⬅️ Вернуться к выбору глав", callback_data="back_to_chapter_selection")]
            ]
            reply_markup = InlineKeyboardMarkup(keyboard)
            await query.message.edit_text(success_message, reply_markup=reply_markup, parse_mode=ParseMode.MARKDOWN)
        else:
            # Стандартное поведение
            await query.message.edit_text(success_message, parse_mode=ParseMode.MARKDOWN)
        
async def handle_settings_apikeys(query, user_state):
    """Обработчик настроек API ключей"""
    # Получаем текущие ключи
    current_keys = user_state.api_keys
    
    # Создаем сообщение
    message = "🔑 **Управление API ключами**\n\n"
    
    if current_keys:
        message += f"📋 Настроено {len(current_keys)} ключей:\n"
        for i, key in enumerate(current_keys, 1):
            # Показываем только первые и последние символы ключа для безопасности
            masked_key = key[:5] + "..." + key[-3:] if len(key) > 10 else "***"
            message += f"{i}. `{masked_key}`\n"
    else:
        message += "⚠️ Нет настроенных API ключей.\n"
    
    message += "\nДля управления ключами используйте команды:\n"
    message += "• `/addkey ВАШ_КЛЮЧ` - добавить новый ключ\n"
    message += "• `/removekey НОМЕР` - удалить ключ по номеру\n"
    message += "• `/clearkeys` - удалить все ключи\n"
    message += "• `/rotation on/off` - включить/выключить ротацию"
    
    await query.message.edit_text(message, parse_mode=ParseMode.MARKDOWN)

async def handle_settings_model(query, user_state):
    """Обработчик настроек модели перевода"""
    # Создаем клавиатуру с доступными моделями
    keyboard = []
    for model_name in MODELS:
        keyboard.append([InlineKeyboardButton(model_name, callback_data=f"set_model_{model_name}")])
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.message.edit_text(
        f"🧠 **Выбор модели перевода**\n\n"
        f"Текущая модель: **{user_state.model}**\n\n"
        f"Выберите новую модель:",
        reply_markup=reply_markup,
        parse_mode=ParseMode.MARKDOWN
    )

async def handle_settings_temperature(query, user_state):
    """Обработчик настроек температуры"""
    # Создаем клавиатуру с вариантами температуры
    keyboard = [
        [
            InlineKeyboardButton("0.0 (детермин.)", callback_data="set_temp_0.0"),
            InlineKeyboardButton("0.5 (низкая)", callback_data="set_temp_0.5")
        ],
        [
            InlineKeyboardButton("0.7 (средняя)", callback_data="set_temp_0.7"),
            InlineKeyboardButton("1.0 (стандарт)", callback_data="set_temp_1.0")
        ],
        [
            InlineKeyboardButton("1.5 (творческая)", callback_data="set_temp_1.5"),
            InlineKeyboardButton("2.0 (максим.)", callback_data="set_temp_2.0")
        ]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await query.message.edit_text(
        f"🌡️ **Настройка температуры модели**\n\n"
        f"Текущая температура: **{user_state.temperature}**\n\n"
        f"Чем выше температура, тем более творческим будет перевод:\n"
        f"• 0.0 - максимально детерминированный\n"
        f"• 1.0 - стандартный (рекомендуется)\n"
        f"• 2.0 - максимально творческий\n\n"
        f"Выберите новое значение:",
        reply_markup=reply_markup,
        parse_mode=ParseMode.MARKDOWN
    )

async def handle_settings_prompt(query, user_state):
    """Обработчик настроек промпта"""
    # Создаем клавиатуру
    keyboard = [
        [InlineKeyboardButton("Установить новый промпт", callback_data="set_custom_prompt")],
        [InlineKeyboardButton("Использовать стандартный промпт", callback_data="reset_prompt")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    prompt_status = "Используется кастомный промпт" if user_state.custom_prompt else "Используется стандартный промпт"
    prompt_preview = ""
    if user_state.custom_prompt:
        # Показываем первые 200 символов промпта
        prompt_preview = "\n\n**Текущий промпт:**\n" + user_state.custom_prompt[:200]
        if len(user_state.custom_prompt) > 200:
            prompt_preview += "..."
    
    await query.message.edit_text(
        f"📝 **Настройка промпта перевода**\n\n"
        f"Статус: **{prompt_status}**{prompt_preview}\n\n"
        f"Выберите действие:",
        reply_markup=reply_markup,
        parse_mode=ParseMode.MARKDOWN
    )

async def handle_settings_rotation(query, user_state):
    """Обработчик настроек ротации ключей"""
    # Создаем клавиатуру
    keyboard = [
        [InlineKeyboardButton(
            "Выключить ротацию" if user_state.use_key_rotation else "Включить ротацию", 
            callback_data="toggle_rotation"
        )]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    status = "включена" if user_state.use_key_rotation else "выключена"
    keys_info = f"Настроено ключей: {len(user_state.api_keys)}"
    recommendation = ""
    
    if user_state.use_key_rotation and len(user_state.api_keys) < 2:
        recommendation = "\n\n⚠️ Для эффективной ротации рекомендуется добавить больше API ключей."
    
    await query.message.edit_text(
        f"🔄 **Настройка ротации API ключей**\n\n"
        f"Текущий статус: **Ротация {status}**\n"
        f"{keys_info}{recommendation}\n\n"
        f"При включенной ротации система будет автоматически переключаться между ключами "
        f"при превышении лимитов API или возникновении ошибок.\n\n"
        f"Выберите действие:",
        reply_markup=reply_markup,
        parse_mode=ParseMode.MARKDOWN
    )

async def handle_settings_glossary(query, user_state):
    """Обработчик настроек глоссария"""
    try:
        logger.info("Показ настроек глоссария")
        
        # Создаем клавиатуру
        keyboard = [
            [InlineKeyboardButton("Загрузить глоссарий (JSON)", callback_data="upload_glossary")]
        ]
        
        if user_state.glossary_data:
            keyboard.append([InlineKeyboardButton("Удалить текущий глоссарий", callback_data="remove_glossary")])
        
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        glossary_status = "Не настроен" if not user_state.glossary_data else f"Загружен ({len(user_state.glossary_data)} терминов)"
        
        await query.message.edit_text(
            f"🔍 **Управление глоссарием**\n\n"
            f"Текущий статус: **{glossary_status}**\n\n"
            f"Глоссарий позволяет задать единообразный перевод определенных терминов в документе.\n"
            f"Загрузите файл JSON с терминами в формате:\n"
            f"```\n{{\n  \"term1\": \"перевод1\",\n  \"term2\": \"перевод2\"\n}}\n```\n\n"
            f"Выберите действие:",
            reply_markup=reply_markup,
            parse_mode=ParseMode.MARKDOWN
        )
    except Exception as e:
        logger.error(f"Ошибка показа настроек глоссария: {e}", exc_info=True)
        try:
            await query.message.edit_text(
                "❌ Ошибка при загрузке настроек глоссария",
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("⬅️ Назад", callback_data="back_to_chapter_selection")]])
            )
        except Exception as e2:
            logger.error(f"Ошибка отправки сообщения об ошибке: {e2}")

async def handle_settings_proxy(query, user_state):
    """Обработчик настроек прокси"""
    # Создаем клавиатуру
    keyboard = [
        [InlineKeyboardButton("Настроить прокси", callback_data="set_proxy")]
    ]
    
    if user_state.proxy_string:
        keyboard.append([InlineKeyboardButton("Отключить прокси", callback_data="reset_proxy")])
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    proxy_status = f"Настроен: `{user_state.proxy_string}`" if user_state.proxy_string else "Не используется"
    
    await query.message.edit_text(
        f"🌐 **Настройка прокси**\n\n"
        f"Текущий статус: **{proxy_status}**\n\n"
        f"Поддерживаются HTTP(S) и SOCKS5 прокси.\n"
        f"Формат: `http(s)://user:pass@host:port` или `socks5(h)://host:port`\n\n"
        f"Выберите действие:",
        reply_markup=reply_markup,
        parse_mode=ParseMode.MARKDOWN
    )

async def handle_glossary_upload(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обрабатывает загрузку файла глоссария"""
    user_id = update.effective_user.id
    user_state = get_user_state(user_id)
    
    # Проверяем, что пользователь в правильном состоянии
    if user_state.step != "waiting_glossary":
        return
    
    # Проверяем, что получен файл
    if not update.message.document:
        await update.message.reply_text("⚠️ Пожалуйста, отправьте файл глоссария в формате JSON.")
        return
    
    document = update.message.document
    file_name = document.file_name
    
    # Проверяем расширение файла
    if not file_name.lower().endswith('.json'):
        await update.message.reply_text("⚠️ Пожалуйста, отправьте файл с расширением .json")
        return
    
    # Скачиваем файл
    file = await context.bot.get_file(document.file_id)
    file_path = f"temp_glossary_{user_id}.json"
    await file.download_to_drive(file_path)
    
    try:
        # Загружаем JSON
        with open(file_path, 'r', encoding='utf-8') as f:
            glossary_data = json.load(f)
        
        # Проверяем формат
        if not isinstance(glossary_data, dict):
            await update.message.reply_text("⚠️ Некорректный формат глоссария. Должен быть словарь.")
            return
        
        # Сохраняем глоссарий
        user_state.glossary_data = glossary_data
        user_state.step = "waiting_file"  # Возвращаем к начальному состоянию
        
        await update.message.reply_text(
            f"✅ Глоссарий успешно загружен!\n"
            f"📋 Добавлено {len(glossary_data)} терминов."
        )
        
    except json.JSONDecodeError:
        await update.message.reply_text("⚠️ Не удалось прочитать JSON файл. Проверьте формат.")
    except Exception as e:
        await update.message.reply_text(f"⚠️ Ошибка при загрузке глоссария: {e}")
    finally:
        # Удаляем временный файл
        if os.path.exists(file_path):
            os.remove(file_path)

def get_possible_output_formats_old(input_format: str) -> list:
    """Старая функция - заменена на версию с OUTPUT_FORMATS"""
    if input_format in ['txt', 'docx', 'html', 'xml']:
        return ['txt', 'docx', 'html']
    elif input_format == 'epub':
        return ['txt', 'docx', 'html', 'epub']
    else:
        return ['txt']

def determine_input_format(file_extension: str) -> str:
    """Определяет входной формат файла по расширению"""
    for fmt, extensions in SUPPORTED_FORMATS.items():
        if file_extension in extensions:
            return fmt
    return 'txt'  # fallback

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик команды /start"""
    user_id = update.effective_user.id
    reset_user_state(user_id)
    
    welcome_message = """
🤖 **Добро пожаловать в бот переводчика файлов!**

Я могу переводить различные форматы файлов используя Google Gemini AI.

**Поддерживаемые форматы:**
• TXT - текстовые файлы
• DOCX - документы Word
• HTML - веб-страницы
• EPUB - электронные книги
• XML - XML документы

**Как использовать:**
1. Отправьте файл для перевода
2. Выберите выходной формат
3. Введите API ключ Google Gemini
4. Получите переведенный файл

Отправьте файл чтобы начать! 📁
    """
    
    await update.message.reply_text(
        welcome_message,
        parse_mode=ParseMode.MARKDOWN
    )

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик загруженных файлов"""
    user_id = update.effective_user.id
    state = get_user_state(user_id)
    
    # Проверяем, не в режиме ли ожидания глоссария пользователь
    if state.step == "waiting_glossary":
        document = update.message.document
        file_name = document.file_name
        
        # Проверяем расширение файла
        if not file_name.lower().endswith('.json'):
            await update.message.reply_text("⚠️ Пожалуйста, отправьте файл с расширением .json")
            return
        
        # Скачиваем файл
        file = await context.bot.get_file(document.file_id)
        file_path = f"temp_glossary_{user_id}.json"
        await file.download_to_drive(file_path)
        
        try:
            # Загружаем JSON
            with open(file_path, 'r', encoding='utf-8') as f:
                glossary_data = json.load(f)
            
            # Проверяем формат
            if not isinstance(glossary_data, dict):
                await update.message.reply_text("⚠️ Некорректный формат глоссария. Должен быть словарь.")
                return
            
            # Сохраняем глоссарий
            state.glossary_data = glossary_data
            
            # Проверяем, был ли предыдущий шаг выбором глав
            previous_step = state.session_data.get("previous_step")
            
            if previous_step == "chapter_selection":
                # Если мы пришли из выбора глав, возвращаемся туда
                state.step = "chapter_selection"
                
                # Отображаем кнопку возврата к выбору глав
                keyboard = [
                    [InlineKeyboardButton("⬅️ Вернуться к выбору глав", callback_data="back_to_chapter_selection")]
                ]
            else:
                # Возвращаем к начальному состоянию
                state.step = "waiting_file"
                
                # Отображаем клавиатуру с настройками
                keyboard = [
                    [InlineKeyboardButton("⚙️ Настройки", callback_data="settings_main")]
                ]
            
            reply_markup = InlineKeyboardMarkup(keyboard)
            
            await update.message.reply_text(
                f"✅ Глоссарий успешно загружен!\n"
                f"📋 Добавлено {len(glossary_data)} терминов.",
                reply_markup=reply_markup
            )
            
        except json.JSONDecodeError:
            await update.message.reply_text("⚠️ Не удалось прочитать JSON файл. Проверьте формат.")
        except Exception as e:
            await update.message.reply_text(f"⚠️ Ошибка при загрузке глоссария: {e}")
        finally:
            # Удаляем временный файл
            if os.path.exists(file_path):
                os.remove(file_path)
        
        return  # Прекращаем дальнейшую обработку
    
    # Стандартная обработка файла для перевода
    if state.step != "waiting_file":
        await update.message.reply_text("Пожалуйста, завершите текущий процесс или используйте /start для начала заново.")
        return
    
    document = update.message.document
    if not document:
        await update.message.reply_text("Пожалуйста, отправьте файл.")
        return
    
    # Проверяем размер файла (максимум 20MB для Telegram)
    if document.file_size > 20 * 1024 * 1024:
        await update.message.reply_text("Файл слишком большой. Максимальный размер: 20MB")
        return
    
    # Определяем формат файла по расширению
    file_name = document.file_name
    file_extension = Path(file_name).suffix.lower()
    
    # Проверяем поддерживаемые форматы
    supported_extensions = {'.txt', '.docx', '.html', '.htm', '.epub', '.xml'}
    if file_extension not in supported_extensions:
        await update.message.reply_text(
            f"Неподдерживаемый формат файла: {file_extension}\n"
            f"Поддерживаемые форматы: {', '.join(supported_extensions)}"
        )
        return
    
    try:
        # Скачиваем файл
        file = await context.bot.get_file(document.file_id)
        
        # Создаем temporary файл
        temp_dir = tempfile.mkdtemp()
        file_path = os.path.join(temp_dir, file_name)
        
        await file.download_to_drive(file_path)
        
        # Проверяем, что файл действительно скачался
        if not os.path.exists(file_path):
            await update.message.reply_text("❌ Ошибка: Не удалось скачать файл")
            return
        
        # Определяем формат файла
        state.file_format = determine_input_format(file_extension)
        
        # Сохраняем информацию о файле
        state.file_path = file_path
        state.file_name = file_name
        state.step = "format_selection"
        
        # Показываем варианты выходного формата
        await show_format_selection(update, state)
        
    except Exception as e:
        logger.error(f"Ошибка при обработке файла: {e}")
        await update.message.reply_text("Произошла ошибка при обработке файла. Попробуйте еще раз.")

async def show_format_selection(update, state: UserState):
    """Показывает выбор между созданием глоссария и переводом"""
    keyboard = [
        [InlineKeyboardButton("📖 Перевести файл", callback_data="action_translate")],
        [InlineKeyboardButton("📚 Создать глоссарий", callback_data="action_glossary")]
    ]
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    message_text = (
        f"📁 Файл получен: `{state.file_name}`\n"
        f"📄 Формат: `{state.file_format.upper()}`\n\n"
        f"Выберите действие:"
    )
    
    try:
        if hasattr(update, 'edit_message_text'):
            # Это CallbackQuery
            await update.edit_message_text(
                message_text,
                reply_markup=reply_markup,
                parse_mode=ParseMode.MARKDOWN
            )
        else:
            # Это Update, отправляем новое сообщение
            await update.message.reply_text(
                message_text,
                reply_markup=reply_markup,
                parse_mode=ParseMode.MARKDOWN
            )
    except Exception as e:
        logger.error(f"Ошибка отправки сообщения: {e}")

async def handle_format_selection(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик выбора действия (перевод или глоссарий)"""
    query = update.callback_query
    user_id = query.from_user.id
    state = get_user_state(user_id)
    
    if state.step != "format_selection":
        await query.answer("Неверный шаг процесса")
        return
    
    callback_data = query.data
    
    if callback_data == "action_translate":
        # Переходим к выбору выходного формата для перевода
        await query.answer("Выбран перевод файла")
        await show_output_format_selection(query, state)
        
    elif callback_data == "action_glossary":
        # Переходим к созданию глоссария
        await query.answer("Выбрано создание глоссария")
        state.action_type = "glossary"
        state.step = "api_key"
        await show_api_key_request(query, state)
    else:
        await query.answer("Неверные данные")

async def show_api_key_request(update: Update, state: UserState):
    """Показывает запрос API ключа"""
    action_text = "создания глоссария" if getattr(state, 'action_type', '') == "glossary" else "перевода"
    
    try:
        if hasattr(update, 'edit_message_text'):
            # Это CallbackQuery
            await update.edit_message_text(
                f"🔑 **API ключ Google Gemini**\n\n"
                f"📁 Файл: `{state.file_name}`\n"
                f"🎯 Действие: {action_text}\n\n"
                f"Для работы с Google Gemini API необходим API ключ.\n\n"
                f"**Получить ключ:**\n"
                f"1. Перейдите на [Google AI Studio](https://aistudio.google.com/app/apikey)\n"
                f"2. Создайте новый API ключ\n"
                f"3. Скопируйте и отправьте его боту\n\n"
                f"🔒 Ваш ключ будет использован только для этого запроса",
                parse_mode=ParseMode.MARKDOWN
            )
        else:
            # Это Update, отправляем новое сообщение
            await update.message.reply_text(
                f"🔑 **API ключ Google Gemini**\n\n"
                f"📁 Файл: `{state.file_name}`\n"
                f"🎯 Действие: {action_text}\n\n"
                f"Для работы с Google Gemini API необходим API ключ.\n\n"
                f"**Получить ключ:**\n"
                f"1. Перейдите на [Google AI Studio](https://aistudio.google.com/app/apikey)\n"
                f"2. Создайте новый API ключ\n"
                f"3. Скопируйте и отправьте его боту\n\n"
                f"🔒 Ваш ключ будет использован только для этого запроса",
                parse_mode=ParseMode.MARKDOWN
            )
    except Exception as e:
        logger.error(f"Ошибка отправки запроса API ключа: {e}")

async def show_output_format_selection(update: Update, state: UserState):
    """Показывает выбор выходного формата для перевода"""
    # Получаем возможные выходные форматы для данного входного формата
    possible_formats = get_possible_output_formats(state.file_format)
    
    keyboard = []
    for display_name, format_code in possible_formats:
        keyboard.append([InlineKeyboardButton(display_name, callback_data=f"format_{format_code}")])
    
    # Добавляем кнопку "Назад"
    keyboard.append([InlineKeyboardButton("⬅️ Назад", callback_data="back_to_action_selection")])
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    try:
        await update.edit_message_text(
            f"📁 Файл: `{state.file_name}`\n"
            f"📄 Формат: `{state.file_format.upper()}`\n\n"
            f"Выберите выходной формат для перевода:",
            reply_markup=reply_markup,
            parse_mode=ParseMode.MARKDOWN
        )
    except Exception as e:
        logger.error(f"Ошибка обновления сообщения: {e}")

async def handle_output_format_selection(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик выбора выходного формата"""
    query = update.callback_query
    user_id = query.from_user.id
    state = get_user_state(user_id)
    
    callback_data = query.data
    
    if callback_data == "back_to_action_selection":
        await query.answer()
        state.step = "format_selection"
        await show_format_selection(query, state)
        return
    
    if not callback_data.startswith("format_"):
        await query.answer("Неверные данные")
        return
    
    selected_format = callback_data.replace("format_", "")
    state.output_format = selected_format
    state.action_type = "translate"
    state.step = "api_key"
    
    await query.answer()
    await show_api_key_request(query, state)

async def handle_text_input(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Универсальный обработчик текстовых сообщений"""
    user_id = update.effective_user.id
    state = get_user_state(user_id)
    
    if state.step == "api_key":
        await handle_api_key(update, context)
    elif state.step == "chapter_input":
        await handle_chapter_input(update, context)
    elif state.step == "waiting_custom_prompt":
        # Обработка пользовательского промпта
        prompt_text = update.message.text.strip()
        
        # Проверяем наличие плейсхолдера {text}
        if "{text}" not in prompt_text:
            await update.message.reply_text(
                "⚠️ В промпте должен быть плейсхолдер `{text}` для указания места вставки переводимого текста.\n"
                "Пожалуйста, отправьте промпт снова или /cancel для отмены.",
                parse_mode=ParseMode.MARKDOWN
            )
            return
        
        # Сохраняем промпт
        state.custom_prompt = prompt_text
        state.step = "waiting_file"  # Возвращаем к начальному состоянию
        
        # Для безопасности пытаемся удалить сообщение с промптом
        try:
            await update.message.delete()
        except Exception as e:
            logger.warning(f"Не удалось удалить сообщение с промптом: {e}")
        
        await update.message.reply_text("✅ Пользовательский промпт сохранен!")
    
    elif state.step == "waiting_proxy":
        # Обработка прокси
        proxy_text = update.message.text.strip().lower()
        
        if proxy_text == "none":
            state.proxy_string = None
            state.step = "waiting_file"
            await update.message.reply_text("✅ Прокси отключен!")
            return
            
        # Проверяем формат прокси
        if (not proxy_text.startswith(("http://", "https://", "socks4://", "socks5://", "socks5h://")) or 
            "://" not in proxy_text):
            await update.message.reply_text(
                "⚠️ Некорректный формат URL прокси.\n"
                "Используйте формат: `http(s)://user:pass@host:port` или `socks5(h)://host:port`\n"
                "Для отключения отправьте `none` или /cancel для отмены.",
                parse_mode=ParseMode.MARKDOWN
            )
            return
        
        # Сохраняем прокси
        state.proxy_string = proxy_text
        state.step = "waiting_file"
        
        # Для безопасности удаляем сообщение с прокси
        try:
            await update.message.delete()
        except Exception as e:
            logger.warning(f"Не удалось удалить сообщение с прокси: {e}")
        
        await update.message.reply_text(f"✅ Прокси настроен: {proxy_text.split('@')[-1]}")
    
    elif state.step == "waiting_glossary":
        # Если пользователь отправил текст вместо файла глоссария
        await update.message.reply_text(
            "⚠️ Пожалуйста, отправьте файл глоссария в формате JSON.\n"
            "Используйте /cancel для отмены."
        )
    
    else:
        # Неожиданное текстовое сообщение
        await update.message.reply_text(
            "🤔 Я не понимаю, что вы хотите сделать.\n"
            "Используйте /start для начала работы или /help для справки."
        )

async def validate_api_key(api_key: str) -> tuple[bool, str]:
    """Проверяет валидность API ключа через Google Gemini API"""
    try:
        if not genai or not google_exceptions:
            return False, "Google API библиотеки не установлены"
            
        # Настраиваем API
        genai.configure(api_key=api_key)
        
        # Пытаемся получить список моделей для проверки ключа
        models = genai.list_models()
        
        # Проверяем, есть ли доступные модели Gemini
        gemini_models = [m for m in models if m.name.startswith("models/")]
        
        if gemini_models:
            return True, "API ключ действителен."
        else:
            return False, "Ключ принят API, но не найдено доступных моделей Gemini."
            
    except google_exceptions.Unauthenticated as e:
        return False, f"Ошибка аутентификации (неверный ключ): {str(e)}"
    except Exception as e:
        return False, f"Ошибка проверки API ключа: {str(e)}"

async def handle_api_key(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик API ключа"""
    user_id = update.effective_user.id
    state = get_user_state(user_id)
    
    if state.step != "api_key":
        return
    
    api_key = update.message.text.strip()
    
    # Базовая проверка API ключа
    if len(api_key) < 30 or not api_key.startswith('AI'):
        await update.message.reply_text(
            "❌ Неверный формат API ключа.\n"
            "API ключ должен начинаться с 'AI' и быть достаточно длинным.\n"
            "Попробуйте еще раз."
        )
        return
    
    # Отправляем сообщение о проверке ключа
    checking_message = await update.message.reply_text(
        "🔍 **Проверяю API ключ...**\n\n"
        "⏳ Выполняю тестовый запрос к Google Gemini API...",
        parse_mode=ParseMode.MARKDOWN
    )
    
    # Проверяем валидность API ключа
    is_valid, validation_message = await validate_api_key(api_key)
    
    if is_valid:
        # Ключ действителен
        state.api_key = api_key
        
        # Проверяем, что за действие выбрано
        if getattr(state, 'action_type', '') == "glossary":
            # Для создания глоссария переходим к выбору модели
            state.step = "glossary_model_selection"
            
            await checking_message.edit_text(
                "✅ **API ключ действителен!**\n\n"
                "🔑 Ключ успешно проверен\n"
                "🤖 Выберите модель для создания глоссария...",
                parse_mode=ParseMode.MARKDOWN
            )
            
            # Показываем выбор модели для глоссария
            await show_glossary_model_selection(update, state)
        else:
            # Обычный перевод
            state.step = "chapter_selection"
            
            await checking_message.edit_text(
                "✅ **API ключ действителен!**\n\n"
                "🔑 Ключ успешно проверен\n"
                "📝 Анализирую файл для определения глав...",
                parse_mode=ParseMode.MARKDOWN
            )
            
            # Анализируем файл для определения количества глав
            await analyze_file_chapters(update, state)
        
    else:
        # Ключ недействителен
        await checking_message.edit_text(
            "❌ **API ключ недействителен**\n\n"
            f"🚫 {validation_message}\n\n"
            "**Как получить API ключ:**\n"
            "1. Перейдите на https://aistudio.google.com/\n"
            "2. Войдите в аккаунт Google\n"
            "3. Создайте новый API ключ\n"
            "4. Отправьте его мне\n\n"
            "🔐 Отправьте корректный API ключ:",
            parse_mode=ParseMode.MARKDOWN
        )

async def show_glossary_model_selection(update: Update, state: UserState):
    """Показывает выбор модели для создания глоссария"""
    keyboard = []
    
    # Добавляем модели для глоссария
    for model_name in MODELS.keys():
        # Создаем короткие названия для кнопок
        short_name = model_name.replace("Gemini ", "").replace("gemma", "Gemma")
        if len(short_name) > 25:  # Обрезаем слишком длинные названия
            short_name = short_name[:22] + "..."
        
        keyboard.append([InlineKeyboardButton(
            f"🤖 {short_name}", 
            callback_data=f"glossary_model_{model_name}"
        )])
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    try:
        await update.message.reply_text(
            f"🤖 **Выбор модели для глоссария**\n\n"
            f"📁 Файл: `{state.file_name}`\n"
            f"📚 Действие: Создание глоссария\n\n"
            f"**Доступные модели:**\n"
            f"• **Gemini 2.5** - Новейшие модели (рекомендуется)\n"
            f"• **Gemini 2.0** - Быстрые и эффективные\n"
            f"• **Gemini 1.5** - Проверенные временем\n\n"
            f"Выберите модель для создания глоссария:",
            reply_markup=reply_markup,
            parse_mode=ParseMode.MARKDOWN
        )
    except Exception as e:
        logger.error(f"Ошибка показа выбора модели для глоссария: {e}")

async def start_glossary_creation(update: Update, state: UserState):
    """Запускает процесс создания глоссария по логике Worker.py"""
    try:
        # Промпт из Launcher.py (точная копия)
        glossary_prompt = """Ты профессиональный лингвист-терминолог. Твоя задача - создать глоссарий терминов для последовательного перевода книги.

ИНСТРУКЦИИ:
1. Найди в тексте ВСЕ:
   - Имена персонажей (включая прозвища, титулы)
   - Названия мест, организаций, техник, артефактов
   - Специфические термины и понятия мира произведения
   - Устойчивые словосочетания и титулы

2. Для каждого термина предложи ОДИН лучший вариант перевода на русский язык.

3. Учитывай контекст и жанр произведения при переводе.

4. НЕ включай в глоссарий:
   - Обычные слова без специального значения
   - Термины, встречающиеся только 1 раз (если это не ключевое имя/название)

ФОРМАТ ВЫВОДА (строго JSON):
{{
  "термин_на_оригинале": "перевод_на_русский",
  "Son Goku": "Сон Гоку",
  "Kamehameha": "Камехамеха"
}}

ВАЖНО: 
- Выводи ТОЛЬКО JSON без дополнительного текста
- Сохраняй оригинальное написание (регистр букв)
- Для имён используй благозвучную транслитерацию
- Для терминов предпочитай осмысленный перевод транслитерации

Текст для анализа:
{text}"""
        
        # Извлекаем главы из файла (как в Worker.py)
        chapters = await extract_chapters_from_file(state.file_path, state.file_format)
        
        if not chapters:
            await update.message.reply_text(
                "❌ Не удалось извлечь главы из файла для создания глоссария.",
                parse_mode=ParseMode.MARKDOWN
            )
            return
        
        # Отправляем сообщение о начале создания
        progress_message = await update.message.reply_text(
            f"📚 Создание глоссария\n\n"
            f"📁 Файл: {state.file_name}\n"
            f"📄 Формат: {state.file_format.upper()}\n"
            f"📊 Найдено глав: {len(chapters)}\n\n"
            "🔄 Анализирую главы и создаю глоссарий терминов...\n"
            "⏳ Это может занять несколько минут..."
        )
        
        # Создаем глоссарий через API
        import google.generativeai as genai
        genai.configure(api_key=state.api_key)
        
        # Используем выбранную модель или модель по умолчанию
        model_name = getattr(state, 'glossary_model', "models/gemini-2.5-flash")
        model = genai.GenerativeModel(model_name)
        
        # Накопительный глоссарий (как в Worker.py)
        current_glossary = {}
        processed_chapters = 0
        
        # Обрабатываем каждую главу отдельно (логика Worker.py)
        for i, (chapter_name, chapter_text) in enumerate(chapters, 1):
            try:
                # Показываем прогресс
                await progress_message.edit_text(
                    f"📚 Создание глоссария\n\n"
                    f"📁 Файл: {state.file_name}\n"
                    f"📄 Формат: {state.file_format.upper()}\n"
                    f"📊 Обработано глав: {i-1}/{len(chapters)}\n"
                    f"📖 Текущая глава: {chapter_name}\n"
                    f"🔄 Найдено терминов: {len(current_glossary)}\n\n"
                    "⏳ Анализирую главу..."
                )
                
                # Ограничиваем текст главы (как в Worker.py)
                limited_text = chapter_text[:30000]
                prompt = glossary_prompt.format(text=limited_text)
                
                # Делаем API запрос (с retry логикой как в Worker.py)
                response = await asyncio.get_event_loop().run_in_executor(
                    None, 
                    lambda: generate_content_with_retry(model, prompt, chapter_name)
                )
                
                if response and response.text:
                    # Парсим ответ (используем логику Worker.py)
                    chapter_terms = parse_glossary_response(response)
                    
                    if chapter_terms:
                        # Накапливаем термины (логика Worker.py: if term not in current_glossary)
                        for term, definition in chapter_terms.items():
                            if term not in current_glossary:
                                current_glossary[term] = definition
                        
                        logger.info(f"Обработана глава {chapter_name}, добавлено терминов: {len(chapter_terms)}, всего: {len(current_glossary)}")
                
                processed_chapters += 1
                
                # Добавляем небольшую задержку между запросами
                await asyncio.sleep(1)
                        
            except Exception as e:
                logger.warning(f"Ошибка обработки главы {chapter_name}: {e}")
                continue
        
        # Финальный результат
        if current_glossary:
            # Сохраняем глоссарий в состоянии пользователя
            state.glossary_data = current_glossary
            
            # Создаем файл глоссария для скачивания
            import json
            glossary_json = json.dumps(current_glossary, ensure_ascii=False, indent=2)
            
            # Отправляем глоссарий как файл
            import io
            glossary_file = io.BytesIO(glossary_json.encode('utf-8'))
            glossary_file.name = f"glossary_{state.file_name.split('.')[0]}.json"
            
            # Экранируем специальные символы в имени файла для Markdown
            safe_filename = state.file_name.replace('_', '\\_').replace('*', '\\*').replace('[', '\\[').replace(']', '\\]').replace('(', '\\(').replace(')', '\\)').replace('~', '\\~').replace('`', '\\`').replace('>', '\\>').replace('#', '\\#').replace('+', '\\+').replace('-', '\\-').replace('=', '\\=').replace('|', '\\|').replace('{', '\\{').replace('}', '\\}').replace('.', '\\.')
            
            await progress_message.edit_text(
                "✅ Глоссарий создан успешно!\n\n"
                f"📁 Файл: {state.file_name}\n"
                f"📊 Обработано глав: {processed_chapters}/{len(chapters)}\n"
                f"📚 Найдено терминов: {len(current_glossary)}\n\n"
                "📥 Файл глоссария будет отправлен отдельным сообщением.\n"
                "💾 Глоссарий также сохранен в памяти бота для будущих переводов."
            )
            
            await update.message.reply_document(
                document=glossary_file,
                caption=f"📚 Глоссарий для файла: {state.file_name}\n📊 Количество терминов: {len(current_glossary)}"
            )
            
            # Возвращаем пользователя в начальное состояние
            state.step = "waiting_file"
        else:
            await progress_message.edit_text(
                f"❌ Ошибка создания глоссария\n\n"
                f"Не удалось извлечь термины из {processed_chapters} обработанных глав.\n"
                "Попробуйте еще раз с другим файлом."
            )
            
    except Exception as e:
        logger.error(f"Ошибка создания глоссария: {e}", exc_info=True)
        await update.message.reply_text(
            f"❌ Ошибка создания глоссария\n\n"
            f"Произошла ошибка: {str(e)}"
        )

def generate_content_with_retry(model, prompt, chapter_name):
    """Генерирует контент с retry логикой как в Worker.py"""
    import time
    import random
    from google.api_core.exceptions import ResourceExhausted, DeadlineExceeded
    
    max_retries = 5
    base_delay = 5
    for attempt in range(max_retries):
        try:
            response = model.generate_content(prompt, request_options={"timeout": 120})
            return response
        except ResourceExhausted as e:
            if attempt < max_retries - 1:
                delay = base_delay * (2 ** attempt) + random.uniform(0, 1)
                logger.warning(f"Rate limit hit for chapter {chapter_name}. Retrying in {delay:.2f} seconds... (Attempt {attempt + 1}/{max_retries})")
                time.sleep(delay)
            else:
                logger.error(f"API limit reached for chapter {chapter_name} after {max_retries} attempts.")
                raise e
        except DeadlineExceeded as e:
            logger.error(f"API timeout for chapter {chapter_name}: {str(e)}")
            return None

    return None

def parse_glossary_response(response) -> dict:
    """Парсит ответ AI и извлекает глоссарий (по логике Worker.py)"""
    import json
    
    try:
        if not response or not hasattr(response, 'text') or not response.text:
            return {}
            
        cleaned_text = response.text.strip()
        if not cleaned_text:
            logger.warning("Received empty response from API.")
            return {}

        if cleaned_text.startswith("```json"):
            cleaned_text = cleaned_text[7:]
        if cleaned_text.endswith("```"):
            cleaned_text = cleaned_text[:-3]
        
        return json.loads(cleaned_text)

    except json.JSONDecodeError:
        logger.error(f"Failed to decode JSON from API response: {response.text[:200]}")
        return {}
    except Exception as e:
        logger.error(f"An unexpected error occurred in parse_glossary_response: {str(e)}")
        return {}

async def handle_glossary_model_selection(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик выбора модели для глоссария"""
    query = update.callback_query
    user_id = query.from_user.id
    state = get_user_state(user_id)
    
    if state.step != "glossary_model_selection":
        await query.answer("Неверный шаг процесса")
        return
    
    callback_data = query.data
    
    if not callback_data.startswith("glossary_model_"):
        await query.answer("Неверные данные")
        return
    
    # Получаем выбранную модель
    selected_model = callback_data.replace("glossary_model_", "")
    if selected_model not in MODELS:
        await query.answer("Неизвестная модель")
        return
    
    # Сохраняем ID модели для API (как в Worker.py)
    state.glossary_model = MODELS[selected_model]["id"]
    state.step = "glossary_ready"
    
    await query.answer(f"Выбрана модель: {selected_model}")
    
    # Показываем кнопку "Начать создание глоссария"
    await show_glossary_start_options(query, state)

async def show_glossary_start_options(update: Update, state: UserState):
    """Показывает опции для начала создания глоссария"""
    keyboard = [
        [InlineKeyboardButton("📚 Начать создание глоссария", callback_data="start_glossary_creation")],
        [InlineKeyboardButton("🤖 Изменить модель", callback_data="change_glossary_model")],
        [InlineKeyboardButton("⬅️ Назад к выбору действия", callback_data="back_to_action_selection")]
    ]
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    try:
        await update.edit_message_text(
            f"🔧 **Настройки глоссария**\n\n"
            f"📁 Файл: `{state.file_name}`\n"
            f"📄 Формат: `{state.file_format.upper()}`\n"
            f"🤖 Модель: `{state.model}`\n\n"
            f"Готов к созданию глоссария.\n"
            f"Глоссарий будет содержать:\n"
            f"• Имена персонажей\n"
            f"• Названия мест и организаций\n"
            f"• Специфические термины\n"
            f"• Техники и артефакты\n\n"
            f"Выберите действие:",
            reply_markup=reply_markup,
            parse_mode=ParseMode.MARKDOWN
        )
    except Exception as e:
        logger.error(f"Ошибка показа опций глоссария: {e}")

async def handle_glossary_options(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик опций глоссария"""
    query = update.callback_query
    user_id = query.from_user.id
    state = get_user_state(user_id)
    
    callback_data = query.data
    
    if callback_data == "start_glossary_creation":
        await query.answer("Начинаю создание глоссария")
        state.step = "creating_glossary"
        await start_glossary_creation(query, state)
        
    elif callback_data == "change_glossary_model":
        await query.answer()
        state.step = "glossary_model_selection"
        await show_glossary_model_selection(query, state)
        
    elif callback_data == "back_to_action_selection":
        await query.answer()
        state.step = "format_selection"
        await show_format_selection(query, state)

async def extract_chapters_from_file(file_path: str, file_format: str) -> list:
    """Извлекает главы из файла (логика Worker.py)"""
    try:
        chapters = []
        
        if file_format == 'epub':
            # Для EPUB файлов - извлекаем главы как в Worker.py
            import zipfile
            from bs4 import BeautifulSoup
            import ebooklib
            from ebooklib import epub
            
            try:
                # Читаем EPUB как в Worker.py
                book = epub.read_epub(file_path)
                epub_chapters = [item for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT)]
                
                for chapter in epub_chapters:
                    chapter_name = chapter.get_name()
                    
                    # Извлекаем текст главы (как в Worker.py)
                    soup = BeautifulSoup(chapter.get_content(), "lxml")
                    chapter_text = soup.get_text(separator=" ", strip=True)
                    
                    # Фильтруем слишком короткие главы
                    if chapter_text and len(chapter_text) > 100:
                        chapters.append((chapter_name, chapter_text))
                        
                logger.info(f"Извлечено {len(chapters)} глав из EPUB файла")
                return chapters
                
            except Exception as e:
                logger.error(f"Ошибка чтения EPUB через ebooklib: {e}")
                # Fallback - пробуем через zipfile (старый метод)
                with zipfile.ZipFile(file_path, 'r') as epub_zip:
                    html_files = [
                        name for name in epub_zip.namelist()
                        if name.lower().endswith(('.html', '.xhtml', '.htm'))
                        and not name.startswith(('__MACOSX', 'META-INF/'))
                    ]
                    
                    for html_file in html_files:
                        try:
                            content = epub_zip.read(html_file).decode('utf-8', errors='ignore')
                            soup = BeautifulSoup(content, 'lxml')
                            text = soup.get_text(separator=' ', strip=True)
                            if text and len(text) > 100:
                                chapters.append((html_file, text))
                        except Exception as e:
                            logger.warning(f"Ошибка чтения файла {html_file}: {e}")
                            continue
                
                return chapters
                
        elif file_format == 'txt':
            # Для текстовых файлов - разбиваем на псевдо-главы
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                full_text = f.read()
            
            # Пробуем найти разделители глав
            chapter_patterns = [
                r'\n\s*Chapter\s+\d+',
                r'\n\s*CHAPTER\s+\d+', 
                r'\n\s*Глава\s+\d+',
                r'\n\s*ГЛАВА\s+\d+',
                r'\n\s*\d+\.\s*',
                r'\n\s*\*\*\*\s*\n',
                r'\n\s*---\s*\n'
            ]
            
            import re
            for pattern in chapter_patterns:
                splits = re.split(pattern, full_text)
                if len(splits) > 3:  # Если нашли разделение на главы
                    for i, chapter_text in enumerate(splits[1:], 1):  # Пропускаем первую часть до первой главы
                        if len(chapter_text.strip()) > 500:  # Минимальная длина главы
                            chapters.append((f"Chapter_{i}", chapter_text.strip()))
                    if chapters:
                        logger.info(f"Разбит на {len(chapters)} глав по паттерну: {pattern}")
                        return chapters
            
            # Если паттерны не сработали - разбиваем на куски по размеру
            chunk_size = 10000  # 10K символов на главу
            for i in range(0, len(full_text), chunk_size):
                chunk = full_text[i:i + chunk_size]
                if len(chunk.strip()) > 500:
                    chapters.append((f"Part_{i//chunk_size + 1}", chunk.strip()))
            
            return chapters
                
        elif file_format == 'docx':
            # Для DOCX файлов - пробуем разбить на главы по параграфам
            import docx
            doc = docx.Document(file_path)
            
            current_chapter = ""
            chapter_num = 1
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                
                # Проверяем, не начинается ли новая глава
                if any(keyword in text.lower() for keyword in ['chapter', 'глава']) and len(text) < 100:
                    # Сохраняем предыдущую главу
                    if current_chapter and len(current_chapter) > 500:
                        chapters.append((f"Chapter_{chapter_num}", current_chapter.strip()))
                        chapter_num += 1
                    current_chapter = ""
                else:
                    current_chapter += text + "\n"
            
            # Добавляем последнюю главу
            if current_chapter and len(current_chapter) > 500:
                chapters.append((f"Chapter_{chapter_num}", current_chapter.strip()))
            
            # Если глав мало - разбиваем на куски
            if len(chapters) < 3:
                full_text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
                chunk_size = 8000
                chapters = []
                for i in range(0, len(full_text), chunk_size):
                    chunk = full_text[i:i + chunk_size]
                    if len(chunk.strip()) > 500:
                        chapters.append((f"Part_{i//chunk_size + 1}", chunk.strip()))
            
            return chapters
            
        return []
        
    except Exception as e:
        logger.error(f"Ошибка извлечения глав из файла: {e}")
        return []

async def extract_chapters_for_glossary(file_path: str, file_format: str) -> str:
    """Извлекает текст из файла для создания глоссария"""
    try:
        if file_format == 'epub':
            import zipfile
            from bs4 import BeautifulSoup
            
            with zipfile.ZipFile(file_path, 'r') as epub_zip:
                # Получаем все HTML файлы
                html_files = [
                    name for name in epub_zip.namelist()
                    if name.lower().endswith(('.html', '.xhtml', '.htm'))
                    and not name.startswith(('__MACOSX', 'META-INF/'))
                ]
                
                all_text = []
                for html_file in html_files[:20]:  # Берем первые 20 файлов для анализа (было 10)
                    try:
                        content = epub_zip.read(html_file).decode('utf-8', errors='ignore')
                        soup = BeautifulSoup(content, 'lxml')
                        text = soup.get_text(separator=' ', strip=True)
                        if text and len(text) > 100:  # Фильтруем слишком короткие файлы
                            all_text.append(text)
                    except Exception as e:
                        logger.warning(f"Ошибка чтения файла {html_file}: {e}")
                        continue
                
                return ' '.join(all_text)
                
        elif file_format == 'txt':
            # Для текстовых файлов
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
                
        elif file_format == 'docx':
            # Для DOCX файлов
            import docx
            doc = docx.Document(file_path)
            return ' '.join([paragraph.text for paragraph in doc.paragraphs])
            
        return ""
        
    except Exception as e:
        logger.error(f"Ошибка извлечения текста из файла: {e}")
        return ""

async def analyze_file_chapters(update: Update, state: UserState):
    """Анализирует файл для определения количества глав"""
    try:
        chapters_found = 0
        
        # Сначала пробуем использовать точную логику TransGemini
        if state.file_format == 'epub':
            transgemini_info = await get_transgemini_chapters_info(state.file_path, state.file_format)
            if transgemini_info['total_content'] > 0:
                chapters_found = transgemini_info['total_content']
                # Сохраняем детальную информацию в состоянии для дальнейшего использования
                state.chapters_info = transgemini_info
                logger.info(f"Использована точная логика TransGemini: найдено {chapters_found} глав")
            else:
                # Fallback к старой логике
                chapters_found = await count_chapters_in_file(state.file_path, state.file_format)
        else:
            # Для других форматов используем стандартную логику
            chapters_found = await count_chapters_in_file(state.file_path, state.file_format)
        
        state.total_chapters = max(1, chapters_found)
        
        # Показываем опции выбора глав
        await show_chapter_selection(update, state)
        
    except Exception as e:
        logger.error(f"Ошибка при анализе глав: {e}")
        # В случае ошибки - сразу к настройкам перевода
        state.step = "translating"
        await show_translation_options(update, state)

async def get_transgemini_chapters_info(file_path: str, file_format: str) -> dict:
    """Получает информацию о главах используя точную логику TransGemini"""
    try:
        logger.info(f"Анализ файла: {file_path}, формат: {file_format}")
        
        if file_format == 'epub':
            with zipfile.ZipFile(file_path, 'r') as epub_zip:
                # Получаем все HTML файлы (как в TransGemini)
                html_files = sorted([
                    name for name in epub_zip.namelist()
                    if name.lower().endswith(('.html', '.xhtml', '.htm'))
                    and not name.startswith(('__MACOSX', 'META-INF/'))
                ])
                
                logger.info(f"Найдено HTML файлов: {len(html_files)}")
                
                # Пытаемся найти NAV файл (как в TransGemini)
                nav_path = None
                for html_file in html_files:
                    if 'nav' in Path(html_file).name.lower():
                        nav_path = html_file
                        break
                
                chapters_info = {
                    'all_files': [],
                    'content_files': [],
                    'skip_files': [],
                    'nav_file': nav_path,
                    'total_all': len(html_files),
                    'total_content': 0,
                    'original_path': file_path  # Добавляем путь к оригинальному EPUB
                }
                
                TRANSLATED_SUFFIX = '_translated'  # Константа из TransGemini
                
                for file_path_in_epub in html_files:
                    try:
                        file_info = epub_zip.getinfo(file_path_in_epub)
                        file_size = file_info.file_size
                    except:
                        file_size = 0
                    
                    is_nav = (nav_path and file_path_in_epub == nav_path)
                    is_translated = Path(file_path_in_epub).stem.endswith(TRANSLATED_SUFFIX)
                    
                    file_data = {
                        'path': file_path_in_epub,
                        'name': Path(file_path_in_epub).name,
                        'size': file_size,
                        'is_nav': is_nav,
                        'is_translated': is_translated,
                        'is_selected': False,
                        'category': 'unknown'
                    }
                    
                    if is_nav:
                        file_data['category'] = 'nav'
                        file_data['is_selected'] = False
                        chapters_info['skip_files'].append(file_data)
                    else:
                        # Применяем точную логику TransGemini для определения типа файла
                        item_text_lower = file_path_in_epub.lower()
                        path = Path(item_text_lower)
                        filename_lower = path.name
                        filename_base = path.stem.split('.')[0]  # Get stem before first dot
                        
                        skip_indicators = ['toc', 'nav', 'ncx', 'cover', 'title', 'index', 'copyright', 'about', 'meta', 'opf',
                                          'masthead', 'colophon', 'imprint', 'acknowledgments', 'dedication',
                                          'glossary', 'bibliography', 'notes', 'annotations', 'epigraph', 'halftitle',
                                          'frontmatter', 'backmatter', 'preface', 'introduction', 'appendix', 'biography',
                                          'isbn', 'legal', 'notice', 'otherbooks', 'prelims', 'team', 'promo', 'bonus']
                        content_indicators = ['chapter', 'part', 'section', 'content', 'text', 'page', 'body', 'main', 'article',
                                            'chp', 'chap', 'prt', 'sec', 'glava', 'prologue', 'epilogue']
                        
                        is_likely_skip = any(skip in filename_base for skip in skip_indicators)
                        parent_dir_lower = str(path.parent).lower()
                        is_likely_skip = is_likely_skip or any(skip in parent_dir_lower for skip in ['toc', 'nav', 'meta', 'frontmatter', 'backmatter', 'index', 'notes'])
                        is_likely_content = any(indicator in filename_base for indicator in content_indicators)
                        is_chapter_like = (re.fullmatch(r'(ch|gl|chap|chapter|part|section|sec|glava)[\d_-]+.*', filename_base) or 
                                          re.fullmatch(r'[\d]+', filename_base) or 
                                          re.match(r'^[ivxlcdm]+$', filename_base))
                        
                        if not is_likely_skip and (is_likely_content or is_chapter_like):
                            file_data['category'] = 'content'
                            file_data['is_selected'] = True
                            chapters_info['content_files'].append(file_data)
                        else:
                            if not is_likely_skip and 'text' in filename_base:
                                file_data['category'] = 'text'
                                file_data['is_selected'] = True
                                chapters_info['content_files'].append(file_data)
                            else:
                                file_data['category'] = 'skip'
                                file_data['is_selected'] = False
                                chapters_info['skip_files'].append(file_data)
                    
                    chapters_info['all_files'].append(file_data)
                
                chapters_info['total_content'] = len(chapters_info['content_files'])
                logger.info(f"TransGemini анализ EPUB: {chapters_info['total_content']} глав из {chapters_info['total_all']} файлов")
                return chapters_info
                
        return {'total_all': 0, 'total_content': 0, 'all_files': [], 'content_files': [], 'skip_files': [], 'nav_file': None}
        
    except Exception as e:
        logger.error(f"Ошибка TransGemini анализа: {e}")
        return {'total_all': 0, 'total_content': 0, 'all_files': [], 'content_files': [], 'skip_files': [], 'nav_file': None}

async def get_chapters_info(file_path: str, file_format: str) -> dict:
    """Получает детальную информацию о главах в файле"""
    try:
        if file_format == 'epub':
            with zipfile.ZipFile(file_path, 'r') as epub_zip:
                # Получаем все HTML файлы
                all_html_files = sorted([
                    name for name in epub_zip.namelist()
                    if name.lower().endswith(('.html', '.xhtml', '.htm'))
                    and not name.startswith(('__MACOSX', 'META-INF/'))
                ])
                
                # Анализируем каждый файл
                chapters_info = {
                    'all_files': [],
                    'content_files': [],
                    'skip_files': [],
                    'total_all': len(all_html_files),
                    'total_content': 0,
                    'original_path': file_path  # Добавляем путь к оригинальному EPUB
                }
                
                skip_indicators = [
                    'toc', 'nav', 'ncx', 'cover', 'title', 'index', 'copyright', 'about', 'meta', 'opf',
                    'masthead', 'colophon', 'imprint', 'acknowledgments', 'dedication',
                    'glossary', 'bibliography', 'notes', 'annotations', 'epigraph', 'halftitle',
                    'frontmatter', 'backmatter', 'preface', 'introduction', 'appendix', 'biography',
                    'isbn', 'legal', 'notice', 'otherbooks', 'prelims', 'team', 'promo', 'bonus'
                ]
                
                for html_file in all_html_files:
                    filename_base = Path(html_file).stem.split('.')[0].lower()
                    
                    try:
                        file_info = epub_zip.getinfo(html_file)
                        file_size = file_info.file_size
                    except:
                        file_size = 0
                    
                    # Проверяем, является ли файл служебным
                    is_skip_file = any(skip_word in filename_base for skip_word in skip_indicators)
                    is_translated = filename_base.endswith('_translated')
                    
                    file_data = {
                        'path': html_file,
                        'name': Path(html_file).name,
                        'size': file_size,
                        'is_skip': is_skip_file,
                        'is_translated': is_translated
                    }
                    
                    chapters_info['all_files'].append(file_data)
                    
                    if not is_skip_file and not is_translated and file_size > 500:
                        chapters_info['content_files'].append(file_data)
                    else:
                        chapters_info['skip_files'].append(file_data)
                
                chapters_info['total_content'] = len(chapters_info['content_files'])
                logger.info(f"Анализ завершен: {chapters_info['total_content']} глав для перевода из {chapters_info['total_all']} файлов")
                return chapters_info
                
        logger.warning(f"Неподдерживаемый формат файла: {file_format}")
        return {'total_all': 0, 'total_content': 0, 'all_files': [], 'content_files': [], 'skip_files': []}
        
    except Exception as e:
        logger.error(f"Ошибка анализа глав: {e}", exc_info=True)
        return {'total_all': 0, 'total_content': 0, 'all_files': [], 'content_files': [], 'skip_files': []}

async def count_chapters_in_file(file_path: str, file_format: str) -> int:
    """Подсчитывает количество глав в файле"""
    try:
        if file_format == 'txt':
            # Пробуем разные кодировки
            content = ""
            encodings = ['utf-8', 'cp1251', 'latin-1', 'cp866']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                # Если ничего не помогло, читаем в бинарном режиме
                with open(file_path, 'rb') as f:
                    raw_content = f.read()
                    content = raw_content.decode('utf-8', errors='ignore')
            
            # Ищем заголовки глав
            import re
            patterns = [
                r'^\s*(Глава|Chapter|ГЛАВА|CHAPTER)\s+\d+',
                r'^\s*(Часть|Part|ЧАСТЬ|PART)\s+\d+',
                r'^\s*\d+\.\s*[А-ЯA-Z]',
                r'^#{1,3}\s+',  # Markdown заголовки
            ]
            
            total_matches = 0
            for pattern in patterns:
                matches = re.findall(pattern, content, re.MULTILINE | re.IGNORECASE)
                total_matches = max(total_matches, len(matches))
            
            return max(1, total_matches)
                
        elif file_format == 'docx':
            # Для DOCX используем python-docx
            try:
                from docx import Document
                doc = Document(file_path)
                chapter_count = 0
                
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text and (
                        text.lower().startswith(('глава', 'chapter', 'часть', 'part')) or
                        para.style.name.startswith('Heading')
                    ):
                        chapter_count += 1
                
                return max(1, chapter_count)
            except ImportError:
                return 5  # Fallback если docx не установлен
                
        elif file_format == 'html':
            # Пробуем разные кодировки для HTML
            content = ""
            encodings = ['utf-8', 'cp1251', 'latin-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                with open(file_path, 'rb') as f:
                    raw_content = f.read()
                    content = raw_content.decode('utf-8', errors='ignore')
            
            import re
            # Ищем HTML заголовки
            headers = re.findall(r'<h[1-6][^>]*>(.*?)</h[1-6]>', content, re.IGNORECASE | re.DOTALL)
            return max(1, len(headers))
        
        elif file_format == 'epub':
            # Используем точную логику TransGemini.py для EPUB файлов
            try:
                chapter_count = 0
                with zipfile.ZipFile(file_path, 'r') as epub_zip:
                    # Получаем HTML файлы так же, как в TransGemini
                    html_files_in_epub = sorted([
                        name for name in epub_zip.namelist()
                        if name.lower().endswith(('.html', '.xhtml', '.htm'))
                        and not name.startswith(('__MACOSX', 'META-INF/'))  # Исключаем служебные папки
                    ])
                    
                    if not html_files_in_epub:
                        logger.warning(f"В EPUB файле не найдено HTML/XHTML файлов")
                        return 5
                    
                    logger.info(f"Всего HTML файлов в EPUB: {len(html_files_in_epub)}")
                    for html_file in html_files_in_epub:
                        logger.debug(f"HTML файл: {html_file}")
                    
                    # Фильтруем главы, исключая служебные файлы (как в TransGemini)
                    content_files = []
                    for file_path_in_epub in html_files_in_epub:
                        filename_lower = Path(file_path_in_epub).name.lower()
                        filename_base = Path(file_path_in_epub).stem.split('.')[0].lower()
                        
                        # Список служебных файлов (как в TransGemini)
                        skip_indicators = [
                            'toc', 'nav', 'ncx', 'cover', 'title', 'index', 'copyright', 'about', 'meta', 'opf',
                            'masthead', 'colophon', 'imprint', 'acknowledgments', 'dedication',
                            'glossary', 'bibliography', 'notes', 'annotations', 'epigraph', 'halftitle',
                            'frontmatter', 'backmatter', 'preface', 'introduction', 'appendix', 'biography',
                            'isbn', 'legal', 'notice', 'otherbooks', 'prelims', 'team', 'promo', 'bonus'
                        ]
                        
                        # Список индикаторов контента (главы)
                        content_indicators = [
                            'chapter', 'part', 'section', 'content', 'text', 'page', 'body', 'main', 'article',
                            'chp', 'chap', 'prt', 'sec', 'glava', 'prologue', 'epilogue'
                        ]
                        
                        # Проверяем, является ли файл служебным
                        is_skip_file = any(skip_word in filename_base for skip_word in skip_indicators)
                        is_content_file = any(content_word in filename_base for content_word in content_indicators)
                        
                        # Также проверяем файлы с суффиксом _translated
                        is_translated = filename_base.endswith('_translated')
                        
                        # Дополнительная проверка: исключаем очень короткие HTML файлы (менее 1KB)
                        try:
                            file_info = epub_zip.getinfo(file_path_in_epub)
                            file_size = file_info.file_size
                        except:
                            file_size = 0
                        
                        # Если файл не служебный и не переведенный, и имеет достаточный размер
                        if not is_skip_file and not is_translated and file_size > 1000:
                            content_files.append(file_path_in_epub)
                            logger.debug(f"Найдена глава: {file_path_in_epub} (размер: {file_size} байт)")
                        else:
                            logger.debug(f"Пропущен файл: {file_path_in_epub} (служебный: {is_skip_file}, переведенный: {is_translated}, размер: {file_size})")
                    
                    # Если после фильтрации осталось мало файлов, используем менее строгую фильтрацию
                    if len(content_files) < 3:
                        logger.info("Мало глав после строгой фильтрации, применяем более мягкие критерии")
                        content_files = []
                        for file_path_in_epub in html_files_in_epub:
                            filename_lower = Path(file_path_in_epub).name.lower()
                            filename_base = Path(file_path_in_epub).stem.split('.')[0].lower()
                            
                            # Более короткий список служебных файлов
                            skip_indicators_short = ['toc', 'nav', 'cover', 'title', 'copyright', 'meta', 'opf']
                            is_skip_file = any(skip_word in filename_base for skip_word in skip_indicators_short)
                            is_translated = filename_base.endswith('_translated')
                            
                            try:
                                file_info = epub_zip.getinfo(file_path_in_epub)
                                file_size = file_info.file_size
                            except:
                                file_size = 0
                            
                            if not is_skip_file and not is_translated and file_size > 500:
                                content_files.append(file_path_in_epub)
                    
                    chapter_count = len(content_files)
                    logger.info(f"EPUB анализ: найдено {chapter_count} глав из {len(html_files_in_epub)} HTML файлов")
                    
                    return max(1, chapter_count)
                    
            except Exception as e:
                logger.error(f"Ошибка анализа EPUB: {e}")
                # Fallback: пробуем прочитать как обычный файл
                try:
                    # Некоторые EPUB читаются как текст
                    encodings = ['utf-8', 'cp1251', 'latin-1']
                    for encoding in encodings:
                        try:
                            with open(file_path, 'r', encoding=encoding) as f:
                                content = f.read()
                                patterns = [
                                    r'(?:chapter|глава|часть)\s*\d+',
                                    r'<h[1-6][^>]*>.*?</h[1-6]>',
                                ]
                                total_matches = 0
                                for pattern in patterns:
                                    matches = re.findall(pattern, content, re.IGNORECASE)
                                    total_matches = max(total_matches, len(matches))
                                if total_matches > 0:
                                    return total_matches
                            break
                        except UnicodeDecodeError:
                            continue
                except:
                    pass
                
                return 20  # Разумное значение по умолчанию для книг
        
        return 5  # Значение по умолчанию
        
    except Exception as e:
        logger.error(f"Ошибка при подсчете глав: {e}")
        return 5

async def show_chapter_selection(update: Update, state: UserState):
    """Показывает опции выбора глав"""
    keyboard = [
        [InlineKeyboardButton("📖 Все главы", callback_data="chapters_all")],
        [InlineKeyboardButton("🔢 Выбрать диапазон", callback_data="chapters_range")],
        [InlineKeyboardButton("📋 Показать все главы", callback_data="show_all_chapters")],
        [InlineKeyboardButton("🔍 Настроить глоссарий", callback_data="setup_glossary_from_chapter_selection")],
        [InlineKeyboardButton("▶️ Перейти к настройкам", callback_data="skip_chapters")]
    ]
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    chapter_info = ""
    if state.total_chapters > 1:
        chapter_info = f"📊 В файле обнаружено примерно **{state.total_chapters} глав/разделов**\n\n"
    
    message_text = (
        f"⚙️ **Настройка перевода**\n\n"
        f"{chapter_info}"
        f"📁 Файл: `{state.file_name}`\n"
        f"📄 Формат: `{state.output_format.upper()}`\n\n"
        f"Выберите опцию:"
    )
    
    # Проверяем тип объекта update и используем соответствующий метод
    try:
        if hasattr(update, 'edit_message_text'):
            # Это CallbackQuery
            await update.edit_message_text(
                message_text,
                reply_markup=reply_markup,
                parse_mode=ParseMode.MARKDOWN
            )
        else:
            # Это Update, отправляем новое сообщение
            await update.message.reply_text(
                message_text,
                reply_markup=reply_markup,
                parse_mode=ParseMode.MARKDOWN
            )
    except BadRequest as e:
        if "Message is not modified" not in str(e):
            raise

async def show_all_chapters(update: Update, state: UserState):
    """Показывает все найденные главы в файле"""
    try:
        logger.info(f"Показ всех глав для файла: {state.file_name}")
        
        # Используем сохраненную информацию или получаем новую
        chapters_info = getattr(state, 'chapters_info', None)
        if not chapters_info:
            logger.info("Получаем информацию о главах через TransGemini")
            chapters_info = await get_transgemini_chapters_info(state.file_path, state.file_format)
            state.chapters_info = chapters_info
        
        logger.info(f"Информация о главах: {chapters_info}")
        
        if chapters_info['total_all'] == 0:
            logger.warning("Не найдено ни одной главы")
            try:
                await update.edit_message_text(
                    "❌ Не удалось проанализировать главы в файле.",
                    reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("⬅️ Назад", callback_data="back_to_chapter_selection")]])
                )
            except BadRequest as e:
                if "Message is not modified" not in str(e):
                    raise
            return
        
        # Формируем текст со списком глав
        message_text = f"📋 **Анализ глав (TransGemini логика)**\n\n"
        message_text += f"📁 Файл: `{state.file_name}`\n\n"
        message_text += f"📊 **Статистика:**\n"
        message_text += f"• Всего HTML файлов: `{chapters_info['total_all']}`\n"
        message_text += f"• Главы для перевода: `{chapters_info['total_content']}`\n"
        message_text += f"• Служебные файлы: `{len(chapters_info['skip_files'])}`\n"
        if chapters_info['nav_file']:
            message_text += f"• NAV файл (оглавление): `{Path(chapters_info['nav_file']).name}`\n"
        message_text += "\n"
        
        # Показываем главы содержания (те, что будут переведены)
        if chapters_info['content_files']:
            message_text += f"✅ **Главы для перевода ({len(chapters_info['content_files'])}):**\n"
            for i, file_data in enumerate(chapters_info['content_files'][:20], 1):  # Показываем первые 20
                size_kb = file_data['size'] // 1024 if file_data['size'] > 0 else 0
                category_emoji = {"content": "📖", "text": "📄"}.get(file_data['category'], "📄")
                message_text += f"{i}. {category_emoji} `{file_data['name']}` ({size_kb}KB)\n"
            
            if len(chapters_info['content_files']) > 20:
                message_text += f"... и еще {len(chapters_info['content_files']) - 20} глав\n"
        
        # Показываем служебные файлы (первые несколько)
        if chapters_info['skip_files']:
            message_text += f"\n🚫 **Служебные файлы ({len(chapters_info['skip_files'])}):**\n"
            for file_data in chapters_info['skip_files'][:8]:  # Показываем первые 8
                size_kb = file_data['size'] // 1024 if file_data['size'] > 0 else 0
                category_emoji = {"nav": "🧭", "skip": "⏭️"}.get(file_data['category'], "❓")
                reason = {"nav": "навигация", "skip": "служебный"}.get(file_data['category'], "неизвестно")
                message_text += f"• {category_emoji} `{file_data['name']}` ({size_kb}KB) - {reason}\n"
            
            if len(chapters_info['skip_files']) > 8:
                message_text += f"... и еще {len(chapters_info['skip_files']) - 8} файлов\n"
        
        # Обновляем точное количество глав в состоянии только для EPUB файлов
        # Для других форматов сохраняем первоначальное значение
        if state.file_format == 'epub' and chapters_info['total_content'] > 0:
            state.total_chapters = chapters_info['total_content']
        # Для не-EPUB файлов или если анализ EPUB не дал результатов, сохраняем исходное значение
        
        keyboard = [
            [InlineKeyboardButton(f"📖 Перевести все {chapters_info['total_content']} глав", callback_data="chapters_all")],
            [InlineKeyboardButton("🔢 Выбрать диапазон", callback_data="chapters_range")],
            [InlineKeyboardButton("⬅️ Назад", callback_data="back_to_chapter_selection")]
        ]
        
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        try:
            await update.edit_message_text(
                message_text,
                reply_markup=reply_markup,
                parse_mode=ParseMode.MARKDOWN
            )
        except BadRequest as e:
            if "Message is not modified" not in str(e):
                raise
        
    except Exception as e:
        logger.error(f"Ошибка показа глав: {e}")
        try:
            await update.edit_message_text(
                f"❌ Ошибка при анализе глав: {str(e)}",
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("⬅️ Назад", callback_data="back_to_chapter_selection")]])
            )
        except BadRequest as e2:
            if "Message is not modified" not in str(e2):
                raise

async def handle_chapter_selection(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик выбора глав"""
    query = update.callback_query
    user_id = query.from_user.id
    state = get_user_state(user_id)
    
    logger.info(f"Получен callback_data: {query.data}, шаг: {state.step}")
    
    if state.step != "chapter_selection":
        logger.warning(f"Неверный шаг процесса. Текущий: {state.step}, ожидается: chapter_selection")
        await query.answer("Неверный шаг процесса")
        return
    
    callback_data = query.data
    logger.info(f"Обрабатываем callback_data: {callback_data}")
    
    if callback_data == "chapters_all":
        logger.info("Выбраны все главы")
        state.start_chapter = 1
        state.chapter_count = 0  # 0 = все главы
        state.step = "translating"
        
        await query.answer("Выбраны все главы")
        await show_translation_options(query, state)
        
    elif callback_data == "chapters_range":
        logger.info("Переход к выбору диапазона")
        await query.answer()
        await show_chapter_range_input(query, state)
        
    elif callback_data == "show_all_chapters":
        logger.info("Показ всех глав")
        await query.answer()
        await show_all_chapters(query, state)
        
    elif callback_data == "setup_glossary_from_chapter_selection":
        logger.info("Переход к настройке глоссария")
        # Запоминаем, что мы были в выборе глав
        state.session_data["previous_step"] = state.step
        # Показываем меню глоссария
        await query.answer("Переход к настройке глоссария")
        await handle_settings_glossary(query, state)
        
    elif callback_data == "skip_chapters":
        logger.info("Пропуск выбора глав")
        state.step = "translating"
        await query.answer()
        await show_translation_options(query, state)
        
    elif callback_data == "back_to_chapter_selection":
        logger.info("Возврат к выбору глав")
        await query.answer()
        await show_chapter_selection(query, state)
    else:
        logger.warning(f"Неизвестный callback_data: {callback_data}")

async def show_chapter_range_input(update: Update, state: UserState):
    """Показывает интерфейс для ввода диапазона глав"""
    keyboard = []
    
    # Быстрые варианты выбора
    if state.total_chapters >= 5:
        keyboard.extend([
            [
                InlineKeyboardButton("1-5 глав", callback_data="range_1_5"),
                InlineKeyboardButton("6-10 глав", callback_data="range_6_10")
            ],
            [
                InlineKeyboardButton("11-15 глав", callback_data="range_11_15"),
                InlineKeyboardButton("16-20 глав", callback_data="range_16_20")
            ]
        ])
    
    # Кнопки для ручного ввода
    keyboard.extend([
        [InlineKeyboardButton("✏️ Ввести вручную", callback_data="range_manual")],
        [InlineKeyboardButton("⬅️ Назад", callback_data="back_to_chapters")]
    ])
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.edit_message_text(
        f"🔢 **Выбор диапазона глав**\n\n"
        f"📊 Всего глав в файле: `{state.total_chapters}`\n\n"
        f"Выберите диапазон или введите вручную:",
        reply_markup=reply_markup,
        parse_mode=ParseMode.MARKDOWN
    )

async def handle_chapter_range_selection(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик выбора диапазона глав"""
    query = update.callback_query
    user_id = query.from_user.id
    state = get_user_state(user_id)
    
    callback_data = query.data
    
    if callback_data == "back_to_chapters":
        await query.answer()
        await show_chapter_selection(query, state)
        return
    elif callback_data.startswith("range_"):
        if callback_data == "range_manual":
            state.step = "chapter_input"
            await query.answer()
            await query.edit_message_text(
                f"✏️ **Ручной ввод диапазона**\n\n"
                f"📊 Всего глав: `{state.total_chapters}`\n\n"
                f"Введите диапазон в одном из форматов:\n"
                f"• `5` - только 5-я глава\n"
                f"• `3-8` - главы с 3 по 8\n"
                f"• `10+5` - начиная с 10-й, всего 5 глав\n\n"
                f"Отправьте ваш выбор:",
                parse_mode=ParseMode.MARKDOWN
            )
            return
        
        # Обработка быстрых вариантов
        range_parts = callback_data.replace("range_", "").split("_")
        if len(range_parts) == 2:
            start_ch = int(range_parts[0])
            end_ch = int(range_parts[1])
            
            state.start_chapter = start_ch
            state.chapter_count = end_ch - start_ch + 1
            state.step = "translating"
            
            await query.answer(f"Выбраны главы {start_ch}-{end_ch}")
            await show_translation_options(query, state)

async def handle_chapter_input(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик ручного ввода диапазона глав"""
    user_id = update.effective_user.id
    state = get_user_state(user_id)
    
    if state.step != "chapter_input":
        return
    
    input_text = update.message.text.strip()
    
    try:
        # Парсим различные форматы ввода
        if "-" in input_text:
            # Формат "3-8"
            parts = input_text.split("-")
            start_ch = int(parts[0])
            end_ch = int(parts[1])
            
            if start_ch < 1 or end_ch > state.total_chapters or start_ch > end_ch:
                raise ValueError("Неверный диапазон")
                
            state.start_chapter = start_ch
            state.chapter_count = end_ch - start_ch + 1
            
        elif "+" in input_text:
            # Формат "10+5"
            parts = input_text.split("+")
            start_ch = int(parts[0])
            count = int(parts[1])
            
            if start_ch < 1 or start_ch > state.total_chapters or count < 1:
                raise ValueError("Неверный диапазон")
                
            state.start_chapter = start_ch
            state.chapter_count = min(count, state.total_chapters - start_ch + 1)
            
        else:
            # Формат "5" - одна глава
            chapter_num = int(input_text)
            
            if chapter_num < 1 or chapter_num > state.total_chapters:
                raise ValueError("Неверный номер главы")
                
            state.start_chapter = chapter_num
            state.chapter_count = 1
        
        state.step = "translating"
        
        # Показываем подтверждение
        end_chapter = min(state.start_chapter + state.chapter_count - 1, state.total_chapters)
        range_text = f"глава {state.start_chapter}" if state.chapter_count == 1 else f"главы {state.start_chapter}-{end_chapter}"
        
        await update.message.reply_text(
            f"✅ Выбран диапазон: **{range_text}**\n\n"
            f"Переходим к настройкам перевода...",
            parse_mode=ParseMode.MARKDOWN
        )
        
        await show_translation_options(update, state)
        
    except (ValueError, IndexError):
        await update.message.reply_text(
            f"❌ Неверный формат ввода!\n\n"
            f"Используйте один из форматов:\n"
            f"• `5` - только 5-я глава\n"
            f"• `3-8` - главы с 3 по 8\n"
            f"• `10+5` - начиная с 10-й, всего 5 глав\n\n"
            f"Максимум глав в файле: `{state.total_chapters}`\n"
            f"Попробуйте еще раз:",
            parse_mode=ParseMode.MARKDOWN
        )

async def show_translation_options(update: Update, state: UserState):
    """Показывает дополнительные опции перевода"""
    keyboard = [
        [InlineKeyboardButton("🇷🇺 Русский", callback_data="lang_русский")],
        [InlineKeyboardButton("🇺🇸 English", callback_data="lang_английский")],
        [InlineKeyboardButton("🇩🇪 Deutsch", callback_data="lang_немецкий")],
        [InlineKeyboardButton("🇫🇷 Français", callback_data="lang_французский")],
        [InlineKeyboardButton("🇪🇸 Español", callback_data="lang_испанский")],
        [InlineKeyboardButton("🤖 Выбрать модель", callback_data="select_model")],
        [InlineKeyboardButton("▶️ Начать перевод", callback_data="start_translation")],
        [InlineKeyboardButton("⬅️ Назад к выбору глав", callback_data="back_to_translation_options")]
    ]
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    # Информация о выбранных главах
    chapter_info = ""
    if hasattr(state, 'total_chapters') and state.total_chapters > 0:
        if state.chapter_count == 0:  # Все главы
            chapter_info = f"📖 Главы: все ({state.total_chapters})\n"
        elif state.chapter_count == 1:
            chapter_info = f"📖 Глава: {state.start_chapter}\n"
        else:
            end_chapter = min(state.start_chapter + state.chapter_count - 1, state.total_chapters)
            chapter_info = f"� Главы: {state.start_chapter}-{end_chapter}\n"
    
    try:
        # Проверяем тип объекта update и используем соответствующий метод
        if hasattr(update, 'edit_message_text'):
            # Это CallbackQuery
            await update.edit_message_text(
                f"🔧 **Настройки перевода**\n\n"
                f"📁 Файл: `{state.file_name}`\n"
                f"📄 Формат: `{state.output_format.upper()}`\n"
                f"{chapter_info}"
                f"🌍 Язык: `{state.target_language}`\n"
                f"🤖 Модель: `{state.model}`\n\n"
                f"Выберите язык перевода или начните перевод:",
                reply_markup=reply_markup,
                parse_mode=ParseMode.MARKDOWN
            )
        else:
            # Это Update, отправляем новое сообщение
            await update.message.reply_text(
                f"🔧 **Настройки перевода**\n\n"
                f"📁 Файл: `{state.file_name}`\n"
                f"📄 Формат: `{state.output_format.upper()}`\n"
                f"{chapter_info}"
                f"🌍 Язык: `{state.target_language}`\n"
                f"🤖 Модель: `{state.model}`\n\n"
                f"Выберите язык перевода или начните перевод:",
                reply_markup=reply_markup,
                parse_mode=ParseMode.MARKDOWN
            )
    except Exception as e:
        # Если не удалось обновить сообщение, отправляем новое
        if "Message is not modified" in str(e):
            logger.warning("Сообщение не изменилось, пропускаем обновление")
        else:
            logger.error(f"Ошибка обновления сообщения: {e}")
            # Определяем где отправить сообщение
            if hasattr(update, 'message') and update.message:
                await update.message.reply_text(
                    f"🔧 **Настройки перевода**\n\n"
                    f"📁 Файл: `{state.file_name}`\n"
                    f"📄 Формат: `{state.output_format.upper()}`\n"
                    f"{chapter_info}"
                    f"🌍 Язык: `{state.target_language}`\n"
                    f"🤖 Модель: `{state.model}`\n\n"
                    f"Выберите язык перевода или начните перевод:",
                    reply_markup=reply_markup,
                    parse_mode=ParseMode.MARKDOWN
                )
            elif hasattr(update, 'from_user'):
                # Для CallbackQuery используем bot.send_message
                await update.get_bot().send_message(
                    chat_id=update.message.chat_id,
                    text=f"🔧 **Настройки перевода**\n\n"
                         f"📁 Файл: `{state.file_name}`\n"
                         f"📄 Формат: `{state.output_format.upper()}`\n"
                         f"{chapter_info}"
                         f"🌍 Язык: `{state.target_language}`\n"
                         f"🤖 Модель: `{state.model}`\n\n"
                         f"Выберите язык перевода или начните перевод:",
                    reply_markup=reply_markup,
                    parse_mode=ParseMode.MARKDOWN
                )

async def show_model_selection(update: Update, state: UserState):
    """Показывает выбор модели Gemini"""
    # Получаем доступные модели из TransGemini.py
    keyboard = []
    models_per_row = 1  # По одной модели в ряду для лучшей читаемости
    
    model_buttons = []
    for model_name in MODELS.keys():
        # Создаем короткие названия для кнопок
        short_name = model_name.replace("Gemini ", "").replace("gemma", "Gemma")
        if len(short_name) > 25:  # Обрезаем слишком длинные названия
            short_name = short_name[:22] + "..."
        
        model_buttons.append(InlineKeyboardButton(
            f"🤖 {short_name}", 
            callback_data=f"model_{model_name}"
        ))
    
    # Группируем кнопки по рядам
    for i in range(0, len(model_buttons), models_per_row):
        keyboard.append(model_buttons[i:i + models_per_row])
    
    # Добавляем кнопку "Назад"
    keyboard.append([InlineKeyboardButton("⬅️ Назад к настройкам", callback_data="back_to_translation_options")])
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    await update.edit_message_text(
        f"🤖 **Выбор модели Gemini**\n\n"
        f"📁 Файл: `{state.file_name}`\n"
        f"🎯 Текущая модель: `{state.model}`\n\n"
        f"**Доступные модели:**\n"
        f"• **Gemini 2.5** - Новейшие модели (рекомендуется)\n"
        f"• **Gemini 2.0** - Быстрые и эффективные\n"
        f"• **Gemini 1.5** - Проверенные временем\n\n"
        f"Выберите модель для перевода:",
        reply_markup=reply_markup,
        parse_mode=ParseMode.MARKDOWN
    )

async def handle_translation_options(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик опций перевода"""
    query = update.callback_query
    user_id = query.from_user.id
    state = get_user_state(user_id)
    
    callback_data = query.data
    
    if callback_data.startswith("lang_"):
        # Изменение языка
        language = callback_data.replace("lang_", "")
        state.target_language = language
        
        await query.answer(f"Выбран язык: {language}")
        await show_translation_options(query, state)
        
    elif callback_data == "select_model":
        # Показать выбор модели
        await query.answer()
        await show_model_selection(query, state)
        
    elif callback_data.startswith("model_"):
        # Изменение модели
        model_name = callback_data.replace("model_", "")
        if model_name in MODELS:
            state.model = model_name
            await query.answer(f"Выбрана модель: {model_name}")
            await show_translation_options(query, state)
        else:
            await query.answer("Неизвестная модель")
            
    elif callback_data == "back_to_translation_options":
        # Вернуться к выбору глав
        await query.answer()
        state.step = "chapter_selection"  # Возвращаем правильный шаг
        await show_chapter_selection(query, state)
        
    elif callback_data == "start_translation":
        await query.answer()
        await start_translation(query, state)

async def start_translation(update: Update, state: UserState):
    import time
    start_time = time.time()
    logger.info(f"⏳ Перевод запущен в {time.strftime('%H:%M:%S', time.localtime(start_time))}")
    """Запускает процесс перевода используя TransGemini.py"""
    # Отправляем сообщение о начале перевода
    await update.edit_message_text(
        f"🔄 **Начинаю перевод...**\n\n"
        f"📁 Файл: `{state.file_name}`\n"
        f"🌍 Язык: `{state.target_language}`\n"
        f"📄 Формат: `{state.output_format.upper()}`\n\n"
        f"⏳ Это может занять несколько минут...",
        parse_mode=ParseMode.MARKDOWN
    )
    
    try:
        # Проверяем, что file_path не None
        if not state.file_path:
            await update.callback_query.edit_message_text(
                "❌ Ошибка: Путь к файлу не найден. Попробуйте загрузить файл заново."
            )
            return
            
        # Создаем выходной файл
        input_path = Path(state.file_path)
        output_dir = input_path.parent
        output_name = f"{input_path.stem}_translated.{state.output_format}"
        output_path = output_dir / output_name
        
        logger.info(f"Начинаю перевод файла: {state.file_path}")
        logger.info(f"Выходной файл: {output_path}")
        logger.info(f"Формат входной: {state.file_format}, выходной: {state.output_format}")
        logger.info(f"Язык: {state.target_language}, Модель: {state.model}")
        logger.info(f"Главы: начиная с {getattr(state, 'start_chapter', 1)}, количество: {getattr(state, 'chapter_count', 0)}")
        logger.info(f"Ротация API ключей: {'Включена' if state.use_key_rotation and len(state.api_keys) > 1 else 'Выключена'}")
        
        # Проверяем, нужно ли использовать ротацию ключей
        if state.use_key_rotation and len(state.api_keys) > 1:
            # Используем run_translation_with_auto_restart для автоматической ротации ключей
            logger.info(f"Используем ротацию с {len(state.api_keys)} API ключами")
            
            # Подготавливаем настройки для функции
            settings = state.get_settings_dict()
            
            # Настраиваем папку вывода и имя файла
            settings['output_folder'] = str(output_dir)
            settings['output_format'] = state.output_format
            
            # Создаем отдельный поток для запуска функции с автоматической ротацией
            # Это необходимо, так как run_translation_with_auto_restart - синхронная функция
            def run_translation_thread():
                try:
                    run_translation_with_auto_restart(settings)
                    logger.info("Перевод с ротацией ключей завершен успешно")
                except Exception as e:
                    logger.error(f"Ошибка при переводе с ротацией ключей: {e}")
            
            # Запускаем перевод в отдельном потоке
            translation_thread = threading.Thread(target=run_translation_thread)
            translation_thread.start()
            
            # Ждем завершения перевода (с таймаутом для безопасности)
            max_wait_time = 3600  # максимальное время ожидания в секундах (1 час)
            translation_thread.join(timeout=max_wait_time)
            
            # Проверяем, успешно ли завершился перевод
            if translation_thread.is_alive():
                logger.warning("Перевод с ротацией превысил максимальное время ожидания")
                success = False
                error_message = "Превышено максимальное время ожидания перевода"
            else:
                # Проверяем наличие выходного файла
                if output_path.exists():
                    success = True
                    error_message = None
                else:
                    success = False
                    error_message = "Не удалось найти выходной файл после перевода с ротацией ключей"
        else:
            # Используем стандартный метод перевода
            success, error_message = await translate_file_with_transgemini(
                input_file=state.file_path,
                output_file=str(output_path),
                input_format=state.file_format,
                output_format=state.output_format,
                target_language=state.target_language,
                api_key=state.api_key,
                model_name=state.model,
                progress_callback=None,  # Отключаем прогресс-бар
                start_chapter=getattr(state, 'start_chapter', 1),
                chapter_count=getattr(state, 'chapter_count', 0),
                chapters_info=getattr(state, 'chapters_info', None),  # Передаем информацию о главах
                glossary_data=getattr(state, 'glossary_data', None)  # Передаем глоссарий
            )
        
        end_time = time.time()
        duration = end_time - start_time
        if success and output_path.exists():
            logger.info(f"✅ Перевод успешно завершен, файл создан: {output_path}")
            logger.info(f"⏱️ Время перевода: {duration:.1f} сек. (завершено в {time.strftime('%H:%M:%S', time.localtime(end_time))})")
            await send_translated_file(update, state, str(output_path))
        else:
            # Добавляем детальную диагностику
            logger.error(f"❌ Проблема с результатом перевода:")
            logger.error(f"   success: {success}")
            logger.error(f"   output_path: {output_path}")
            logger.error(f"   output_path.exists(): {output_path.exists() if output_path else 'N/A'}")
            logger.error(f"   output_path.parent: {output_path.parent if output_path else 'N/A'}")
            logger.error(f"   error_message: {error_message}")
            
            # Проверяем, что находится в выходной папке
            if output_path and output_path.parent.exists():
                try:
                    files_in_output_dir = list(output_path.parent.iterdir())
                    logger.info(f"📁 Файлы в выходной директории {output_path.parent}:")
                    for file in files_in_output_dir:
                        logger.info(f"   - {file.name} (размер: {file.stat().st_size if file.is_file() else 'dir'})")
                except Exception as e:
                    logger.error(f"Ошибка при просмотре выходной директории: {e}")
            
            # Проверяем возможные альтернативные местоположения файла
            if output_path:
                possible_locations = [
                    output_path.parent / f"{output_path.stem}_translated{output_path.suffix}",
                    output_path.parent / f"{output_path.stem}.txt",
                    output_path.parent / f"{Path(state.file_name).stem}_translated.txt",
                ]
                
                logger.info("🔍 Проверяем возможные местоположения переведенного файла:")
                for possible_path in possible_locations:
                    exists = possible_path.exists()
                    logger.info(f"   {possible_path}: {'✅ НАЙДЕН' if exists else '❌ НЕТ'}")
                    if exists and possible_path.is_file():
                        logger.info(f"      Размер: {possible_path.stat().st_size} байт")
                        # Если нашли файл, попробуем его отправить
                        try:
                            await send_translated_file(update, state, str(possible_path))
                            return  # Успешно отправили файл
                        except Exception as send_error:
                            logger.error(f"Ошибка отправки найденного файла {possible_path}: {send_error}")
            
            error_text = "❌ **Ошибка при переводе**\n\n"
            if error_message and "успешно завершен" not in error_message.lower():
                error_text += f"Детали ошибки: `{error_message}`\n\n"
            else:
                error_text += "Перевод был завершен, но переведенный файл не найден.\n\n"
            
            error_text += "Возможные причины:\n"
            error_text += "• Worker создал файл в неожиданном месте\n"
            error_text += "• Проблема с правами доступа к файлу\n"
            error_text += "• Неверный API ключ Google Gemini\n"
            error_text += "• Превышен лимит запросов API\n"
            error_text += "• Проблемы с интернет-соединением\n"
            error_text += "• Файл слишком большой для обработки\n\n"
            error_text += "Попробуйте еще раз через несколько минут."
            
            await update.edit_message_text(
                error_text,
                parse_mode=ParseMode.MARKDOWN
            )
            
    except Exception as e:
        logger.error(f"Критическая ошибка при переводе: {e}", exc_info=True)
        await update.edit_message_text(
            f"❌ **Критическая ошибка при переводе**\n\n"
            f"Произошла неожиданная ошибка: `{str(e)}`\n\n"
            f"Обратитесь к разработчику или попробуйте еще раз позже.",
            parse_mode=ParseMode.MARKDOWN
        )
    
    # Очищаем состояние пользователя
    user_id = update.from_user.id if hasattr(update, 'from_user') else update.effective_user.id
    reset_user_state(user_id)

def extract_body_content_from_html(html_content: str) -> str:
    """
    Извлекает содержимое <body> из HTML, удаляя CSS стили и оставляя только контент
    Решает проблему попадания CSS стилей в тело EPUB файла
    """
    if not html_content or not html_content.strip():
        return ""
    
    try:
        from bs4 import BeautifulSoup
        
        # Специальная обработка для случаев, где CSS стили попадают в начало файла
        # как текст с названием главы (например: "0002_Chapter_2_Bom__Spring_1 <br />body { font-family...")
        if '<br />body {' in html_content and 'font-family' in html_content:
            logger.info("🧹 Обнаружены CSS стили в тексте, выполняем специальную очистку...")
            
            # Разделяем по <br /> и ищем CSS блок
            parts = html_content.split('<br />')
            
            # Ищем часть с CSS стилями и удаляем её
            clean_parts = []
            css_block_started = False
            
            for part in parts:
                part_stripped = part.strip()
                
                # Проверяем, является ли эта часть CSS стилем
                if ('body {' in part_stripped or 
                    'font-family' in part_stripped or
                    'line-height' in part_stripped or
                    'margin:' in part_stripped or
                    'padding:' in part_stripped or
                    'color:' in part_stripped or
                    part_stripped.endswith('}') and any(css_prop in part_stripped for css_prop in ['font-size', 'border', 'background'])):
                    logger.info(f"   Удаляем CSS фрагмент: {part_stripped[:100]}...")
                    continue
                
                # Пропускаем пустые части
                if not part_stripped:
                    continue
                    
                clean_parts.append(part)
            
            # Соединяем очищенные части
            html_content = '<br />'.join(clean_parts)
            logger.info(f"✅ Специальная очистка завершена, осталось {len(clean_parts)} частей")
        
        # Парсим HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Находим тег <body>
        body_tag = soup.find('body')
        if body_tag:
            # Извлекаем содержимое body, убирая сам тег <body>
            body_content = ""
            for element in body_tag.contents:
                body_content += str(element)
            
            # Преобразуем HTML в Markdown-like формат для TransGemini
            from bs4 import BeautifulSoup
            clean_soup = BeautifulSoup(body_content, 'html.parser')
            
            # Заменяем HTML теги на Markdown/текст
            markdown_content = ""
            
            for element in clean_soup.find_all():
                if element.name == 'p':
                    # Параграфы разделяем двумя переносами строк
                    text = element.get_text().strip()
                    if text:
                        markdown_content += text + "\n\n"
                elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    # Заголовки преобразуем в Markdown
                    level = int(element.name[1])
                    text = element.get_text().strip()
                    if text:
                        markdown_content += '#' * level + ' ' + text + "\n\n"
                elif element.name == 'br':
                    markdown_content += "\n"
                elif element.name in ['strong', 'b']:
                    text = element.get_text().strip()
                    if text:
                        markdown_content += f"**{text}**"
                elif element.name in ['em', 'i']:
                    text = element.get_text().strip()
                    if text:
                        markdown_content += f"*{text}*"
            
            # Если не удалось разобрать структуру, просто извлекаем текст
            if not markdown_content.strip():
                markdown_content = clean_soup.get_text()
            
            # Нормализуем переносы строк и пробелы
            markdown_content = re.sub(r'\n\s*\n\s*\n+', '\n\n', markdown_content)
            markdown_content = re.sub(r'[ \t]+', ' ', markdown_content)
            
            logger.info(f"✅ Извлечено и преобразовано содержимое body в Markdown ({len(markdown_content)} символов)")
            return markdown_content.strip()
        else:
            # Если нет тега body, возвращаем весь контент, но убираем стили и HTML теги
            logger.warning("⚠️ Тег <body> не найден, используем весь контент")
            
            # Убираем теги <head>, <style>, <html>, и DOCTYPE
            content = re.sub(r'<!DOCTYPE[^>]*>', '', html_content, flags=re.IGNORECASE)
            content = re.sub(r'<html[^>]*>', '', content, flags=re.IGNORECASE)
            content = re.sub(r'</html>', '', content, flags=re.IGNORECASE)
            content = re.sub(r'<head[^>]*>.*?</head>', '', content, flags=re.DOTALL | re.IGNORECASE)
            content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL | re.IGNORECASE)
            content = re.sub(r'<\?xml[^>]*\?>', '', content, flags=re.IGNORECASE)
            
            # Дополнительная очистка от CSS стилей, которые могли попасть как текст
            content = re.sub(r'body\s*\{[^}]*\}', '', content, flags=re.DOTALL | re.IGNORECASE)
            content = re.sub(r'[a-zA-Z\-]+\s*\{[^}]*\}', '', content, flags=re.DOTALL)
            
            # Преобразуем основные HTML теги в текст с сохранением структуры
            content = re.sub(r'<p[^>]*>', '\n', content, flags=re.IGNORECASE)
            content = re.sub(r'</p>', '\n\n', content, flags=re.IGNORECASE)
            content = re.sub(r'<br\s*/?>', '\n', content, flags=re.IGNORECASE)
            content = re.sub(r'<h[1-6][^>]*>(.*?)</h[1-6]>', r'# \1\n\n', content, flags=re.IGNORECASE | re.DOTALL)
            
            # Убираем оставшиеся HTML теги
            content = re.sub(r'<[^>]+>', '', content)
            
            # Убираем множественные пустые строки и нормализуем пробелы
            content = re.sub(r'\n\s*\n\s*\n+', '\n\n', content)
            content = re.sub(r'[ \t]+', ' ', content)
            
            return content.strip()
            
    except Exception as e:
        logger.error(f"❌ Ошибка извлечения body контента: {e}")
        logger.info("   Возвращаем оригинальный контент")
        return html_content

def format_glossary_for_prompt(glossary_data: dict, text_content: str = None, use_dynamic_glossary: bool = True) -> str:
    """Форматирует глоссарий для промпта с возможной фильтрацией"""
    if not glossary_data:
        logger.info("📚 Глоссарий пуст, пропускаем")
        return ""
        
    logger.info(f"📚 Исходный глоссарий содержит {len(glossary_data)} терминов")
    
    # Применяем динамическую фильтрацию, если указан текст и включена опция
    glossary_to_use = glossary_data
    if text_content and use_dynamic_glossary and glossary_data:
        glossary_to_use = DynamicGlossaryFilter.filter_glossary_for_text(
            glossary_data, text_content
        )
        logger.info(f"🔍 После динамической фильтрации остался {len(glossary_to_use)} терминов")
        
    if not glossary_to_use:
        logger.info("📚 После фильтрации глоссарий пуст")
        return ""
        
    glossary_lines = []
    for original, translation in glossary_to_use.items():
        glossary_lines.append(f"  {original} = {translation}")
        
    glossary_text = f"\n\n**ГЛОССАРИЙ:**\n" + "\n".join(glossary_lines)
    logger.info(f"📚 Сформирован глоссарий для промпта: {len(glossary_lines)} терминов")
    
    return glossary_text

class DynamicGlossaryFilterBot:
    """Улучшенный фильтр глоссария для бота с дополнительными возможностями"""
    
    @staticmethod
    def filter_glossary_for_text(full_glossary, text, min_word_length=3):
        """
        Возвращает только те термины из глоссария, которые встречаются в тексте
        
        Args:
            full_glossary: полный словарь глоссария {оригинал: перевод}
            text: текст для анализа
            min_word_length: минимальная длина термина для поиска
        """
        if not full_glossary or not text:
            return {}
            
        filtered_glossary = {}
        text_lower = text.lower()
        
        for original, translation in full_glossary.items():
            # Пропускаем слишком короткие термины
            if len(original) < min_word_length:
                continue
                
            # Быстрая проверка наличия в тексте
            if original.lower() in text_lower:
                # Дополнительная проверка на границы слов для точности
                import re
                # Создаем паттерн для поиска целого слова/фразы
                pattern = r'\b' + re.escape(original) + r'\b'
                if re.search(pattern, text, re.IGNORECASE):
                    filtered_glossary[original] = translation
                    
        return filtered_glossary

class DynamicWorker(Worker):
    """Расширенный Worker с поддержкой динамической фильтрации глоссария"""
    
    def __init__(self, *args, **kwargs):
        # Извлекаем дополнительные параметры для динамического глоссария
        self.glossary_data = kwargs.pop('glossary_data', None)
        self.use_dynamic_glossary = kwargs.pop('use_dynamic_glossary', True)
        
        # Инициализируем базовый Worker
        super().__init__(*args, **kwargs)
        
        logger.info(f"🔧 DynamicWorker создан:")
        logger.info(f"   📚 Глоссарий: {'Да' if self.glossary_data else 'Нет'}")
        logger.info(f"   🔄 Динамическая фильтрация: {'Включена' if self.use_dynamic_glossary else 'Выключена'}")
        if self.glossary_data:
            logger.info(f"   📊 Размер глоссария: {len(self.glossary_data)} терминов")
    
    def prepare_chunk_prompt(self, chunk_text, base_prompt_template):
        """
        Подготавливает промпт для чанка с учетом динамической фильтрации глоссария
        Переопределяем этот метод для добавления динамического глоссария
        """
        try:
            if not self.glossary_data or not self.use_dynamic_glossary:
                # Если нет глоссария или динамическая фильтрация отключена - используем базовый промпт
                return base_prompt_template.replace("{text}", chunk_text)
            
            # Применяем динамическую фильтрацию глоссария для этого чанка
            filtered_glossary = DynamicGlossaryFilterBot.filter_glossary_for_text(
                self.glossary_data, chunk_text
            )
            
            original_count = len(self.glossary_data)
            filtered_count = len(filtered_glossary)
            logger.info(f"🔍 Динамическая фильтрация: {original_count} → {filtered_count} терминов для чанка")
            
            # Формируем промпт с динамическим глоссарием для этого чанка
            if filtered_glossary:
                dynamic_glossary_text = format_glossary_for_prompt(
                    filtered_glossary, use_dynamic_glossary=False  # Уже отфильтрован
                )
                
                # Создаем промпт с динамическим глоссарием
                chunk_prompt = f"""{dynamic_glossary_text}

ТЕКСТ ДЛЯ ПЕРЕВОДА:
{chunk_text}"""
                logger.info(f"📚 Применен динамический глоссарий: {filtered_count} терминов")
            else:
                # Если нет релевантных терминов, используем обычный промпт
                chunk_prompt = chunk_text
                logger.info(f"📚 Нет релевантных терминов для данного чанка")
            
            return chunk_prompt
            
        except Exception as e:
            logger.error(f"❌ Ошибка подготовки промпта с динамическим глоссарием: {e}")
            # Fallback к стандартному промпту
            return base_prompt_template.replace("{text}", chunk_text)

async def translate_file_with_transgemini(input_file: str, output_file: str, 
                                        input_format: str, output_format: str,
                                        target_language: str, api_key: str, 
                                        model_name: str, progress_callback=None,
                                        start_chapter: int = 1, chapter_count: int = 0,
                                        chapters_info: dict = None, 
                                        glossary_data: dict = None) -> tuple[bool, str]:
    """
    Асинхронная обертка для TransGemini.py Worker класса
    Использует точно такую же логику как TransGemini для сохранения структуры файлов
    """
    
    logger.info(f"🚀 translate_file_with_transgemini: Начинаем перевод")
    logger.info(f"📁 Входной файл: {input_file}")
    logger.info(f"📄 Формат: {input_format} -> {output_format}")
    logger.info(f"📚 Глоссарий передан: {'Да' if glossary_data else 'Нет'}")
    if glossary_data:
        logger.info(f"📚 Размер глоссария: {len(glossary_data)} терминов")
    logger.info(f"🤖 Модель: {model_name}")
    
    start_time = datetime.datetime.now()
    
    def run_worker():
        """Запускает Worker в синхронном режиме"""
        try:
            # Импортируем необходимые компоненты из TransGemini
            from TransGemini import Worker, MODELS
            
            # Получаем конфигурацию модели
            model_config = MODELS.get(model_name, MODELS.get("Gemini 2.0 Flash", MODELS[list(MODELS.keys())[0]]))
            logger.info(f"🤖 Используем модель: {model_name} с конфигурацией: {model_config}")
            
            # Определяем, использовать ли динамический глоссарий
            use_dynamic_glossary = bool(glossary_data)
            
            # Определяем prompt на основе целевого языка
            # При динамическом глоссарии НЕ добавляем глоссарий в базовый промпт
            if target_language.lower() in ['русский', 'russian', 'ru']:
                base_prompt = """Переведи следующий текст на русский язык. Сохрани исходное форматирование, структуру диалогов и разбивку на абзацы. Не добавляй никаких комментариев или пояснений к переводу."""
                
                # Добавляем статический глоссарий к промпту только если динамическая фильтрация отключена
                if glossary_data and not use_dynamic_glossary:
                    glossary_text = format_glossary_for_prompt(glossary_data, use_dynamic_glossary=False)
                    base_prompt += glossary_text
                    logger.info("📚 Статический глоссарий добавлен к промпту (RU)")
                elif glossary_data and use_dynamic_glossary:
                    logger.info("📚 Динамический глоссарий будет применяться к каждому чанку (RU)")
                
                prompt_template = base_prompt + "\n\n{text}"
            else:
                base_prompt = f"""Translate the following text to {target_language}. Preserve the original formatting, dialogue structure, and paragraph breaks. Do not add any comments or explanations to the translation."""
                
                # Добавляем статический глоссарий к промпту только если динамическая фильтрация отключена
                if glossary_data and not use_dynamic_glossary:
                    glossary_text = format_glossary_for_prompt(glossary_data, use_dynamic_glossary=False)
                    base_prompt += glossary_text
                    logger.info("📚 Статический глоссарий добавлен к промпту (EN)")
                elif glossary_data and use_dynamic_glossary:
                    logger.info("📚 Динамический глоссарий будет применяться к каждому чанку (EN)")
                
                prompt_template = base_prompt + "\n\n{text}"
            
            # Определяем выходную директорию из переданного output_file
            output_dir = os.path.dirname(output_file)
            if not output_dir:
                output_dir = os.path.dirname(input_file)
            
            # Подготавливаем данные о файлах для обработки в формате TransGemini
            # TransGemini ожидает список кортежей: (input_type, filepath, epub_html_path_or_none)
            input_type = input_format.lower()
            files_to_process_data = []
            
            if input_type == 'epub':
                # Для EPUB файлов нужно получить список HTML файлов внутри
                try:
                    with zipfile.ZipFile(input_file, 'r') as epub_zip:
                        # Получаем все HTML файлы как в TransGemini
                        html_files_in_epub = sorted([
                            name for name in epub_zip.namelist()
                            if name.lower().endswith(('.html', '.xhtml', '.htm'))
                            and not name.startswith(('__MACOSX', 'META-INF/'))
                        ])
                        
                        # Фильтруем главы, исключая служебные файлы (как в анализе)
                        content_files = []
                        for file_path_in_epub in html_files_in_epub:
                            filename_base = Path(file_path_in_epub).stem.split('.')[0].lower()
                            
                            # Список служебных файлов (как в TransGemini)
                            skip_indicators = ['toc', 'nav', 'ncx', 'cover', 'title', 'index', 'copyright', 'about', 'meta', 'opf',
                                              'masthead', 'colophon', 'imprint', 'acknowledgments', 'dedication',
                                              'glossary', 'bibliography', 'notes', 'annotations', 'epigraph', 'halftitle',
                                              'frontmatter', 'backmatter', 'preface', 'introduction', 'appendix', 'biography',
                                              'isbn', 'legal', 'notice', 'otherbooks', 'prelims', 'team', 'promo', 'bonus']
                            
                            is_skip_file = any(skip_word in filename_base for skip_word in skip_indicators)
                            is_translated = filename_base.endswith('_translated')
                            
                            # Проверяем размер файла
                            try:
                                file_info = epub_zip.getinfo(file_path_in_epub)
                                file_size = file_info.file_size
                            except:
                                file_size = 0
                            
                            # Если файл не служебный и не переведенный, и имеет достаточный размер
                            if not is_skip_file and not is_translated and file_size > 500:
                                content_files.append(file_path_in_epub)
                        
                        logger.info(f"📝 Найдено {len(content_files)} HTML файлов для обработки в EPUB")
                        
                        # Ограничиваем количество файлов если указано
                        if chapter_count > 0:
                            # Берем файлы начиная с start_chapter
                            start_idx = max(0, start_chapter - 1)
                            end_idx = min(len(content_files), start_idx + chapter_count)
                            selected_files = content_files[start_idx:end_idx]
                            logger.info(f"📝 Выбрано {len(selected_files)} файлов (главы {start_chapter}-{start_chapter + len(selected_files) - 1})")
                        else:
                            selected_files = content_files
                            logger.info(f"📝 Выбраны все {len(selected_files)} файлов")
                        
                        # Добавляем каждый HTML файл как отдельную задачу
                        for html_file in selected_files:
                            files_to_process_data.append(('epub', input_file, html_file))
                        
                        if not files_to_process_data:
                            logger.error("❌ Не найдено HTML файлов для обработки в EPUB")
                            return False, "В EPUB файле не найдено подходящих HTML файлов для перевода"
                            
                except Exception as e:
                    logger.error(f"❌ Ошибка анализа EPUB файла: {e}")
                    return False, f"Ошибка анализа EPUB файла: {str(e)}"
            else:
                # Для других форматов используем прямую обработку
                files_to_process_data = [(input_type, input_file, None)]
            
            logger.info(f"📝 Подготовленные файлы для обработки: {files_to_process_data}")
            
            # Создаем DynamicWorker с поддержкой динамического глоссария
            # Для EPUB файлов используем промежуточный формат HTML, затем соберем EPUB отдельно
            worker_output_format = 'html' if output_format == 'epub' else output_format
            
            worker = DynamicWorker(
                api_key=api_key,
                out_folder=output_dir,
                prompt_template=prompt_template,
                files_to_process_data=files_to_process_data,
                model_config=model_config,
                max_concurrent_requests=1,  # Последовательная обработка для стабильности Telegram бота
                output_format=worker_output_format,  # Используем промежуточный формат для EPUB
                chunking_enabled_gui=True,
                chunk_limit=900000,  # Максимальный размер чанка
                chunk_window=500,
                temperature=0.1,
                chunk_delay_seconds=0.5,  # Уменьшенная задержка между чанками для быстрого перевода
                proxy_string=None,
                # Дополнительные параметры для динамического глоссария
                glossary_data=glossary_data,
                use_dynamic_glossary=use_dynamic_glossary
            )
            
            logger.info(f"DynamicWorker создан с динамической фильтрацией: {'Включена' if use_dynamic_glossary else 'Выключена'}")
            
            # Добавляем обработчик для захвата логов Worker'а
            worker_logs = []
            worker_errors = []
            
            def capture_worker_log(message):
                worker_logs.append(message)
                logger.info(f"Worker Log: {message}")
                # Логируем прогресс по API активности
                if any(word in message for word in ['[API START]', '[API CALL]', '[API RESPONSE]', 'Обработка', 'Начинаю обработку', 'Получен ответ', 'Отправляю запрос']):
                    logger.info(f"[PROGRESS] {message}")
                if any(keyword in message.lower() for keyword in ['error', 'failed', 'exception', 'ошибка']):
                    worker_errors.append(message)
            
            def extract_progress_info(log_message: str) -> str:
                """Извлекает информацию о прогрессе из лог-сообщения Worker'а"""
                try:
                    # Паттерны для различных этапов обработки
                    patterns = [
                        # API активность - новые паттерны
                        (r'\[API START\]\s*([^:]+):\s*Начинаем\s+API\s+запрос', 
                         lambda m: f"🔄 **{m.group(1)}**\n🚀 Начинаем API запрос..."),
                        
                        (r'\[API CALL\]\s*([^:]+):\s*Отправляем\s+запрос\s+к\s+API', 
                         lambda m: f"📡 **{m.group(1)}**\n🤖 Отправляем в Gemini API..."),
                        
                        (r'\[API RESPONSE\]\s*([^:]+):\s*Получен\s+ответ\s+от\s+API', 
                         lambda m: f"✅ **{m.group(1)}**\n📝 Получен ответ от API!"),
                        
                        # Обработка EPUB файлов
                        (r'\[INFO\]\s*([^:]+):\s*Обработка\s+(\d+)/(\d+)\s+чанков', 
                         lambda m: f"📄 **{m.group(1)}**\n⏳ Обрабатываю чанк {m.group(2)} из {m.group(3)}"),
                        
                        # Начало обработки файла
                        (r'\[INFO\]\s*([^:]+):\s*Начинаю\s+обработку', 
                         lambda m: f"🚀 **Начинаю обработку**\n📄 {m.group(1)}"),
                        
                        # Обработка HTML файлов из EPUB
                        (r'\[INFO\]\s*([^:]+):\s*Контент\s+\(([^)]+)\).*разделяем', 
                         lambda m: f"📖 **{m.group(1)}**\n🔄 Разделяю контент ({m.group(2)})"),
                        
                        # Отправка в API (старый формат)
                        (r'\[INFO\]\s*([^:]+):\s*Отправляю\s+запрос\s+в\s+API', 
                         lambda m: f"🤖 **{m.group(1)}**\n📡 Отправляю в Gemini API..."),
                        
                        # Успешный ответ API (старый формат)
                        (r'\[INFO\]\s*([^:]+):\s*Получен\s+ответ.*символов', 
                         lambda m: f"✅ **{m.group(1)}**\n📝 Получен перевод от Gemini"),
                        
                        # Задержка между запросами
                        (r'\[INFO\]\s*([^:]+):\s*Применяем\s+задержку\s+([\d.]+)\s+сек', 
                         lambda m: f"⏰ **{m.group(1)}**\n⏳ Ожидание {m.group(2)} сек..."),
                        
                        # Завершение обработки файла
                        (r'\[INFO\]\s*([^:]+):\s*Обработка\s+завершена', 
                         lambda m: f"✅ **{m.group(1)}**\n🎉 Обработка завершена!"),
                        
                        # Прогресс чанков
                        (r'Chunk\s+(\d+)/(\d+)', 
                         lambda m: f"📝 **Обработка чанка**\n🔢 {m.group(1)} из {m.group(2)}"),
                    ]
                    
                    for pattern, formatter in patterns:
                        match = re.search(pattern, log_message, re.IGNORECASE)
                        if match:
                            try:
                                return formatter(match)
                            except Exception as e:
                                logger.error(f"Ошибка форматирования прогресса: {e}")
                                return f"📋 {log_message}"
                    
                    return None
                    
                except Exception as e:
                    logger.error(f"Ошибка извлечения прогресса: {e}")
                    return None
            
            # Подключаем обработчик логов
            worker.log_message.connect(capture_worker_log)
            
            # Подключаем callback для прогресса если передан
            if progress_callback:
                def on_log_message(message):
                    progress_info = extract_progress_info(message)
                    if progress_info:
                        try:
                            # Запускаем callback асинхронно
                            asyncio.run_coroutine_threadsafe(
                                progress_callback(progress_info), 
                                asyncio.get_event_loop()
                            )
                        except Exception as e:
                            logger.error(f"Ошибка вызова progress_callback: {e}")
                
                worker.log_message.connect(on_log_message)
            
            # Запускаем обработку
            logger.info("🏃 Запускаем Worker.run()...")
            worker.run()
            
            # Проверяем результаты
            if worker_errors:
                error_msg = f"Обнаружены ошибки во время перевода: {'; '.join(worker_errors[:3])}"
                logger.error(f"❌ {error_msg}")
                return False, error_msg

            # Специальная обработка для EPUB - создаем полный EPUB с заменой только переведенных глав
            if output_format == 'epub' and input_format == 'epub':
                logger.info("📚 Начинаем сборку полного EPUB файла с частичным переводом (как в TransGemini)...")
                
                # Логируем состояние выходной директории
                logger.info(f"📁 Файлы в выходной директории {output_dir}:")
                for file in os.listdir(output_dir):
                    file_path = os.path.join(output_dir, file)
                    if os.path.isfile(file_path):
                        file_size = os.path.getsize(file_path)
                        logger.info(f"   - {file} (размер: {file_size})")
                
                logger.info(f"📝 Список файлов для обработки:")
                for i, (ftype, fpath, html_path) in enumerate(files_to_process_data):
                    logger.info(f"   {i+1}. Type: {ftype}, File: {Path(fpath).name}, HTML: {html_path}")
                
                try:
                    # Импортируем функцию сборки EPUB из TransGemini
                    from TransGemini import write_to_epub
                    
                    # Создаем маппинг переведенных файлов
                    translation_mapping = {}
                    translated_files_found = []
                    
                    logger.info(f"🔍 Поиск переведенных HTML файлов в {output_dir}...")
                    
                    for file in os.listdir(output_dir):
                        if file.endswith('_translated.html'):
                            file_path = os.path.join(output_dir, file)
                            # Извлекаем оригинальное имя без _translated.html
                            base_name = file.replace('_translated.html', '')
                            
                            logger.info(f"🔍 Обрабатываем переведенный файл: {file} (базовое имя: {base_name})")
                            
                            # Читаем переведенный контент
                            with open(file_path, 'r', encoding='utf-8') as f:
                                raw_translated_content = f.read()
                            
                            # Извлекаем только содержимое <body> из переведенного HTML, 
                            # удаляя CSS стили, которые должны быть в <head>
                            translated_content = extract_body_content_from_html(raw_translated_content)
                            
                            # Добавляем детальное логирование переведенного контента
                            raw_content_preview = raw_translated_content[:300].replace('\n', ' ') if raw_translated_content else "пустой"
                            logger.info(f"🔍 Глава переведена: {file}")
                            logger.info(f"📝 Размер исходного контента: {len(raw_translated_content) if raw_translated_content else 0} символов")
                            logger.info(f"📖 Превью исходного: {raw_content_preview}...")
                            
                            content_preview = translated_content[:300].replace('\n', ' ') if translated_content else "пустой"
                            logger.info(f"📝 Размер очищенного контента: {len(translated_content) if translated_content else 0} символов")
                            logger.info(f"📖 Превью очищенного: {content_preview}...")
                            
                            if len(translated_content) < 100:
                                logger.warning(f"⚠️ Подозрительно короткий переведенный контент в {file}: {translated_content}")
                            
                            # Проверяем, что в очищенном контенте нет CSS стилей и HTML тегов
                            if 'font-family' in translated_content or 'line-height' in translated_content:
                                logger.warning(f"⚠️ В очищенном контенте всё ещё есть CSS стили: {file}")
                                logger.info(f"   Начало: {translated_content[:500]}")
                            elif '<p>' in translated_content or '<div>' in translated_content or '<br' in translated_content:
                                logger.warning(f"⚠️ В очищенном контенте всё ещё есть HTML теги: {file}")
                                logger.info(f"   Начало: {translated_content[:500]}")
                            else:
                                logger.info(f"✅ Очищенный контент не содержит CSS стилей и HTML тегов")
                            
                            # Ищем соответствующий HTML файл в списке обработанных
                            matched_original_path = None
                            
                            # Сначала ищем точное соответствие по stem имени
                            for (ftype, fpath, html_path) in files_to_process_data:
                                if html_path:
                                    html_stem = Path(html_path).stem
                                    if html_stem == base_name:
                                        matched_original_path = html_path
                                        logger.info(f"✅ Найдено точное соответствие: {base_name} -> {html_path}")
                                        break
                            
                            # Если точное соответствие не найдено, ищем по более гибким критериям
                            if not matched_original_path:
                                logger.info(f"⚠️ Точное соответствие для {base_name} не найдено, ищем по гибким критериям...")
                                
                                for (ftype, fpath, html_path) in files_to_process_data:
                                    if html_path:
                                        html_stem = Path(html_path).stem.lower()
                                        base_name_lower = base_name.lower()
                                        
                                        # Проверяем различные варианты сопоставления
                                        if (html_stem == base_name_lower or 
                                            html_stem.replace('_', '-') == base_name_lower.replace('_', '-') or
                                            html_stem.replace('-', '_') == base_name_lower.replace('-', '_') or
                                            base_name_lower in html_stem or 
                                            html_stem in base_name_lower):
                                            matched_original_path = html_path
                                            logger.info(f"✅ Найдено гибкое соответствие: {base_name} -> {html_path}")
                                            break
                            
                            if matched_original_path:
                                translation_mapping[matched_original_path] = translated_content
                                translated_files_found.append(file)
                                
                                # Добавляем логирование переведенного контента
                                content_preview = translated_content[:200].replace('\n', ' ') if translated_content else "пустой"
                                logger.info(f"✅ Сопоставлен переведенный файл: {file} -> {matched_original_path}")
                                logger.info(f"🔍 Глава переведена: {len(translated_content)} символов")
                                logger.info(f"📝 Превью перевода: {content_preview}...")
                            else:
                                logger.warning(f"⚠️ Не найдено соответствие для переведенного файла: {file}")
                                logger.info(f"   Доступные HTML пути в files_to_process_data:")
                                for (ftype, fpath, html_path) in files_to_process_data:
                                    if html_path:
                                        logger.info(f"     - {html_path} (stem: {Path(html_path).stem})")
                    
                    logger.info(f"🔍 Результат создания маппинга переводов:")
                    logger.info(f"   Найдено переведенных файлов: {len(translation_mapping)}")
                    for orig_path, content in translation_mapping.items():
                        content_preview = content[:100].replace('\n', ' ') if content else "пустой"
                        logger.info(f"   - {orig_path} ({len(content)} символов): {content_preview}...")
                    
                    if not translation_mapping:
                        logger.error("❌ Не найдено переведенных файлов для сборки EPUB!")
                        logger.info("📋 Проверьте, что Worker успешно создал переведенные HTML файлы")
                        return False, "Не найдено переведенных файлов для сборки EPUB"
                    
                    if translation_mapping:
                        # Теперь создаем полный EPUB файл: берем ВСЕ файлы из оригинального EPUB
                        # и заменяем только те, что были переведены
                        
                        logger.info("📖 Создаем полный EPUB с частичным переводом...")
                        
                        # Получаем ВСЕ HTML файлы из оригинального EPUB
                        all_epub_parts = []
                        with zipfile.ZipFile(input_file, 'r') as epub_zip:
                            all_html_files = sorted([
                                name for name in epub_zip.namelist()
                                if name.lower().endswith(('.html', '.xhtml', '.htm'))
                                and not name.startswith(('__MACOSX', 'META-INF/'))
                            ])
                            
                            logger.info(f"📚 Обрабатываем все HTML файлы в EPUB ({len(all_html_files)} файлов):")
                            
                            for html_file in all_html_files:
                                try:
                                    # Проверяем, есть ли перевод для этого файла
                                    if html_file in translation_mapping:
                                        # Используем переведенный контент (строка)
                                        content = translation_mapping[html_file]
                                        status = "ПЕРЕВЕДЕН"
                                        logger.info(f"   ✅ {html_file} - {status} ({len(content)} символов)")
                                    else:
                                        # Используем оригинальный контент (байты для TransGemini)
                                        with epub_zip.open(html_file) as f:
                                            content = f.read()  # Оставляем как bytes для оригинального контента
                                        status = "ОРИГИНАЛ"
                                        logger.info(f"   📄 {html_file} - {status} ({len(content)} байт)")
                                    
                                    # Проверяем, переведенный файл или оригинальный
                                    is_translated = html_file in translation_mapping
                                    
                                    all_epub_parts.append({
                                        'original_filename': html_file,
                                        'content_to_write': content,  # str для переводов, bytes для оригиналов
                                        'image_map': {},  # Пустая карта изображений
                                        'is_original_content': not is_translated,
                                        'translation_warning': None
                                    })
                                    
                                except Exception as e:
                                    logger.error(f"❌ Ошибка обработки файла {html_file}: {e}")
                                    # В случае ошибки, пропускаем файл (но лучше бы добавить как оригинал)
                                    try:
                                        with epub_zip.open(html_file) as f:
                                            content = f.read()  # Оставляем как bytes для оригинального контента
                                        all_epub_parts.append({
                                            'original_filename': html_file,
                                            'content_to_write': content,  # bytes для оригинального контента
                                            'image_map': {},
                                            'is_original_content': True,
                                            'translation_warning': f"Ошибка обработки: {e}"
                                        })
                                        logger.info(f"   📄 {html_file} - ОРИГИНАЛ (после ошибки, {len(content)} байт)")
                                    except Exception as e2:
                                        logger.error(f"❌ Критическая ошибка с файлом {html_file}: {e2}")
                        
                        logger.info(f"📦 Подготовлен полный EPUB с {len(all_epub_parts)} HTML файлами:")
                        logger.info(f"   Переведенных файлов: {len(translation_mapping)}")
                        logger.info(f"   Оригинальных файлов: {len(all_epub_parts) - len(translation_mapping)}")
                        # Безопасный подсчет размера для строк и байтов
                        total_content_size = 0
                        for part in all_epub_parts:
                            content = part.get('content_to_write', '')
                            if isinstance(content, bytes):
                                total_content_size += len(content)
                            elif isinstance(content, str):
                                total_content_size += len(content)
                        logger.info(f"   Общий размер контента: {total_content_size} символов/байт")
                        
                        # Создаем финальный EPUB файл
                        final_epub_path = os.path.join(output_dir, f"{Path(input_file).stem}_translated.epub")
                        
                        # Подготавливаем метаданные для сборки EPUB - читаем из оригинального файла
                        build_metadata = extract_epub_metadata(input_file)
                        
                        # Дополняем/переопределяем нужные поля
                        build_metadata.update({
                            'title': f"{Path(input_file).stem}_translated",
                            'author': 'TransGemini Bot',
                            'combined_image_map': {}  # Пустая карта изображений
                        })
                        
                        # Вызываем функцию сборки EPUB с полным набором файлов
                        logger.info("🔧 Вызываем write_to_epub для создания полного EPUB...")
                        
                        # Добавляем детальное логирование как вы предложили
                        logger.info(f"📄 processed_epub_parts: {[p['original_filename'] for p in all_epub_parts]}")
                        logger.info(f"📊 Всего частей: {len(all_epub_parts)}")
                        logger.info(f"📁 Оригинальный EPUB: {input_file}")
                        logger.info(f"📁 Выходной файл: {final_epub_path}")
                        logger.info(f"🔧 build_metadata: {build_metadata}")
                        
                        success, error_msg = write_to_epub(
                            out_path=final_epub_path,
                            processed_epub_parts=all_epub_parts,  # ВСЕ файлы (переведенные + оригинальные)
                            original_epub_path=input_file,
                            build_metadata=build_metadata,
                            book_title_override=None
                        )
                        
                        if success and os.path.exists(final_epub_path):
                            # Перемещаем финальный EPUB файл в нужное место
                            if final_epub_path != output_file:
                                try:
                                    os.makedirs(os.path.dirname(output_file), exist_ok=True)
                                    shutil.move(final_epub_path, output_file)
                                    final_output_path = output_file
                                except Exception as e:
                                    logger.warning(f"⚠️ Не удалось переместить EPUB: {e}")
                                    final_output_path = final_epub_path
                            else:
                                final_output_path = final_epub_path
                            
                            # Очищаем промежуточные HTML файлы
                            for file in translated_files_found:
                                try:
                                    file_path = os.path.join(output_dir, file)
                                    os.remove(file_path)
                                    logger.info(f"🗑️ Удален промежуточный файл: {file}")
                                except Exception as e:
                                    logger.warning(f"⚠️ Не удалось удалить промежуточный файл {file}: {e}")
                            
                            file_size = os.path.getsize(final_output_path)
                            end_time = datetime.datetime.now()
                            duration = end_time - start_time
                            
                            logger.info(f"✅ Полный EPUB файл с частичным переводом успешно создан!")
                            logger.info(f"📁 Выходной файл: {final_output_path}")
                            logger.info(f"📊 Размер файла: {file_size} байт ({file_size // 1024} KB)")
                            logger.info(f"📖 Переведено глав: {len(translation_mapping)}")
                            logger.info(f"📄 Всего глав в EPUB: {len(all_epub_parts)}")
                            logger.info(f"⏱️ Время выполнения: {duration}")
                            
                            return True, f"EPUB перевод завершен. Файл сохранен: {os.path.basename(final_output_path)} ({file_size} байт)"
                        else:
                            logger.error(f"❌ Ошибка сборки EPUB: {error_msg}")
                            return False, f"Ошибка сборки EPUB: {error_msg}"
                    else:
                        logger.error("❌ Не найдено переведенных HTML файлов для сборки EPUB")
                        
                        # Подробная диагностика состояния
                        logger.info("🔍 Диагностика состояния для отладки:")
                        logger.info(f"   📁 Выходная директория: {output_dir}")
                        logger.info(f"   📄 Количество файлов для обработки: {len(files_to_process_data)}")
                        
                        # Показываем все файлы в выходной директории
                        logger.info("   📁 Содержимое выходной директории:")
                        for file in os.listdir(output_dir):
                            file_path = os.path.join(output_dir, file)
                            if os.path.isfile(file_path):
                                file_size = os.path.getsize(file_path)
                                logger.info(f"     - {file} ({file_size} байт)")
                        
                        # Показываем список файлов для обработки
                        logger.info("   📋 Список files_to_process_data:")
                        for i, (ftype, fpath, html_path) in enumerate(files_to_process_data):
                            logger.info(f"     {i+1}. Type: {ftype}, File: {Path(fpath).name}, HTML: {html_path}")
                        
                        # Проверяем возможные альтернативные имена файлов
                        logger.info("   🔍 Поиск альтернативных переведенных файлов:")
                        alternative_patterns = [
                            "*_translated.html",
                            "*_translated.txt", 
                            "*_translated.*",
                            f"{Path(input_file).stem}*"
                        ]
                        
                        import glob
                        for pattern in alternative_patterns:
                            pattern_path = os.path.join(output_dir, pattern)
                            matches = glob.glob(pattern_path)
                            if matches:
                                logger.info(f"     Паттерн '{pattern}': найдено {len(matches)} файлов")
                                for match in matches[:5]:  # Показываем первые 5
                                    match_size = os.path.getsize(match) if os.path.exists(match) else 0
                                    logger.info(f"       - {Path(match).name} ({match_size} байт)")
                            else:
                                logger.info(f"     Паттерн '{pattern}': файлы не найдены")
                        
                        return False, "Не найдено переведенных HTML файлов для сборки EPUB. Проверьте логи для диагностики."
                        
                except Exception as e:
                    logger.error(f"❌ Ошибка при сборке EPUB: {e}", exc_info=True)
                    return False, f"Ошибка сборки EPUB: {str(e)}"
            
            # Обычная обработка для не-EPUB файлов
            worker_output_ext = 'html' if output_format == 'epub' else output_format
            
            # Ищем созданный файл в выходной директории
            # Worker создает файлы с суффиксом _translated
            input_name = Path(input_file).stem
            expected_output_name = f"{input_name}_translated.{worker_output_ext}"
            expected_output_path = os.path.join(output_dir, expected_output_name)
            
            if os.path.exists(expected_output_path):
                # Если нужно переименовать файл
                final_output_path = expected_output_path
                if expected_output_path != output_file:
                    try:
                        # Создаем директорию для целевого файла
                        os.makedirs(os.path.dirname(output_file), exist_ok=True)
                        # Перемещаем файл
                        shutil.move(expected_output_path, output_file)
                        logger.info(f"✅ Файл перемещен с {expected_output_path} на {output_file}")
                        final_output_path = output_file
                    except Exception as e:
                        logger.warning(f"⚠️ Не удалось переместить файл: {e}, используем оригинальный путь")
                        final_output_path = expected_output_path
                
                file_size = os.path.getsize(final_output_path)
                end_time = datetime.datetime.now()
                duration = end_time - start_time
                
                logger.info(f"✅ Перевод завершен успешно!")
                logger.info(f"📁 Выходной файл: {final_output_path}")
                logger.info(f"📊 Размер файла: {file_size} байт")
                logger.info(f"⏱️ Время выполнения: {duration}")
                
                return True, f"Перевод завершен. Файл сохранен: {os.path.basename(final_output_path)} ({file_size} байт)"
            else:
                # Ищем любые созданные файлы с _translated
                created_files = []
                for file in os.listdir(output_dir):
                    if '_translated' in file and file.endswith(f'.{output_format}'):
                        created_files.append(os.path.join(output_dir, file))
                
                if created_files:
                    # Берем первый найденный файл
                    actual_output = created_files[0]
                    logger.info(f"✅ Найден созданный файл: {actual_output}")
                    
                    # Перемещаем к нужному имени
                    final_output_path = actual_output
                    if actual_output != output_file:
                        try:
                            shutil.move(actual_output, output_file)
                            logger.info(f"✅ Файл перемещен на {output_file}")
                            final_output_path = output_file
                        except Exception as e:
                            logger.warning(f"⚠️ Не удалось переместить файл: {e}")
                            final_output_path = actual_output
                    
                    file_size = os.path.getsize(final_output_path)
                    return True, f"Перевод завершен. Файл сохранен: {os.path.basename(final_output_path)} ({file_size} байт)"
                else:
                    error_msg = f"Файл не был создан. Ожидался: {expected_output_path}"
                    logger.error(f"❌ {error_msg}")
                    return False, error_msg
                    
        except Exception as e:
            logger.error(f"❌ Ошибка при запуске Worker: {e}", exc_info=True)
            return False, f"Ошибка обработки: {str(e)}"
    
    # Запускаем Worker в отдельном потоке
    try:
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(None, run_worker)
        return result
    except Exception as e:
        logger.error(f"❌ Ошибка выполнения: {e}", exc_info=True)
        return False, f"Ошибка выполнения: {str(e)}"


async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Команда помощи"""
    help_text = """
🤖 **TransGemini Telegram Bot** - переводчик файлов

📖 **Поддерживаемые форматы:**
• EPUB (электронные книги)
• TXT (текстовые файлы)
• DOCX (документы Word)
• HTML (веб-страницы)

🔧 **Как использовать:**
1. Отправьте файл боту
2. Выберите выходной формат
3. Введите API ключ Gemini
4. Выберите главы для перевода
5. Настройте параметры перевода
6. Получите переведенный файл

✨ **Новые функции:**
• `/settings` - управление всеми настройками бота
• `/apikeys` - управление несколькими API ключами
• `/addkey ВАШ_КЛЮЧ` - добавить новый API ключ
• `/removekey НОМЕР` - удалить API ключ по номеру
• `/clearkeys` - удалить все API ключи
• `/rotation on/off` - включить/выключить автоматическую ротацию ключей

🔄 **Автоматическая ротация ключей:**
При включенной ротации бот будет автоматически переключаться между вашими API ключами при достижении лимитов или возникновении ошибок.

⚙️ **Команды:**
/start - начать работу
/help - показать справку
/cancel - отменить текущую операцию

🔑 **Получение API ключа:**
1. Перейдите на ai.google.dev
2. Создайте проект или выберите существующий
3. Получите API ключ для Gemini
4. Введите ключ в боте

💡 **Совет:** Бот использует TransGemini.py для высококачественного перевода с сохранением структуры документов.
    """
    await update.message.reply_text(help_text, parse_mode='Markdown')


async def cancel_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Команда отмены"""
    user_id = update.effective_user.id
    reset_user_state(user_id)
    
    await update.message.reply_text(
        "❌ Текущая операция отменена.\n"
        "📎 Отправьте новый файл для начала перевода или используйте /start",
        reply_markup=InlineKeyboardMarkup([])
    )


async def send_translated_file(update: Update, state: UserState, translated_file_path: str):
    """Отправляет переведенный файл пользователю"""
    try:
        file_path = Path(translated_file_path)
        
        if not file_path.exists():
            # Определяем через какой метод отправить сообщение
            if hasattr(update, 'edit_message_text'):
                # Это CallbackQuery
                await update.edit_message_text("❌ Переведенный файл не найден!")
            else:
                # Это обычный Update
                await update.message.reply_text("❌ Переведенный файл не найден!")
            return
        
        # Проверяем размер файла (ограничение Telegram - 50MB)
        file_size = file_path.stat().st_size
        if file_size > 50 * 1024 * 1024:  # 50 MB
            error_msg = (f"❌ Файл слишком большой ({file_size / 1024 / 1024:.1f} MB). "
                        f"Максимальный размер для Telegram: 50 MB")
            
            if hasattr(update, 'edit_message_text'):
                await update.edit_message_text(error_msg)
            else:
                await update.message.reply_text(error_msg)
            return
        
        # Определяем объект для отправки файла
        if hasattr(update, 'message') and update.message:
            # CallbackQuery с message
            message_obj = update.message
        elif hasattr(update, 'callback_query') and update.callback_query and update.callback_query.message:
            # Update с callback_query
            message_obj = update.callback_query.message
        else:
            # Прямой Update
            message_obj = update.message
        
        # Отправляем файл
        with open(file_path, 'rb') as f:
            await message_obj.reply_document(
                document=f,
                filename=file_path.name,
                caption=f"✅ Перевод завершен!\n"
                       f"📄 Файл: {file_path.name}\n"
                       f"📊 Размер: {file_size / 1024:.1f} KB\n"
                       f"🎯 Переведено с помощью TransGemini"
            )
        
        # Сбрасываем состояние пользователя
        user_id = update.effective_user.id if hasattr(update, 'effective_user') else update.from_user.id
        reset_user_state(user_id)
        
        logger.info(f"✅ Файл отправлен пользователю {user_id}: {file_path.name}")
        
    except Exception as e:
        logger.error(f"❌ Ошибка отправки файла: {e}", exc_info=True)
        
        error_msg = f"❌ Ошибка при отправке файла: {str(e)}"
        
        try:
            if hasattr(update, 'edit_message_text'):
                await update.edit_message_text(error_msg)
            elif hasattr(update, 'callback_query') and update.callback_query.message:
                await update.callback_query.message.reply_text(error_msg)
            else:
                await update.message.reply_text(error_msg)
        except Exception as send_error:
            logger.error(f"❌ Не удалось отправить сообщение об ошибке: {send_error}")


def load_env_file():
    """Загружает переменные из .env файла"""
    env_path = Path('.env')
    if env_path.exists():
        with open(env_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#') and '=' in line:
                    key, value = line.split('=', 1)
                    os.environ[key.strip()] = value.strip()

def main():
    """Основная функция запуска бота"""
    print("🤖 Запуск Telegram бота переводчика файлов")
    print("🔧 Использует TransGemini.py для высококачественного перевода")
    print("=" * 60)
    
    # Загружаем .env файл если существует
    load_env_file()
    
    # Получаем токен бота из переменной окружения
    bot_token = os.getenv('TELEGRAM_BOT_TOKEN')
    
    if not bot_token:
        print("❌ Ошибка: Не найден токен бота!")
        print("Создайте бота через @BotFather в Telegram и получите токен")
        print("=" * 40)
        
        # Попросим пользователя ввести токен
        bot_token = input("Введите токен бота: ").strip()
        if not bot_token:
            print("Токен не введен. Выход.")
            sys.exit(1)
    
    # Создаем приложение
    application = Application.builder().token(bot_token).build()
    
    # Добавляем основные обработчики
    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("help", help_command))
    application.add_handler(CommandHandler("cancel", cancel_command))
    
    # Добавляем новые обработчики команд для управления API ключами
    application.add_handler(CommandHandler("apikeys", handle_apikeys_command))
    application.add_handler(CommandHandler("addkey", handle_addkey_command))
    application.add_handler(CommandHandler("removekey", handle_removekey_command))
    application.add_handler(CommandHandler("clearkeys", handle_clearkeys_command))
    application.add_handler(CommandHandler("rotation", handle_rotation_command))
    
    # Добавляем обработчик команды настроек
    application.add_handler(CommandHandler("settings", handle_settings_command))
    
    # Обработчик файлов
    application.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    
    # Обработчик API ключа и ввода диапазона глав (текстовые сообщения)
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text_input))
    
    # Обработчики callback кнопок
    application.add_handler(CallbackQueryHandler(handle_format_selection, pattern=r"^action_"))
    application.add_handler(CallbackQueryHandler(handle_output_format_selection, pattern=r"^(format_.*|back_to_action_selection)$"))
    application.add_handler(CallbackQueryHandler(handle_chapter_selection, pattern=r"^(chapters_|skip_chapters|back_to_chapter_selection|show_all_chapters|setup_glossary_from_chapter_selection)"))
    application.add_handler(CallbackQueryHandler(handle_chapter_range_selection, pattern=r"^(range_|back_to_chapters)"))
    application.add_handler(CallbackQueryHandler(handle_translation_options, pattern=r"^(lang_|select_model$|model_|back_to_translation_options$|start_translation$)"))
    
    # Обработчики для глоссария
    application.add_handler(CallbackQueryHandler(handle_glossary_model_selection, pattern=r"^glossary_model_"))
    application.add_handler(CallbackQueryHandler(handle_glossary_options, pattern=r"^(start_glossary_creation|change_glossary_model|back_to_action_selection)$"))
    
    # Новые обработчики callback кнопок для расширенных настроек
    application.add_handler(CallbackQueryHandler(handle_keys_callback, pattern=r"^(confirm_clear_keys|cancel_clear_keys)"))
    application.add_handler(CallbackQueryHandler(handle_settings_callback, pattern=r"^(settings_|set_model_|set_temp_|toggle_rotation|set_custom_prompt|reset_prompt|upload_glossary|remove_glossary|set_proxy|reset_proxy)"))
    
    print(f"✅ Бот настроен с токеном: {bot_token[:10]}...")
    print("🚀 Бот запущен! Нажмите Ctrl+C для остановки.")
    print("🎯 Отправьте боту файл для начала перевода!")
    
    # Запускаем бота
    application.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == '__main__':
    main()
